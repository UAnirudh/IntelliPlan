import flask
import sys as _sys

# ── Dev-mode fix: when this file is run via `python App.py` it loads as
# `__main__`. Any subsequent `from App import ...` (e.g. from auth_api's
# lazy helpers) would re-execute this whole file as a SECOND `App`
# module, creating duplicate `db`/`User`/Flask-app instances and
# producing "current Flask app is not registered with this 'SQLAlchemy'
# instance" errors when blueprints touch the DB. Aliasing __main__ as
# `App` makes both import paths resolve to the same module object.
if __name__ == "__main__":
    _sys.modules.setdefault("App", _sys.modules[__name__])
from flask import render_template, request, redirect, session, url_for
import requests
import os
from dotenv import load_dotenv
from datetime import datetime, timezone, timedelta
import time
from time_utils import utcnow
import desktop_auth
from studentvue_helper import (
    test_login,
    get_assignments as get_sv_assignments,
    get_missing_assignments,
    normalize_district_url,
    _compute_priority as compute_priority,
)
from ai_provider import ai_available, chat as ai_chat, vision as ai_vision, transcribe_audio, chat_json as ai_chat_json
import re
import html as _html_mod
import unicodedata as _unicodedata
import json
import uuid
import base64
import hashlib
import io
import functools
import random
import sentry_sdk
from sentry_sdk.integrations.flask import FlaskIntegration
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from auth_api import auth_bp, verify_token
from chatbot_api import chatbot_bp
from plani_agent import plani_agent_bp
from werkzeug.utils import secure_filename
import secrets as secrets_module
import urllib.parse
from flask import jsonify, send_from_directory
from datetime import datetime, timedelta, date
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import inspect, text
from flask_login import (
    LoginManager, UserMixin, login_user, logout_user,
    current_user
)
from flask_bcrypt import Bcrypt
from werkzeug.middleware.proxy_fix import ProxyFix
import streak_engine
import pet_engine
import analytics as app_analytics
import scheduler_engine
import scheduler_clarify

# ── FIX: Use database-backed sessions so Railway container restarts
#         don't wipe the OAuth state between redirect hops.
#         flask-session with "sqlalchemy" keeps state in the DB itself.
from flask_session import Session

try:
    from google_calendar_helper import (
        get_auth_url, exchange_code_for_token,
        get_upcoming_events, add_schedule_to_calendar, find_free_slots,
        compute_free_hours, merge_token_data, has_calendar_scope
    )
    GCAL_AVAILABLE = True
except Exception as e:
    print(f"Google Calendar not available: {e}")
    GCAL_AVAILABLE = False

try:
    from notion_helper import (
        test_notion_token, test_notion_token_detail, get_notion_databases,
        get_shared_pages, create_intelliplan_database,
        get_notion_tasks, create_notion_task,
        update_notion_task, complete_notion_task,
        get_notion_auth_url, exchange_notion_code,
        get_upcoming_notion_tasks, add_schedule_to_notion,
    )
    NOTION_AVAILABLE = True
except Exception as e:
    print(f"Notion not available: {e}")
    NOTION_AVAILABLE = False

try:
    from canvas_oauth import (
        get_canvas_auth_url, exchange_canvas_code,
        refresh_canvas_token, revoke_canvas_token,
        oauth_is_configured as canvas_oauth_configured,
        DEFAULT_CANVAS_BASE,
    )
    CANVAS_OAUTH_AVAILABLE = True
except Exception as e:
    print(f"Canvas OAuth not available: {e}")
    CANVAS_OAUTH_AVAILABLE = False

if os.getenv("SENTRY_DSN"):
    sentry_sdk.init(
        dsn=os.getenv("SENTRY_DSN"),
        integrations=[FlaskIntegration()],
        traces_sample_rate=0.1
    )

limiter = Limiter(key_func=get_remote_address)

load_dotenv()

app = flask.Flask(
    __name__,
    template_folder="Main_Project/templates",
)

app.secret_key = os.getenv("SECRET_KEY", "intelliplan-dev-key")
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_port=1)

# ── Response compression ──────────────────────────────────────────────
#   Every page ships ~200KB of HTML (the shared design system is inlined
#   into base.html), so without compression a first paint on a real
#   network takes ~2s. gzip/br cuts text responses ~85-90% (200KB → ~30KB),
#   which is the single biggest win for perceived load time. Defensive
#   import so the app still boots if the package isn't installed yet.
try:
    from flask_compress import Compress

    # Only compress payloads big enough to beat the CPU cost; the shared
    # CSS/JS/HTML and JSON API responses all clear this easily.
    app.config.setdefault("COMPRESS_MIN_SIZE", 1024)
    app.config.setdefault("COMPRESS_LEVEL", 6)
    app.config.setdefault("COMPRESS_MIMETYPES", [
        "text/html", "text/css", "text/xml", "text/plain",
        "application/json", "application/javascript", "application/xml",
        "application/rss+xml", "image/svg+xml",
    ])
    Compress(app)
except Exception as _compress_err:  # pragma: no cover - optional dependency
    print(f"[startup] response compression disabled: {_compress_err}")

APP_BASE_URL = os.getenv("APP_BASE_URL", "https://intelliplan.tech").rstrip("/")
APP_DOMAIN = APP_BASE_URL.replace("https://", "").replace("http://", "").split("/", 1)[0]
LEGACY_ALLOWED_ORIGINS = [
    origin.strip().rstrip("/")
    for origin in os.getenv("LEGACY_ALLOWED_ORIGINS", "").split(",")
    if origin.strip()
]
ALLOWED_WEB_ORIGINS = [
    APP_BASE_URL,
    f"https://www.{APP_DOMAIN}" if not APP_DOMAIN.startswith("www.") else APP_BASE_URL,
    *LEGACY_ALLOWED_ORIGINS,
]

# ── FIX: Switch SESSION_TYPE from "filesystem" to "sqlalchemy".
#   Filesystem sessions are stored in /tmp on Railway — ephemeral containers
#   can spin up a NEW instance to serve the OAuth callback, which has an empty
#   /tmp and therefore loses oauth_state, causing the IPE-XXXXXXXX 500 error.
#   Storing sessions in the same Postgres/SQLite DB as the app makes them
#   survive across instances and restarts.
#   We configure this BEFORE db = SQLAlchemy(app) because flask-session
#   needs the app config ready, and we wire it up after db is created below.
app.config["SESSION_TYPE"] = "sqlalchemy"
app.config["SESSION_PERMANENT"] = True
# ── Session cookie SameSite ───────────────────────────────────────
# Default is Lax, NOT None, and the difference is why sign-in was failing
# on phones.
#
# SameSite=None marks a cookie as usable in a cross-site context. Safari's
# Intelligent Tracking Prevention — on by default on every iPhone — treats
# such cookies as tracking cookies and will refuse or evict them. The
# symptom is exactly what was reported: the password is accepted,
# login_user() runs, the redirect fires, and the very next request arrives
# with no session, so @login_required bounces the student back to the login
# page. It looks like the password was wrong. It wasn't; the cookie was
# never kept.
#
# None was chosen originally because IntelliPlan is embedded in an iframe by
# Lotus, and a cookie in a cross-site frame genuinely does need
# SameSite=None. Both requirements are real, so this is resolved per
# response rather than per deployment: Lax by default, upgraded to None
# only for requests that actually arrive inside a frame. See
# _samesite_for_embeds() below.
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_HTTPONLY"] = True
# Secure cookies are HTTPS-only. Over plain http (local development) the
# browser silently discards them, which presents as "login does nothing".
app.config["SESSION_COOKIE_SECURE"] = APP_BASE_URL.startswith("https://")
app.config["REMEMBER_COOKIE_SECURE"] = APP_BASE_URL.startswith("https://")
app.config["REMEMBER_COOKIE_HTTPONLY"] = True
app.config["REMEMBER_COOKIE_SAMESITE"] = "Lax"
app.config["PREFERRED_URL_SCHEME"] = "https" if APP_BASE_URL.startswith("https://") else "http"
app.permanent_session_lifetime = timedelta(days=7)

app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv("DATABASE_URL", "sqlite:///intelliplan.db")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024
app.config["NOTES_UPLOAD_FOLDER"] = os.path.join(app.root_path, "uploads", "course_notes")
os.makedirs(app.config["NOTES_UPLOAD_FOLDER"], exist_ok=True)

db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = "login"

# ── FIX: Point flask-session at the SQLAlchemy db so sessions are durable.
app.config["SESSION_SQLALCHEMY"] = db
Session(app)

@app.teardown_appcontext
def shutdown_session(exception=None):
    if exception:
        db.session.rollback()
    db.session.remove()

app.register_blueprint(auth_bp)
app.register_blueprint(chatbot_bp)
app.register_blueprint(plani_agent_bp)
from extra_features import extras_bp
app.register_blueprint(extras_bp)

#: Cookies whose SameSite attribute we manage. Anything else Flask or an
#: extension sets is left exactly as it was.
_SAMESITE_MANAGED_COOKIES = ("session", "remember_token")


def _request_is_framed() -> bool:
    """Whether this request is being made from inside an iframe.

    `Sec-Fetch-Dest: iframe` is set by the browser itself and cannot be
    forged by page script, which makes it the right signal here — the whole
    point is to relax a cookie restriction, so the input has to be one the
    embedding page does not control.

    Browsers without Sec-Fetch-* (older Safari) fall through to False and
    therefore get Lax. That is the safe direction: a student signing in
    directly keeps working, and only the embed degrades on browsers that
    were already unreliable for third-party cookies.
    """
    dest = (request.headers.get("Sec-Fetch-Dest") or "").lower()
    if dest in ("iframe", "frame", "embed", "object"):
        return True
    # Explicit opt-in for the embed product's own entry points, so the
    # widget routes work even when Sec-Fetch-Dest is absent.
    return (request.args.get("embed") == "1") or bool(session.get("_embedded"))


def _upgrade_samesite(cookie: str) -> str:
    """Rewrite one Set-Cookie value to SameSite=None; Secure.

    SameSite=None without Secure is rejected outright by every current
    browser, so Secure is added rather than assumed.
    """
    if "samesite=" in cookie.lower():
        cookie = re.sub(r"SameSite\s*=\s*\w+", "SameSite=None", cookie, flags=re.IGNORECASE)
    else:
        cookie += "; SameSite=None"
    if "secure" not in cookie.lower():
        cookie += "; Secure"
    return cookie


class _EmbedSameSiteMiddleware:
    """Upgrade session cookies to SameSite=None only inside a frame.

    This is WSGI middleware and not an ``after_request`` hook, which is
    where the first attempt at this went wrong: Flask runs after_request
    handlers *before* it saves the session, so the session cookie does not
    exist yet at that point and the hook silently rewrote nothing. By the
    time the response reaches WSGI, every Set-Cookie — Flask's session,
    flask-session's, and Flask-Login's remember token — is present.

    Varying per response is the whole point. SESSION_COOKIE_SAMESITE is
    application-wide config read at serialisation time, and both needs are
    real: a student on an iPhone needs Lax so Safari keeps the cookie, and
    the Lotus iframe needs None so the cookie is sent at all.
    """

    def __init__(self, wsgi_app, managed):
        self._app = wsgi_app
        self._managed = managed

    def __call__(self, environ, start_response):
        framed = self._is_framed(environ)

        def _start(status, headers, exc_info=None):
            if framed:
                headers = [
                    (k, _upgrade_samesite(v))
                    if k.lower() == "set-cookie" and v.split("=", 1)[0].strip() in self._managed
                    else (k, v)
                    for k, v in headers
                ]
            return start_response(status, headers, exc_info)

        return self._app(environ, _start)

    @staticmethod
    def _is_framed(environ) -> bool:
        dest = (environ.get("HTTP_SEC_FETCH_DEST") or "").lower()
        if dest in ("iframe", "frame", "embed", "object"):
            return True
        return "embed=1" in (environ.get("QUERY_STRING") or "")


app.wsgi_app = _EmbedSameSiteMiddleware(app.wsgi_app, _SAMESITE_MANAGED_COOKIES)


@app.after_request
def add_cors_headers(response):
    origin = request.headers.get("Origin", "")
    is_extension = origin.startswith("chrome-extension://")
    is_local = origin.startswith("http://localhost") or origin.startswith("http://127.0.0.1")
    # HTML documents do not need CORS — only XHR/fetch endpoints. Limiting
    # CORS to /api/* removes a wildcard from every HTML response while
    # keeping the programmatic surfaces working.
    try:
        path = (request.path or "")
    except Exception:
        path = ""
    needs_cors = path.startswith("/api/")
    if needs_cors:
        if origin and (origin.rstrip("/") in ALLOWED_WEB_ORIGINS or is_extension or is_local):
            response.headers["Access-Control-Allow-Origin"] = origin
            response.headers["Vary"] = "Origin"
        else:
            response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization, X-Extension-Token"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    else:
        # Make sure no upstream layer leaks a wildcard CORS header on HTML.
        for h in ("Access-Control-Allow-Origin", "Access-Control-Allow-Headers",
                  "Access-Control-Allow-Methods", "Access-Control-Allow-Credentials"):
            response.headers.pop(h, None)
    # Sensitive auth/account pages must NEVER be embeddable — an attacker who
    # can iframe the login or settings page can clickjack credentials and
    # account changes. These pages are never part of the embed product, so
    # locking them down doesn't affect Lotus/other embedders, which only embed
    # the study tools.
    try:
        path = (request.path or "").rstrip("/")
    except Exception:
        path = ""
    # ── CSP: hardened directives, scoped frame-ancestors per surface ──
    # Auth/account pages can never be iframed (clickjacking risk). All
    # other surfaces stay iframeable for Lotus and the embeddable widgets.
    #
    # Set CSP_ENFORCE=1 in env to flip from Report-Only to enforced mode;
    # default ships Report-Only so violations are surfaced without breaking
    # the live site. After 48h of clean reports, flip the env flag.
    _common_csp = (
        "default-src 'self'; "
        # 'unsafe-inline' on script-src is required today because most
        # templates ship inline <script> blocks. Migrate to nonces in a
        # follow-up; CSP_ENFORCE stays off until then.
        "script-src 'self' 'unsafe-inline' 'unsafe-eval' "
            "https://www.googletagmanager.com https://www.google-analytics.com "
            "https://www.clarity.ms https://*.clarity.ms "
            "https://browser.sentry-cdn.com https://*.sentry.io "
            "https://js.stripe.com https://challenges.cloudflare.com "
            "https://meet.jit.si; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
        "img-src 'self' data: blob: https:; "
        "font-src 'self' data: https://fonts.gstatic.com; "
        "connect-src 'self' https://www.google-analytics.com "
            "https://*.clarity.ms https://*.sentry.io https://api.stripe.com "
            "https://*.googleapis.com https://meet.jit.si; "
        "media-src 'self' blob:; "
        "object-src 'none'; "
        "base-uri 'self'; "
        "form-action 'self'; "
    )
    if _is_frame_sensitive(path):
        response.headers["X-Frame-Options"] = "DENY"
        csp_value = _common_csp + "frame-src 'none'; frame-ancestors 'none'"
    else:
        # Embeddable surface — Lotus + our own origin + local dev.
        response.headers.pop("X-Frame-Options", None)
        csp_value = _common_csp + (
            "frame-src https://meet.jit.si https://challenges.cloudflare.com; "
            "frame-ancestors 'self' https://lotus-72e3e.web.app "
            "https://intelliplan.tech http://localhost:5000"
        )
    csp_header_name = "Content-Security-Policy" if os.getenv("CSP_ENFORCE") == "1" else "Content-Security-Policy-Report-Only"
    response.headers[csp_header_name] = csp_value
    response.headers["X-Permitted-Cross-Domain-Policies"] = "none"
    return response

# Paths that must never be iframed (clickjacking protection). Matched as a
# prefix against the request path so sub-routes (e.g. /login/account) are
# covered too.
_FRAME_SENSITIVE_PREFIXES = (
    "/login", "/register", "/settings", "/connect", "/account",
)

def _is_frame_sensitive(path):
    p = (path or "").rstrip("/")
    return any(p == pre or p.startswith(pre + "/") for pre in _FRAME_SENSITIVE_PREFIXES)

@app.after_request
def add_static_cache_headers(response):
    """Cache static assets (icons, images, fonts, css/js) for repeat visits and
    better Core Web Vitals. Never caches the service worker or manifest, which
    must revalidate so updates ship.

    Flask's static handler sets ``Cache-Control: no-cache`` by default, so this
    hook must OVERRIDE that value rather than only fill it in when absent —
    otherwise every asset revalidates on every navigation (a round-trip per
    icon/css/js on each page). Media (icons/logos/images/fonts) get a long
    7-day cache; CSS/JS aren't content-hashed, so they get a shorter TTL with
    stale-while-revalidate so a deploy's new styles propagate within a day
    while repeat views still serve instantly from cache."""
    try:
        p = request.path or ""
        no_cache_assets = ("/static/sw.js", "/static/manifest.json")
        if (p.startswith("/static/")
                and p not in no_cache_assets
                and response.status_code in (200, 304)):
            if p.endswith((".css", ".js")):
                response.headers["Cache-Control"] = "public, max-age=86400, stale-while-revalidate=604800"
            else:
                response.headers["Cache-Control"] = "public, max-age=604800, immutable"
    except Exception:
        pass
    return response

def _default_limit_exempt():
    """Exempt safe, idempotent GET/HEAD/OPTIONS requests from the global
    default limits. Public pages (landing, /login, /register, legal, blog,
    etc.) are read-only and shouldn't be throttled by the per-IP default —
    that wrongly blocks shared-NAT users and CI runners. Sensitive
    state-changing endpoints keep their explicit @limiter.limit decorators,
    which still apply regardless of this exemption."""
    return request.method in ("GET", "HEAD", "OPTIONS")

limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    default_limits_exempt_when=_default_limit_exempt,
)

# ── MODELS ────────────────────────────────────────────────────
class User(UserMixin, db.Model):
    __tablename__ = "users"
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(255), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False, default="")
    google_id = db.Column(db.String(255), unique=True, nullable=True)
    name = db.Column(db.String(255), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    # ── Referral columns ──
    referral_code = db.Column(db.String(16), unique=True, nullable=True)
    referred_by_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    # ── Phone + reminder opt-in ──
    phone = db.Column(db.String(32), nullable=True)            # E.164 (+15551234567) or blank
    sms_reminders_opt_in = db.Column(db.Boolean, default=False)
    push_reminders_opt_in = db.Column(db.Boolean, default=False)
    reminder_lead_minutes = db.Column(db.Integer, default=60)  # how far ahead to remind
    sms_carrier = db.Column(db.String(32), default="tmobile")  # SMS-over-email gateway key
    # ── Notification preferences (see intelliplan/notifications/) ──
    email_reminders_opt_in = db.Column(db.Boolean, default=False)
    # Minutes east of UTC. Every timestamp here is naive UTC, so this is the
    # only way to reach the student's wall clock — without it, "don't text me
    # after 10pm" silently means 10pm wherever the server happens to live.
    utc_offset_minutes = db.Column(db.Integer, default=0)
    quiet_hours_enabled = db.Column(db.Boolean, default=True)
    quiet_hours_start = db.Column(db.Integer, default=22)
    quiet_hours_end = db.Column(db.Integer, default=7)
    # Comma-separated EventKind values; empty means "the defaults".
    notification_kinds = db.Column(db.String(512), nullable=True)
    # ── COPPA: under-13 gating ──
    birth_year = db.Column(db.Integer, nullable=True)          # collected at signup
    parent_email = db.Column(db.String(255), nullable=True)    # for under-13 accounts
    parent_consent_granted = db.Column(db.Boolean, default=False)
    parent_consent_token = db.Column(db.String(64), nullable=True)  # signed verification token
    # JSON: {"grade_source":"active|canvas|...", "assignment_sources":["active","google_classroom",...]}
    lms_preferences = db.Column(db.Text, default="{}")
    # ── AI personalization opt-in. When True, IntelliPlan injects the
    # student's grades, performance patterns, and history into AI prompts
    # for the scheduler, tutor, and other AI features. Default OFF to
    # respect privacy — the toggle lives in Settings → Privacy.
    ai_personalization_opt_in = db.Column(db.Boolean, default=False)
    # ── Marketing email: onboarding sequence and newsletters, sent from
    # an external tool rather than by this app.
    #
    # Separate from email_reminders_opt_in on purpose. That one is
    # transactional — the deadline reminders the student signed up for.
    # This is marketing, and conflating the two means an unsubscribe from
    # a newsletter silently kills the reminders the planner exists to
    # send, or worse, a reminders opt-in is read as permission to market.
    #
    # Default OFF, and it stays off unless the student ticked the box on
    # a form. A pre-ticked box is not consent under GDPR, and CAN-SPAM
    # makes the sender responsible either way.
    marketing_emails_opt_in = db.Column(db.Boolean, default=False)
    # When they agreed. Consent you cannot date is consent you cannot
    # evidence, and "when did this person opt in" is the first question
    # asked in any complaint.
    marketing_opt_in_at = db.Column(db.DateTime, nullable=True)
    # ── Role: student | teacher | parent. Drives /teacher and /parent
    # dashboards plus the StudentLink consent flow.
    role = db.Column(db.String(16), default="student")

    # ── Focus enforcement (Active study) ──────────────────────────────
    #: What happens when the camera check-in decides the student has
    #: genuinely drifted: "off" (a dismissible nudge, the old behaviour),
    #: "alarm", "takeover", or "stakes". Chosen during onboarding and
    #: changeable in Settings. Default is off — an app that starts blaring
    #: at someone who never asked it to is an app they uninstall.
    focus_enforcement = db.Column(db.String(16), default="off")
    #: Path under /uploads/focus_alarms for a sound the student uploaded,
    #: or empty for the built-in tone.
    focus_alarm_file = db.Column(db.String(255), nullable=True)
    #: Seconds of continuous distraction before enforcement fires. A short
    #: grace period is the difference between "you looked away" and "you
    #: left" — without it the alarm goes off every time someone reaches for
    #: a textbook.
    focus_grace_seconds = db.Column(db.Integer, default=25)
    linked_accounts = db.relationship("LinkedAccount", backref="user", lazy=True, cascade="all, delete-orphan")
    dismissed = db.relationship("DismissedAssignment", backref="user", lazy=True, cascade="all, delete-orphan")
    descriptions = db.relationship("CustomDescription", backref="user", lazy=True, cascade="all, delete-orphan")

class UserIdentity(db.Model):
    """Student identity profile collected at onboarding and editable in settings.

    Drives chatbot/tutor personalization — grade level, academic focus areas, and
    goals/priorities are injected into the Plani/Tutor system prompt so replies
    match the student's curriculum and ambitions.
    """
    __tablename__ = "user_identities"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), unique=True, nullable=False)
    grade_level = db.Column(db.String(32), nullable=True)        # e.g. "11th grade"
    focus_areas = db.Column(db.Text, default="[]")                # JSON list of subjects
    goals = db.Column(db.Text, default="")                        # free-text priorities
    completed = db.Column(db.Boolean, default=False)              # questionnaire finished
    availability = db.Column(db.Text, default="{}")               # JSON: day -> time range
    weekly_commitments = db.Column(db.Text, default="")           # free-text extracurriculars
    class_schedule = db.Column(db.Text, default="[]")             # JSON list of class slots
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    def focus_list(self):
        try:
            value = json.loads(self.focus_areas or "[]")
            return value if isinstance(value, list) else []
        except (TypeError, json.JSONDecodeError):
            return []

    def avail_dict(self):
        try:
            v = json.loads(self.availability or "{}")
            return v if isinstance(v, dict) else {}
        except Exception:
            return {}

    def class_list(self):
        try:
            v = json.loads(self.class_schedule or "[]")
            return v if isinstance(v, list) else []
        except Exception:
            return []

    def to_dict(self):
        return {
            "grade_level": self.grade_level or "",
            "focus_areas": self.focus_list(),
            "goals": self.goals or "",
            "completed": bool(self.completed),
            "availability": self.avail_dict(),
            "weekly_commitments": self.weekly_commitments or "",
            "class_schedule": self.class_list(),
        }


class LinkedAccount(db.Model):
    __tablename__ = "linked_accounts"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    profile_id = db.Column(db.String(16), unique=True, default=lambda: str(uuid.uuid4())[:8])
    name = db.Column(db.String(255), default="My Account")
    login_type = db.Column(db.String(32), nullable=False)
    credentials = db.Column(db.Text, nullable=False)
    is_active = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def get_credentials(self):
        return json.loads(self.credentials)

    def set_credentials(self, creds_dict):
        self.credentials = json.dumps(creds_dict)

class DismissedAssignment(db.Model):
    __tablename__ = "dismissed_assignments"
    id = db.Column(db.Integer, primary_key=True)
    # Indexed: looked up by owner on every dashboard / command-center load.
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True, index=True)
    guest_session_id = db.Column(db.String(64), nullable=True, index=True)
    title = db.Column(db.String(512), nullable=False)
    data = db.Column(db.Text, default="{}")

class TestMark(db.Model):
    __tablename__ = "test_marks"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True, index=True)
    guest_session_id = db.Column(db.String(64), nullable=True, index=True)
    title = db.Column(db.String(512), nullable=False)
    data = db.Column(db.Text, default="{}")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class CustomDescription(db.Model):
    __tablename__ = "custom_descriptions"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    assignment_title = db.Column(db.String(512), nullable=False)
    description = db.Column(db.Text, nullable=False)

class GoogleIntegration(db.Model):
    __tablename__ = "google_integrations"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    token_data = db.Column(db.Text, nullable=False)
    # Multi-account support — each row is one connected Google Account.
    # `is_active=True` marks the account currently used for calendar sync.
    account_email = db.Column(db.String(255), nullable=True)
    account_name = db.Column(db.String(255), nullable=True)
    is_active = db.Column(db.Boolean, default=True)
    connected_at = db.Column(db.DateTime, default=datetime.utcnow)

class DesktopAuthCode(db.Model):
    """A one-time ticket letting the desktop app claim a browser sign-in.

    Minted at the end of a Google callback that the desktop client started,
    spent once by /api/desktop/auth/exchange, and worthless after two
    minutes. See desktop_auth.py for why each column is here — in short,
    only the hash of the code is kept, and redeeming it also requires the
    PKCE verifier that never leaves the app.
    """
    __tablename__ = "desktop_auth_codes"
    id = db.Column(db.Integer, primary_key=True)
    # Unique so a repeat insert collides rather than quietly creating a
    # second live ticket for the same code.
    code_hash = db.Column(db.String(64), nullable=False, unique=True, index=True)
    code_challenge = db.Column(db.String(64), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)
    # Set the moment it is spent. Presence here, not deletion, is what
    # makes a replay fail: the row has to stick around long enough to say
    # "already used" rather than "never existed".
    used_at = db.Column(db.DateTime, nullable=True)


class NotionIntegration(db.Model):
    __tablename__ = "notion_integrations"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    token = db.Column(db.String(512), nullable=False)
    database_id = db.Column(db.String(256), nullable=True)
    # New OAuth metadata — populated by /oauth/notion/callback. Older
    # rows from the manual-token flow leave these as NULL.
    auth_type = db.Column(db.String(16), default="manual")  # "oauth" | "manual"
    workspace_id = db.Column(db.String(64), nullable=True)
    workspace_name = db.Column(db.String(256), nullable=True)
    workspace_icon = db.Column(db.String(512), nullable=True)
    bot_id = db.Column(db.String(64), nullable=True)
    connected_at = db.Column(db.DateTime, default=datetime.utcnow)


class CanvasIntegration(db.Model):
    """Canvas LMS OAuth tokens. Separate from the legacy token-paste flow,
    which still writes to LinkedAccount.credentials for backwards compat."""
    __tablename__ = "canvas_integrations"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    canvas_base = db.Column(db.String(256), nullable=False)
    access_token = db.Column(db.String(2048), nullable=False)
    refresh_token = db.Column(db.String(2048), nullable=True)
    token_expires_at = db.Column(db.DateTime, nullable=True)
    canvas_user_id = db.Column(db.String(64), nullable=True)
    canvas_user_name = db.Column(db.String(256), nullable=True)
    connected_at = db.Column(db.DateTime, default=datetime.utcnow)

class ClassroomIntegration(db.Model):
    """Google Classroom OAuth tokens. Separate from GoogleIntegration (which is
    for Calendar) because Classroom uses its own client credentials and scope
    set, and a user may connect Classroom independently of Calendar."""
    __tablename__ = "classroom_integrations"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    access_token = db.Column(db.String(2048), nullable=False)
    refresh_token = db.Column(db.String(2048), nullable=True)
    token_expires_at = db.Column(db.DateTime, nullable=True)
    account_email = db.Column(db.String(255), nullable=True)
    account_name = db.Column(db.String(255), nullable=True)
    connected_at = db.Column(db.DateTime, default=datetime.utcnow)

class BlackboardIntegration(db.Model):
    """Blackboard Learn OAuth tokens. Blackboard is per-institution: each row
    stores the institution's Learn URL (e.g. https://learn.school.edu) along
    with the access/refresh tokens and user identity from that institution."""
    __tablename__ = "blackboard_integrations"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    institution_url = db.Column(db.String(512), nullable=False)
    access_token = db.Column(db.String(2048), nullable=False)
    refresh_token = db.Column(db.String(2048), nullable=True)
    token_expires_at = db.Column(db.DateTime, nullable=True)
    bb_user_id = db.Column(db.String(64), nullable=True)
    bb_username = db.Column(db.String(255), nullable=True)
    connected_at = db.Column(db.DateTime, default=datetime.utcnow)

class MoodleIntegration(db.Model):
    """Moodle web-services token. No OAuth: each Moodle instance is self-hosted
    and authenticates via a per-user web-service token the user generates in
    their Moodle preferences. We store the institution's Moodle URL plus that
    token, plus the user identity returned by core_webservice_get_site_info."""
    __tablename__ = "moodle_integrations"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    moodle_url = db.Column(db.String(512), nullable=False)
    ws_token = db.Column(db.String(512), nullable=False)
    moodle_user_id = db.Column(db.String(64), nullable=True)
    moodle_username = db.Column(db.String(255), nullable=True)
    moodle_fullname = db.Column(db.String(255), nullable=True)
    connected_at = db.Column(db.DateTime, default=datetime.utcnow)

class ManualTask(db.Model):
    __tablename__ = "manual_tasks"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    title = db.Column(db.String(512), nullable=False)
    due_date = db.Column(db.String(32), default="")
    priority = db.Column(db.String(16), default="Medium")
    course = db.Column(db.String(256), default="Personal")
    estimated_time = db.Column(db.Integer, default=60)
    notes = db.Column(db.Text, default="")
    done = db.Column(db.Boolean, default=False)
    notion_page_id = db.Column(db.String(256), nullable=True)
    # Provenance tags — populated when this task was created by the CSV
    # importer, smart-paste, or the unsupported-LMS extension scraper so we
    # can refresh-replace cleanly instead of duplicating on every sync.
    import_source = db.Column(db.String(32), default="")   # "csv"|"paste"|"scraper:<lms>"|""
    import_batch_id = db.Column(db.String(64), default="") # uuid for the import session
    external_id = db.Column(db.String(128), default="")    # LMS-side id when scraper provides one
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class ImportedGrade(db.Model):
    """Course grades imported from CSV, smart-paste, or extension scraper.

    Exists so users on unsupported LMSes can still see their grades on the
    /grades page and have them feed into AI personalization. Refreshed
    in-place by the extension auto-sync via import_batch_id matching.
    """
    __tablename__ = "imported_grades"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    course = db.Column(db.String(256), nullable=False)
    percentage = db.Column(db.Float, nullable=True)
    letter = db.Column(db.String(4), default="")
    teacher = db.Column(db.String(256), default="")
    period = db.Column(db.String(64), default="")
    source = db.Column(db.String(32), default="csv")       # "csv"|"paste"|"scraper:<lms>"
    source_label = db.Column(db.String(64), default="")    # human-friendly e.g. "Aeries (Riverside)"
    last_synced = db.Column(db.DateTime, default=datetime.utcnow)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class SavedSchedule(db.Model):
    __tablename__ = "saved_schedules"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    name = db.Column(db.String(256), default="My Schedule")
    schedule_data = db.Column(db.Text, nullable=False)
    # Interactive View progress ({block_id: {done, checked, ...}}), synced
    # from the client so checked-off blocks follow the student across devices.
    progress_json = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    is_active = db.Column(db.Boolean, default=True)

class SchedulerPreset(db.Model):
    """A saved answer to a clarifying question, keyed by normalized task title.

    Students re-add the same vague task constantly ("Study", every week), so
    the answers they gave last time are worth offering back instead of asking
    again. One row per (owner, task key).
    """
    __tablename__ = "scheduler_presets"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True, index=True)
    guest_session_id = db.Column(db.String(64), nullable=True, index=True)
    task_key = db.Column(db.String(96), nullable=False)      # scheduler_clarify.preset_key()
    label = db.Column(db.String(200), default="")            # title as the student typed it
    answers_json = db.Column(db.Text, nullable=False, default="{}")
    times_used = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_used_at = db.Column(db.DateTime, default=datetime.utcnow)

    def answers(self):
        try:
            v = json.loads(self.answers_json or "{}")
            return v if isinstance(v, dict) else {}
        except (TypeError, ValueError):
            return {}

    def to_dict(self):
        return {
            "task_key": self.task_key, "label": self.label,
            "answers": self.answers(), "times_used": self.times_used or 0,
            "last_used_at": self.last_used_at.isoformat() if self.last_used_at else None,
        }


class ManualPlanPreset(db.Model):
    """A hand-built day, saved to be laid down again.

    Distinct from SchedulerPreset above, which remembers the *answers* a
    student gave about one vague task. This remembers a whole shape of a
    day — "Weekday evening", "Saturday catch-up" — as a list of blocks
    with clock times and no dates. Applying it stamps those times onto a
    date the student picks.

    Dateless on purpose: a preset is a routine, and a routine that
    remembered it was authored on the 14th would be useless on the 15th.
    """
    __tablename__ = "manual_plan_presets"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True, index=True)
    guest_session_id = db.Column(db.String(64), nullable=True, index=True)
    name = db.Column(db.String(120), nullable=False, default="My day")
    blocks_json = db.Column(db.Text, nullable=False, default="[]")
    times_used = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_used_at = db.Column(db.DateTime, default=datetime.utcnow)

    def blocks(self):
        try:
            v = json.loads(self.blocks_json or "[]")
            return v if isinstance(v, list) else []
        except (TypeError, ValueError):
            return []

    def to_dict(self):
        blocks = self.blocks()
        return {
            "id": self.id,
            "name": self.name,
            "blocks": blocks,
            "block_count": len(blocks),
            "total_minutes": sum(int(b.get("duration_minutes") or 0) for b in blocks),
            "times_used": self.times_used or 0,
            "created_at": self.created_at.isoformat() if self.created_at else None,
            "last_used_at": self.last_used_at.isoformat() if self.last_used_at else None,
        }


class DayArchive(db.Model):
    """Day-by-day snapshots — schedules, resources, notes, and anything else."""
    __tablename__ = "day_archives"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    archive_date = db.Column(db.Date, nullable=False, index=True)
    item_type = db.Column(db.String(64), nullable=False)
    title = db.Column(db.String(256), default="")
    payload = db.Column(db.Text, nullable=False)
    meta_json = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Task:
    def __init__(self, name, deadline, duration, priority_weight=1, difficulty=1):
        self.name = name
        self.deadline = deadline
        self.duration = duration
        self.priority_weight = priority_weight
        self.difficulty = difficulty

class TaskFeedback(db.Model):
    __tablename__ = "task_feedback"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    title = db.Column(db.String(512), nullable=False)
    course = db.Column(db.String(256), default="")
    estimated_time = db.Column(db.Integer, default=60)
    actual_time = db.Column(db.Integer, nullable=True)
    difficulty = db.Column(db.String(16), default="Medium")
    priority = db.Column(db.String(16), default="Medium")
    completed_at = db.Column(db.DateTime, default=datetime.utcnow)
    day_of_week = db.Column(db.String(16), default="")
    time_of_day = db.Column(db.String(16), default="")

class PushSubscription(db.Model):
    """One row per *browser*, not per student.

    A student with a phone and a laptop has two subscriptions, and both have
    to survive: the endpoint is what a push service addresses, and it is
    unique per browser install. Keying these rows on user_id alone meant
    enabling notifications on the second device overwrote the first, so the
    phone went quiet the moment the laptop subscribed — with nothing in the
    UI to suggest it had happened.

    ``endpoint`` is denormalised out of ``subscription_json`` so the upsert
    can match on it without parsing every row.
    """

    __tablename__ = "push_subscriptions"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    #: Nullable for rows written before this column existed; backfilled by
    #: _migrate_push_subscription_endpoints() at boot.
    endpoint = db.Column(db.String(512), nullable=True, index=True)
    subscription_json = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class ReminderSent(db.Model):
    """Dedupe row — one per (user, task, channel) so a reminder is sent
    at most once per task, regardless of how often the cron runs."""
    __tablename__ = "reminders_sent"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    task_key = db.Column(db.String(256), nullable=False)  # "manual:<id>" or "canvas:<id>"
    channel = db.Column(db.String(16), nullable=False)    # "sms" or "push"
    sent_at = db.Column(db.DateTime, default=datetime.utcnow)

class CourseNote(db.Model):
    __tablename__ = "course_notes"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    course_name = db.Column(db.String(255), nullable=False)
    course_id = db.Column(db.String(128), nullable=True)
    course_source = db.Column(db.String(32), nullable=True)
    note_date = db.Column(db.String(32), nullable=False)
    title = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=True)
    stored_filename = db.Column(db.String(255), nullable=True)
    text_content = db.Column(db.Text, default="")
    summary_cache = db.Column(db.Text, default="")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class ExtensionToken(db.Model):
    __tablename__ = "extension_tokens"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    token = db.Column(db.String(64), unique=True, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class ApiKey(db.Model):
    """A credential for the public REST API, plus the application it came from.

    One row covers the whole lifecycle: the developer applies (status
    'pending'), a human or the auto-approver moves it to 'active' and the
    secret is minted, and the developer or an admin can 'revoked' it. The
    application answers stay on the row because they are what an admin
    reviews, and what we go back to when a key starts behaving oddly.

    The secret itself is never stored — only its SHA-256 and the leading
    `key_prefix`, which is what the dashboard shows so a developer can tell
    two keys apart without us being able to reconstruct either one.
    """
    __tablename__ = "api_keys"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False, index=True)

    # ── The application ──
    app_name = db.Column(db.String(120), nullable=False)
    app_url = db.Column(db.String(512), default="")
    use_case = db.Column(db.Text, default="")            # what they're building
    expected_volume = db.Column(db.String(32), default="low")   # low | medium | high
    contact_email = db.Column(db.String(255), default="")
    accepted_terms_at = db.Column(db.DateTime, nullable=True)

    # ── The credential ──
    scopes = db.Column(db.Text, default="")              # space-separated
    key_prefix = db.Column(db.String(24), default="", index=True)
    key_hash = db.Column(db.String(64), default="", index=True)

    # ── Lifecycle ──
    status = db.Column(db.String(16), default="pending", index=True)  # pending|active|revoked|denied
    review_note = db.Column(db.String(500), default="")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    approved_at = db.Column(db.DateTime, nullable=True)
    revoked_at = db.Column(db.DateTime, nullable=True)
    last_used_at = db.Column(db.DateTime, nullable=True)
    request_count = db.Column(db.Integer, default=0)
    rate_limit_per_min = db.Column(db.Integer, default=60)

    def scope_list(self):
        return [s for s in (self.scopes or "").split() if s]

    def to_dict(self):
        return {
            "id": self.id,
            "app_name": self.app_name,
            "app_url": self.app_url or "",
            "use_case": self.use_case or "",
            "expected_volume": self.expected_volume or "low",
            "scopes": self.scope_list(),
            "key_prefix": self.key_prefix or "",
            "status": self.status,
            "review_note": self.review_note or "",
            "rate_limit_per_min": self.rate_limit_per_min or 60,
            "request_count": self.request_count or 0,
            "created_at": self.created_at.isoformat() if self.created_at else None,
            "approved_at": self.approved_at.isoformat() if self.approved_at else None,
            "revoked_at": self.revoked_at.isoformat() if self.revoked_at else None,
            "last_used_at": self.last_used_at.isoformat() if self.last_used_at else None,
        }


class Lesson(db.Model):
    """Uploaded lesson recording (audio or video) + AI-generated summary."""
    __tablename__ = "lessons"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    title = db.Column(db.String(255), nullable=False)
    course = db.Column(db.String(128), default="")
    tags = db.Column(db.Text, default="[]")            # JSON list of tag strings
    media_kind = db.Column(db.String(16), default="video")  # video | audio
    original_filename = db.Column(db.String(255), default="")
    stored_filename = db.Column(db.String(255), default="")
    mime_type = db.Column(db.String(64), default="")
    duration_seconds = db.Column(db.Integer, default=0)
    transcript = db.Column(db.Text, default="")
    summary = db.Column(db.Text, default="")
    summary_status = db.Column(db.String(16), default="pending")  # pending | ready | failed
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def tag_list(self):
        try:
            v = json.loads(self.tags or "[]")
            return v if isinstance(v, list) else []
        except Exception:
            return []

    def to_dict(self):
        return {
            "id": self.id,
            "title": self.title,
            "course": self.course or "",
            "tags": self.tag_list(),
            "media_kind": self.media_kind,
            "stream_url": f"/lessons/{self.id}/stream",
            "original_filename": self.original_filename,
            "duration_seconds": int(self.duration_seconds or 0),
            "summary": self.summary or "",
            "summary_status": self.summary_status or "pending",
            "created_at": self.created_at.isoformat() if self.created_at else None,
        }


class StudyGroup(db.Model):
    """A small study group: shared notes, schedule, video room link."""
    __tablename__ = "study_groups"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), nullable=False)
    topic = db.Column(db.String(120), default="")           # e.g. "AP Calc BC", "SAT Math"
    level = db.Column(db.String(32), default="any")         # beginner | intermediate | advanced | any
    style = db.Column(db.String(32), default="any")         # focused | discussion | quizzing | flashcards | any
    visibility = db.Column(db.String(16), default="public") # public | private
    description = db.Column(db.Text, default="")
    shared_notes = db.Column(db.Text, default="")
    meeting_url = db.Column(db.String(512), default="")     # generated Jitsi room
    next_meeting_at = db.Column(db.DateTime, nullable=True)
    next_meeting_topic = db.Column(db.String(255), default="")
    suggested_plan = db.Column(db.Text, default="")          # AI-generated study plan
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class StudyGroupMember(db.Model):
    __tablename__ = "study_group_members"
    id = db.Column(db.Integer, primary_key=True)
    group_id = db.Column(db.Integer, db.ForeignKey("study_groups.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    role = db.Column(db.String(16), default="member")  # owner | member
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)


class StudyGroupTask(db.Model):
    """Collaborative task inside a study group. Any member can create,
    claim, complete, or delete (creator-only) a task."""
    __tablename__ = "study_group_tasks"
    id = db.Column(db.Integer, primary_key=True)
    group_id = db.Column(db.Integer, db.ForeignKey("study_groups.id"), nullable=False, index=True)
    title = db.Column(db.String(200), nullable=False)
    notes = db.Column(db.Text, nullable=True)
    due_date = db.Column(db.DateTime, nullable=True)
    done = db.Column(db.Boolean, default=False)
    claimed_by_user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    created_by_user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    completed_at = db.Column(db.DateTime, nullable=True)


class VoiceSeat(db.Model):
    """One member's presence in a study group's voice room.

    A row exists only while someone is in voice. It is presence, not
    membership — leaving voice deletes the row, and staying in it is what
    `last_seen_at` proves.

    That column is the whole design. A browser tab that crashes, loses
    Wi-Fi, or is closed by swiping the app away never sends a "leave", so
    presence tracked purely by join/leave drifts into a room that claims
    six people and has one. Instead the client says "still here" on a
    heartbeat and anything that has gone quiet for longer than
    VOICE_STALE_SECONDS is simply not in the room any more. No cleanup
    job has to run for that to be true.
    """
    __tablename__ = "voice_seats"
    id = db.Column(db.Integer, primary_key=True)
    group_id = db.Column(db.Integer, db.ForeignKey("study_groups.id"), nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False, index=True)
    display_name = db.Column(db.String(120), default="")
    is_muted = db.Column(db.Boolean, default=True)      # everyone joins muted
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_seen_at = db.Column(db.DateTime, default=datetime.utcnow, index=True)

    __table_args__ = (
        # One seat per person per room. Without this a double-click on
        # "Join" seats the same student twice and the roster shows a ghost.
        db.UniqueConstraint("group_id", "user_id", name="uq_voice_seat_member"),
    )


class LMSToken(db.Model):
    """OAuth / API tokens for a user's connection to one LMS provider.

    `provider` is the registry key (google_classroom, blackboard, moodle,
    powerschool). `tokens_json` stores access/refresh tokens and any
    provider-specific metadata (base_url for Moodle, etc.).
    """
    __tablename__ = "lms_tokens"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False, index=True)
    provider = db.Column(db.String(32), nullable=False, index=True)
    tokens_json = db.Column(db.Text, default="{}")
    last_synced_at = db.Column(db.DateTime, nullable=True)
    last_sync_count = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class StudentLink(db.Model):
    """Teacher/parent ↔ student consent link.

    Students must accept a link before any data is exposed to the
    linker. The ``relationship`` column doubles as audit log of WHO
    requested access (teacher vs. parent)."""
    __tablename__ = "student_links"
    id = db.Column(db.Integer, primary_key=True)
    linker_user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False, index=True)
    student_user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False, index=True)
    relationship = db.Column(db.String(16), default="teacher")  # teacher | parent
    invite_token = db.Column(db.String(64), nullable=True)
    accepted_at = db.Column(db.DateTime, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class FeatureFlag(db.Model):
    """Admin-controlled kill switches with optional percentage rollout.

    Default behaviour for an unknown flag is "enabled" — so apps still
    work if a flag row is missing or the table fails to load.

    When ``rollout_percentage`` is set (0-100), the flag uses a
    deterministic hash on user_id so assignment is stable and doesn't
    shift as other users sign up. 100 = everyone, 0 = nobody."""
    __tablename__ = "feature_flags"
    id = db.Column(db.Integer, primary_key=True)
    key = db.Column(db.String(64), unique=True, nullable=False)
    enabled = db.Column(db.Boolean, default=True)
    rollout_percentage = db.Column(db.Integer, default=100)
    description = db.Column(db.String(255), default="")
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class LiveSession(db.Model):
    """A real-time study room — Jitsi-backed. Anyone with the link can
    join. Owner can toggle audio/video defaults and pin materials."""
    __tablename__ = "live_sessions"
    id = db.Column(db.Integer, primary_key=True)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    title = db.Column(db.String(160), nullable=False, default="Study session")
    topic = db.Column(db.String(160), default="")
    room_slug = db.Column(db.String(48), unique=True, nullable=False)
    audio_only = db.Column(db.Boolean, default=False)
    video_enabled = db.Column(db.Boolean, default=True)
    audio_enabled = db.Column(db.Boolean, default=True)
    materials = db.Column(db.Text, default="")
    is_open = db.Column(db.Boolean, default=True)
    group_id = db.Column(db.Integer, db.ForeignKey("study_groups.id"), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class SavedMeeting(db.Model):
    """User-saved recurring or one-off class meeting links (Teams, Zoom, Meet, etc.)."""
    __tablename__ = "saved_meetings"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    name = db.Column(db.String(120), nullable=False)
    url = db.Column(db.String(512), nullable=False)
    platform = db.Column(db.String(32), default="other")
    schedule_text = db.Column(db.String(200), default="")
    is_recurring = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class SyllabusRecord(db.Model):
    """Stored syllabus imports — one row per course PDF."""
    __tablename__ = "syllabus_records"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    course_name = db.Column(db.String(160), default="")
    filename = db.Column(db.String(255), default="")
    assignments_json = db.Column(db.Text, default="[]")
    imported_count = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class SessionMessage(db.Model):
    """Chat messages for live study sessions and study groups."""
    __tablename__ = "session_messages"
    id = db.Column(db.Integer, primary_key=True)
    context_type = db.Column(db.String(16), nullable=False)  # live | group
    context_id = db.Column(db.Integer, nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    author_name = db.Column(db.String(120), default="Guest")
    body = db.Column(db.Text, nullable=False)
    saved_to_library = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class StudySession(db.Model):
    __tablename__ = "study_sessions"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    mode = db.Column(db.String(16), default="casual")
    questions_total = db.Column(db.Integer, default=0)
    questions_correct = db.Column(db.Integer, default=0)
    points_earned = db.Column(db.Integer, default=0)
    duration_seconds = db.Column(db.Integer, default=0)
    completed = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class StudyPoints(db.Model):
    __tablename__ = "study_points"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    total_points = db.Column(db.Integer, default=0)
    spark_balance = db.Column(db.Integer, default=0)
    sparks_earned_total = db.Column(db.Integer, default=0)
    level = db.Column(db.Integer, default=1)
    streak_count = db.Column(db.Integer, default=0)
    streak_freeze_count = db.Column(db.Integer, default=0)
    freeze_capacity = db.Column(db.Integer, default=2)
    last_active_date = db.Column(db.String(16), default="")
    repair_last_used = db.Column(db.String(16), default="")
    repair_eligible_until = db.Column(db.DateTime, nullable=True)
    streak_history = db.Column(db.Text, default="[]")
    session_history = db.Column(db.Text, default="[]")
    badges = db.Column(db.Text, default="[]")
    active_booster = db.Column(db.Text, default="null")
    active_cosmetics = db.Column(db.Text, default="{}")
    weekly_quests = db.Column(db.Text, default="{}")
    shop_purchases = db.Column(db.Text, default="[]")
    longest_streak = db.Column(db.Integer, default=0)
    total_sessions = db.Column(db.Integer, default=0)
    last_daily_claim = db.Column(db.String(16), default="")   # YYYY-MM-DD of last daily-chest claim
    updated_at = db.Column(db.DateTime, default=datetime.utcnow)

class StudyMastery(db.Model):
    __tablename__ = "study_mastery"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True)
    guest_session_id = db.Column(db.String(64), nullable=True)
    question_key = db.Column(db.String(512), nullable=False)
    question_text = db.Column(db.Text, default="")
    answer_text = db.Column(db.Text, default="")
    topic = db.Column(db.String(256), default="")
    mastery_level = db.Column(db.Integer, default=0)
    times_seen = db.Column(db.Integer, default=0)
    times_correct = db.Column(db.Integer, default=0)
    times_partial = db.Column(db.Integer, default=0)
    last_seen = db.Column(db.String(16), default="")
    next_review = db.Column(db.String(16), default="")
    easiness_factor = db.Column(db.Float, default=2.5)
    interval_days = db.Column(db.Integer, default=1)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class UserStreak(db.Model):
    """Task-completion streak — separate from the study-session streak in
    StudyPoints. Qualifying actions: completing a task OR viewing the
    dashboard (plan review). All date logic uses the user's local timezone."""
    __tablename__ = "user_streaks"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), unique=True, nullable=False)
    current_streak = db.Column(db.Integer, default=0)
    longest_streak = db.Column(db.Integer, default=0)
    last_qualifying_action_at = db.Column(db.DateTime(timezone=True), nullable=True)
    last_qualifying_local_date = db.Column(db.String(16), default="")
    freezes_available = db.Column(db.Integer, default=2)
    freezes_used_total = db.Column(db.Integer, default=0)
    timezone = db.Column(db.String(64), default="")
    nudge_shown_date = db.Column(db.String(16), default="")
    qualified_dates_json = db.Column(db.Text, default="[]")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class PlaniPet(db.Model):
    """Virtual creature that grows with site usage. Duolingo-style mascot.

    XP is awarded for: daily visits, task completion, schedule generation,
    streak milestones, study sessions, tutor chats, and grade imports.
    Stage is derived from cumulative XP via pet_engine.stage_for_xp().

    Care actions (feed/play/pet/study_with) are cooldown-gated. We track
    last-action timestamps and a chest-streak so a perfect-week chest
    pays more than a one-off check-in.
    """
    __tablename__ = "plani_pets"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), unique=True, nullable=False)
    name = db.Column(db.String(40), default="Plani")
    xp = db.Column(db.Integer, default=0)
    last_visit_local_date = db.Column(db.String(16), default="")
    hatched_at = db.Column(db.DateTime, default=datetime.utcnow)
    # Care actions cooldowns
    last_fed_at = db.Column(db.DateTime, nullable=True)
    last_played_at = db.Column(db.DateTime, nullable=True)
    last_petted_at = db.Column(db.DateTime, nullable=True)
    last_studied_at = db.Column(db.DateTime, nullable=True)
    # Daily chest (separate from streak)
    last_chest_local_date = db.Column(db.String(16), default="")
    chest_streak_days = db.Column(db.Integer, default=0)
    # Perfect-week tracking — week iso "YYYY-Www" we already paid out
    perfect_week_paid = db.Column(db.String(16), default="")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class FeatureRequest(db.Model):
    """User-submitted feature ideas surfaced on /features.

    Vote count is materialized for cheap list rendering and recomputed
    whenever a vote is toggled. Status: open / planned / in_progress /
    shipped / declined — admin only updates this.
    """
    __tablename__ = "feature_requests"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    title = db.Column(db.String(140), nullable=False)
    body = db.Column(db.Text, default="")
    category = db.Column(db.String(40), default="general")
    status = db.Column(db.String(20), default="open")
    vote_count = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class FeatureRequestVote(db.Model):
    __tablename__ = "feature_request_votes"
    id = db.Column(db.Integer, primary_key=True)
    request_id = db.Column(db.Integer, db.ForeignKey("feature_requests.id"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    __table_args__ = (db.UniqueConstraint("request_id", "user_id", name="uq_feature_vote"),)


class SiteFeedback(db.Model):
    """User-submitted feedback: bug reports, feature ideas, praise, or
    general comments. Powers the floating feedback widget available on
    every authenticated page."""
    __tablename__ = "site_feedback"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False, index=True)
    category = db.Column(db.String(24), nullable=False, default="general")
    mood = db.Column(db.Integer, nullable=True)
    message = db.Column(db.Text, nullable=False)
    page_url = db.Column(db.String(512), default="")
    status = db.Column(db.String(16), default="new")
    admin_note = db.Column(db.Text, default="")
    # JSON blob captured by the client bug-report dialog: route, viewport,
    # theme, connectivity and the last few JS failures. Empty for feedback
    # sent through the plain widget.
    diagnostics = db.Column(db.Text, default="")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class ClientErrorLog(db.Model):
    """JavaScript failures reported by the browser.

    Rows are collapsed on `fingerprint` (kind + message + source + line) so a
    render loop firing the same TypeError a thousand times produces one row
    with a count, not a thousand rows. `first_seen` / `last_seen` bracket the
    occurrence window, which is what tells us whether a bug is still live
    after a deploy.
    """
    __tablename__ = "client_error_logs"
    id = db.Column(db.Integer, primary_key=True)
    # Nullable: anonymous visitors hit JS errors too, and those are often the
    # most useful ones (they are on the pages with the least testing).
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=True, index=True)
    fingerprint = db.Column(db.String(64), nullable=False, index=True)
    kind = db.Column(db.String(32), default="error")
    message = db.Column(db.String(512), default="")
    stack = db.Column(db.Text, default="")
    source = db.Column(db.String(512), default="")
    line = db.Column(db.Integer, default=0)
    page_url = db.Column(db.String(512), default="")
    user_agent = db.Column(db.String(300), default="")
    viewport = db.Column(db.String(24), default="")
    context = db.Column(db.Text, default="")
    count = db.Column(db.Integer, default=1)
    resolved = db.Column(db.Boolean, default=False)
    first_seen = db.Column(db.DateTime, default=datetime.utcnow)
    last_seen = db.Column(db.DateTime, default=datetime.utcnow, index=True)


class MediaBalanceSession(db.Model):
    """Lightweight digital-wellbeing log. Records the *minutes* a user
    spent in IntelliPlan per local date — no per-event tracking, no
    restrictions, no lockouts. Used to surface gentle insights on /balance.
    """
    __tablename__ = "media_balance_sessions"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False)
    local_date = db.Column(db.String(16), nullable=False)  # YYYY-MM-DD
    minutes = db.Column(db.Integer, default=0)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    __table_args__ = (db.UniqueConstraint("user_id", "local_date", name="uq_balance_day"),)


class MediaBalancePrefs(db.Model):
    """Per-user opt-in awareness reminders. Awareness only — never blocks."""
    __tablename__ = "media_balance_prefs"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), unique=True, nullable=False)
    reminders_enabled = db.Column(db.Boolean, default=False)
    reminder_minutes = db.Column(db.Integer, default=45)   # gentle nudge cadence
    daily_goal_minutes = db.Column(db.Integer, default=60) # awareness target
    night_nudges_enabled = db.Column(db.Boolean, default=True)  # sleep reminders after night_start_hour
    night_start_hour = db.Column(db.Integer, default=22)        # local hour (0-23) when sleep mode begins
    night_cadence_minutes = db.Column(db.Integer, default=10)   # night reminder repeat cadence
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class AccessibilityPrefs(db.Model):
    """Per-user accessibility settings, server-stored so they follow the
    user across devices. The same set is also mirrored to localStorage on
    the client so the page can apply them before first paint."""
    __tablename__ = "accessibility_prefs"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), unique=True, nullable=False)
    dyslexia_font = db.Column(db.Boolean, default=False)
    text_scale = db.Column(db.Integer, default=100)       # 90 / 100 / 115 / 130 / 150
    line_spacing = db.Column(db.Integer, default=100)     # 100 / 130 / 160 / 200
    high_contrast = db.Column(db.Boolean, default=False)
    reduced_motion = db.Column(db.Boolean, default=False)
    underline_links = db.Column(db.Boolean, default=False)
    focus_ring_bold = db.Column(db.Boolean, default=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class EmailSend(db.Model):
    """One row per lifecycle email per user — the deduplication ledger.

    The unique constraint on (user_id, email_key) is the whole point: two
    cron fires racing each other both try to insert, one wins, the other
    takes an IntegrityError and skips. Doing this with a SELECT-then-send
    would leave a window where both reads miss and both sends go out, and
    the failure mode of that bug is a student getting the same email twice.

    The row is written *before* the provider call, so a crash mid-send
    leaves a stale ``pending`` row and the user simply never gets that one
    email. That is the right way round: a missing welcome is a small loss,
    a duplicate is a support ticket and an unsubscribe.
    """
    __tablename__ = "email_sends"
    __table_args__ = (
        db.UniqueConstraint("user_id", "email_key", name="uq_email_sends_user_key"),
    )
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id"), nullable=False, index=True)
    #: "welcome", "feedback_v1", "newsletter_2026_08" — the campaign identity.
    #: Versioned in the key so a second feedback ask can ship later without
    #: colliding with the first.
    email_key = db.Column(db.String(64), nullable=False)
    sent_at = db.Column(db.DateTime, default=datetime.utcnow)
    #: pending | sent | failed
    status = db.Column(db.String(16), default="pending")
    provider_message_id = db.Column(db.String(128), nullable=True)


class EmailSuppression(db.Model):
    """Addresses that must never receive marketing, keyed on the address.

    Deliberately not keyed on user_id. Someone who unsubscribes, deletes
    their account and signs up again is the same person with the same
    inbox, and a user-id suppression would quietly start mailing them
    again. The address is the thing that received the mail and the thing
    the complaint would come from.
    """
    __tablename__ = "email_suppressions"
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(255), unique=True, nullable=False, index=True)
    #: "unsubscribe" | "bounce" | "complaint" | "manual"
    reason = db.Column(db.String(32), default="unsubscribe")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


STREAK_TIERS = [
    {"id": "spark", "name": "Spark", "min": 1, "max": 6, "bonus": 5, "freeze_cap": 2, "color": "#ef4444"},
    {"id": "flame", "name": "Flame", "min": 7, "max": 13, "bonus": 10, "freeze_cap": 3, "color": "#f97316"},
    {"id": "blaze", "name": "Blaze", "min": 14, "max": 20, "bonus": 15, "freeze_cap": 3, "color": "#f59e0b"},
    {"id": "inferno", "name": "Inferno", "min": 21, "max": 29, "bonus": 20, "freeze_cap": 5, "color": "#dc2626"},
    {"id": "wildfire", "name": "Wildfire", "min": 30, "max": 59, "bonus": 30, "freeze_cap": 5, "color": "#f97316"},
    {"id": "firestorm", "name": "Firestorm", "min": 60, "max": 99, "bonus": 50, "freeze_cap": 5, "color": "#06b6d4"},
    {"id": "legendary", "name": "Legendary", "min": 100, "max": 364, "bonus": 75, "freeze_cap": 5, "color": "#8b5cf6"},
    {"id": "eternal", "name": "Eternal", "min": 365, "max": 99999, "bonus": 150, "freeze_cap": 5, "color": "#22c55e"},
]

LEVEL_TITLE_TIERS = [
    (1, "Learner"),
    (5, "Student"),
    (10, "Scholar"),
    (15, "Researcher"),
    (20, "Expert"),
    (25, "Veteran"),
    (30, "Mentor"),
    (40, "Master"),
    (50, "Legend"),
]

def level_title_for(level):
    title = LEVEL_TITLE_TIERS[0][1]
    for threshold, candidate in LEVEL_TITLE_TIERS:
        if level >= threshold:
            title = candidate
    return title

LEVELS = [
    (level, level_title_for(level), 0 if level == 1 else int(round((35 * ((level - 1) ** 2)) / 25) * 25))
    for level in range(1, 51)
]

STREAK_MILESTONES = {
    3: {"sparks": 25, "freezes": 0, "badge": "first_flame", "title": None},
    7: {"sparks": 75, "freezes": 1, "badge": "week_warrior", "title": None},
    14: {"sparks": 100, "freezes": 1, "badge": "fortnight_fighter", "title": None},
    21: {"sparks": 150, "freezes": 0, "badge": "inferno_initiate", "title": "Grinder"},
    30: {"sparks": 250, "freezes": 2, "badge": "monthly_master", "title": "Monthly Master"},
    60: {"sparks": 500, "freezes": 2, "badge": "sixty_strong", "title": None},
    100: {"sparks": 1000, "freezes": 3, "badge": "century_club", "title": "Legendary"},
    365: {"sparks": 2000, "freezes": 3, "badge": "year_of_fire", "title": "Eternal"},
}

BADGE_CATALOG = {
    "first_flame": {"name": "First Flame", "kind": "streak"},
    "week_warrior": {"name": "Week Warrior", "kind": "streak"},
    "fortnight_fighter": {"name": "Fortnight Fighter", "kind": "streak"},
    "inferno_initiate": {"name": "Inferno Initiate", "kind": "streak"},
    "monthly_master": {"name": "Monthly Master", "kind": "streak"},
    "sixty_strong": {"name": "Sixty Strong", "kind": "streak"},
    "century_club": {"name": "Century Club", "kind": "streak"},
    "year_of_fire": {"name": "Year of Fire", "kind": "streak"},
    "first_session": {"name": "First Session", "kind": "session"},
    "getting_serious": {"name": "Getting Serious", "kind": "session"},
    "dedicated": {"name": "Dedicated", "kind": "session"},
    "committed": {"name": "Committed", "kind": "session"},
    "unstoppable": {"name": "Unstoppable", "kind": "session"},
    "sharp": {"name": "Sharp", "kind": "accuracy"},
    "precise": {"name": "Precise", "kind": "accuracy"},
    "flawless": {"name": "Flawless", "kind": "accuracy"},
    "perfect_week": {"name": "Perfect Week", "kind": "special"},
    "speed_demon": {"name": "Speed Demon", "kind": "special"},
    "night_owl": {"name": "Night Owl", "kind": "special"},
    "early_bird": {"name": "Early Bird", "kind": "special"},
    "comeback_kid": {"name": "Comeback Kid", "kind": "special"},
    "quest_starter": {"name": "Quest Starter", "kind": "quest"},
    "quest_finisher": {"name": "Quest Finisher", "kind": "quest"},
    "weekly_champion": {"name": "Weekly Champion", "kind": "quest"},
    "shopper": {"name": "Spark Shopper", "kind": "shop"},
    "deal_hunter": {"name": "Deal Hunter", "kind": "shop"},
    "freeze_ready": {"name": "Freeze Ready", "kind": "protection"},
    "booster_pilot": {"name": "Booster Pilot", "kind": "shop"},
    "spark_saver": {"name": "Spark Saver", "kind": "currency"},
    "style_setter": {"name": "Style Setter", "kind": "cosmetic"},
}

SHOP_ITEMS = {
    "streak_freeze": {"name": "Streak Freeze", "price": 200, "kind": "protection", "value": 1, "description": "Blocks one missed day."},
    "freeze_pack": {"name": "Freeze Pack", "price": 500, "kind": "protection", "value": 3, "description": "Three freezes at a discount."},
    "weekend_shield": {"name": "Weekend Shield", "price": 350, "kind": "protection", "value": 2, "description": "Adds two freezes for busy weekends or travel days."},
    "repair_token": {"name": "Repair Token", "price": 900, "kind": "inventory", "field": "repair_credits", "value": 1, "description": "Cuts the next streak repair cost in half."},
    "booster_2x": {"name": "2x Sparks", "price": 100, "kind": "booster", "multiplier": 2, "uses": 1, "description": "Doubles Sparks in the next session."},
    "booster_3x": {"name": "3x Sparks", "price": 250, "kind": "booster", "multiplier": 3, "uses": 1, "description": "Triples Sparks in the next session."},
    "daily_booster": {"name": "Daily Booster", "price": 350, "kind": "booster", "multiplier": 1.5, "hours": 24, "description": "+50% Sparks for 24 hours."},
    "focus_booster": {"name": "Focus Booster", "price": 180, "kind": "booster", "multiplier": 1.25, "hours": 72, "description": "+25% Sparks for the next three days."},
    "skip_pack": {"name": "Question Skip Pack", "price": 75, "kind": "inventory", "field": "skips", "value": 5, "description": "Skip five questions without losing momentum."},
    "hint_pack": {"name": "Hint Token Pack", "price": 120, "kind": "inventory", "field": "hints", "value": 10, "description": "Use AI hints on hard questions."},
    "color_gold": {"name": "Streak Color: Gold", "price": 400, "kind": "cosmetic", "slot": "streak_color", "value": "gold", "description": "Gold streak glow."},
    "color_neon": {"name": "Streak Color: Neon", "price": 400, "kind": "cosmetic", "slot": "streak_color", "value": "neon", "description": "Neon streak glow."},
    "color_forest": {"name": "Streak Color: Forest", "price": 400, "kind": "cosmetic", "slot": "streak_color", "value": "forest", "description": "Calm green streak glow."},
    "title_scholar": {"name": "Profile Title: Scholar", "price": 300, "kind": "cosmetic", "slot": "title", "value": "Scholar", "description": "Display Scholar by your name."},
    "title_grinder": {"name": "Profile Title: Grinder", "price": 300, "kind": "cosmetic", "slot": "title", "value": "Grinder", "description": "Display Grinder by your name."},
    "title_comeback": {"name": "Profile Title: Comeback Kid", "price": 450, "kind": "cosmetic", "slot": "title", "value": "Comeback Kid", "description": "Display Comeback Kid by your name."},
    "calendar_dark": {"name": "Calendar Theme: Dark", "price": 500, "kind": "cosmetic", "slot": "calendar_theme", "value": "dark", "description": "Dark calendar grid."},
    "frame_ember": {"name": "Streak Frame: Ember", "price": 600, "kind": "cosmetic", "slot": "streak_frame", "value": "ember", "description": "Animated ember frame."},
    "frame_aurora": {"name": "Streak Frame: Aurora", "price": 800, "kind": "cosmetic", "slot": "streak_frame", "value": "aurora", "description": "Rare aurora frame."},
    "frame_cosmic": {"name": "Streak Frame: Cosmic", "price": 950, "kind": "cosmetic", "slot": "streak_frame", "value": "cosmic", "description": "Premium cosmic streak frame."},
}

QUEST_POOL = [
    {"id": "study_5_days", "title": "Study 5 days this week", "metric": "study_days", "target": 5, "reward": 80},
    {"id": "answer_75_correct", "title": "Answer 75 questions correctly", "metric": "correct_answers", "target": 75, "reward": 60},
    {"id": "serious_3_sessions", "title": "Complete 3 Serious or Extreme sessions", "metric": "focus_sessions", "target": 3, "reward": 70},
    {"id": "maintain_7_days", "title": "Maintain your streak all 7 days", "metric": "study_days", "target": 7, "reward": 120},
    {"id": "master_3_concepts", "title": "Master 3 new concepts", "metric": "mastered_concepts", "target": 3, "reward": 60},
    {"id": "earn_200_session", "title": "Earn 200 Sparks in a single session", "metric": "single_session_sparks", "target": 200, "reward": 50},
    {"id": "perfect_session", "title": "Complete a perfect session", "metric": "perfect_sessions", "target": 1, "reward": 90},
    {"id": "study_30_minutes", "title": "Spend 30+ minutes studying this week", "metric": "study_minutes", "target": 30, "reward": 70},
]

def safe_json_load(raw, fallback):
    try:
        if raw in (None, ""):
            return fallback
        return json.loads(raw)
    except Exception:
        return fallback

def current_study_tier(streak_count):
    streak_count = int(streak_count or 0)
    for tier in STREAK_TIERS:
        if tier["min"] <= streak_count <= tier["max"]:
            return tier
    return STREAK_TIERS[0]

def repair_cost_for(streak_count):
    streak_count = int(streak_count or 0)
    if streak_count <= 6:
        return 150
    if streak_count <= 13:
        return 300
    if streak_count <= 29:
        return 500
    if streak_count <= 59:
        return 800
    return 1200

def level_for_sparks(sparks_total):
    sparks_total = int(sparks_total or 0)
    current = LEVELS[0]
    for level in LEVELS:
        if sparks_total >= level[2]:
            current = level
    return {"level": current[0], "title": current[1], "next": next(({"level": l, "title": t, "required": req} for l, t, req in LEVELS if req > sparks_total), None)}

def add_badges(p, badge_ids):
    badges = safe_json_load(p.badges, [])
    changed = []
    for badge_id in badge_ids:
        if badge_id and badge_id not in badges:
            badges.append(badge_id)
            changed.append(badge_id)
    if changed:
        p.badges = json.dumps(badges)
    return changed

def grant_sparks(p, amount, reason="", apply_booster=True):
    amount = max(0, int(round(float(amount or 0))))
    if amount <= 0:
        return {"awarded": 0, "base": 0, "multiplier": 1, "level_up": None}
    multiplier = 1
    booster = safe_json_load(p.active_booster, None)
    now = utcnow()
    if apply_booster and isinstance(booster, dict):
        expires_at = booster.get("expires_at")
        if expires_at:
            try:
                if datetime.fromisoformat(expires_at) < now:
                    booster = None
            except Exception:
                booster = None
        if booster:
            multiplier = float(booster.get("multiplier", 1) or 1)
            if booster.get("uses") is not None:
                booster["uses"] = max(0, int(booster.get("uses", 0)) - 1)
                p.active_booster = json.dumps(booster) if booster["uses"] > 0 else "null"
    awarded = int(round(amount * multiplier))
    old_level = int(p.level or 1)
    p.spark_balance = int(p.spark_balance or 0) + awarded
    p.sparks_earned_total = int(p.sparks_earned_total or p.total_points or 0) + awarded
    p.total_points = int(p.sparks_earned_total or 0)
    new_level = level_for_sparks(p.sparks_earned_total)["level"]
    level_up = None
    if new_level > old_level:
        p.level = new_level
        p.spark_balance += 50
        p.sparks_earned_total += 50
        p.total_points = p.sparks_earned_total
        level_up = {"from": old_level, "to": new_level, "bonus": 50}
    p.updated_at = utcnow()
    return {"awarded": awarded, "base": amount, "multiplier": multiplier, "level_up": level_up}

def active_week_id(dt=None):
    dt = dt or datetime.now()
    iso = dt.isocalendar()
    return f"{iso.year}-W{iso.week:02d}"

def ensure_weekly_quests(p):
    week_id = active_week_id()
    data = safe_json_load(p.weekly_quests, {})
    if data.get("week_id") == week_id and data.get("quests"):
        return data
    seed = f"{week_id}:{p.user_id or p.guest_session_id or p.id or 'study'}"
    rng = random.Random(seed)
    quests = [dict(q) for q in rng.sample(QUEST_POOL, 3)]
    data = {
        "week_id": week_id,
        "quests": quests,
        "progress": {q["id"]: 0 for q in quests},
        "completed": [],
        "completion_bonus_claimed": False,
        "study_dates": [],
    }
    p.weekly_quests = json.dumps(data)
    return data

def update_quest_progress(p, session_data):
    quests = ensure_weekly_quests(p)
    progress = quests.setdefault("progress", {})
    completed = set(quests.get("completed", []))
    study_date = session_data.get("date") or datetime.now().strftime("%Y-%m-%d")
    study_dates = set(quests.get("study_dates", []))
    study_dates.add(study_date)
    quests["study_dates"] = sorted(study_dates)
    metrics = {
        "study_days": len(study_dates),
        "correct_answers": int(session_data.get("correct", 0) or 0),
        "focus_sessions": 1 if session_data.get("mode") in ("serious", "extreme") else 0,
        "mastered_concepts": int(session_data.get("mastered_concepts", 0) or 0),
        "single_session_sparks": int(session_data.get("sparks", 0) or 0),
        "perfect_sessions": 1 if int(session_data.get("questions", 0) or 0) > 0 and int(session_data.get("correct", 0) or 0) >= int(session_data.get("questions", 0) or 0) else 0,
        "study_minutes": int(round((int(session_data.get("duration", 0) or 0)) / 60)),
    }
    rewards = []
    quest_badges = []
    for q in quests.get("quests", []):
        qid = q["id"]
        metric = q["metric"]
        if metric == "study_days":
            progress[qid] = len(study_dates)
        elif metric == "single_session_sparks":
            progress[qid] = max(int(progress.get(qid, 0) or 0), metrics[metric])
        else:
            progress[qid] = int(progress.get(qid, 0) or 0) + metrics.get(metric, 0)
        if progress[qid] >= q["target"] and qid not in completed:
            completed.add(qid)
            rewards.append({"quest_id": qid, "title": q["title"], "sparks": q["reward"]})
            grant_sparks(p, q["reward"], f"quest:{qid}", apply_booster=False)
            quest_badges.append("quest_starter")
    quests["completed"] = sorted(completed)
    if len(completed) >= 3 and not quests.get("completion_bonus_claimed"):
        quests["completion_bonus_claimed"] = True
        rewards.append({"quest_id": "weekly_completion", "title": "Weekly Completion Bonus", "sparks": 150, "freezes": 1})
        grant_sparks(p, 150, "weekly_completion", apply_booster=False)
        p.streak_freeze_count = min(int(p.freeze_capacity or 2), int(p.streak_freeze_count or 0) + 1)
        quest_badges.extend(["quest_finisher", "weekly_champion"])
    if quest_badges:
        add_badges(p, quest_badges)
    p.weekly_quests = json.dumps(quests)
    return rewards

def passive_freezes_due(p, today=None):
    today = today or datetime.now()
    if int(p.streak_count or 0) < 30:
        return 0
    history = safe_json_load(p.shop_purchases, [])
    week_id = active_week_id(today)
    passive_key = f"passive_freeze:{week_id}"
    if any(item.get("id") == passive_key for item in history if isinstance(item, dict)):
        return 0
    amount = 2 if int(p.streak_count or 0) >= 100 else 1
    history.append({"id": passive_key, "item_id": "passive_freeze", "qty": amount, "created_at": utcnow().isoformat()})
    p.shop_purchases = json.dumps(history[-200:])
    return amount

def reconcile_missed_streak(p):
    if not p.last_active_date or int(p.streak_count or 0) <= 0:
        return None
    today = datetime.now().date()
    try:
        last = datetime.strptime(p.last_active_date[:10], "%Y-%m-%d").date()
    except Exception:
        return None
    if last >= today - timedelta(days=1):
        return None
    history = safe_json_load(p.streak_history, [])
    cosmetics = safe_json_load(p.active_cosmetics, {})
    missed_day = (last + timedelta(days=1)).strftime("%Y-%m-%d")
    tier = current_study_tier(p.streak_count)
    p.freeze_capacity = tier["freeze_cap"]
    if int(p.streak_freeze_count or 0) > 0:
        p.streak_freeze_count = max(0, int(p.streak_freeze_count or 0) - 1)
        p.last_active_date = missed_day
        if missed_day not in history:
            history.append(missed_day)
            p.streak_history = json.dumps(sorted(history)[-90:])
        return {"type": "freeze_consumed", "message": "Streak Freeze used - your streak is safe. Back tomorrow."}
    if not p.repair_eligible_until or p.repair_eligible_until < utcnow():
        p.repair_eligible_until = utcnow() + timedelta(hours=48)
        cosmetics["broken_streak_count"] = int(p.streak_count or 0)
        cosmetics["broken_last_active_date"] = p.last_active_date
        p.active_cosmetics = json.dumps(cosmetics)
        p.streak_count = 0
        return {"type": "repair_window", "message": "Your streak broke, but you can still save it."}
    return {"type": "repair_window", "message": "Your streak broke, but you can still save it."}

def apply_study_schema_migrations():
    inspector = inspect(db.engine)
    if "study_points" not in inspector.get_table_names():
        return
    existing = {col["name"] for col in inspector.get_columns("study_points")}
    dialect = db.engine.dialect.name
    text_type = "TEXT"
    dt_type = "TIMESTAMP" if dialect != "sqlite" else "DATETIME"
    columns = {
        "spark_balance": "INTEGER DEFAULT 0",
        "sparks_earned_total": "INTEGER DEFAULT 0",
        "level": "INTEGER DEFAULT 1",
        "freeze_capacity": "INTEGER DEFAULT 2",
        "repair_last_used": f"{text_type} DEFAULT ''",
        "repair_eligible_until": dt_type,
        "badges": f"{text_type} DEFAULT '[]'",
        "active_booster": f"{text_type} DEFAULT 'null'",
        "active_cosmetics": f"{text_type} DEFAULT '{{}}'",
        "weekly_quests": f"{text_type} DEFAULT '{{}}'",
        "shop_purchases": f"{text_type} DEFAULT '[]'",
        "longest_streak": "INTEGER DEFAULT 0",
        "total_sessions": "INTEGER DEFAULT 0",
    }
    for name, ddl in columns.items():
        if name not in existing:
            db.session.execute(text(f"ALTER TABLE study_points ADD COLUMN {name} {ddl}"))
    db.session.execute(text("UPDATE study_points SET spark_balance = COALESCE(spark_balance, total_points, 0)"))
    db.session.execute(text("UPDATE study_points SET sparks_earned_total = COALESCE(sparks_earned_total, total_points, 0)"))
    db.session.execute(text("UPDATE study_points SET level = COALESCE(level, 1), freeze_capacity = COALESCE(freeze_capacity, 2), longest_streak = COALESCE(longest_streak, streak_count, 0), total_sessions = COALESCE(total_sessions, 0)"))
    db.session.execute(text("UPDATE study_points SET badges = COALESCE(badges, '[]'), active_booster = COALESCE(active_booster, 'null'), active_cosmetics = COALESCE(active_cosmetics, '{}'), weekly_quests = COALESCE(weekly_quests, '{}'), shop_purchases = COALESCE(shop_purchases, '[]')"))
    db.session.commit()

# ── Command Center models (additive — see docs/command-center/) ────
# Register the three new tables (briefing_cache, health_snapshots,
# student_signals) against the existing ``db`` instance. The
# ``register(db)`` callback pattern avoids any circular import with the
# ``intelliplan`` package while keeping the model classes accessible on
# the App.py namespace for the upcoming /api/today handler.
from intelliplan.models import command_center as _cc_models
from intelliplan.models import learning_graph as _lg_models
from intelliplan.models import active_session as _as_models
from intelliplan.migrations import (
    apply_active_session_migrations,
    apply_notification_migrations,
    apply_command_center_migrations,
    apply_email_migrations,
    apply_learning_graph_migrations,
    apply_media_balance_migrations,
    apply_sync_migrations,
)
BriefingCache, HealthSnapshot, StudentSignal = _cc_models.register(db)
StudentProfile, ConceptMastery, LearningEvent = _lg_models.register(db)
# Active study sessions — the table the scheduler's estimation model learns
# from. See intelliplan/models/active_session.py for the privacy contract
# covering the focus-sample rows.
ActiveSession, ActiveFocusSample = _as_models.register(db)
# Notification outbox — durable queue with dedupe, retries, and expiry.
from intelliplan.notifications import models as _notif_models
NotificationOutbox = _notif_models.register(db)
# Single-runner lease for the in-process notification timer. Registered
# here so create_all() builds it; see notifications_glue.start_ticker for
# why more than one worker sweeping at once is not safe.
CronLease = _notif_models.register_lease(db)
# Offline replay ledger. Registered here, alongside the other models, so
# the create_all() below builds the table and its unique index; the
# request hooks that use it are installed further down, once `current_user`
# is available to scope ops by owner.
from intelliplan.sync import models as _sync_models
_sync_models.register(db)

with app.app_context():
    db.create_all()
    apply_study_schema_migrations()
    apply_media_balance_migrations(db)
    apply_command_center_migrations(db)
    apply_learning_graph_migrations(db)
    apply_active_session_migrations(db)
    apply_notification_migrations(db)
    apply_sync_migrations(db)
    apply_email_migrations(db)

@login_manager.user_loader
def load_user(user_id):
    # Defensive: if the SELECT for User fails (e.g. a column from a new
    # migration doesn't exist yet on a freshly-deployed DB), DO NOT
    # propagate the exception. Flask-Login bubbles it up into every
    # template render, and the error page itself extends base.html
    # which calls is_logged_in() → load_user → the same exception, so
    # users see only the raw "Server Error" fallback with no way out.
    try:
        return db.session.get(User, int(user_id))
    except Exception as _e:
        try:
            db.session.rollback()
        except Exception:
            pass
        print(f"[load_user] failed for {user_id}: {_e}")
        return None

@login_manager.request_loader
def load_user_from_request(req):
    auth_header = req.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        return None
    token = auth_header.split(" ", 1)[1].strip()
    if not token:
        return None
    payload = verify_token(token)
    if not payload:
        return None
    user_id = payload.get("user_id")
    if not user_id:
        return None
    try:
        return db.session.get(User, int(user_id))
    except Exception:
        return None

# ── CONSTANTS ─────────────────────────────────────────────────
PRIORITY_COLORS = {
    "High": "#ef4444",
    "Medium": "#f59e0b",
    "Low": "#22c55e",
}

DIFFICULTY_COLORS = {
    "Easy": "#86efac",
    "Medium": "#60a5fa",
    "Hard": "#8b5cf6",
}

WORKLOAD_COLORS = {
    "light": "#dcfce7",
    "moderate": "#fef3c7",
    "heavy": "#fee2e2",
}

API_ERROR_MESSAGES = {
    "ai": "AI scheduling is temporarily unavailable. Please try again in a moment.",
    "groq": "AI scheduling is temporarily unavailable. Please try again in a moment.",
    "canvas": "Canvas connection failed. Check your API token in Settings.",
    "studentvue": "StudentVue connection failed. Check your credentials in Settings.",
    "google_calendar": "Google Calendar sync is temporarily unavailable.",
    "notion": "Notion connection failed. Try reconnecting in Integrations.",
    "generic": "Service temporarily unavailable. Please try again later."
}


# Exception types whose text is always internal detail, never something a
# student can act on. A SQLAlchemy error's str() contains the failing SQL
# statement and the table's column list; an OSError contains filesystem
# paths. Returning either to an API caller is a free schema dump.
_OPAQUE_ERROR_TYPES = (
    "OperationalError", "ProgrammingError", "IntegrityError", "DataError",
    "InternalError", "InvalidRequestError", "StatementError", "DBAPIError",
    "SQLAlchemyError", "OSError", "IOError", "AttributeError", "TypeError",
    "KeyError", "IndexError", "ImportError", "ModuleNotFoundError",
    "RecursionError", "MemoryError",
)

# Fragments that mean the text is internal even when the exception type
# looked harmless — a wrapped driver error, a traceback fragment, a path.
_LEAKY_FRAGMENTS = re.compile(
    r"(?:\bSELECT\b|\bINSERT\b|\bUPDATE\b|\bDELETE\b|\bFROM\s+\w+|\bTable\b|"
    r"\bcolumn\b|Traceback|File \"|/usr/|[A-Za-z]:\\\\|site-packages|"
    r"psycopg|sqlite3|sqlalchemy|\bDSN\b|password=|token=|secret=)",
    re.IGNORECASE,
)


def safe_error_message(exc, fallback=None, limit=180):
    """A message safe to send to an API caller.

    Most of this codebase's handlers returned ``str(e)`` directly. That is
    fine for a hand-written ``ValueError("Due date must be in the future")``
    and disastrous for a database error, whose text carries the SQL and the
    schema. Rather than blanking every message — several are genuinely
    useful and the UI shows them — this keeps short, human-sounding text
    and replaces anything that looks like machinery.

    The full exception still goes to the server log at every call site; the
    only thing narrowed is what crosses the network.
    """
    fallback = fallback or API_ERROR_MESSAGES["generic"]
    if exc is None:
        return fallback
    if type(exc).__name__ in _OPAQUE_ERROR_TYPES:
        return fallback
    text = str(exc).strip()
    if not text or len(text) > 300:
        return fallback
    if _LEAKY_FRAGMENTS.search(text):
        return fallback
    if "\n" in text:
        text = text.split("\n", 1)[0].strip()
    return text[:limit]

# ── ERROR HELPERS ─────────────────────────────────────────────
def make_error_id():
    return "IPE-" + str(uuid.uuid4())[:8].upper()

# ── HELPERS ───────────────────────────────────────────────────
def get_guest_session_id():
    if "guest_id" not in session:
        session["guest_id"] = str(uuid.uuid4())
    return session["guest_id"]


# ── Plan cache ────────────────────────────────────────────────────────
# Building a plan fits the student's estimation model over their history
# and then runs the optimizer, which is cheap in absolute terms but not
# free, and the scheduler page reads it on every navigation. Cache per
# identity, and — this is the load-bearing half — invalidate the moment a
# study session ends, because "finishing a session changes the next plan"
# is the entire promise of the feedback loop.

_PLAN_CACHE = {}
_PLAN_CACHE_TTL = 120          # seconds
_PLAN_CACHE_MAX = 500          # bound the process's memory


def _plan_cache_key(user_id, guest_id):
    return f"u{user_id}" if user_id else (f"g{guest_id}" if guest_id else None)


def get_cached_plan(user_id, guest_id):
    """Return a cached plan payload, or None when absent or stale."""
    key = _plan_cache_key(user_id, guest_id)
    if not key:
        return None
    entry = _PLAN_CACHE.get(key)
    if not entry:
        return None
    stamped, payload = entry
    if (time.monotonic() - stamped) > _PLAN_CACHE_TTL:
        _PLAN_CACHE.pop(key, None)
        return None
    return payload


def set_cached_plan(user_id, guest_id, payload):
    key = _plan_cache_key(user_id, guest_id)
    if not key:
        return payload
    _PLAN_CACHE[key] = (time.monotonic(), payload)
    overflow = len(_PLAN_CACHE) - _PLAN_CACHE_MAX
    if overflow > 0:
        for stale_key, _ in sorted(_PLAN_CACHE.items(), key=lambda kv: kv[1][0])[:overflow]:
            _PLAN_CACHE.pop(stale_key, None)
    return payload


def invalidate_schedule_cache(user_id=None, guest_id=None):
    """Drop one student's cached plan. Safe to call with nothing to drop."""
    key = _plan_cache_key(user_id, guest_id)
    if key:
        _PLAN_CACHE.pop(key, None)

def is_logged_in():
    if current_user.is_authenticated:
        return True
    return "login_type" in session

def _default_lms_prefs():
    return {
        "grade_source": "active",
        "assignment_sources": [
            "active", "google_classroom", "blackboard", "moodle", "notion", "manual",
        ],
    }


def get_user_lms_prefs():
    if not current_user.is_authenticated:
        return _default_lms_prefs()
    raw = getattr(current_user, "lms_preferences", None) or "{}"
    try:
        prefs = json.loads(raw) if isinstance(raw, str) else (raw or {})
    except Exception:
        prefs = {}
    out = _default_lms_prefs()
    if prefs.get("grade_source"):
        out["grade_source"] = prefs["grade_source"]
    if isinstance(prefs.get("assignment_sources"), list) and prefs["assignment_sources"]:
        out["assignment_sources"] = prefs["assignment_sources"]
    return out


def _linked_account_by_type(login_type):
    if not current_user.is_authenticated or not login_type:
        return None
    acct = LinkedAccount.query.filter_by(
        user_id=current_user.id, login_type=login_type
    ).order_by(LinkedAccount.is_active.desc(), LinkedAccount.id.desc()).first()
    if not acct:
        return None
    creds = acct.get_credentials()
    creds["login_type"] = acct.login_type
    return creds


def get_grade_account():
    """Account used for grades/gradebook — respects user LMS preference."""
    prefs = get_user_lms_prefs()
    src = (prefs.get("grade_source") or "active").strip().lower()
    if src == "active":
        return get_active_account()
    if src in ("canvas", "studentvue", "schoology"):
        return _linked_account_by_type(src) or get_active_account()
    return get_active_account()


def get_active_account():
    if current_user.is_authenticated:
        acct = LinkedAccount.query.filter_by(user_id=current_user.id, is_active=True).first()
        if acct:
            creds = acct.get_credentials()
            creds["login_type"] = acct.login_type
            return creds
        return None
    login_type = session.get("login_type")
    if not login_type:
        return None
    if login_type == "canvas":
        return {
            "login_type": "canvas",
            "canvas_token": session.get("canvas_token"),
            "canvas_url": session.get("canvas_url"),
        }
    if login_type == "studentvue":
        return {
            "login_type": "studentvue",
            "sv_username": session.get("sv_username"),
            "sv_password": session.get("sv_password"),
            "sv_district_url": session.get("sv_district_url"),
        }
    if login_type == "schoology":
        return {
            "login_type": "schoology",
            "schoology_key": session.get("schoology_key"),
            "schoology_secret": session.get("schoology_secret"),
        }
    return None

def _norm_title(title):
    """Normalize an assignment title into a stable matching key.

    Completion state is persisted keyed by the assignment's *display
    title*. Those strings drift between an LMS sync and a later render —
    HTML entities (``&amp;`` vs ``&``), smart quotes, NBSPs, trailing
    whitespace, and casing all vary — and any drift made a previously
    completed assignment reappear ("it resets when I revisit the page").

    Collapsing to a normalized key (unescape entities → NFKC →
    collapse whitespace → casefold) makes the match survive that drift
    without needing a stable per-assignment id from every provider.
    """
    if not title:
        return ""
    s = _html_mod.unescape(str(title))
    s = _unicodedata.normalize("NFKC", s)
    # Fold typographic variants NFKC leaves alone — curly quotes and
    # dashes routinely differ between an LMS feed and a later render.
    s = s.translate(_TYPO_FOLD)
    s = re.sub(r"\s+", " ", s).strip().casefold()
    return s


# Smart quotes / dashes → ASCII equivalents (NFKC doesn't fold these).
_TYPO_FOLD = {
    0x2018: "'", 0x2019: "'", 0x201A: "'", 0x201B: "'",
    0x201C: '"', 0x201D: '"', 0x201E: '"', 0x201F: '"',
    0x2013: "-", 0x2014: "-", 0x2015: "-", 0x2212: "-",
}


class _DismissedSet:
    """Set-like view over completed-assignment titles that compares by
    normalized key (see :func:`_norm_title`).

    Drop-in for the old ``{r.title for r in rows}`` set: every existing
    ``title in dismissed`` / ``title not in dismissed`` call site keeps
    working unchanged but now matches robustly across title drift.
    """
    __slots__ = ("_keys",)

    def __init__(self, titles):
        self._keys = {_norm_title(t) for t in titles}

    def __contains__(self, title):
        return _norm_title(title) in self._keys

    def __iter__(self):
        return iter(self._keys)

    def __len__(self):
        return len(self._keys)

    def __bool__(self):
        return bool(self._keys)


def get_dismissed_titles():
    if current_user.is_authenticated:
        rows = DismissedAssignment.query.filter_by(user_id=current_user.id).all()
    else:
        gid = get_guest_session_id()
        rows = DismissedAssignment.query.filter_by(guest_session_id=gid).all()
    return _DismissedSet(r.title for r in rows)

def get_dismissed_rows():
    if current_user.is_authenticated:
        return DismissedAssignment.query.filter_by(user_id=current_user.id).all()
    gid = get_guest_session_id()
    return DismissedAssignment.query.filter_by(guest_session_id=gid).all()

def save_dismissed(title, data_dict):
    # Dedupe on the normalized key (not the raw string) so the same
    # assignment dismissed under a slightly-drifted title doesn't create
    # a second row. The original title is still stored for display.
    key = _norm_title(title)
    if current_user.is_authenticated:
        rows = DismissedAssignment.query.filter_by(user_id=current_user.id).all()
        if not any(_norm_title(r.title) == key for r in rows):
            db.session.add(DismissedAssignment(user_id=current_user.id, title=title, data=json.dumps(data_dict)))
    else:
        gid = get_guest_session_id()
        rows = DismissedAssignment.query.filter_by(guest_session_id=gid).all()
        if not any(_norm_title(r.title) == key for r in rows):
            db.session.add(DismissedAssignment(guest_session_id=gid, title=title, data=json.dumps(data_dict)))
    db.session.commit()

def delete_dismissed(title):
    # Match on the normalized key so un-completing works even when the
    # title sent now differs from the stored one by entity/whitespace/
    # case drift. Also sweeps up any historical duplicate rows.
    key = _norm_title(title)
    if current_user.is_authenticated:
        rows = DismissedAssignment.query.filter_by(user_id=current_user.id).all()
    else:
        gid = get_guest_session_id()
        rows = DismissedAssignment.query.filter_by(guest_session_id=gid).all()
    for r in rows:
        if _norm_title(r.title) == key:
            db.session.delete(r)
    db.session.commit()

def get_test_titles():
    if current_user.is_authenticated:
        rows = TestMark.query.filter_by(user_id=current_user.id).all()
    else:
        gid = get_guest_session_id()
        rows = TestMark.query.filter_by(guest_session_id=gid).all()
    return {r.title for r in rows}

def get_test_marks():
    if current_user.is_authenticated:
        return TestMark.query.filter_by(user_id=current_user.id).order_by(TestMark.created_at.desc()).all()
    gid = get_guest_session_id()
    return TestMark.query.filter_by(guest_session_id=gid).order_by(TestMark.created_at.desc()).all()

def save_test_mark(title, data_dict):
    if current_user.is_authenticated:
        existing = TestMark.query.filter_by(user_id=current_user.id, title=title).first()
        if not existing:
            db.session.add(TestMark(user_id=current_user.id, title=title, data=json.dumps(data_dict)))
    else:
        gid = get_guest_session_id()
        existing = TestMark.query.filter_by(guest_session_id=gid, title=title).first()
        if not existing:
            db.session.add(TestMark(guest_session_id=gid, title=title, data=json.dumps(data_dict)))
    db.session.commit()

def delete_test_mark(title):
    if current_user.is_authenticated:
        TestMark.query.filter_by(user_id=current_user.id, title=title).delete()
    else:
        gid = get_guest_session_id()
        TestMark.query.filter_by(guest_session_id=gid, title=title).delete()
    db.session.commit()

def get_custom_description(assignment_title):
    if current_user.is_authenticated:
        row = CustomDescription.query.filter_by(user_id=current_user.id, assignment_title=assignment_title).first()
    else:
        gid = get_guest_session_id()
        row = CustomDescription.query.filter_by(guest_session_id=gid, assignment_title=assignment_title).first()
    return row.description if row else None

def save_custom_description(assignment_title, description):
    if current_user.is_authenticated:
        row = CustomDescription.query.filter_by(user_id=current_user.id, assignment_title=assignment_title).first()
        if row:
            row.description = description
        else:
            db.session.add(CustomDescription(user_id=current_user.id, assignment_title=assignment_title, description=description))
    else:
        gid = get_guest_session_id()
        row = CustomDescription.query.filter_by(guest_session_id=gid, assignment_title=assignment_title).first()
        if row:
            row.description = description
        else:
            db.session.add(CustomDescription(guest_session_id=gid, assignment_title=assignment_title, description=description))
    db.session.commit()

def get_google_token():
    if current_user.is_authenticated:
        # Prefer the active row; fall back to the most recent if none flagged.
        gi = (
            GoogleIntegration.query
            .filter_by(user_id=current_user.id, is_active=True)
            .first()
            or GoogleIntegration.query
            .filter_by(user_id=current_user.id)
            .order_by(GoogleIntegration.id.desc())
            .first()
        )
        if gi:
            try:
                return json.loads(gi.token_data)
            except (TypeError, json.JSONDecodeError):
                db.session.delete(gi)
                db.session.commit()
                return None
    return session.get("google_token")

def get_notion_token_and_db():
    if current_user.is_authenticated:
        ni = NotionIntegration.query.filter_by(user_id=current_user.id).first()
        if ni and ni.token:
            return ni.token, ni.database_id
    return session.get("notion_token"), session.get("notion_database_id")

def get_study_profile(user_id=None, guest_id=None):
    if user_id:
        p = StudyPoints.query.filter_by(user_id=user_id).first()
        if not p:
            p = StudyPoints(user_id=user_id)
            db.session.add(p)
            db.session.commit()
    else:
        p = StudyPoints.query.filter_by(guest_session_id=guest_id).first()
        if not p:
            p = StudyPoints(guest_session_id=guest_id)
            db.session.add(p)
            db.session.commit()
    if p.spark_balance is None:
        p.spark_balance = int(p.total_points or 0)
    if p.sparks_earned_total is None or int(p.sparks_earned_total or 0) < int(p.total_points or 0):
        p.sparks_earned_total = int(p.total_points or 0)
    if not p.level:
        p.level = level_for_sparks(p.sparks_earned_total)["level"]
    if not p.freeze_capacity:
        p.freeze_capacity = current_study_tier(p.streak_count)["freeze_cap"]
    if p.badges in (None, ""):
        p.badges = "[]"
    if p.active_booster in (None, ""):
        p.active_booster = "null"
    if p.active_cosmetics in (None, ""):
        p.active_cosmetics = "{}"
    if p.weekly_quests in (None, ""):
        p.weekly_quests = "{}"
    if p.shop_purchases in (None, ""):
        p.shop_purchases = "[]"
    return p

@functools.lru_cache(maxsize=1)
def _asset_version() -> str:
    """Fingerprint for cache-busting /static URLs.

    Flask serves /static with a long max-age, so a deploy that changes a JS
    or CSS file leaves returning users running the previous one until their
    cache expires — a stale bundle that looks exactly like a bug that will
    not reproduce. Hashing the mtimes of the files we actually version gives
    a value that changes on deploy and stays put in between.

    Cached: this walks the filesystem, and it must not run per request.
    """
    try:
        base = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
        stamps = []
        for sub in ("css", "js"):
            d = os.path.join(base, sub)
            if not os.path.isdir(d):
                continue
            for root, _dirs, files in os.walk(d):
                for f in files:
                    if f.endswith((".css", ".js")):
                        stamps.append(str(int(os.path.getmtime(os.path.join(root, f)))))
        if not stamps:
            return "0"
        return hashlib.md5("|".join(sorted(stamps)).encode()).hexdigest()[:10]
    except Exception:
        # A missing fingerprint must never take down a page render.
        return "0"


@app.context_processor
def inject_asset_version():
    # In debug the fingerprint must not be cached: the lru_cache above is
    # what makes this cheap in production (one filesystem walk per
    # process), but it also means an edited stylesheet keeps the previous
    # ?v= for the life of the dev server — so the browser serves the old
    # file from cache and the edit appears not to have worked at all.
    if app.debug:
        _asset_version.cache_clear()
    return dict(asset_v=_asset_version())


@app.context_processor
def inject_auth():
    # Defensive: if load_user blows up
    # (e.g. mid-migration DB schema), we still need every template
    # render — including error.html — to succeed.
    try:
        return dict(logged_in=is_logged_in())
    except Exception:
        return dict(logged_in=False)

# ── SCHEDULE LOGIC ────────────────────────────────────────────
def infer_task_difficulty(points_possible, priority, due_date_str):
    score = float(points_possible or 0)
    try:
        due_date = datetime.fromisoformat(str(due_date_str)[:10])
        days_until_due = (due_date.date() - datetime.now().date()).days
    except ValueError:
        days_until_due = 7
    if priority == "High":
        score += 35
    elif priority == "Medium":
        score += 15
    if days_until_due <= 2:
        score += 20
    elif days_until_due <= 5:
        score += 10
    if score >= 110:
        return "Hard"
    if score >= 55:
        return "Medium"
    return "Easy"

def get_energy_profile(preferred_time):
    preference = (preferred_time or "evening").lower()
    profiles = {
        "morning": {"label": "morning", "summary": "Front-load harder work first.", "recommended_start_hour": 7, "hard_task_window": "7:00 AM - 11:00 AM", "light_task_window": "11:00 AM - 1:00 PM"},
        "afternoon": {"label": "afternoon", "summary": "Place demanding work first.", "recommended_start_hour": 1, "hard_task_window": "1:00 PM - 4:00 PM", "light_task_window": "4:00 PM - 6:00 PM"},
        "evening": {"label": "evening", "summary": "Begin with highest-focus work in early evening.", "recommended_start_hour": 6, "hard_task_window": "6:00 PM - 8:30 PM", "light_task_window": "8:30 PM - 10:30 PM"},
    }
    return profiles.get(preference, profiles["evening"])

def parse_time_slot_start(time_slot):
    if not time_slot or " - " not in time_slot:
        return None
    start_text = time_slot.split(" - ", 1)[0].strip()
    for fmt in ("%I:%M %p", "%I %p", "%H:%M", "%H"):
        try:
            return datetime.strptime(start_text, fmt)
        except ValueError:
            continue
    return None

def infer_block_energy_level(time_slot, preferred_time, difficulty):
    start_dt = parse_time_slot_start(time_slot)
    if start_dt is None:
        return "steady"
    hour = start_dt.hour
    preference = (preferred_time or "evening").lower()
    if preference == "morning":
        if hour < 10: return "peak"
        if hour < 13: return "steady"
        return "wind-down"
    if preference == "afternoon":
        if 13 <= hour < 16: return "peak"
        if 11 <= hour < 18: return "steady"
        return "wind-down"
    if 18 <= hour < 21: return "peak"
    if 16 <= hour < 22: return "steady"
    if difficulty == "Hard": return "steady"
    return "wind-down"

def build_daily_tip(workload_level, preferred_time, high_priority_count, hard_task_count):
    preference = (preferred_time or "evening").lower()
    if workload_level == "heavy":
        return f"Today is a heavier {preference} workload — protect your focus for the first block."
    if hard_task_count >= 2:
        return "Multiple demanding tasks today — clean starts and no distractions before each block."
    if high_priority_count >= 1:
        return "Knock out the urgent task first while your attention is strongest."
    return "Balanced day — finish each block fully and keep your momentum steady."

def classify_block_kind(title, course=""):
    """Infer the *type* of work in a block from its title + course.
    Drives the per-block icon, the checklist template, and the
    "Open in …" deep-link in the Interactive View.

    Returns one of:
      writing | math | reading | exam_prep | review | research |
      project | language | coding | general
    """
    t = f"{title or ''} {course or ''}".lower()
    pairs = [
        ("exam_prep", ("test", "exam", "midterm", "final", "quiz", "ap exam", "study guide")),
        ("writing",   ("essay", "paper", "writing", "draft", "thesis", "argument", "narrative", "reflection", "journal")),
        ("math",      ("math", "algebra", "geometry", "calc", "calculus", "trig", "statistics", "problem set", "equation", "homework problems")),
        ("coding",    ("code", "coding", "program", "python", "java ", "javascript", "leetcode", "debug", "cs ", "computer science")),
        ("language",  ("spanish", "french", "german", "latin", "chinese", "japanese", "vocab", "vocabulary", "conjugat", "translation")),
        ("reading",   ("read", "chapter", "novel", "textbook", "annotate", "literature", "passage")),
        ("research",  ("research", "sources", "annotated bibliography", "citation", "library", "data collection")),
        ("project",   ("project", "presentation", "slides", "poster", "lab report", "lab ", "design", "build")),
        ("review",    ("review", "revise", "go over", "flashcards", "spaced", "recap")),
    ]
    for kind, needles in pairs:
        if any(n in t for n in needles):
            return kind
    return "general"


# Deep-link map: Interactive View uses this to send the student to the
# right tool when they click "Open in …" on a checklist.
BLOCK_KIND_REDIRECT = {
    "writing":   {"label": "Open Writing Assistant", "href": "/writing",   "icon": "i-pencil"},
    "math":      {"label": "Open Math Explainer",    "href": "/math",      "icon": "i-modeler"},
    "reading":   {"label": "Open Lessons",           "href": "/lessons",   "icon": "i-lightbulb"},
    "exam_prep": {"label": "Open Tests Tracker",     "href": "/tests",     "icon": "i-test"},
    "review":    {"label": "Open Learn (spaced)",    "href": "/learn",     "icon": "i-study"},
    "research":  {"label": "Ask the Tutor",          "href": "/tutor",     "icon": "i-tutor"},
    "project":   {"label": "Open Deep Study",        "href": "/study",     "icon": "i-rocket"},
    "language":  {"label": "Open Learn (spaced)",    "href": "/learn",     "icon": "i-study"},
    "coding":    {"label": "Ask the Tutor",          "href": "/tutor",     "icon": "i-tutor"},
    "general":   {"label": "Start Focus Timer",      "href": "/focus",     "icon": "i-clock"},
}


def build_block_checklist(block, kind, assignment_meta):
    """Generate a summary + detailed checklist for a block. Templates are
    keyed by block kind so a writing block gets writing-shaped steps, a
    math block gets practice/check steps, etc. We seed step 1 with the
    block's own title so the checklist feels grounded in their task,
    not generic advice."""
    title = (block.get("assignment") or "this task").strip()
    duration = int(block.get("duration_minutes") or 25)
    short_title = title if len(title) <= 64 else title[:61] + "…"
    course = (block.get("course") or "").strip()
    # Each detailed step is { step, why } so the user understands intent.
    if block.get("is_break"):
        return {
            "summary": ["Step away from your screen", "Hydrate / quick stretch", "Reset your workspace"],
            "detailed": [
                {"step": "Stand up and look 20 feet away for 20 seconds",
                 "why":  "The 20-20-20 rule cuts eye strain so the next block has clean focus."},
                {"step": "Drink a glass of water",
                 "why":  "Mild dehydration reduces working memory measurably."},
                {"step": "Tidy your desk so only the next task's materials are visible",
                 "why":  "Visual clutter taxes attention before you even start."},
            ],
        }
    templates = {
        "writing": {
            "summary": [
                f"Re-read the prompt for '{short_title}'",
                "Outline 3 key points before drafting",
                "Draft without editing — momentum first",
                "Read aloud, mark awkward sentences",
                "Edit one pass for clarity, one pass for grammar",
            ],
            "detailed": [
                {"step": f"Open the prompt for '{short_title}' and rewrite it in your own words",
                 "why":  "Forces real comprehension before you write a single word."},
                {"step": "Bullet 3-5 main ideas and pick the 3 strongest",
                 "why":  "Picking before drafting prevents mid-paragraph thesis drift."},
                {"step": "Set a timer for half the block and draft without backspacing",
                 "why":  "Separating drafting from editing roughly doubles writing speed."},
                {"step": "Read the draft aloud — mark every sentence that trips your tongue",
                 "why":  "Awkward speech = awkward prose; your ear catches what your eye misses."},
                {"step": "Do one pass for argument flow, then one pass for grammar/typos",
                 "why":  "Single-purpose editing passes catch ~30% more issues than mixed passes."},
            ],
        },
        "math": {
            "summary": [
                "Re-do 1-2 example problems first",
                "Work the set, pencil only",
                "Check answers, circle wrong ones",
                "Redo wrong ones from scratch",
                "Note the rule you missed in a flashcard",
            ],
            "detailed": [
                {"step": "Pick 1-2 worked examples and solve them WITHOUT looking at the solution",
                 "why":  "Replicating a known solution from memory primes the same patterns for new problems."},
                {"step": "Work the actual problem set on paper — no calculator unless required",
                 "why":  "Writing math by hand activates the muscle memory you need on tests."},
                {"step": "Check answers; circle every wrong one (don't fix yet)",
                 "why":  "Batching errors lets you see if you're missing one concept across many problems."},
                {"step": "For each wrong problem, redo it from scratch on a clean line",
                 "why":  "Re-deriving > erasing — it forces you to find your real misstep."},
                {"step": "Write the missed rule on a flashcard / Learn deck for tomorrow",
                 "why":  "Spaced repetition is the only thing that locks math facts long-term."},
            ],
        },
        "reading": {
            "summary": [
                f"Skim the section of '{short_title}' first",
                "Annotate margin notes as you read",
                "Summarize each subsection in one line",
                "Write 2 questions the reading raised",
                "5-min recall: close the book and explain it aloud",
            ],
            "detailed": [
                {"step": "Skim headings, bold terms, and the first sentence of each paragraph",
                 "why":  "Building a mental map first triples retention on the careful read."},
                {"step": "Annotate margins with one-word tags (claim, evidence, question, ?)",
                 "why":  "Tagging forces active reading — you can't tag what you didn't process."},
                {"step": "After each subsection, write one sentence that summarizes it",
                 "why":  "Forces synthesis in the moment, when the context is still loaded."},
                {"step": "Write 2 questions the reading raised but didn't answer",
                 "why":  "Open questions become the best discussion / essay material."},
                {"step": "Close the book and recall the main points aloud or in writing",
                 "why":  "Free recall, even with errors, beats re-reading for memory by a wide margin."},
            ],
        },
        "exam_prep": {
            "summary": [
                "List topics you'll likely be tested on",
                "Self-test the weakest 2 topics first",
                "Redo 2-3 past problems / FRQs",
                "Make 1 cheat sheet (then put it away)",
                "Recall everything blind for 5 min",
            ],
            "detailed": [
                {"step": "Write all topics you expect on the exam, then rank each as Strong / Shaky / Weak",
                 "why":  "Studying what you already know feels good but does nothing for your score."},
                {"step": "Open a past test or FRQ for your two weakest topics — work them blind",
                 "why":  "Practice under the format of the real test transfers ~3× better than re-reading notes."},
                {"step": "Score yourself honestly; mark exactly which step broke",
                 "why":  "Most missed points come from one of three repeated process errors — find yours."},
                {"step": "Make a one-page cheat sheet by hand (you won't use it, that's the point)",
                 "why":  "Compression forces you to decide what's load-bearing knowledge."},
                {"step": "Close everything and brain-dump for 5 minutes",
                 "why":  "Retrieval practice the night before correlates strongly with exam performance."},
            ],
        },
        "review":  {
            "summary": [
                "Open your most recent notes / flashcards",
                "Do 10-15 active-recall reps",
                "Mark anything you guessed",
                "Rework the gaps in a fresh sentence",
                "Schedule the gaps for tomorrow",
            ],
            "detailed": [
                {"step": "Open the deck or notes for the topic — start cold, no warm-up",
                 "why":  "Friction at the start of recall is where the actual learning happens."},
                {"step": "Run 10-15 active-recall reps (Anki, Quizlet, or hand-quizzed)",
                 "why":  "Recalling is ~2× more effective per minute than re-reading."},
                {"step": "Mark every card you got but had to guess on",
                 "why":  "Guessed-right = not learned; don't let it pass."},
                {"step": "Rewrite each gap as one sentence, in your own words",
                 "why":  "If you can phrase it yourself, you'll recognize it on the test."},
                {"step": "Tag those cards for review again tomorrow",
                 "why":  "Spaced repetition only works if you actually space it."},
            ],
        },
        "research": {
            "summary": [
                "Define the question in one sentence",
                "Pull 2-3 credible sources",
                "Take notes in your own words",
                "Cite sources as you go (not at the end)",
                "Summarize what you found",
            ],
            "detailed": [
                {"step": "Write the question you're researching in one specific sentence",
                 "why":  "Broad questions waste hours; narrow ones finish themselves."},
                {"step": "Pull 2-3 sources from a library database or Google Scholar — not the open web",
                 "why":  "Curated sources are higher signal and easier to cite cleanly."},
                {"step": "Take notes in your own words only — never copy verbatim",
                 "why":  "Paraphrasing now is the cheapest plagiarism insurance later."},
                {"step": "Add a citation in your draft the moment you pull a fact",
                 "why":  "Reconstructing citations at the end is where most bibliographies break."},
                {"step": "End with a 3-bullet summary of what you actually found",
                 "why":  "If you can't summarize it, you don't have it yet."},
            ],
        },
        "project": {
            "summary": [
                "Look at the rubric / requirements",
                "Pick the next concrete deliverable",
                "Work the deliverable, not the project",
                "Save / commit progress at the end",
                "Note next step for tomorrow",
            ],
            "detailed": [
                {"step": "Re-read the rubric or assignment description top to bottom",
                 "why":  "Most lost points are for ignored rubric items, not bad work."},
                {"step": "Pick ONE concrete deliverable for this block (a slide, a section, a feature)",
                 "why":  "Projects stall on 'work on the project' — they move on shipped pieces."},
                {"step": "Work only the deliverable — defer every shiny side-quest to a notes file",
                 "why":  "Scope creep is the #1 reason projects miss deadlines."},
                {"step": "Save / export / commit before you close the block",
                 "why":  "Future-you should not have to remember what unsaved work was about."},
                {"step": "Write tomorrow's first step in one sentence",
                 "why":  "Starting the next block is the slowest part of any project."},
            ],
        },
        "language": {
            "summary": [
                "Warm up with 5 min of audio",
                "Run vocab flashcards (recall, not recognition)",
                "Practice one grammar pattern in 5 sentences",
                "Read or listen to a real text",
                "Speak / write 3 sentences using today's pattern",
            ],
            "detailed": [
                {"step": "Listen to 5 minutes of native audio (podcast, song with lyrics, news clip)",
                 "why":  "Tunes your ear before active practice — measurable comprehension boost."},
                {"step": "Run today's vocab as RECALL (target → English), not recognition",
                 "why":  "Recognition flatters you; recall is what tests measure."},
                {"step": "Write 5 sentences using today's grammar focus",
                 "why":  "Generating sentences cements rules faster than translating them."},
                {"step": "Read or listen to one paragraph of a real text",
                 "why":  "Untextbook input shows you how the language actually behaves."},
                {"step": "Say / write 3 sentences about your own day using today's pattern",
                 "why":  "Personal content is the easiest to remember — it sticks."},
            ],
        },
        "coding": {
            "summary": [
                "Re-read the problem / spec",
                "Sketch a plan before typing",
                "Write the test first if possible",
                "Implement smallest passing version",
                "Refactor and add edge cases",
            ],
            "detailed": [
                {"step": "Re-read the prompt / spec twice, in your own words once",
                 "why":  "Most coding mistakes start as misread requirements."},
                {"step": "Sketch a 3-line plan on paper before touching the keyboard",
                 "why":  "Coding without a plan is debugging with a plan — pay now or pay later."},
                {"step": "Write a failing test (or a sample call you expect to work)",
                 "why":  "A test makes 'done' a binary, not a feeling."},
                {"step": "Implement the dumbest version that could pass — no abstractions yet",
                 "why":  "Premature abstractions are the #1 source of bugs in student code."},
                {"step": "Add 2 edge cases (empty, big, weird) and refactor only after they pass",
                 "why":  "Refactoring red code makes both problems harder."},
            ],
        },
        "general": {
            "summary": [
                f"Re-read what '{short_title}' actually asks for",
                "Pick the first concrete sub-step",
                "Work it cleanly with no tabs open",
                "Capture progress at the end",
                "Note tomorrow's first move",
            ],
            "detailed": [
                {"step": f"Re-read '{short_title}' top to bottom — note what 'done' actually means",
                 "why":  "Most procrastination is fear of an unclear target."},
                {"step": "Pick a concrete first sub-step you can finish in 10 minutes",
                 "why":  "Starting is the costly part; small first steps make starting cheap."},
                {"step": f"Set a {min(duration, 45)}-min timer and work with only the needed tab open",
                 "why":  "Time-boxing + single-window keeps the block honest."},
                {"step": "At the end, save progress and tag where you stopped",
                 "why":  "Re-orienting next session is what kills momentum."},
                {"step": "Write the next first action in one sentence",
                 "why":  "Tomorrow's start is decided today or it's decided by mood."},
            ],
        },
    }
    return templates.get(kind, templates["general"])


def reallocate_days(schedule_data, availability, preferred_time,
                    commitments, dna, today=None):
    """Decide which day each block belongs on, deterministically.

    The model is good at naming the work, sizing it roughly, and writing
    the "what to actually do" notes. It is bad at the calendar: it cannot
    see how many free minutes each day holds, so it stacks several
    assignments on the first day, leaves later days empty, and puts long
    work the night before it is due. That is the whole reason plans did
    not feel spaced out.

    So the model's *day* assignment is discarded and recomputed here from
    real capacity, deadlines, distributed practice and load balance —
    while everything else it produced is kept intact and travels with the
    block. See scheduler_engine.allocate_across_days.

    Operates in place on ``schedule_data``. Returns the number of blocks
    that could not be placed inside the student's real free time.

    No-ops when there is no schedule, no day has a parseable date, or the
    student has no availability signal at all — in that last case there is
    no capacity model to allocate against, and the legacy behaviour is
    better than a guess.
    """
    days = schedule_data.get("schedule") or []
    if not days:
        return 0

    today = today or datetime.now().date()

    # Map each day entry to a real date, keeping the model's own ordering
    # as the fallback for entries it left undated.
    dated = []
    for idx, day in enumerate(days):
        d = _parse_schedule_day_date(day.get("date"), idx)
        if d is None:
            return 0
        dated.append((d, day))

    horizon = [d for d, _ in dated]
    # Real calendar events, not just the recurring commitments typed into
    # settings. Same source the per-day placement uses — if the two
    # disagreed, the allocator would budget time the placement pass then
    # refuses to fill, and the work would silently overflow.
    try:
        busy_by_date = _planner_busy_by_date()
    except Exception:
        busy_by_date = {}

    capacity = scheduler_engine.plan_capacity(
        min(horizon), (max(horizon) - min(horizon)).days + 1,
        availability, preferred_time, commitments,
        busy_by_date=busy_by_date,
    )
    if not any(capacity.values()):
        return 0

    # Flatten. Breaks are dropped: they are a property of how a day is laid
    # out, and the day is about to be laid out again — keeping them would
    # carry yesterday's rhythm onto a different set of blocks.
    tasks = []
    for _d, day in dated:
        for block in day.get("blocks") or []:
            if block.get("is_break"):
                continue
            task = dict(block)
            task.setdefault("title", block.get("assignment") or "")
            task["est_minutes"] = int(block.get("duration_minutes") or 30)
            tasks.append(task)
    if not tasks:
        return 0

    placed, unplaced = scheduler_engine.allocate_across_days(
        tasks, capacity, today, dna,
    )

    by_date = {d.isoformat(): day for d, day in dated}
    for iso, blocks in placed.items():
        day = by_date.get(iso)
        if day is None:
            continue
        rebuilt = []
        for b in blocks:
            out = dict(b)
            out["assignment"] = out.get("assignment") or out.get("title") or ""
            # "Essay (part 2 of 3)" reads as progress; "Essay" three times
            # reads as the planner repeating itself. Wording and fields match
            # split_oversized_blocks() so both split paths look the same.
            if out.get("part_total", 1) > 1:
                out["assignment"] = (
                    f"{out['parent_title']} (part {out['part_index']} "
                    f"of {out['part_total']})"
                )
            out.pop("est_minutes", None)
            out["time_slot"] = ""          # place_day_blocks assigns the clock
            rebuilt.append(out)
        day["blocks"] = rebuilt

    # Days the allocator gave nothing to must be emptied, not left holding
    # the model's original guess.
    for iso, day in by_date.items():
        if iso not in placed:
            day["blocks"] = []

    return len(unplaced)


def humanize_schedule(schedule_data, preferred_time, hours_per_day,
                      availability=None, commitments=None, dna=None):
    """Make the AI output look and feel like a real human study plan:
      - Enforce minimum transition gaps between blocks (no back-to-back).
      - Prevent two Hard tasks in a row — inject a break or a lighter task.
      - Cap the streak of work blocks before a real break.
      - Re-time blocks into the student's *real* free windows.
      - Attach kind / redirect / checklist data for the Interactive View.

    ``availability`` / ``commitments`` / ``dna`` are the personalization
    inputs. When all three are absent — guests, users who never filled in
    Settings → Availability — placement falls back to the legacy fixed-anchor
    behaviour so nothing regresses for them.

    Operates in place on schedule_data and returns it."""
    from datetime import timedelta as _td
    schedule = schedule_data.get("schedule", []) or []
    personalized = bool(availability) or dna is not None
    carry = []          # blocks pushed forward because a day ran out of time
    overflow_notes = []
    # Hours-per-day pacing → tighter buffers if the student only has 1h,
    # roomier ones if they have a long evening.
    base_gap = 5 if hours_per_day and hours_per_day <= 1.5 else 10
    # Minutes of continuous work before a 15-minute reset is inserted. Derived
    # from the student's own measured focus length rather than fixed at 90:
    # dna is right here and knows it, and asking someone who finishes
    # 25-minute blocks to work an hour and a half straight is the opposite of
    # planning around their habits. Falls back to 90 with no measured history.
    long_break_after = scheduler_engine.long_break_after_for(dna)

    # Take the calendar away from the model before laying anything out.
    # It chose which day each block sits on, and it chose badly — it has no
    # view of how many free minutes a day actually holds, so it front-loads
    # the week and leaves long work until the night before it is due.
    # reallocate_days() recomputes that from real capacity, deadlines,
    # spacing and load; the loop below then only has to arrange each day.
    if personalized:
        try:
            spilled = reallocate_days(
                schedule_data, availability, preferred_time, commitments, dna,
            )
            if spilled:
                overflow_notes.append(
                    f"{spilled} block(s) had no room in your free time this week."
                )
            schedule = schedule_data.get("schedule", []) or []
        except Exception as e:
            # Never let allocation take down plan generation: the model's
            # own day assignment is worse, but it is not nothing.
            print(f"[scheduler] day reallocation failed (non-fatal): {e}")

    next_block_id = 1
    for day_idx, day in enumerate(schedule):
        blocks = day.get("blocks", []) or []
        # 1. Anti-clustering: no two Hard work blocks back-to-back.
        i = 1
        while i < len(blocks):
            prev, cur = blocks[i - 1], blocks[i]
            if (not prev.get("is_break") and not cur.get("is_break")
                    and prev.get("difficulty") == "Hard" and cur.get("difficulty") == "Hard"):
                blocks.insert(i, {
                    "assignment": "Stretch break",
                    "course": "",
                    "duration_minutes": 10,
                    "time_slot": "",
                    "notes": "Two demanding tasks in a row — reset before the next one.",
                    "is_break": True,
                })
                i += 1
            i += 1
        # 2. Long-work-streak rule: more than long_break_after minutes of
        #    continuous study without a break → inject a 15-min break.
        #    Skipped on the personalized path — place_day_blocks() does this
        #    itself, against real clock time, and running both double-breaks.
        i = 0
        run = 0
        while i < len(blocks) and not personalized:
            b = blocks[i]
            if b.get("is_break"):
                run = 0
            else:
                run += int(b.get("duration_minutes") or 0)
                if run >= long_break_after and i + 1 < len(blocks) and not blocks[i + 1].get("is_break"):
                    blocks.insert(i + 1, {
                        "assignment": "Long break",
                        "course": "",
                        "duration_minutes": 15,
                        "time_slot": "",
                        "notes": "You've worked a solid stretch. Step away, eat, walk.",
                        "is_break": True,
                    })
                    run = 0
                    i += 1
            i += 1
        # 3. Re-time everything. We ignore the LLM's exact time strings —
        #    they're often inconsistent (e.g. "8 PM-8:45") — and place blocks
        #    ourselves. When we know the student's real availability we lay
        #    them inside those windows; otherwise we fall back to a fixed
        #    anchor derived from preferred_time.
        day_date = _parse_schedule_day_date(day.get("date"), day_idx)
        placed_ok = False
        if personalized and day_date is not None:
            try:
                windows = scheduler_engine.windows_for_date(
                    day_date, availability, preferred_time, commitments,
                )
                # Anything that didn't fit yesterday tries again today, ahead
                # of today's own work — a dropped task is worse than a late one.
                placed, spilled = scheduler_engine.place_day_blocks(
                    carry + blocks, windows, dna, long_break_after=long_break_after,
                )
                blocks, carry = placed, spilled
                placed_ok = True
                if spilled and windows:
                    overflow_notes.append(
                        f"{len(spilled)} block(s) didn't fit in your free time on "
                        f"{day_date:%a %b %d} — moved to the next available day."
                    )
                elif not windows:
                    # The student has no free time left on this date (the slots
                    # they marked have already passed, or are fully booked by a
                    # commitment). Carrying forward respects what they told us;
                    # falling back to an anchor hour would silently contradict it.
                    overflow_notes.append(
                        f"No free time on {day_date:%a %b %d} — that day's work "
                        f"moved to your next available day."
                    )
            except Exception as pe:
                print(f"[scheduler] personalized placement failed, using anchor: {pe}")
        if not placed_ok:
            profile = get_energy_profile(preferred_time)
            start_hour = profile["recommended_start_hour"]
            # `recommended_start_hour` is stored in 12-hour terms for the
            # afternoon (1 → 1 PM) and evening (6 → 6 PM) profiles. Without
            # this, every evening plan anchored day 2 onward at 6:00 AM.
            if preferred_time in ("afternoon", "evening") and start_hour <= 6:
                start_hour += 12
            cursor = datetime.now().replace(hour=start_hour, minute=0, second=0, microsecond=0)
            # If today's first day is *today*, push start forward to "now + 15min" rounded up.
            if day_idx == 0:
                now = datetime.now()
                soonest = now + _td(minutes=15)
                soonest = soonest.replace(minute=(soonest.minute // 5) * 5, second=0, microsecond=0)
                if soonest > cursor:
                    cursor = soonest
            def _fmt12(dt):
                # Cross-platform 12-hour formatting with no leading zero on the hour.
                return dt.strftime("%I:%M %p").lstrip("0") or "12:00 AM"
            for b_idx, b in enumerate(blocks):
                dur = int(b.get("duration_minutes") or 25)
                end = cursor + _td(minutes=dur)
                b["time_slot"] = f"{_fmt12(cursor)} - {_fmt12(end)}"
                b["start_iso"] = cursor.isoformat()
                b["end_iso"] = end.isoformat()
                # Next block starts after a transition gap (longer after a long block).
                gap = base_gap
                if dur >= 60: gap = max(gap, 10)
                cursor = end + _td(minutes=gap)
        # 4. Attach metadata used by the Interactive View.
        for b in blocks:
            b["block_id"] = f"d{day_idx + 1}-b{next_block_id}"
            next_block_id += 1
            kind = "break" if b.get("is_break") else classify_block_kind(b.get("assignment", ""), b.get("course", ""))
            b["kind"] = kind
            b["redirect"] = BLOCK_KIND_REDIRECT.get(kind) if kind != "break" else None
            b["checklist"] = build_block_checklist(b, kind, {})
        day["blocks"] = blocks
    # Anything still unplaced after the last day genuinely has nowhere to go
    # in the student's stated free time. Surface it instead of silently
    # dropping it — the honest answer is "you don't have room for this".
    if carry and schedule:
        for b in carry:
            b["unplaced"] = True
            b["time_slot"] = "No free time — needs rescheduling"
            b["block_id"] = f"unplaced-{next_block_id}"
            next_block_id += 1
            b.setdefault("kind", "study")
            b.setdefault("checklist", [])
            b["redirect"] = None
        schedule[-1].setdefault("blocks", []).extend(carry)
        overflow_notes.append(
            f"{len(carry)} block(s) could not fit anywhere in your available hours. "
            f"Add availability in Settings, or push a due date."
        )
    # Say which assignments got broken into sittings. The plan shows "Lab
    # report 1/3" without ever explaining where the other two came from, and a
    # student who cannot tell why their schedule changed shape stops trusting
    # it. One line per assignment, not per sitting.
    split_sizes = {}
    for day in schedule:
        for b in day.get("blocks", []) or []:
            # Every sitting of one split carries the same part_total, so a
            # plain assignment collapses them to one entry.
            if b.get("parent_title") and b.get("part_total"):
                split_sizes[b["parent_title"]] = b["part_total"]
    for title, parts in split_sizes.items():
        overflow_notes.append(
            f'"{title}" is longer than one sitting, so it is split into {parts} '
            f"shorter blocks that each match how long you actually focus."
        )

    if overflow_notes:
        schedule_data["placement_notes"] = overflow_notes
    # Placement can move blocks between days, so any per-day totals computed
    # before this point (enrich_schedule_data runs first) are now stale.
    if personalized:
        _recompute_day_totals(schedule_data, hours_per_day)
    return schedule_data


def _recompute_day_totals(schedule_data, hours_per_day=2):
    """Refresh per-day totals and workload labels after blocks were re-placed.

    enrich_schedule_data() runs before placement, so once blocks move between
    days its minute counts and "heavy/moderate/light" labels describe a plan
    that no longer exists.
    """
    total_study = 0
    for day in schedule_data.get("schedule", []) or []:
        blocks = day.get("blocks", []) or []
        study = sum(int(b.get("duration_minutes") or 0)
                    for b in blocks if not b.get("is_break") and not b.get("unplaced"))
        brk = sum(int(b.get("duration_minutes") or 0) for b in blocks if b.get("is_break"))
        total_study += study
        day["study_minutes"] = study
        day["break_minutes"] = brk
        day["total_minutes"] = study + brk
        day["total_hours"] = round((study + brk) / 60, 1)
        if study >= max(int(hours_per_day * 60 * 0.85), 150):
            level = "heavy"
        elif study >= max(int(hours_per_day * 60 * 0.55), 90):
            level = "moderate"
        else:
            level = "light"
        day["workload_level"] = level
        day["color_theme"] = WORKLOAD_COLORS[level]
    schedule_data["total_study_time"] = f"{total_study // 60} hours {total_study % 60} minutes"
    return schedule_data


def _repair_truncated_json(raw):
    """Best-effort close of a JSON document that was cut off mid-write.

    Walks the text tracking string/escape state and bracket depth, rewinds to
    the last structurally complete element, then closes whatever is still
    open. Returns None if there's nothing salvageable.

    This is a net, not a plan: with the thinking budget capped, truncation
    should be rare. But a schedule missing its last day beats an error page.
    """
    if not raw:
        return None
    in_string = escaped = False
    stack = []
    # Index just past the last point where we were at depth>=1 and had just
    # finished an element — a safe place to cut.
    safe_cut = None
    for i, ch in enumerate(raw):
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
        elif ch in "{[":
            stack.append(ch)
        elif ch in "}]":
            if stack:
                stack.pop()
            if len(stack) >= 1:
                safe_cut = i + 1
    if not stack:
        return None            # not actually truncated
    if safe_cut is None:
        return None            # nothing complete to keep
    head = raw[:safe_cut]
    # Re-derive what's still open at the cut point.
    in_string = escaped = False
    open_stack = []
    for ch in head:
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
        elif ch in "{[":
            open_stack.append(ch)
        elif ch in "}]":
            if open_stack:
                open_stack.pop()
    closers = "".join("}" if c == "{" else "]" for c in reversed(open_stack))
    try:
        return json.loads(head + closers)
    except json.JSONDecodeError:
        return None


def _parse_schedule_json(raw):
    """Turn a model response into a schedule dict, or raise.

    Handles, in order: code fences, plain JSON, a JSON object embedded in
    prose, and finally a truncated document. Raises ValueError if what comes
    back has no usable schedule, so the caller can retry rather than hand the
    student an empty plan.
    """
    text = re.sub(r"```json\n?", "", raw or "")
    text = re.sub(r"```\n?", "", text).strip()

    data = None
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        m = re.search(r"\{[\s\S]*\}", text)
        if m:
            try:
                data = json.loads(m.group(0))
            except json.JSONDecodeError:
                data = None
        if data is None:
            data = _repair_truncated_json(text)
            if data is not None:
                print("[scheduler] recovered a truncated schedule by repairing JSON")
    if data is None:
        raise json.JSONDecodeError("Could not parse a schedule", text[:200], 0)
    if not isinstance(data, dict):
        raise ValueError(f"Schedule response was {type(data).__name__}, expected object")
    days = data.get("schedule")
    if not isinstance(days, list) or not days:
        raise ValueError("Schedule response had no 'schedule' array")
    # A day list where nothing has blocks is not a plan.
    if not any((d or {}).get("blocks") for d in days if isinstance(d, dict)):
        raise ValueError("Schedule response had no blocks in any day")
    return data


def _parse_schedule_day_date(raw, day_idx):
    """Resolve a schedule day's "date" field to a real ``date``.

    The model is asked for YYYY-MM-DD but occasionally returns something else.
    Falls back to today + day_idx so placement still works on a malformed day
    rather than dropping the whole personalized path.
    """
    if raw:
        for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%Y/%m/%d"):
            try:
                return datetime.strptime(str(raw).strip()[:10], fmt).date()
            except ValueError:
                continue
    try:
        return (datetime.now() + timedelta(days=int(day_idx))).date()
    except Exception:
        return None


def reflow_schedule(schedule_data, availability=None, commitments=None, dna=None,
                    preferred_time="evening", hours_per_day=2):
    """Re-time a schedule the student has rearranged by hand.

    Unlike humanize_schedule() this preserves block order exactly as given —
    the student just dragged these into the arrangement they want, so
    reordering them again would fight the user. All this does is recompute
    clock times against their real free windows, flag anything that no longer
    fits, and refresh day totals.

    Returns the same dict, mutated, with ``placement_notes`` and per-block
    ``conflict`` flags describing what could not be honoured.
    """
    schedule = schedule_data.get("schedule", []) or []
    carry, notes = [], []
    personalized = bool(availability) or dna is not None
    for day_idx, day in enumerate(schedule):
        # Drop breaks we injected last time — the break rule runs again below,
        # so keeping them would grow the plan by one break per drag.
        blocks = scheduler_engine.strip_auto_breaks(
            [b for b in (day.get("blocks") or []) if b]
        )
        for b in blocks:
            b.pop("conflict", None)
            b.pop("unplaced", None)
        day_date = _parse_schedule_day_date(day.get("date"), day_idx)
        if not personalized or day_date is None:
            day["blocks"] = carry + blocks
            carry = []
            continue
        try:
            windows = scheduler_engine.windows_for_date(
                day_date, availability, preferred_time, commitments,
            )
            placed, spilled = scheduler_engine.place_day_blocks(
                carry + blocks, windows, dna, preserve_order=True,
                # Same interval the plan was generated with. Without this the
                # reflow used the module default, so a single drag re-spaced
                # every break in the day to a cadence the student never saw.
                long_break_after=scheduler_engine.long_break_after_for(dna),
            )
        except Exception as e:
            print(f"[reflow] placement failed for day {day_idx}: {e}")
            day["blocks"] = carry + blocks
            carry = []
            continue
        day["blocks"] = placed
        carry = spilled
        if spilled:
            for b in spilled:
                b["conflict"] = "no_room"
            label = "You have no free time" if not windows else "There isn't room"
            notes.append(
                f"{label} on {day_date:%a %b %d} for "
                f"{len(spilled)} block(s) — moved to your next available day."
            )
    if carry and schedule:
        for b in carry:
            b["conflict"] = "unplaceable"
            b["unplaced"] = True
            b["time_slot"] = "No free time — needs rescheduling"
        schedule[-1].setdefault("blocks", []).extend(carry)
        notes.append(
            f"{len(carry)} block(s) don't fit anywhere in your available hours. "
            f"Add availability in Settings, or move them to a different day."
        )
    schedule_data["placement_notes"] = notes
    if personalized:
        _recompute_day_totals(schedule_data, hours_per_day)
    return schedule_data


def enrich_schedule_data(schedule_data, assignments, preferred_time, hours_per_day):
    assignment_lookup = {item["title"]: item for item in assignments if isinstance(item, dict) and item.get("title")}
    schedule = schedule_data.get("schedule", [])
    total_study_minutes = 0
    for day in schedule:
        study_minutes = sum(block.get("duration_minutes", 0) for block in day.get("blocks", []) if not block.get("is_break"))
        break_minutes = sum(block.get("duration_minutes", 0) for block in day.get("blocks", []) if block.get("is_break"))
        total_minutes = study_minutes + break_minutes
        total_study_minutes += study_minutes
        if study_minutes >= max(int(hours_per_day * 60 * 0.85), 150):
            workload_level = "heavy"
        elif study_minutes >= max(int(hours_per_day * 60 * 0.55), 90):
            workload_level = "moderate"
        else:
            workload_level = "light"
        high_priority_count = 0
        hard_task_count = 0
        for block in day.get("blocks", []):
            if block.get("is_break"):
                block["color"] = "#cbd5e1"
                block["energy_level"] = "reset"
                block["difficulty"] = "Break"
                continue
            # Match on the parent title first. The planner labels a split
            # sitting "APUSH essay (part 2 of 3)", which matches no assignment,
            # so every split block used to fall through to Medium priority, a
            # re-inferred difficulty, and no due date — quietly discarding the
            # planner's own priority and the deadline the clock placer orders
            # the day by.
            assignment_meta = (
                assignment_lookup.get(block.get("parent_title") or "")
                or assignment_lookup.get(block.get("assignment", ""))
                or {}
            )
            priority = assignment_meta.get("priority", "Medium")
            difficulty = assignment_meta.get("difficulty") or infer_task_difficulty(assignment_meta.get("points_possible"), priority, assignment_meta.get("due_date"))
            energy_level = infer_block_energy_level(block.get("time_slot"), preferred_time, difficulty)
            if priority == "High": high_priority_count += 1
            if difficulty == "Hard": hard_task_count += 1
            # The planner's 0..100 score is finer-grained than the three
            # buckets the UI colours by. Keep both rather than overwriting one
            # with the other.
            if isinstance(block.get("priority"), (int, float)):
                block.setdefault("priority_score", int(block["priority"]))
            block["priority"] = priority
            block["difficulty"] = difficulty
            # The real deadline off the matched assignment, not the model's
            # opinion of it. place_day_blocks() orders the day by this, so an
            # assignment due tomorrow is never pushed past one due next week.
            # Blocks whose title the model rewrote miss the lookup and get no
            # due date, which the engine treats as "unknown", not "urgent".
            if assignment_meta.get("due_date"):
                block["due_date"] = assignment_meta["due_date"]
            block["energy_level"] = energy_level
            block["color"] = PRIORITY_COLORS.get(priority, "#60a5fa")
            block["accent_color"] = DIFFICULTY_COLORS.get(difficulty, "#60a5fa")
        day["workload_level"] = workload_level
        day["study_minutes"] = study_minutes
        day["break_minutes"] = break_minutes
        day["total_minutes"] = total_minutes
        day["color_theme"] = WORKLOAD_COLORS[workload_level]
        # Keep the model's day-specific tip — it can name the actual assignment
        # that matters today. build_daily_tip() is a four-branch template and
        # reads identically every day, so it's a fallback, not an override.
        ai_tip = (day.get("daily_tip") or "").strip()
        day["daily_tip"] = ai_tip or build_daily_tip(
            workload_level, preferred_time, high_priority_count, hard_task_count
        )
        if not day.get("total_hours"):
            day["total_hours"] = round(total_minutes / 60, 1)
    schedule_data["energy_profile"] = get_energy_profile(preferred_time)
    schedule_data["total_study_time"] = f"{total_study_minutes // 60} hours {total_study_minutes % 60} minutes"
    return schedule_data

# ── PAGE ROUTES ───────────────────────────────────────────────

@app.route("/api/stats")
def public_stats():
    """Returns live site stats for the landing page. No auth required."""
    try:
        total_users = User.query.count()
        total_assignments = DismissedAssignment.query.count()  # assignments completed
        total_study_sessions = StudySession.query.filter_by(completed=True).count()
        
        # Estimate hours saved: avg 8 min saved per assignment (not having to manually sort)
        hours_saved = round((total_assignments * 8) / 60)
        
        # Pad slightly for display (early traction looks better rounded up)
        display_users = max(total_users, 50)         # show at least 50
        display_assignments = max(total_assignments, 200)
        display_hours = max(hours_saved, 120)
        
        return jsonify({
            "students": display_users,
            "assignments_tracked": display_assignments,
            "hours_saved": display_hours,
            "study_sessions": total_study_sessions,
        })
    except Exception as e:
        print(f"Stats error: {e}")
        return jsonify({
            "students": 50,
            "assignments_tracked": 200,
            "hours_saved": 120,
            "study_sessions": 0,
        })

@app.route("/")
def landing():
    return render_template("landing.html", active_page="landing")

# ── SEO / AI-crawler static files ─────────────────────────────
@app.route("/robots.txt")
def robots_txt():
    return send_from_directory(app.static_folder, "robots.txt")

def _sitemap_lastmod_for(template_name: str | None, fallback: str) -> str:
    """Return an ISO-8601 lastmod for a sitemap URL.

    Prefer the linked template's mtime — it's the closest proxy for "when
    did this page actually change?". Falls back to ``fallback`` when the
    template can't be stat'd (e.g. dynamic routes, missing files)."""
    if template_name:
        try:
            path = os.path.join(app.template_folder or "", template_name)
            if os.path.exists(path):
                ts = os.path.getmtime(path)
                return datetime.utcfromtimestamp(ts).strftime("%Y-%m-%d")
        except Exception:
            pass
    return fallback


# Canonical sitemap. Each entry: (path, template_or_None, default_lastmod,
# changefreq, priority). Onboarding is intentionally absent — it's
# Disallowed in robots.txt, so listing it in the sitemap is a conflict.
_SITEMAP_ENTRIES = [
    # Public marketing
    ("/",                                "landing.html",                  "2026-06-22", "weekly",  "1.0"),
    ("/tutor",                           "tutor.html",                    "2026-06-22", "weekly",  "0.9"),
    ("/faq",                             "faq.html",                      "2026-06-22", "monthly", "0.9"),
    ("/pricing",                         "pricing.html",                  "2026-06-22", "monthly", "0.8"),
    ("/compare",                         "compare.html",                  "2026-06-22", "monthly", "0.8"),
    ("/about",                           "about.html",                    "2026-06-22", "monthly", "0.7"),
    ("/contact",                         "contact.html",                  "2026-06-22", "yearly",  "0.5"),
    ("/install",                         "install.html",                  "2026-06-22", "monthly", "0.6"),
    ("/download",                        "download.html",                 "2026-08-14", "weekly",  "0.8"),
    ("/legal",                           "legal.html",                    "2026-06-22", "yearly",  "0.3"),
    ("/ambassador",                      "ambassador.html",               "2026-06-22", "monthly", "0.7"),
    ("/schools",                         "schools.html",                  "2026-06-22", "monthly", "0.7"),
    ("/uk",                              "uk.html",                       "2026-06-22", "monthly", "0.7"),
    ("/api-docs",                        "api_docs.html",                 "2026-06-22", "monthly", "0.5"),
    # Blog
    ("/blog",                            "blog_index.html",               "2026-06-22", "weekly",  "0.8"),
    ("/blog/how-to-use-canvas-with-a-study-planner", "blog_canvas.html",  "2026-06-22", "monthly", "0.7"),
    ("/blog/studentvue-study-planner",   "blog_studentvue.html",          "2026-06-22", "monthly", "0.7"),
    ("/blog/how-to-prioritize-assignments", "blog_prioritize.html",       "2026-06-22", "monthly", "0.7"),
    ("/blog/best-ai-study-planner",      "blog_best_ai_planner.html",     "2026-06-22", "monthly", "0.8"),
    ("/blog/best-student-planner-app",   "blog_student_planner_app.html", "2026-06-22", "monthly", "0.8"),
    ("/blog/ai-notetaker-for-students",  "blog_ai_notetaker.html",        "2026-06-22", "monthly", "0.8"),
    ("/blog/what-is-a-good-gpa",         "blog_good_gpa.html",            "2026-06-22", "monthly", "0.8"),
    ("/blog/ap-study-planner",           "blog_ap_study_planner.html",    "2026-06-22", "monthly", "0.8"),
    # Compare
    ("/compare/intelliplan-vs-notion",       "compare_notion.html",       "2026-06-22", "monthly", "0.8"),
    ("/compare/intelliplan-vs-myhomework",   "compare_myhomework.html",   "2026-06-22", "monthly", "0.8"),
    ("/compare/intelliplan-vs-turbo-ai",     "compare_turbo.html",        "2026-06-22", "monthly", "0.8"),
    ("/compare/intelliplan-vs-quizlet",      "compare_quizlet.html",      "2026-06-22", "monthly", "0.8"),
    ("/compare/intelliplan-vs-mystudylife",  "compare_mystudylife.html",  "2026-06-22", "monthly", "0.8"),
    # Tools
    ("/tools/final-grade-calculator",    "tool_final_grade.html",         "2026-06-22", "monthly", "0.9"),
    ("/tools/gpa-calculator",            "tool_gpa.html",                 "2026-06-22", "monthly", "0.9"),
    ("/tools/grade-calculator",          "tool_grade.html",               "2026-06-22", "monthly", "0.9"),
    ("/tools/finals-countdown",          "tool_countdown.html",           "2026-06-22", "monthly", "0.7"),
    ("/tools/test-grade-calculator",     "tool_test_grade.html",          "2026-06-22", "monthly", "0.9"),
    ("/tools/study-schedule-maker",      "tool_schedule_maker.html",      "2026-06-22", "monthly", "0.9"),
    ("/tools/text-dissector",            "text_dissector.html",           "2026-06-22", "monthly", "0.8"),
    # App pages
    ("/olympiad",                        "olympiad.html",                 "2026-06-22", "monthly", "0.85"),
    ("/dashboard",                       "dashboard.html",                "2026-06-22", "daily",   "0.9"),
    ("/scheduler",                       "scheduler.html",                "2026-06-22", "daily",   "0.8"),
    ("/study",                           "study.html",                    "2026-06-22", "weekly",  "0.8"),
    ("/learn",                           "study.html",                    "2026-06-22", "weekly",  "0.8"),
    ("/study-and-learn",                 "study.html",                    "2026-06-22", "weekly",  "0.8"),
    ("/grademodel",                      "grademodel.html",               "2026-06-22", "weekly",  "0.8"),
    ("/grades",                          "gradebook.html",                "2026-06-22", "weekly",  "0.7"),
    ("/gradebook",                       "gradebook.html",                "2026-06-22", "weekly",  "0.7"),
    ("/classes",                         None,                            "2026-06-22", "weekly",  "0.7"),
    ("/streak",                          "streak.html",                   "2026-06-22", "weekly",  "0.7"),
    ("/focus",                           "focus.html",                    "2026-06-22", "weekly",  "0.7"),
    ("/lessons",                         None,                            "2026-06-22", "weekly",  "0.7"),
    ("/groups",                          "groups.html",                   "2026-06-22", "weekly",  "0.7"),
    ("/meetings",                        "meetings.html",                 "2026-06-22", "weekly",  "0.7"),
    ("/priority",                        None,                            "2026-06-22", "weekly",  "0.6"),
    ("/writing",                         None,                            "2026-06-22", "weekly",  "0.7"),
    ("/math",                            None,                            "2026-06-22", "weekly",  "0.7"),
    ("/extractor",                       None,                            "2026-06-22", "weekly",  "0.6"),
    ("/tests",                           None,                            "2026-06-22", "weekly",  "0.7"),
    ("/memories",                        None,                            "2026-06-22", "weekly",  "0.6"),
    ("/library",                         "library.html",                  "2026-06-22", "weekly",  "0.7"),
]


@app.route("/sitemap.xml")
def sitemap_xml():
    base = APP_BASE_URL.rstrip("/")
    parts = ['<?xml version="1.0" encoding="UTF-8"?>',
             '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">']
    for path, tpl, default_lm, changefreq, priority in _SITEMAP_ENTRIES:
        lastmod = _sitemap_lastmod_for(tpl, default_lm)
        parts.append(
            f"  <url><loc>{base}{path}</loc>"
            f"<lastmod>{lastmod}</lastmod>"
            f"<changefreq>{changefreq}</changefreq>"
            f"<priority>{priority}</priority></url>"
        )
    parts.append("</urlset>")
    xml = "\n".join(parts)
    return flask.Response(xml, mimetype="application/xml")

@app.route("/llms.txt")
def llms_txt():
    return send_from_directory(app.static_folder, "llms.txt", mimetype="text/plain")

INDEXNOW_KEY = os.getenv("INDEXNOW_KEY", "15d38c49db0d48efa4ec2ad2635b43c9").strip()
INDEXNOW_ENDPOINT = os.getenv("INDEXNOW_ENDPOINT", "https://api.indexnow.org/indexnow").strip()


def _indexnow_key_location():
    return os.getenv("INDEXNOW_KEY_LOCATION") or f"{APP_BASE_URL}/{INDEXNOW_KEY}.txt"


def _indexnow_host():
    parsed = urllib.parse.urlparse(APP_BASE_URL)
    return parsed.netloc or APP_DOMAIN


def _indexnow_sitemap_urls(limit=10000):
    """IndexNow source-of-truth: the same URL list the dynamic sitemap emits.

    Previously we re-read static/sitemap.xml; now that the sitemap is
    generated at request time we pull straight from _SITEMAP_ENTRIES so
    the two surfaces can never drift."""
    base = APP_BASE_URL.rstrip("/")
    urls = [f"{base}{path}" for path, *_ in _SITEMAP_ENTRIES]
    return urls[:limit]


def _indexnow_normalize_urls(urls):
    host = _indexnow_host().lower()
    out = []
    seen = set()
    for raw in urls or []:
        url = str(raw or "").strip()
        if not url:
            continue
        if url.startswith("/"):
            url = APP_BASE_URL.rstrip("/") + url
        parsed = urllib.parse.urlparse(url)
        if parsed.scheme not in {"http", "https"} or parsed.netloc.lower() != host:
            continue
        clean = urllib.parse.urlunparse((parsed.scheme, parsed.netloc, parsed.path or "/", "", parsed.query, ""))
        if clean not in seen:
            seen.add(clean)
            out.append(clean)
    return out[:10000]


def _indexnow_response(ok, http_status, submitted, message, method="POST", urls=None):
    return {
        "status": "ok" if ok else "error",
        "http_status": http_status,
        "submitted": submitted,
        "method": method,
        "endpoint": INDEXNOW_ENDPOINT,
        "keyLocation": _indexnow_key_location(),
        "host": _indexnow_host(),
        "urls": (urls or [])[:5],
        "message": message,
    }


def _submit_indexnow_single_url(url):
    """Submit one URL via IndexNow GET (?url=&key=&keyLocation=)."""
    urls = _indexnow_normalize_urls([url])
    if not INDEXNOW_KEY:
        return {"status": "error", "message": "IndexNow key is not configured.", "submitted": 0}
    if not urls:
        return {"status": "error", "message": "No valid URL for this host.", "submitted": 0}
    target = urls[0]
    params = {
        "url": target,
        "key": INDEXNOW_KEY,
        "keyLocation": _indexnow_key_location(),
    }
    try:
        r = requests.get(INDEXNOW_ENDPOINT, params=params, timeout=20)
        ok = 200 <= r.status_code < 300
        return _indexnow_response(
            ok,
            r.status_code,
            1 if ok else 0,
            "URL submitted to IndexNow." if ok else (r.text[:300] or "IndexNow rejected the request."),
            method="GET",
            urls=[target],
        )
    except Exception as e:
        return {"status": "error", "message": safe_error_message(e), "submitted": 0, "endpoint": INDEXNOW_ENDPOINT}


def _submit_indexnow_urls(urls):
    """Submit up to 10,000 URLs via IndexNow POST JSON bulk API."""
    urls = _indexnow_normalize_urls(urls)
    if not INDEXNOW_KEY:
        return {"status": "error", "message": "IndexNow key is not configured.", "submitted": 0}
    if not urls:
        return {"status": "error", "message": "No valid URLs for this host.", "submitted": 0}
    payload = {
        "host": _indexnow_host(),
        "key": INDEXNOW_KEY,
        "keyLocation": _indexnow_key_location(),
        "urlList": urls,
    }
    try:
        r = requests.post(
            INDEXNOW_ENDPOINT,
            json=payload,
            headers={"Content-Type": "application/json; charset=utf-8"},
            timeout=20,
        )
        ok = 200 <= r.status_code < 300
        return _indexnow_response(
            ok,
            r.status_code,
            len(urls) if ok else 0,
            "URLs submitted to IndexNow." if ok else (r.text[:300] or "IndexNow rejected the request."),
            method="POST",
            urls=urls,
        )
    except Exception as e:
        return {"status": "error", "message": safe_error_message(e), "submitted": 0, "endpoint": INDEXNOW_ENDPOINT}


def notify_indexnow(urls):
    """Notify search engines about changed URLs (single GET, bulk POST)."""
    if not urls:
        return {"status": "error", "message": "No URLs provided.", "submitted": 0}
    if isinstance(urls, str):
        urls = [urls]
    if len(urls) == 1:
        return _submit_indexnow_single_url(urls[0])
    return _submit_indexnow_urls(urls)


@app.route(f"/{INDEXNOW_KEY}.txt")
def indexnow_key_file():
    return flask.Response(INDEXNOW_KEY + "\n", mimetype="text/plain; charset=utf-8")


@app.route("/indexnow")
def indexnow_docs_page():
    return render_template(
        "indexnow.html",
        active_page="indexnow",
        indexnow_key=INDEXNOW_KEY,
        indexnow_key_location=_indexnow_key_location(),
        indexnow_host=_indexnow_host(),
        indexnow_endpoint=INDEXNOW_ENDPOINT,
        app_base_url=APP_BASE_URL,
    )


# ── SEO: /schedule is a permanent alias for /scheduler ──────────
# Google's Search Console previously flagged /schedule as "Page with
# redirect" because it 302'd or 404'd inconsistently. A clean 301 to
# the canonical /scheduler resolves the duplicate and lets the link
# equity flow through.
@app.route("/schedule")
def schedule_alias_redirect():
    return redirect("/scheduler", code=301)


# ── /privacy and /terms are canonical aliases for /legal ────────
# IntelliPlan keeps the Privacy Policy and Terms of Service together
# on a single /legal page (anchored at #privacy and #terms). Old
# inbound links (bookmarks, Google results, footers in saved emails)
# still hit /privacy or /terms, so we 301 them to the anchored URL
# on /legal — Google then consolidates the link equity onto /legal.
@app.route("/privacy")
def privacy_alias_redirect():
    return redirect("/legal#privacy", code=301)


@app.route("/terms")
def terms_alias_redirect():
    return redirect("/legal#terms", code=301)


# ── SEO: canonical + noindex helpers ────────────────────────────
# Public canonical host. Lock it to the apex HTTPS domain so any
# www./http variants are consolidated by the <link rel="canonical">
# emitted in base.html.
CANONICAL_HOST = "https://intelliplan.tech"

# Paths that should NEVER appear in Google's index. Headers (X-Robots-
# Tag) are set on responses for these, in addition to robots.txt
# Disallow rules, so even if a URL gets crawled before robots.txt
# is fetched (e.g. via an external backlink), Google honors noindex.
_NOINDEX_PREFIXES = (
    "/api/", "/push/", "/notifications/", "/cron/", "/oauth/",
    "/calendar/", "/debug/", "/feedback/", "/assignment/",
    "/admin", "/logout", "/live/", "/archive/",
)
_NOINDEX_EXACT = {
    "/login", "/register", "/login/account",
    "/login/canvas", "/login/studentvue", "/login/schoology",
    "/onboarding", "/connect", "/settings", "/dismissed", "/profiles",
    "/features", "/accessibility",
}


def _should_noindex(path):
    if not path:
        return False
    if path in _NOINDEX_EXACT:
        return True
    for pref in _NOINDEX_PREFIXES:
        if path.startswith(pref):
            return True
    return False


@app.context_processor
def _seo_context():
    """Make `canonical_url` + `noindex_page` available to every Jinja template
    so base.html can render the right <link rel="canonical"> and noindex meta
    without each view needing to set them explicitly."""
    try:
        path = request.path or "/"
    except Exception:
        path = "/"
    # Strip the trailing slash on canonical for paths beyond "/" so we don't
    # split traffic between e.g. /faq and /faq/.
    canon_path = path.rstrip("/") or "/"
    return {
        "canonical_url": CANONICAL_HOST + canon_path,
        "noindex_page": _should_noindex(path),
    }


@app.after_request
def _seo_headers(response):
    """Send X-Robots-Tag: noindex, nofollow on every authenticated page,
    every API endpoint, and every internal route. Belt-and-suspenders
    alongside robots.txt — header is authoritative even if a crawler
    skipped robots.txt for that path."""
    try:
        path = request.path or ""
    except Exception:
        return response
    if _should_noindex(path):
        # Use 'noindex, nofollow' — we don't want crawlers following links
        # out of authenticated pages (they'd hit /login redirects).
        existing = response.headers.get("X-Robots-Tag", "")
        if "noindex" not in existing.lower():
            response.headers["X-Robots-Tag"] = "noindex, nofollow"

    # ── Baseline security headers (professional hardening) ──────
    # NOTE: We deliberately do NOT set X-Frame-Options or a restrictive
    # Permissions-Policy — IntelliPlan is embedded in iframes (Lotus, etc.)
    # and uses camera/mic (Jitsi live sessions), notifications, and speech.
    # The headers below are safe everywhere and improve our security posture
    # + Google's "HTTPS / best practices" signals without breaking embeds.
    response.headers.setdefault("X-Content-Type-Options", "nosniff")
    response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
    # HSTS: 2 years, opt-in to includeSubDomains/preload via env so we don't
    # commit subdomain HTTPS before every subdomain is actually serving HTTPS.
    # Once verified, set HSTS_PRELOAD=1 in env to unlock preload eligibility.
    hsts = "max-age=63072000"
    if os.getenv("HSTS_INCLUDE_SUBDOMAINS", "1") == "1":
        hsts += "; includeSubDomains"
    if os.getenv("HSTS_PRELOAD") == "1":
        hsts += "; preload"
    response.headers["Strict-Transport-Security"] = hsts
    # Permissions-Policy: deny features we don't use anywhere. Camera/mic/
    # speaker-selection are NOT denied — Jitsi live sessions need them.
    # Payment is denied today; flip to 'self' if we ever ship Stripe Checkout
    # on our own origin.
    response.headers.setdefault("Permissions-Policy", (
        "geolocation=(), interest-cohort=(), payment=(), usb=(), "
        "magnetometer=(), gyroscope=(), accelerometer=(), "
        "ambient-light-sensor=(), battery=(), bluetooth=()"
    ))
    return response

@app.route("/tutor")
def tutor():
    logged_in = bool(session.get('logged_in') or (current_user.is_authenticated if hasattr(current_user, 'is_authenticated') else False))
    return render_template("tutor.html", active_page="tutor", logged_in=logged_in)

# ── Public info pages ─────────────────────────────────────────
@app.route("/faq")
def faq():
    return render_template("faq.html", active_page="faq")

@app.route("/compare")
def compare():
    return render_template("compare.html", active_page="compare")

@app.route("/pricing")
def pricing():
    return render_template("pricing.html", active_page="pricing")

# ── Blog / guides ──────────────────────────────────────────────
@app.route("/blog/how-to-use-canvas-with-a-study-planner")
def blog_canvas():
    return render_template("blog_canvas.html", active_page="blog")

@app.route("/blog/studentvue-study-planner")
def blog_studentvue():
    return render_template("blog_studentvue.html", active_page="blog")

@app.route("/blog/how-to-prioritize-assignments")
def blog_prioritize():
    return render_template("blog_prioritize.html", active_page="blog")

@app.route("/blog/best-ai-study-planner")
def blog_best_ai_planner():
    return render_template("blog_best_ai_planner.html", active_page="blog")

@app.route("/blog/best-student-planner-app")
def blog_student_planner_app():
    return render_template("blog_student_planner_app.html", active_page="blog")

@app.route("/compare/intelliplan-vs-notion")
def compare_notion():
    return render_template("compare_notion.html", active_page="compare")

@app.route("/compare/intelliplan-vs-myhomework")
def compare_myhomework():
    return render_template("compare_myhomework.html", active_page="compare")

@app.route("/compare/intelliplan-vs-turbo-ai")
def compare_turbo():
    return render_template("compare_turbo.html", active_page="compare")

@app.route("/compare/intelliplan-vs-quizlet")
def compare_quizlet():
    return render_template("compare_quizlet.html", active_page="compare")

@app.route("/compare/intelliplan-vs-mystudylife")
def compare_mystudylife():
    return render_template("compare_mystudylife.html", active_page="compare")

@app.route("/ambassador")
def ambassador():
    return render_template("ambassador.html", active_page="ambassador")

@app.route("/ambassador/dashboard")
def ambassador_dashboard():
    if not current_user.is_authenticated:
        return redirect(url_for("login"))
    return render_template("ambassador_dashboard.html", active_page="ambassador")

@app.route("/schools")
def schools():
    return render_template("schools.html", active_page="schools")

@app.route("/uk")
def uk_landing():
    return render_template("uk.html", active_page="uk")

@app.route("/blog/ai-notetaker-for-students")
def blog_ai_notetaker():
    return render_template("blog_ai_notetaker.html", active_page="blog")

@app.route("/blog/what-is-a-good-gpa")
def blog_good_gpa():
    return render_template("blog_good_gpa.html", active_page="blog")

@app.route("/blog/ap-study-planner")
def blog_ap_study_planner():
    return render_template("blog_ap_study_planner.html", active_page="blog")

@app.route("/blog")
def blog_index():
    return render_template("blog_index.html", active_page="blog")

@app.route("/olympiad")
def olympiad_page():
    return render_template("olympiad.html", active_page="olympiad")

@app.route("/tools/final-grade-calculator")
def tool_final_grade():
    return render_template("tool_final_grade.html", active_page="tools")

@app.route("/tools/gpa-calculator")
def tool_gpa():
    return render_template("tool_gpa.html", active_page="tools")

@app.route("/tools/grade-calculator")
def tool_grade():
    return render_template("tool_grade.html", active_page="tools")

@app.route("/tools/finals-countdown")
def tool_countdown():
    return render_template("tool_countdown.html", active_page="tools")

@app.route("/tools/test-grade-calculator")
def tool_test_grade():
    return render_template("tool_test_grade.html", active_page="tools")

@app.route("/tools/study-schedule-maker")
def tool_schedule_maker():
    return render_template("tool_schedule_maker.html", active_page="tools")

@app.route("/contact")
def contact():
    return render_template("contact.html", active_page="contact")

@app.route("/about")
def about():
    return render_template("about.html", active_page="about")

@app.route("/api-docs")
@app.route("/developers")
def api_docs_page():
    return render_template("api_docs.html", active_page="api_docs")

@app.route("/schedule")
def home():
    return redirect(url_for("dashboard"))

@app.route("/priority")
def priority():
    return render_template("priority.html", active_page="priority")

@app.route("/classes")
def classes():
    return render_template("classes.html", active_page="classes")

@app.route("/grades")
def grades():
    return render_template("grades.html", active_page="grades")

@app.route("/scheduler")
def scheduler():
    return render_template("scheduler.html", active_page="scheduler", load_saved=False)

@app.route("/scheduler/saved")
def scheduler_saved():
    return render_template("scheduler.html", active_page="scheduler", load_saved=True)

@app.route("/grademodel")
def grademodel():
    return render_template("grademodel.html", active_page="grademodel")

@app.route("/gradebook")
def gradebook():
    return render_template("gradebook.html", active_page="gradebook")

@app.route("/dismissed")
def dismissed_page():
    return render_template("dismissed.html", active_page="dismissed")

@app.route("/profiles")
def profiles():
    return render_template("profiles.html", active_page="profiles")

@app.route("/settings")
def settings():
    identity_dict = {
        "grade_level": "", "focus_areas": [], "goals": "",
        "completed": False, "availability": {},
        "weekly_commitments": "", "class_schedule": [],
    }
    if current_user.is_authenticated:
        try:
            identity = _get_or_create_identity(current_user.id)
            identity_dict = identity.to_dict()
        except Exception as e:
            print(f"Settings identity error: {e}")
    return render_template(
        "settings.html",
        active_page="settings",
        identity=identity_dict,
        grade_choices=GRADE_LEVEL_CHOICES,
        focus_choices=FOCUS_AREA_CHOICES,
    )

@app.route("/dashboard")
def dashboard():
    # First-run questionnaire: send the student to the dedicated /onboarding
    # page when their identity row hasn't been marked completed. Doing the
    # redirect here (rather than rendering an inline modal) is what makes
    # "Quick vs Customized" work cleanly and lets the AI-assisted flow have
    # a full page to work with. Guests skip — no identity row.
    grade_choices = GRADE_LEVEL_CHOICES if current_user.is_authenticated else []
    focus_choices = FOCUS_AREA_CHOICES if current_user.is_authenticated else []
    identity_dict = None
    if current_user.is_authenticated:
        identity = None
        try:
            identity = _get_or_create_identity(current_user.id)
        except Exception as _e:
            print(f"[dashboard] identity load failed: {_e}")
            try: db.session.rollback()
            except Exception: pass
            try:
                _run_boot_migration_once()
                identity = _get_or_create_identity(current_user.id)
            except Exception as _e2:
                print(f"[dashboard] identity retry failed: {_e2}")
                identity = None
        if identity is not None:
            if not bool(identity.completed):
                # Hand off to the dedicated onboarding page. Pass ?from=signup
                # so the page can show a richer welcome on the first hop.
                return redirect(url_for("onboarding"))
            try:
                identity_dict = identity.to_dict()
            except Exception:
                identity_dict = None
    return render_template(
        "dashboard.html",
        active_page="dashboard",
        needs_onboarding=False,  # legacy modal stays inert; gating is at the redirect now
        identity=identity_dict,
        grade_choices=grade_choices,
        focus_choices=focus_choices,
        phone=(current_user.phone if current_user.is_authenticated else "") or "",
        sms_opt_in=bool(getattr(current_user, "sms_reminders_opt_in", False)) if current_user.is_authenticated else False,
    )

@app.route("/study")
def study():
    """New Deep Study Pipeline — the 3-step setup → encoding → database workflow."""
    return render_template("deep_study.html", active_page="study")


@app.route("/learn")
def learn():
    """Classic Study & Learn experience — quiz generation, flashcards, mastery, etc."""
    return render_template("study.html", active_page="learn")


@app.route("/study-and-learn")
def study_and_learn_hub():
    """Landing page that replaces the cluttered Study & Learn sidebar group.
    Renders a grid of cards with descriptions for every study/learn surface,
    plus an AI chat bar at the top that recommends the best page for the
    user's current need (and can redirect them to it)."""
    return render_template("study_hub.html", active_page="study_hub")


@app.route("/api/study-hub/recommend", methods=["POST"])
def study_hub_recommend():
    """Lightweight intent-router for the Study & Learn hub chat bar.
    Takes a free-text user query and returns the best-matching feature
    plus a one-line reason. Uses keyword matching so it works without
    a paid LLM round-trip — and stays fast for the chat-bar UX."""
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "Not logged in"}), 401

    data  = request.get_json(silent=True) or {}
    query = (data.get("query") or "").strip().lower()
    if not query:
        return flask.jsonify({"status": "error", "message": "Query required"}), 400

    # Each entry: (page slug, route, title, score-weighted keywords)
    catalog = [
        ("tutor",     "/tutor",     "AI Tutor",        ["help","stuck","explain","tutor","question","ask","why","how","confused","understand"]),
        ("priority",  "/priority",  "Priority Queue",  ["priority","urgent","first","next","triage","important","focus","what to do"]),
        ("classes",   "/classes",   "Classes",         ["classes","course","class","subjects","schedule","period"]),
        ("grades",    "/grades",    "Grades",          ["grades","grade","gpa","report card","marks","score"]),
        ("tests",     "/tests",     "Tests & Quizzes", ["test","exam","quiz","midterm","final","prep","practice"]),
        ("dismissed", "/dismissed", "Completed Work",  ["done","completed","finished","dismissed","archive","past"]),
        ("study",     "/study",     "Deep Study",      ["deep","focus","session","intensive","study session","encoding","sprint"]),
        ("learn",     "/learn",     "Learn",           ["learn","flashcards","quiz me","review","mastery","memorize","practice","cards"]),
        ("focus",     "/focus",     "Focus Timer",     ["timer","pomodoro","focus time","25 minutes","countdown","stopwatch"]),
        ("library",   "/library",   "AP Library",      ["library","ap","resources","textbook","notes","reference"]),
        ("olympiad",  "/olympiad",  "Olympiad Preparer", ["olympiad","amc","aime","usamo","science olympiad","usaco","math contest","competition prep"]),
        ("streak",    "/streak",    "Streak",          ["streak","daily","habit","reward","points","consistency"]),
        ("lessons",   "/lessons",   "Lessons",         ["lessons","tutorial","walkthrough","recording","video lesson"]),
        ("writing",   "/writing",   "Writing",         ["writing","essay","paper","grammar","draft","proofread","editor"]),
        ("math",      "/math",      "Math Explainer",  ["math","equation","algebra","calculus","geometry","solve","derivative","integral"]),
        ("extractor", "/extractor", "Task Extractor",  ["extract","upload","syllabus","pdf","screenshot","import tasks","from image"]),
        ("meetings",  "/meetings",  "Meetings",        ["meeting","zoom","teams","google meet","webex","call","conference"]),
        ("groups",    "/groups",    "Study Groups",    ["group","groups","classmates","friends","peers","collaborate","together"]),
    ]

    best = None
    best_score = 0
    for slug, route, title, keywords in catalog:
        score = 0
        for kw in keywords:
            if kw in query:
                # Multi-word keywords score higher than single words.
                score += 2 if " " in kw else 1
        if slug in query or title.lower() in query:
            score += 3
        if score > best_score:
            best_score = score
            best = (slug, route, title)

    if not best:
        # Fallback: send them to the AI Tutor — it can handle any open-ended question.
        best = ("tutor", "/tutor", "AI Tutor")
        reason = "I'm not sure which page fits best — try asking the AI Tutor, it can guide you from there."
    else:
        reason_map = {
            "tutor":     "The AI Tutor can answer that directly.",
            "priority":  "The Priority Queue ranks your tasks so you know what to tackle first.",
            "classes":   "Classes shows everything for each course in one place.",
            "grades":    "Grades has your current scores and trends.",
            "tests":     "Tests & Quizzes has prep materials and practice exams.",
            "dismissed": "Completed Work is your archive of finished assignments.",
            "study":     "Deep Study runs a structured 3-step encoding session.",
            "learn":     "Learn generates flashcards and quizzes for mastery practice.",
            "focus":     "Focus Timer runs a Pomodoro session with break reminders.",
            "library":   "AP Library has curated resources and study guides.",
            "streak":    "Streak rewards consistent daily study habits.",
            "lessons":   "Lessons has uploaded recordings with AI summaries.",
            "writing":   "Writing helps with essays, grammar, and proofreading.",
            "math":      "Math Explainer walks through problems step by step.",
            "extractor": "Task Extractor imports assignments from PDFs and images.",
            "meetings":  "Meetings shows your upcoming Teams, Zoom, and Meet calls.",
            "groups":    "Study Groups lets you collaborate with classmates.",
        }
        reason = reason_map.get(best[0], "This page should help.")

    return flask.jsonify({
        "status":  "ok",
        "slug":    best[0],
        "route":   best[1],
        "title":   best[2],
        "reason":  reason,
        "score":   best_score,
    })

@app.route("/streak")
def streak():
    return render_template("streak.html", active_page="streak")

@app.route("/focus")
def focus():
    return render_template("focus.html", active_page="focus")


@app.route("/library")
def library():
    """AP & Exam content library — curated outlines, FRQs, and ready-made flashcard sets."""
    return render_template("library.html", active_page="library")


@app.route("/meetings")
def meetings():
    """Meetings hub: quick-launch Teams, Zoom, Google Meet, and user-managed class links."""
    return render_template("meetings.html", active_page="meetings")


# ════════════════════════════════════════════════════════════════
# DEEP STUDY PIPELINE — Step 0 (ingest) / Step 2 (voice coach feedback) /
# Step 3 (fact → active-recall card transform). All three endpoints share
# the Gemini/Groq AI backend used by the rest of the AI surface.
# ════════════════════════════════════════════════════════════════
def _deepstudy_extract_text_from_file(f):
    """Best-effort text extraction for the ingest endpoint. PDF/DOCX use
    optional packages with ASCII fallback; .txt and .md read directly."""
    name = (f.filename or "").lower()
    raw = f.read()
    if name.endswith((".txt", ".md")):
        try:
            return raw.decode("utf-8", errors="replace")
        except Exception:
            return raw.decode("latin-1", errors="replace")
    if name.endswith(".pdf"):
        try:
            import pdfplumber, io
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                return "\n".join(p.extract_text() or "" for p in pdf.pages[:25])
        except Exception:
            t = raw.decode("latin-1", errors="replace")
            return re.sub(r"[^\x20-\x7E\n\t]", " ", t)
    if name.endswith(".docx"):
        try:
            import docx, io
            d = docx.Document(io.BytesIO(raw))
            return "\n".join(p.text for p in d.paragraphs)
        except Exception:
            t = raw.decode("latin-1", errors="replace")
            return re.sub(r"[^\x20-\x7E\n\t]", " ", t)
    return raw.decode("utf-8", errors="replace")


@app.route("/api/deepstudy/ingest", methods=["POST"])
def api_deepstudy_ingest():
    """Step 0: extract source text and produce a topic + factual outline.

    Accepts multipart with optional `file` and `text` fields. Runs the
    extracted material through Gemini to produce a short topic label plus
    a list of atomic factual statements that drive Steps 1-3."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401

    text_in = (request.form.get("text") or "").strip()
    source_text = text_in
    if "file" in request.files and request.files["file"].filename:
        file_text = _deepstudy_extract_text_from_file(request.files["file"]).strip()
        if file_text:
            source_text = (source_text + "\n\n" + file_text).strip() if source_text else file_text

    if not source_text:
        return jsonify({"status": "error", "message": "Provide a file or describe your topic."}), 400

    truncated = source_text[:8000]
    if not ai_available():
        # Graceful fallback: split on sentences so the pipeline keeps working.
        sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", truncated) if 12 < len(s.strip()) < 280][:12]
        return jsonify({
            "status": "ok",
            "topic": (text_in or "Study Session")[:60],
            "source_text": source_text,
            "facts": sentences or [truncated[:240]],
        })

    prompt = (
        "You are a study coach. Given the material below, return ONLY valid JSON with two keys:\n"
        "  \"topic\": a short 2-6 word topic label\n"
        "  \"facts\": an array of 8-12 atomic factual statements (each ≤ 35 words) drawn from the material\n\n"
        f"MATERIAL:\n---\n{truncated}\n---\nJSON:"
    )
    try:
        parsed = ai_chat_json(
            [{"role": "user", "content": prompt}],
            tier="fast",
            temperature=0.2,
            max_tokens=1400,
        )
        topic = (parsed.get("topic") or text_in or "Study Session")[:80]
        facts = parsed.get("facts") or []
        if not isinstance(facts, list):
            facts = []
        facts = [str(f).strip() for f in facts if str(f).strip()][:14]
    except Exception as e:
        print(f"[deepstudy ingest] AI error: {e}")
        topic = (text_in or "Study Session")[:60]
        facts = [s.strip() for s in re.split(r"(?<=[.!?])\s+", truncated) if 12 < len(s.strip()) < 280][:12]

    return jsonify({
        "status": "ok",
        "topic": topic,
        "source_text": source_text,
        "facts": facts,
    })


@app.route("/api/deepstudy/feedback", methods=["POST"])
def api_deepstudy_feedback():
    """Step 2: voice coach reply. Reads the user's utterance, the mode
    (feynman vs blurting), and the ingested source text; returns a short
    spoken-style reply that pushes them back into the technique."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401

    data = request.get_json() or {}
    mode = (data.get("mode") or "feynman").lower()
    topic = (data.get("topic") or "this topic").strip()
    src   = (data.get("source_text") or "")[:3500]
    user_text = (data.get("user_text") or "").strip()
    if not user_text:
        return jsonify({"reply": "Keep going — I'm listening."})

    if not ai_available():
        return jsonify({"reply": "Got it — keep explaining. Try to simplify any jargon."})

    if mode == "feynman":
        system_msg = (
            f"You are Plani, a Feynman-technique voice coach. The student is studying '{topic}'. "
            "Your job is to make sure they explain concepts in PLAIN LANGUAGE, as if to a 10-year-old. "
            "If their utterance quotes textbook jargon or hides behind technical terms from the source, "
            "INTERRUPT gently and ask them to re-explain that exact phrase in everyday words. "
            "If they explain well, give one short specific encouragement and ask the next probing question. "
            "Always reply in 1-3 short sentences — this is a spoken voice call, not a lecture."
        )
    else:
        system_msg = (
            f"You are Plani, a Blurting-technique voice coach. The student is studying '{topic}' and just "
            "brain-dumped what they remember. Cross-reference against the SOURCE MATERIAL below. "
            "Reply in 2-4 short spoken-style sentences: name 1-2 specific things they nailed, then name 1-2 "
            "specific concepts FROM THE SOURCE they did not mention. Be encouraging but factual. "
            f"\n\nSOURCE MATERIAL:\n---\n{src}\n---"
        )

    try:
        reply = ai_chat(
            [
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_text},
            ],
            tier="fast",
            temperature=0.6,
            max_tokens=180,
        )
    except Exception as e:
        print(f"[deepstudy feedback] AI error: {e}")
        reply = "Good — keep going. Try to put that in your own words."

    return jsonify({"reply": reply})


@app.route("/api/deepstudy/transform", methods=["POST"])
def api_deepstudy_transform():
    """Step 3: convert raw source facts into active-recall question cards."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401

    data = request.get_json() or {}
    facts = data.get("facts") or []
    topic = (data.get("topic") or "this topic").strip()
    if not facts or not isinstance(facts, list):
        return jsonify({"status": "error", "message": "No facts to transform."}), 400

    facts_clean = [str(f).strip() for f in facts if str(f).strip()][:14]
    if not ai_available():
        # Fallback: heuristic question for each fact
        cards = [{
            "question": f"In your own words, explain: {f[:120]}",
            "answer": f,
        } for f in facts_clean]
        return jsonify({"status": "ok", "cards": cards})

    numbered = "\n".join(f"{i+1}. {f}" for i, f in enumerate(facts_clean))
    prompt = (
        f"You are an active-recall coach. Topic: {topic}.\n"
        "For each numbered passive fact below, write a SINGLE active-testing question that forces a student to "
        "retrieve the underlying mechanism or definition. The question must NOT contain the answer.\n"
        "Return ONLY valid JSON with key \"cards\": an array of objects with \"question\" and \"answer\" fields, "
        "in the same order as the input.\n\n"
        f"FACTS:\n{numbered}\n\nJSON:"
    )
    try:
        parsed = ai_chat_json(
            [{"role": "user", "content": prompt}],
            tier="fast",
            temperature=0.3,
            max_tokens=2000,
        )
        cards = parsed.get("cards") or []
        if not isinstance(cards, list):
            cards = []
        cards = [{
            "question": str(c.get("question", "")).strip(),
            "answer":   str(c.get("answer",   "")).strip(),
        } for c in cards if c.get("question")][:14]
    except Exception as e:
        print(f"[deepstudy transform] AI error: {e}")
        cards = [{"question": f"Explain in your own words: {f[:120]}", "answer": f} for f in facts_clean]

    return jsonify({"status": "ok", "cards": cards})


@app.route("/api/lms/connect/<provider>", methods=["POST"])
def api_lms_connect(provider):
    """Start an OAuth flow for an LMS provider (google_classroom, brightspace, moodle).

    Returns the auth URL if the provider's client credentials are configured,
    otherwise reports the provider as not-yet-enabled so the UI can show a
    coming-soon state instead of a hard error."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401

    provider = (provider or "").strip().lower()
    if provider not in ("google_classroom", "blackboard", "brightspace", "moodle"):
        return jsonify({"status": "error", "message": f"Unknown LMS provider: {provider}"}), 400

    # Moodle does not use OAuth: each institution issues a per-user web-services
    # token the user pastes in. Tell the UI to show the manual-connect form
    # instead of starting an OAuth redirect.
    if provider == "moodle":
        if os.getenv("MOODLE_ENABLED", "1").strip().lower() in {"0", "false", "no", "off"}:
            return jsonify({"status": "pending", "provider": provider,
                            "message": "Moodle support is launching soon."})
        return jsonify({"status": "manual", "provider": "moodle",
                        "form_endpoint": "/api/lms/connect/moodle/manual"})

    # Blackboard Learn is per-institution: each school has its own OAuth host.
    # The UI must POST {"institution_url": "https://learn.school.edu"} so we
    # know where to redirect the user. If the institution URL is missing, ask
    # the UI to prompt for it instead of failing silently.
    if provider == "blackboard":
        client_id = os.getenv("BLACKBOARD_CLIENT_ID")
        client_secret = os.getenv("BLACKBOARD_CLIENT_SECRET")
        if not client_id or not client_secret:
            print(f"[lms waitlist] {current_user.email} → blackboard")
            return jsonify({"status": "pending", "provider": "blackboard",
                            "message": "Blackboard support is launching soon. We'll email you when it's ready."})
        body = request.get_json(silent=True) or {}
        institution = (body.get("institution_url") or "").strip()
        if not institution:
            return jsonify({"status": "need_institution", "provider": "blackboard",
                            "message": "Enter your school's Blackboard URL (e.g. https://learn.myschool.edu)."})
        # Normalize and validate the institution URL.
        if not institution.startswith(("http://", "https://")):
            institution = "https://" + institution
        institution = institution.rstrip("/")
        try:
            parsed = urllib.parse.urlparse(institution)
            if not parsed.netloc or "." not in parsed.netloc:
                raise ValueError("bad host")
        except Exception:
            return jsonify({"status": "error",
                            "message": "That doesn't look like a valid Blackboard URL."}), 400
        redirect_uri = APP_BASE_URL + "/api/lms/callback/blackboard"
        state = secrets_module.token_urlsafe(24)
        session["lms_oauth_state"] = state
        session["lms_oauth_provider"] = "blackboard"
        session["blackboard_institution_url"] = institution
        # Blackboard's standard scope grants read access to the user's courses
        # and gradebook columns (assignments).
        scope = "read"
        auth_url = (
            f"{institution}/learn/api/public/v1/oauth2/authorizationcode"
            f"?client_id={urllib.parse.quote(client_id, safe='')}"
            f"&redirect_uri={urllib.parse.quote(redirect_uri, safe='')}"
            f"&response_type=code"
            f"&scope={urllib.parse.quote(scope, safe='')}"
            f"&state={urllib.parse.quote(state, safe='')}"
        )
        return jsonify({"status": "ok", "url": auth_url})

    # Remaining provider: google_classroom (already implemented above).
    config = {
        "google_classroom": {
            "client_id_env": "GOOGLE_CLASSROOM_CLIENT_ID",
            "scope": "https://www.googleapis.com/auth/classroom.courses.readonly "
                     "https://www.googleapis.com/auth/classroom.coursework.me.readonly",
            "auth_base": "https://accounts.google.com/o/oauth2/v2/auth",
        },
        "brightspace": {"client_id_env": "BRIGHTSPACE_CLIENT_ID", "auth_base": None},
    }[provider]

    client_id = os.getenv(config["client_id_env"])
    if not client_id or not config.get("auth_base"):
        # Record a waitlist signup so we can notify the user once the
        # integration is live (table reuses email_subscribers from the
        # marketing waitlist if available; otherwise we just log).
        print(f"[lms waitlist] {current_user.email} → {provider}")
        return jsonify({
            "status": "pending",
            "provider": provider,
            "message": f"{provider.replace('_', ' ').title()} support is launching soon. We'll email you when it's ready."
        })

    redirect_uri = APP_BASE_URL + f"/api/lms/callback/{provider}"
    state = secrets_module.token_urlsafe(24)
    session["lms_oauth_state"] = state
    session["lms_oauth_provider"] = provider
    auth_url = (
        f"{config['auth_base']}?client_id={urllib.parse.quote(client_id, safe='')}"
        f"&redirect_uri={urllib.parse.quote(redirect_uri, safe='')}"
        f"&response_type=code&access_type=offline&prompt=consent"
        f"&scope={urllib.parse.quote(config['scope'], safe='')}"
        f"&state={urllib.parse.quote(state, safe='')}"
    )
    return jsonify({"status": "ok", "url": auth_url})


# ── GOOGLE CLASSROOM HELPERS ──────────────────────────────────
def _classroom_token_endpoint():
    return "https://oauth2.googleapis.com/token"

def _classroom_exchange_code(code):
    """Exchange an auth code for a Google Classroom access + refresh token.
    Returns the parsed JSON response or raises."""
    client_id = os.getenv("GOOGLE_CLASSROOM_CLIENT_ID", "")
    client_secret = os.getenv("GOOGLE_CLASSROOM_CLIENT_SECRET", "")
    redirect_uri = APP_BASE_URL + "/api/lms/callback/google_classroom"
    r = requests.post(_classroom_token_endpoint(), data={
        "code": code,
        "client_id": client_id,
        "client_secret": client_secret,
        "redirect_uri": redirect_uri,
        "grant_type": "authorization_code",
    }, timeout=20)
    r.raise_for_status()
    return r.json()

def _classroom_refresh_token(refresh_token):
    """Refresh a Google Classroom access token using a stored refresh token."""
    client_id = os.getenv("GOOGLE_CLASSROOM_CLIENT_ID", "")
    client_secret = os.getenv("GOOGLE_CLASSROOM_CLIENT_SECRET", "")
    r = requests.post(_classroom_token_endpoint(), data={
        "refresh_token": refresh_token,
        "client_id": client_id,
        "client_secret": client_secret,
        "grant_type": "refresh_token",
    }, timeout=20)
    r.raise_for_status()
    return r.json()

def _classroom_access_token_for(user_id):
    """Return a valid access token for the given user, refreshing if needed.
    Returns None if no Classroom integration exists or refresh fails."""
    row = ClassroomIntegration.query.filter_by(user_id=user_id).order_by(ClassroomIntegration.id.desc()).first()
    if not row:
        return None, None
    now = utcnow()
    # Refresh if missing expiry or already expired (with 60s safety margin).
    needs_refresh = (not row.token_expires_at) or (row.token_expires_at <= now + timedelta(seconds=60))
    if needs_refresh and row.refresh_token:
        try:
            data = _classroom_refresh_token(row.refresh_token)
            row.access_token = data.get("access_token", row.access_token)
            ttl = int(data.get("expires_in", 3600))
            row.token_expires_at = now + timedelta(seconds=ttl)
            db.session.commit()
        except Exception as e:
            print(f"[classroom] refresh failed for user {user_id}: {e}")
            return None, row
    return row.access_token, row

def _classroom_get_userinfo(access_token):
    """Fetch the connected Google account's email + name."""
    try:
        r = requests.get(
            "https://www.googleapis.com/oauth2/v3/userinfo",
            headers={"Authorization": f"Bearer {access_token}"}, timeout=15,
        )
        if r.status_code == 200:
            return r.json()
    except Exception as e:
        print(f"[classroom] userinfo error: {e}")
    return {}

def _classroom_fetch_assignments(access_token):
    """Fetch active courses and their pending (or overdue) coursework for the
    connected user. Returns a list of dicts shaped to match the unified-tasks
    format used by the rest of IntelliPlan."""
    out = []
    headers = {"Authorization": f"Bearer {access_token}"}
    try:
        cr = requests.get(
            "https://classroom.googleapis.com/v1/courses",
            headers=headers, params={"courseStates": "ACTIVE", "pageSize": 50}, timeout=20,
        )
        if cr.status_code != 200:
            print(f"[classroom] courses HTTP {cr.status_code}: {cr.text[:200]}")
            return out
        courses = cr.json().get("courses", []) or []
    except Exception as e:
        print(f"[classroom] courses fetch failed: {e}")
        return out

    today = date.today() if 'date' in globals() else utcnow().date()
    for c in courses:
        cid = c.get("id"); cname = c.get("name") or "Google Classroom"
        if not cid:
            continue
        try:
            wr = requests.get(
                f"https://classroom.googleapis.com/v1/courses/{cid}/courseWork",
                headers=headers, params={"pageSize": 50}, timeout=20,
            )
            if wr.status_code != 200:
                continue
            items = wr.json().get("courseWork", []) or []
        except Exception as e:
            print(f"[classroom] coursework error for {cid}: {e}")
            continue
        for w in items:
            title = (w.get("title") or "").strip()
            if not title:
                continue
            dd = w.get("dueDate") or {}
            y, m, d = dd.get("year"), dd.get("month"), dd.get("day")
            if not (y and m and d):
                continue
            try:
                due = date(int(y), int(m), int(d))
            except Exception:
                continue
            days = (due - today).days
            if days < -14:
                continue
            pts = w.get("maxPoints") or 0
            priority = compute_priority(days, pts, title)
            est_minutes, description = _lms_row_sizing(w, pts)
            out.append({
                "id": f"gc-{w.get('id', '')}",
                "course_id": str(cid),
                "title": title,
                "course": cname,
                "due_date": due.strftime("%Y-%m-%d"),
                "priority": priority,
                "source": "google_classroom",
                "estimated_time": est_minutes,
                "description": description,
                "difficulty": "Medium",
                "color": PRIORITY_COLORS.get(priority, "#f59e0b"),
            })
    return out


# ── BLACKBOARD LEARN HELPERS ──────────────────────────────────
def _blackboard_basic_auth_header():
    cid = os.getenv("BLACKBOARD_CLIENT_ID", "")
    sec = os.getenv("BLACKBOARD_CLIENT_SECRET", "")
    raw = f"{cid}:{sec}".encode("utf-8")
    return "Basic " + base64.b64encode(raw).decode("ascii")

def _blackboard_exchange_code(institution_url, code):
    """Exchange an auth code for a Blackboard access + refresh token. Blackboard
    uses Basic auth with the client credentials and form-encoded body."""
    redirect_uri = APP_BASE_URL + "/api/lms/callback/blackboard"
    r = requests.post(
        f"{institution_url}/learn/api/public/v1/oauth2/token",
        headers={"Authorization": _blackboard_basic_auth_header(),
                 "Content-Type": "application/x-www-form-urlencoded"},
        data={"code": code, "redirect_uri": redirect_uri, "grant_type": "authorization_code"},
        timeout=20,
    )
    r.raise_for_status()
    return r.json()

def _blackboard_refresh_token(institution_url, refresh_token):
    r = requests.post(
        f"{institution_url}/learn/api/public/v1/oauth2/token",
        headers={"Authorization": _blackboard_basic_auth_header(),
                 "Content-Type": "application/x-www-form-urlencoded"},
        data={"refresh_token": refresh_token, "grant_type": "refresh_token"},
        timeout=20,
    )
    r.raise_for_status()
    return r.json()

def _blackboard_access_token_for(user_id):
    row = BlackboardIntegration.query.filter_by(user_id=user_id).order_by(BlackboardIntegration.id.desc()).first()
    if not row:
        return None, None
    now = utcnow()
    needs_refresh = (not row.token_expires_at) or (row.token_expires_at <= now + timedelta(seconds=60))
    if needs_refresh and row.refresh_token and row.institution_url:
        try:
            data = _blackboard_refresh_token(row.institution_url, row.refresh_token)
            row.access_token = data.get("access_token", row.access_token)
            ttl = int(data.get("expires_in", 3600))
            row.token_expires_at = now + timedelta(seconds=ttl)
            db.session.commit()
        except Exception as e:
            print(f"[blackboard] refresh failed for user {user_id}: {e}")
            return None, row
    return row.access_token, row

def _blackboard_get_userinfo(institution_url, access_token):
    try:
        r = requests.get(
            f"{institution_url}/learn/api/public/v1/users/me",
            headers={"Authorization": f"Bearer {access_token}"}, timeout=15,
        )
        if r.status_code == 200:
            return r.json()
    except Exception as e:
        print(f"[blackboard] userinfo error: {e}")
    return {}

def _blackboard_fetch_assignments(institution_url, access_token, bb_user_id=None):
    """Fetch active courses and their gradebook assignment columns for the user."""
    out = []
    headers = {"Authorization": f"Bearer {access_token}"}
    today = date.today()
    # Get the user's courses (Blackboard returns memberships).
    user_path = f"users/{bb_user_id}" if bb_user_id else "users/me"
    try:
        cr = requests.get(
            f"{institution_url}/learn/api/public/v1/{user_path}/courses",
            headers=headers, timeout=20,
        )
        if cr.status_code != 200:
            print(f"[blackboard] courses HTTP {cr.status_code}: {cr.text[:200]}")
            return out
        memberships = cr.json().get("results", []) or []
    except Exception as e:
        print(f"[blackboard] courses fetch failed: {e}")
        return out

    for m in memberships:
        cid = m.get("courseId") or m.get("course", {}).get("id")
        if not cid:
            continue
        # Fetch the course name (best effort) and gradebook columns.
        cname = "Course"
        try:
            cdetails = requests.get(
                f"{institution_url}/learn/api/public/v3/courses/{cid}",
                headers=headers, timeout=15,
            )
            if cdetails.status_code == 200:
                j = cdetails.json()
                cname = j.get("name") or j.get("displayName") or cname
        except Exception:
            pass
        try:
            gr = requests.get(
                f"{institution_url}/learn/api/public/v2/courses/{cid}/gradebook/columns",
                headers=headers,
                params={
                    "fields": "id,name,displayName,score.possible,grading.due,availability.available",
                    "limit": 100,
                },
                timeout=20,
            )
            if gr.status_code == 404:
                gr = requests.get(
                    f"{institution_url}/learn/api/public/v1/courses/{cid}/gradebook/columns",
                    headers=headers, timeout=20,
                )
            if gr.status_code != 200:
                continue
            cols = gr.json().get("results", []) or []
        except Exception as e:
            print(f"[blackboard] gradebook error for {cid}: {e}")
            continue
        for col in cols:
            availability = col.get("availability") or {}
            if str(availability.get("available") or "").lower() in {"no", "false", "disabled"}:
                continue
            name = (col.get("displayName") or col.get("name") or "").strip()
            grading = col.get("grading") or {}
            due_iso = grading.get("due") or col.get("due")
            if not name or not due_iso:
                continue
            try:
                due = datetime.fromisoformat(due_iso.replace("Z", "+00:00")).date()
            except Exception:
                continue
            days = (due - today).days
            if days < -14:
                continue
            score = col.get("score") or {}
            possible = score.get("possible") or 0
            priority = compute_priority(days, possible, name)
            est_minutes, description = _lms_row_sizing(col, possible)
            out.append({
                "id": f"bb-{cid}-{col.get('id', '')}",
                "course_id": str(cid),
                "title": name,
                "course": cname,
                "due_date": due.strftime("%Y-%m-%d"),
                "priority": priority,
                "source": "blackboard",
                "estimated_time": est_minutes,
                "description": description,
                "difficulty": "Medium",
                "color": PRIORITY_COLORS.get(priority, "#f59e0b"),
            })
    return out


# ── MOODLE HELPERS ────────────────────────────────────────────
def _moodle_call(moodle_url, ws_token, function, **params):
    """Call a Moodle web service function and return parsed JSON.
    Moodle returns a 200 OK with an "exception" field on errors, so we treat
    that as a failure too."""
    base = (moodle_url or "").rstrip("/") + "/webservice/rest/server.php"
    data = {"wstoken": ws_token, "wsfunction": function, "moodlewsrestformat": "json"}
    data.update(params)
    r = requests.post(base, data=data, timeout=20)
    r.raise_for_status()
    j = r.json()
    if isinstance(j, dict) and j.get("exception"):
        raise RuntimeError(f"Moodle error: {j.get('message') or j.get('errorcode')}")
    return j

def _moodle_fetch_assignments(moodle_url, ws_token, moodle_user_id=None):
    """Fetch upcoming Moodle assignments. Tries mod_assign_get_assignments first
    (per-course assignment list); falls back to core_calendar_get_action_events
    if not allowed. Both functions are normally exposed by the default Moodle
    mobile/web service, but each site can restrict them."""
    out = []
    today = date.today()
    # Step 1: list the user's courses so we can fetch their assignments.
    try:
        if moodle_user_id:
            courses = _moodle_call(moodle_url, ws_token,
                                   "core_enrol_get_users_courses", userid=int(moodle_user_id))
        else:
            site = _moodle_call(moodle_url, ws_token, "core_webservice_get_site_info")
            uid = site.get("userid")
            courses = _moodle_call(moodle_url, ws_token,
                                   "core_enrol_get_users_courses", userid=int(uid)) if uid else []
    except Exception as e:
        print(f"[moodle] courses fetch failed: {e}")
        courses = []
    if not isinstance(courses, list) or not courses:
        # Fallback to action events if courses can't be listed.
        try:
            evs = _moodle_call(moodle_url, ws_token,
                               "core_calendar_get_action_events_by_timesort",
                               timesortfrom=int(utcnow().timestamp()) - 14 * 86400,
                               limitnum=100)
            for e in (evs.get("events", []) or []) if isinstance(evs, dict) else []:
                ts = e.get("timesort") or e.get("timestart")
                if not ts:
                    continue
                due = datetime.utcfromtimestamp(int(ts)).date()
                days = (due - today).days
                if days < -14:
                    continue
                title = (e.get("name") or "").strip()
                if not title:
                    continue
                priority = compute_priority(days, 0, title)
                out.append({
                    "id": f"mdl-{e.get('id', '')}",
                    "course_id": str(e.get("course", {}).get("id") or ""),
                    "title": title,
                    "course": (e.get("course") or {}).get("fullname") or "Moodle",
                    "due_date": due.strftime("%Y-%m-%d"),
                    "priority": priority,
                    "source": "moodle",
                    "estimated_time": 60,
                    "difficulty": "Medium",
                    "color": PRIORITY_COLORS.get(priority, "#f59e0b"),
                })
        except Exception as e:
            print(f"[moodle] events fetch failed: {e}")
        return out

    # Step 2: fetch assignments for the courses we found.
    course_ids = []
    course_names = {}
    for c in courses:
        cid = c.get("id")
        if cid:
            course_ids.append(int(cid))
            course_names[int(cid)] = c.get("fullname") or c.get("shortname") or "Moodle"
    try:
        params = {}
        for i, cid in enumerate(course_ids[:50]):
            params[f"courseids[{i}]"] = cid
        result = _moodle_call(moodle_url, ws_token, "mod_assign_get_assignments", **params)
    except Exception as e:
        print(f"[moodle] assignments fetch failed: {e}")
        return out
    if not isinstance(result, dict):
        return out
    for course in result.get("courses", []) or []:
        cid = course.get("id")
        cname = course.get("fullname") or course_names.get(cid) or "Moodle"
        for a in course.get("assignments", []) or []:
            duedate = a.get("duedate") or 0
            if not duedate:
                continue
            try:
                due = datetime.utcfromtimestamp(int(duedate)).date()
            except Exception:
                continue
            days = (due - today).days
            if days < -14:
                continue
            title = (a.get("name") or "").strip()
            if not title:
                continue
            grade = a.get("grade") or 0
            priority = compute_priority(days, grade, title)
            est_minutes, description = _lms_row_sizing(a, grade)
            out.append({
                "id": f"mdl-{a.get('id', '')}",
                "course_id": str(cid or ""),
                "title": title,
                "course": cname,
                "due_date": due.strftime("%Y-%m-%d"),
                "priority": priority,
                "source": "moodle",
                "estimated_time": est_minutes,
                "description": description,
                "difficulty": "Medium",
                "color": PRIORITY_COLORS.get(priority, "#f59e0b"),
            })
    return out


@app.route("/api/lms/connect/moodle/manual", methods=["POST"])
def api_lms_connect_moodle_manual():
    """Manual Moodle connect: user pastes their institution URL + web-service
    token. We validate by calling core_webservice_get_site_info, which both
    confirms the token is valid and returns the user identity."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    body = request.get_json(silent=True) or {}
    moodle_url = (body.get("moodle_url") or "").strip()
    ws_token = (body.get("ws_token") or "").strip()
    if not moodle_url or not ws_token:
        return jsonify({"status": "error", "message": "Both Moodle URL and token are required."}), 400
    if not moodle_url.startswith(("http://", "https://")):
        moodle_url = "https://" + moodle_url
    moodle_url = moodle_url.rstrip("/")
    try:
        info = _moodle_call(moodle_url, ws_token, "core_webservice_get_site_info")
    except Exception as e:
        return jsonify({"status": "error",
                        "message": f"Could not verify with Moodle: {e}"}), 400
    row = MoodleIntegration.query.filter_by(user_id=current_user.id).order_by(MoodleIntegration.id.desc()).first()
    now = utcnow()
    if not row:
        row = MoodleIntegration(user_id=current_user.id, moodle_url=moodle_url, ws_token=ws_token)
        db.session.add(row)
    row.moodle_url = moodle_url
    row.ws_token = ws_token
    row.moodle_user_id = str(info.get("userid") or "") or None
    row.moodle_username = info.get("username") or None
    row.moodle_fullname = info.get("fullname") or None
    row.connected_at = now
    db.session.commit()
    return jsonify({"status": "ok",
                    "fullname": row.moodle_fullname,
                    "username": row.moodle_username,
                    "site": info.get("sitename")})


@app.route("/api/lms/disconnect/blackboard", methods=["POST"])
def api_blackboard_disconnect():
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    try:
        BlackboardIntegration.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        return jsonify({"status": "ok"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": safe_error_message(e)}), 500

@app.route("/api/lms/status/blackboard", methods=["GET"])
def api_blackboard_status():
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    row = BlackboardIntegration.query.filter_by(user_id=current_user.id).order_by(BlackboardIntegration.id.desc()).first()
    if not row:
        return jsonify({"status": "ok", "connected": False})
    return jsonify({"status": "ok", "connected": True,
                    "institution_url": row.institution_url,
                    "username": row.bb_username,
                    "connected_at": row.connected_at.isoformat() if row.connected_at else None})

@app.route("/api/lms/disconnect/moodle", methods=["POST"])
def api_moodle_disconnect():
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    try:
        MoodleIntegration.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        return jsonify({"status": "ok"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": safe_error_message(e)}), 500

@app.route("/api/lms/status/moodle", methods=["GET"])
def api_moodle_status():
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    row = MoodleIntegration.query.filter_by(user_id=current_user.id).order_by(MoodleIntegration.id.desc()).first()
    if not row:
        return jsonify({"status": "ok", "connected": False})
    return jsonify({"status": "ok", "connected": True,
                    "moodle_url": row.moodle_url,
                    "fullname": row.moodle_fullname,
                    "connected_at": row.connected_at.isoformat() if row.connected_at else None})


@app.route("/api/lms/callback/<provider>")
def api_lms_callback(provider):
    """OAuth callback for LMS providers. Stores the access token on the user
    record so subsequent syncs can pull courses and assignments."""
    if not current_user.is_authenticated:
        return redirect(url_for("login"))

    state = request.args.get("state", "")
    expected = session.pop("lms_oauth_state", None)
    if not state or state != expected:
        return render_template("error.html", error_code=400, error_id="LMS-STATE-MISMATCH",
                               message="OAuth state mismatch — please retry the connection."), 400
    code = request.args.get("code", "")
    if not code:
        return redirect("/connect?lms_error=1")

    if provider == "blackboard":
        institution = session.pop("blackboard_institution_url", None)
        if not institution:
            print("[blackboard callback] missing institution_url in session")
            return redirect("/connect?lms_error=1")
        try:
            tok = _blackboard_exchange_code(institution, code)
            access = tok.get("access_token")
            refresh = tok.get("refresh_token")
            ttl = int(tok.get("expires_in", 3600))
            if not access:
                print(f"[blackboard callback] no access_token: {tok}")
                return redirect("/connect?lms_error=1")
            info = _blackboard_get_userinfo(institution, access)
            row = BlackboardIntegration.query.filter_by(user_id=current_user.id).order_by(BlackboardIntegration.id.desc()).first()
            now = utcnow()
            if not row:
                row = BlackboardIntegration(user_id=current_user.id, institution_url=institution, access_token=access)
                db.session.add(row)
            row.institution_url = institution
            row.access_token = access
            if refresh:
                row.refresh_token = refresh
            row.token_expires_at = now + timedelta(seconds=ttl)
            row.bb_user_id = (info.get("id") or info.get("userId") or "") or row.bb_user_id
            row.bb_username = info.get("userName") or info.get("username") or row.bb_username
            row.connected_at = now
            db.session.commit()
            print(f"[blackboard callback] connected {row.bb_username or current_user.email}")
        except Exception as e:
            print(f"[blackboard callback] FAILED: {e}")
            return redirect("/connect?lms_error=1")
        return redirect("/connect?lms_connected=blackboard")

    if provider == "google_classroom":
        try:
            tok = _classroom_exchange_code(code)
            access = tok.get("access_token")
            refresh = tok.get("refresh_token")
            ttl = int(tok.get("expires_in", 3600))
            if not access:
                print(f"[classroom callback] no access_token in response: {tok}")
                return redirect("/connect?lms_error=1")
            info = _classroom_get_userinfo(access)
            row = ClassroomIntegration.query.filter_by(user_id=current_user.id).order_by(ClassroomIntegration.id.desc()).first()
            now = utcnow()
            if not row:
                row = ClassroomIntegration(user_id=current_user.id, access_token=access)
                db.session.add(row)
            row.access_token = access
            if refresh:  # Google only returns refresh on first consent
                row.refresh_token = refresh
            row.token_expires_at = now + timedelta(seconds=ttl)
            row.account_email = info.get("email") or row.account_email
            row.account_name = info.get("name") or row.account_name
            row.connected_at = now
            db.session.commit()
            print(f"[classroom callback] connected {row.account_email or current_user.email}")
        except Exception as e:
            print(f"[classroom callback] FAILED: {e}")
            return redirect("/connect?lms_error=1")
        return redirect("/connect?lms_connected=google_classroom")

    # Brightspace / Moodle still pending full backend.
    print(f"[lms callback] {current_user.email} authorized {provider} (storage pending)")
    return redirect(f"/connect?lms_connected={provider}")


@app.route("/api/lms/disconnect/google_classroom", methods=["POST"])
def api_classroom_disconnect():
    """Disconnect Google Classroom — deletes the stored token row."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    try:
        ClassroomIntegration.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        return jsonify({"status": "ok"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": safe_error_message(e)}), 500


@app.route("/api/lms/status/google_classroom", methods=["GET"])
def api_classroom_status():
    """Lightweight status check so the UI can render Connected / Not connected."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    row = ClassroomIntegration.query.filter_by(user_id=current_user.id).order_by(ClassroomIntegration.id.desc()).first()
    if not row:
        return jsonify({"status": "ok", "connected": False})
    return jsonify({
        "status": "ok",
        "connected": True,
        "account_email": row.account_email,
        "account_name": row.account_name,
        "connected_at": row.connected_at.isoformat() if row.connected_at else None,
    })

@app.route("/legal")
def legal():
    return render_template("legal.html", active_page="legal")

@app.route("/install")
def install():
    return render_template("install.html")

@app.route("/install/ios")
def install_ios():
    return render_template("install_ios.html")


@app.route("/download")
def download_desktop():
    """The desktop download page.

    Asset links come from the latest GitHub Release rather than from
    hardcoded filenames, because the filenames carry the version and would
    be wrong the day after every release. See desktop_releases.py.
    """
    import desktop_releases

    release = desktop_releases.latest_release()
    return render_template(
        "download.html",
        active_page="download",
        release=release,
        grouped=desktop_releases.by_platform(release),
        human_size=desktop_releases.human_size,
    )


@app.route("/api/desktop/latest")
def api_desktop_latest():
    """Machine-readable version of the same thing.

    The page uses this to fill in its download button after detecting the
    visitor's OS client-side, so the HTML itself stays cacheable and
    identical for everyone.
    """
    import desktop_releases

    release = desktop_releases.latest_release()
    return flask.jsonify({
        "status": "ok",
        "version": release.get("version"),
        "notes_url": release.get("notes_url"),
        "published_at": release.get("published_at"),
        "unavailable": bool(release.get("unavailable")),
        "platforms": desktop_releases.by_platform(release),
    })

# ── AUTH ROUTES ───────────────────────────────────────────────
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        return redirect(url_for("login_account"), 307)
    if is_logged_in():
        return redirect("/command-center")
    return render_template("login.html", active_page="login")

@app.route("/register", methods=["GET", "POST"])
def register():
    if current_user.is_authenticated:
        return redirect("/command-center")
    error = None
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "").strip()
        confirm = request.form.get("confirm_password", "").strip()
        if not email or not password:
            error = "Please fill in all fields."
        elif password != confirm:
            error = "Passwords do not match."
        elif len(password) < 8:
            error = "Password must be at least 8 characters."
        else:
            # Same defensive lookup as login — re-run the migration if
            # the SELECT trips on a missing column.
            try:
                _existing = User.query.filter_by(email=email).first()
            except Exception as _e:
                print(f"[register] User lookup failed: {_e}")
                try: db.session.rollback()
                except Exception: pass
                try:
                    _run_boot_migration_once()
                    _existing = User.query.filter_by(email=email).first()
                except Exception:
                    _existing = None
            if _existing:
                error = "An account with that email already exists."
        # ── COPPA: capture birth year and (if needed) parent email ──
        birth_year_raw = request.form.get("birth_year", "").strip()
        parent_email_raw = request.form.get("parent_email", "").strip().lower()
        birth_year_val = None
        age = None
        if not error and birth_year_raw:
            try:
                birth_year_val = int(birth_year_raw)
                current_year = utcnow().year
                if birth_year_val < 1900 or birth_year_val > current_year:
                    error = "Please enter a valid birth year."
                else:
                    age = current_year - birth_year_val
            except ValueError:
                error = "Please enter a valid birth year."
        elif not error:
            error = "Please tell us your birth year so we can comply with COPPA."
        if not error and age is not None and age < 13:
            if not parent_email_raw or "@" not in parent_email_raw:
                error = ("Because you're under 13, a parent or guardian's email is required. "
                         "We'll send them a one-time consent link before activating your account.")
            elif parent_email_raw == email:
                error = "Your parent or guardian's email must be different from your own."

        if not error:
            try:
                pw_hash = bcrypt.generate_password_hash(password).decode("utf-8")
                # Optional phone + SMS opt-in from the register form
                phone_raw = request.form.get("phone", "").strip()
                phone_norm = _normalise_phone(phone_raw) if phone_raw else None
                sms_optin = bool(request.form.get("sms_reminders_opt_in"))
                if not phone_norm:
                    sms_optin = False  # can't reminder without a number
                under_13 = (age is not None and age < 13)
                consent_token = secrets_module.token_urlsafe(24) if under_13 else None
                # Newsletters and onboarding email. Unticked means absent
                # from the form, which is the only reading of "no" a
                # checkbox has.
                marketing_optin = bool(request.form.get("marketing_emails_opt_in"))
                if under_13:
                    # COPPA: marketing to a child needs verifiable parental
                    # consent, and the consent gate below covers using the
                    # planner, not being sold to. Refusing here rather than
                    # deferring keeps a flag from sitting true and being
                    # picked up by an export later. Same shape as sms_optin
                    # above: no prerequisite, no opt-in.
                    marketing_optin = False
                user = User(
                    email=email, password_hash=pw_hash,
                    phone=phone_norm, sms_reminders_opt_in=sms_optin,
                    birth_year=birth_year_val,
                    parent_email=parent_email_raw if under_13 else None,
                    parent_consent_granted=not under_13,  # adults skip the gate
                    parent_consent_token=consent_token,
                    marketing_emails_opt_in=marketing_optin,
                    marketing_opt_in_at=utcnow() if marketing_optin else None,
                )
                db.session.add(user)
                db.session.commit()
                # Fire the parental consent email out-of-band.
                if under_13 and parent_email_raw:
                    try:
                        consent_url = f"{APP_BASE_URL}/parent/consent?token={consent_token}"
                        deny_url = f"{APP_BASE_URL}/parent/deny?token={consent_token}"
                        body = (
                            f"Hi,\n\nYour child ({email}) signed up for IntelliPlan, a free study "
                            f"planner. Because they're under 13, COPPA requires your consent before "
                            f"their account becomes active.\n\n"
                            f"What IntelliPlan does: helps students plan homework, prioritize "
                            f"assignments, and study with an AI tutor. No ads. No data sold. We collect "
                            f"only what's needed to run the planner (email, grade level, assignments) "
                            f"and you can request deletion at any time.\n\n"
                            f"✅ Approve the account:\n{consent_url}\n\n"
                            f"❌ Deny / delete this signup:\n{deny_url}\n\n"
                            f"If you didn't expect this email, you can either click the deny link "
                            f"above to remove the account, or simply ignore this message — the account "
                            f"stays locked and inactive until you approve it.\n\n— IntelliPlan"
                        )
                        # Send via SMTP if SMTP_HOST is set; otherwise log
                        # the link so it can still be delivered manually.
                        _send_email(parent_email_raw, "Consent needed: your child's IntelliPlan account", body)
                        print(f"[coppa] consent link emailed to {parent_email_raw}: {consent_url}")
                    except Exception as _e:
                        print(f"[coppa] consent email failed: {_e}")
                # Apply any pending referral (sets referred_by_id). Safe no-op if there's none.
                try:
                    _grant_referral_bonus(user)
                except Exception as _ref_e:
                    print(f"[referral] grant failed: {_ref_e}")
                # Under-13 accounts stay logged-out until the parent clicks
                # the consent link. Show a friendly "waiting for parent" page.
                if under_13:
                    return render_template(
                        "register.html",
                        active_page="login",
                        error=None,
                        info=(f"Account created for {email}. We've emailed a consent link to "
                              f"{parent_email_raw}. Your account will activate as soon as they approve it."),
                    )
                login_user(user, remember=True)
                # Best-effort welcome SMS so the user sees the integration
                # work the moment they opt in.
                if phone_norm and sms_optin:
                    # Best-effort welcome text via the email-to-SMS gateway.
                    try:
                        _sms_send_email_gateway(
                            phone_norm,
                            "Welcome to IntelliPlan! Open the app to customise your reminder times.",
                            carrier=(user.sms_carrier or "tmobile"),
                        )  # return value intentionally ignored here
                    except Exception: pass
                return redirect("/command-center")
            except Exception as _e:
                print(f"[register] user create failed: {_e}")
                try: db.session.rollback()
                except Exception: pass
                error = "Could not create that account right now — please try again."
    return render_template("register.html", active_page="login", error=error)


GRADE_LEVEL_CHOICES = [
    "6th grade", "7th grade", "8th grade",
    "9th grade", "10th grade", "11th grade", "12th grade",
    "Undergraduate", "Graduate", "Other",
]

FOCUS_AREA_CHOICES = [
    "Math", "Science", "Biology", "Chemistry", "Physics",
    "English / Literature", "History", "Foreign language",
    "Computer Science", "Economics", "Arts",
    "Test prep (SAT / ACT / AP)",
]


def _get_or_create_identity(user_id):
    identity = UserIdentity.query.filter_by(user_id=user_id).first()
    if not identity:
        identity = UserIdentity(user_id=user_id)
        db.session.add(identity)
        db.session.commit()
    return identity


def _ai_personalization_enabled(user=None):
    """True when the user has opted into AI personalization. Default: False.

    Guests and users who haven't toggled it on in Settings never get their
    grade history shipped to the AI provider. This is the privacy gate that
    governs build_student_context().
    """
    u = user if user is not None else (current_user if current_user.is_authenticated else None)
    if not u:
        return False
    return bool(getattr(u, "ai_personalization_opt_in", False))


def _summarize_grade_signals(grades_list, max_courses=12):
    """Distill a grades payload into the signals that actually help an LLM.

    Returns a dict: {"course_grades": [...], "strong": [...], "weak": [...],
    "average": float|None}. Designed to be cheap (no extra API calls) — the
    caller supplies the grades they already fetched, we just compress them
    into something a prompt can reference without leaking PII.
    """
    if not isinstance(grades_list, (list, tuple)) or not grades_list:
        return None
    rows = []
    for g in grades_list[:max_courses]:
        if not isinstance(g, dict):
            continue
        course = (g.get("course") or g.get("class_name") or g.get("name") or "").strip()
        if not course:
            continue
        pct = g.get("percentage")
        if pct is None:
            pct = g.get("grade_percent") or g.get("score")
        try:
            pct_val = float(str(pct).rstrip("%")) if pct is not None else None
        except (TypeError, ValueError):
            pct_val = None
        rows.append({"course": course[:60], "percent": pct_val,
                     "letter": (g.get("letter") or g.get("grade") or "")[:4]})
    if not rows:
        return None
    scored = [r for r in rows if r["percent"] is not None]
    avg = round(sum(r["percent"] for r in scored) / len(scored), 1) if scored else None
    strong = [r["course"] for r in sorted(scored, key=lambda r: -r["percent"])[:3] if r["percent"] >= 88]
    weak = [r["course"] for r in sorted(scored, key=lambda r: r["percent"])[:3] if r["percent"] < 78]
    return {"course_grades": rows, "strong": strong, "weak": weak, "average": avg}


def build_student_context(user_id=None, grades_summary=None, depth="full"):
    """Build the personalization prompt block for AI features.

    Returns an empty string if the user is a guest, hasn't opted in, or has
    no useful signal. Otherwise returns a compact section suitable for
    embedding in a system or user prompt.

    depth:
      - "full"  → identity + grades + class schedule (for scheduler)
      - "tutor" → identity + grades only (for chatbot/tutor)
      - "thin"  → identity only (for low-stakes features like writing tone)
    """
    try:
        uid = user_id if user_id is not None else (current_user.id if current_user.is_authenticated else None)
        if not uid:
            return ""
        u = db.session.get(User, uid)
        if not u or not _ai_personalization_enabled(u):
            return ""
        identity = _get_or_create_identity(uid)
        ident = identity.to_dict()
        bits = []
        if ident.get("grade_level"):
            bits.append(f"Grade level: {ident['grade_level']}")
        if ident.get("focus_areas"):
            bits.append(f"Academic focus: {', '.join(ident['focus_areas'][:6])}")
        if ident.get("goals"):
            bits.append(f"Stated goals: {ident['goals'][:240]}")
        if depth == "full" and ident.get("weekly_commitments"):
            bits.append(f"Weekly commitments: {ident['weekly_commitments'][:200]}")
        if depth == "full" and ident.get("availability"):
            av = "; ".join(f"{d}: {t}" for d, t in ident["availability"].items() if t)
            if av:
                bits.append(f"Availability: {av[:300]}")
        # Grade signals — only if the caller already gathered them, so we
        # never make an extra LMS request inside the prompt path.
        if grades_summary and isinstance(grades_summary, dict):
            if grades_summary.get("average") is not None:
                bits.append(f"Current overall average: {grades_summary['average']}%")
            if grades_summary.get("strong"):
                bits.append(f"Strongest subjects (>= 88%): {', '.join(grades_summary['strong'][:3])}")
            if grades_summary.get("weak"):
                bits.append(f"Subjects needing more time (< 78%): {', '.join(grades_summary['weak'][:3])}")
            cg = grades_summary.get("course_grades") or []
            if cg and depth == "full":
                line = ", ".join(
                    f"{r['course']} {int(r['percent'])}%" if r.get("percent") is not None else r["course"]
                    for r in cg[:8]
                )
                if line:
                    bits.append(f"Course grades: {line}")
        if not bits:
            return ""
        header = (
            "\n=== STUDENT CONTEXT (use to personalize, do NOT echo verbatim) ===\n"
            + "\n".join(f"  - {b}" for b in bits)
            + "\nGuidance: lean on strong subjects to build confidence; allocate more careful "
              "explanation and study time to weaker subjects; align tone to the student's stated goals."
            + "\n=== END STUDENT CONTEXT ===\n"
        )
        return header
    except Exception as _e:
        # Personalization is non-essential — never let it break an AI call.
        print(f"[personalization] build_student_context failed: {_e}")
        return ""


def normalized_custom_task_views(custom_tasks):
    """Give free-text custom tasks the same dict shape as an assignment.

    They arrive as bare title strings, but the clarification pass needs to
    look at course / estimate / due date uniformly across both kinds.
    """
    views = []
    for t in custom_tasks or []:
        if isinstance(t, dict):
            views.append(dict(t))
        else:
            views.append({"title": str(t or "").strip(), "course": "",
                          "estimated_time": 0, "due_date": ""})
    return views


def _preset_query(user_id=None, guest_id=None):
    q = SchedulerPreset.query
    return q.filter_by(user_id=user_id) if user_id else q.filter_by(guest_session_id=guest_id)


def load_scheduler_presets(user_id=None, guest_id=None):
    """Saved clarification answers for this owner, keyed by task key."""
    try:
        return {p.task_key: p.to_dict() for p in _preset_query(user_id, guest_id).all()}
    except Exception as e:
        print(f"[clarify] preset load failed (non-fatal): {e}")
        return {}


def save_scheduler_presets(payload, user_id=None, guest_id=None):
    """Upsert clarification answers so we can offer them back next time."""
    if not payload:
        return 0
    saved = 0
    try:
        for key, entry in payload.items():
            row = _preset_query(user_id, guest_id).filter_by(task_key=key).first()
            if row is None:
                row = SchedulerPreset(user_id=user_id, guest_session_id=guest_id, task_key=key)
                db.session.add(row)
            row.label = str(entry.get("label") or key)[:200]
            # Merge rather than replace: a later partial answer shouldn't wipe
            # fields the student already filled in.
            merged = {**row.answers(), **(entry.get("answers") or {})}
            row.answers_json = json.dumps(merged)
            row.last_used_at = utcnow()
            saved += 1
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"[clarify] preset save failed (non-fatal): {e}")
        return 0
    return saved


def _mark_presets_used(keys, user_id=None, guest_id=None):
    if not keys:
        return
    try:
        for row in _preset_query(user_id, guest_id).filter(SchedulerPreset.task_key.in_(list(keys))).all():
            row.times_used = (row.times_used or 0) + 1
            row.last_used_at = utcnow()
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"[clarify] preset usage bump failed (non-fatal): {e}")


@app.route("/scheduler/presets", methods=["GET"])
def list_scheduler_presets():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    presets = load_scheduler_presets(uid, gid)
    return flask.jsonify({"status": "ok", "presets": presets})


@app.route("/scheduler/presets/delete", methods=["POST"])
def delete_scheduler_preset():
    """Forget a saved answer. Students change classes; presets must be undoable."""
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    key = (request.json or {}).get("task_key")
    if not key:
        return flask.jsonify({"status": "error", "message": "task_key required"}), 400
    try:
        n = _preset_query(uid, gid).filter_by(task_key=key).delete()
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500
    return flask.jsonify({"status": "ok", "deleted": n})


def build_scheduler_personalization(user_id=None, guest_id=None, feedback_limit=400):
    """Gather everything the placement engine needs for one student.

    Returns ``(dna, availability, commitments)``. Any piece may be empty —
    the engine and humanize_schedule() both degrade to the pre-personalization
    behaviour when a signal is missing, so this never has to raise.

    Unlike build_student_context(), this is NOT gated on the AI-personalization
    opt-in: none of it leaves the server. Availability and completion history
    shape *where blocks land locally*; the opt-in gate governs what we ship to
    the AI provider, which is handled separately by StudyDNA.to_prompt().
    """
    dna = scheduler_engine.StudyDNA()
    availability, commitments = {}, ""
    try:
        q = TaskFeedback.query
        q = q.filter_by(user_id=user_id) if user_id else q.filter_by(guest_session_id=guest_id)
        rows = q.order_by(TaskFeedback.id.desc()).limit(feedback_limit).all()
        feedback = [{
            "estimated_time": r.estimated_time, "actual_time": r.actual_time,
            "course": r.course, "day_of_week": r.day_of_week,
            "time_of_day": r.time_of_day,
        } for r in rows]

        sq = SavedSchedule.query
        sq = sq.filter_by(user_id=user_id) if user_id else sq.filter_by(guest_session_id=guest_id)
        recent = sq.order_by(SavedSchedule.created_at.desc()).limit(10).all()
        progress = [
            scheduler_engine.summarize_progress(s.progress_json)
            for s in recent if s.progress_json
        ]

        # Real timed sittings from /active. The app has been measuring how
        # long students actually work, when they work, and when they lose
        # focus — and then planning their week from numbers they typed into
        # an estimate box, because none of it was ever read back here. This
        # is the single largest source of genuine personalization available,
        # and it was on the floor.
        sessions = []
        try:
            asq = ActiveSession.query.filter(ActiveSession.state != "running")
            asq = (asq.filter_by(user_id=user_id) if user_id
                   else asq.filter_by(guest_session_id=guest_id))
            for s in asq.order_by(ActiveSession.started_at.desc()).limit(120).all():
                started = s.started_at
                sessions.append({
                    "planned_minutes": s.planned_minutes or 0,
                    "active_minutes": s.active_minutes,
                    "course": s.course or "",
                    "completed_work": bool(s.completed_work),
                    "distraction_events": s.distraction_events or 0,
                    # Only meaningful when the focus check-in actually ran —
                    # a zero from a session without it is an absence of
                    # measurement, not a measurement of zero.
                    "focus_streak_minutes": (
                        int(round((s.longest_focus_streak or 0) / 60.0))
                        if s.focus_enabled and s.longest_focus_streak else 0
                    ),
                    "day_of_week": started.strftime("%a") if started else "",
                    "time_of_day": (
                        scheduler_engine.slot_for_hour(started.hour) if started else ""
                    ),
                })
        except Exception as e:
            print(f"[scheduler] session history unavailable (non-fatal): {e}")

        dna = scheduler_engine.build_study_dna(feedback, progress, sessions)
    except Exception as e:
        print(f"[scheduler] study DNA build failed (non-fatal): {e}")

    if user_id:
        try:
            identity = _get_or_create_identity(user_id)
            availability = identity.avail_dict() or {}
            commitments = identity.weekly_commitments or ""
        except Exception as e:
            print(f"[scheduler] identity load failed (non-fatal): {e}")
    return dna, availability, commitments


def _fetch_grades_for_personalization():
    """Best-effort grades fetch used by AI endpoints that don't already have them.

    Returns the same shape as /grades/data, or an empty list on any failure.
    Wrapped so caller code can stay one-line: `_summarize_grade_signals(_fetch_grades_for_personalization())`.
    """
    try:
        acct = get_grade_account()
        if acct:
            lt = acct.get("login_type")
            if lt == "studentvue":
                from studentvue_helper import get_grades as _sv
                return _sv(acct["sv_district_url"], acct["sv_username"], acct["sv_password"]) or []
            if lt == "schoology":
                from schoology_helper import get_schoology_grades as _sc
                return _sc(acct["schoology_key"], acct["schoology_secret"]) or []
            if lt == "canvas":
                from canvas_helper import get_grades as _cv
                return _cv(acct.get("canvas_url", "https://canvas.instructure.com"), acct["canvas_token"]) or []
        # No LMS — fall back to imported CSV / paste / scraper grades so AI
        # personalization still works for unsupported-LMS students.
        return _imported_grades_payload()
    except Exception as _e:
        print(f"[personalization] grades fetch failed: {_e}")
    return []


@app.route("/onboarding", methods=["GET", "POST"])
def onboarding():
    if not is_logged_in():
        return redirect(url_for("login"))
    try:
        identity = _get_or_create_identity(current_user.id)
    except Exception as _e:
        # The dashboard modal handles the questionnaire too — if the
        # legacy /onboarding page can't load (e.g. column missing on a
        # not-yet-fully-migrated DB), forward to dashboard and let the
        # modal handle it from there. Don't 500.
        print(f"[onboarding] identity load failed: {_e}")
        try: db.session.rollback()
        except Exception: pass
        try:
            _run_boot_migration_once()
            identity = _get_or_create_identity(current_user.id)
        except Exception as _e2:
            print(f"[onboarding] retry failed: {_e2}")
            return redirect("/command-center")
    if request.method == "POST":
        # Legacy form path — preserved so old browsers/users with cached JS
        # still complete onboarding even when the new SPA-style flow can't
        # run (e.g. JS disabled). Mirrors the new quick path.
        grade = (request.form.get("grade_level") or "").strip()[:32]
        focus = request.form.getlist("focus_areas")
        focus = [f.strip()[:48] for f in focus if f.strip()][:12]
        goals = (request.form.get("goals") or "").strip()[:1000]
        identity.grade_level = grade or None
        identity.focus_areas = json.dumps(focus)
        identity.goals = goals
        identity.completed = True
        db.session.commit()
        # Clear the resume-state so we don't bounce them back into the chat
        # next time they hit /onboarding (a real bug in the prior flow).
        for k in ("onb_mode", "onb_step", "onb_messages"):
            session.pop(k, None)
        session.modified = True
        next_url = request.args.get("next")
        if next_url and next_url.startswith("/"):
            return redirect(next_url)
        return redirect("/command-center")
    return render_template(
        "onboarding.html",
        active_page="onboarding",
        identity=identity.to_dict(),
        grade_choices=GRADE_LEVEL_CHOICES,
        focus_choices=FOCUS_AREA_CHOICES,
    )


@app.route("/identity", methods=["POST"])
def update_identity():
    if not is_logged_in():
        return jsonify({"error": "auth required"}), 401
    payload = request.get_json(silent=True) or {}
    identity = _get_or_create_identity(current_user.id)
    if "grade_level" in payload:
        identity.grade_level = str(payload.get("grade_level") or "").strip()[:32] or None
    if "focus_areas" in payload:
        raw = payload.get("focus_areas") or []
        if isinstance(raw, list):
            identity.focus_areas = json.dumps([str(x).strip()[:48] for x in raw if str(x).strip()][:12])
    if "goals" in payload:
        identity.goals = str(payload.get("goals") or "").strip()[:1000]
    if "availability" in payload:
        av = payload.get("availability") or {}
        if isinstance(av, dict):
            identity.availability = json.dumps(av)
    if "weekly_commitments" in payload:
        identity.weekly_commitments = str(payload.get("weekly_commitments") or "").strip()[:500]
    if "class_schedule" in payload:
        cs = payload.get("class_schedule") or []
        if isinstance(cs, list):
            identity.class_schedule = json.dumps(cs[:50])
    if payload.get("completed"):
        identity.completed = True
    else:
        identity.completed = True
    db.session.commit()
    # When the client signals completion, wipe the resume-state so a future
    # visit to /onboarding starts fresh at "Quick or Custom" rather than
    # silently dumping the user back into the chat they already finished.
    if payload.get("completed"):
        for k in ("onb_mode", "onb_step", "onb_messages"):
            session.pop(k, None)
        session.modified = True
    return jsonify({"ok": True, "identity": identity.to_dict()})


@app.route("/api/identity", methods=["GET"])
def get_identity():
    if not is_logged_in():
        return jsonify({"error": "auth required"}), 401
    identity = _get_or_create_identity(current_user.id)
    return jsonify({"identity": identity.to_dict()})


# ── Onboarding state (resume-from-where-you-left-off) ──────────
@app.route("/api/onboarding/state", methods=["GET", "POST"])
def api_onboarding_state():
    """Get or update onboarding step state. Stored in the session so the user
    can refresh mid-flow and land back on the same step.
    """
    if not current_user.is_authenticated:
        return jsonify({"error": "auth required"}), 401
    if request.method == "GET":
        return jsonify({
            "mode": session.get("onb_mode", ""),       # "quick" | "custom"
            "step": session.get("onb_step", "intro"),  # intro|chat|import|done
            "messages": session.get("onb_messages", []),
        })
    body = request.get_json(silent=True) or {}
    if "mode" in body:
        session["onb_mode"] = (body.get("mode") or "")[:16]
    if "step" in body:
        session["onb_step"] = (body.get("step") or "intro")[:16]
    if "messages" in body and isinstance(body["messages"], list):
        # Keep the last 30 turns to bound session size.
        session["onb_messages"] = body["messages"][-30:]
    session.modified = True
    return jsonify({"ok": True})


@app.route("/api/onboarding/chat", methods=["POST"])
@limiter.limit("30 per hour")
def api_onboarding_chat():
    """AI-assisted onboarding conversation.

    Body: {"messages": [{"role":"user|assistant","content":"..."}, ...]}
    Returns the next assistant reply AND any profile fields it could
    confidently extract from the conversation so far. The caller posts
    those back to /identity so they persist immediately.
    """
    if not current_user.is_authenticated:
        return jsonify({"error": "auth required"}), 401
    if not ai_available():
        return jsonify({"error": "ai unavailable",
                        "reply": "The AI is offline — switch to Quick onboarding for now and refine later in Settings.",
                        "extracted": {}}), 503
    body = request.get_json(silent=True) or {}
    msgs = body.get("messages") or []
    if not isinstance(msgs, list):
        msgs = []
    # Take only the last 20 turns to control prompt size.
    msgs = msgs[-20:]
    identity = _get_or_create_identity(current_user.id)
    so_far = identity.to_dict()
    system = (
        "You are Plani, IntelliPlan's onboarding assistant. Your goal is to learn "
        "enough about the student in under 2 minutes to personalize their study plan.\n\n"
        "ASK FOR (one or two at a time, in this order):\n"
        "  1. Grade level (e.g. '10th grade', 'sophomore in college').\n"
        "  2. Subjects/courses they're currently taking and which feel hardest.\n"
        "  3. Their main goal this semester (raise a grade, prep for a test, build a habit).\n"
        "  4. Their typical free time on weekdays vs weekends.\n"
        "  5. Any weekly commitments (sports, work, family).\n\n"
        "VOICE: warm, concise (1–3 sentences), one focused question per turn. "
        "Never ask all five at once. Acknowledge what they just said before asking the next thing. "
        "If they've covered everything, say so and recommend they tap 'Done'.\n\n"
        # Without this the model invents an origin (it has been answering
        # "students at Michigan" / "at UPenn"). Stated as a fact so it has
        # something true to say instead of guessing.
        "IF ASKED WHO MADE INTELLIPLAN: Anirudh Ulabala built it solo — no company, "
        "no team, no university. Point to /about. Never invent an origin story. "
        "Answer in one short sentence inside 'reply', then continue onboarding.\n\n"
        "OUTPUT FORMAT: return ONLY valid JSON shaped exactly like:\n"
        '{"reply":"...your message to the student (under 240 chars)...",'
        '"extracted":{"grade_level":"","focus_areas":[],"goals":"","weekly_commitments":"","availability":{}},'
        '"complete":false}\n\n'
        "FIELD RULES:\n"
        "  - extracted.focus_areas: at most 8 short course names. Strip teacher names and percentages.\n"
        "  - extracted.weekly_commitments: ONE compact line like 'piano Wed 5:30-6pm; robotics Sat/Sun; tennis Sat 10-11am'.\n"
        "  - extracted.availability: simple {day: 'morning|afternoon|evening|none'} for the 7 weekday keys. "
        "No nested objects. Skip days you're unsure about.\n"
        "  - reply: KEEP IT SHORT. One acknowledgement sentence + one question. Do not echo the student's full input.\n"
        "  - Only include extracted fields you're confident about — leave others as empty string / empty array / empty object.\n"
        "  - Set complete=true only when you have grade_level + focus_areas + goals at minimum.\n\n"
        f"PROFILE SO FAR: {json.dumps(so_far)}"
    )
    chat_messages = [{"role": "system", "content": system}]
    for m in msgs:
        if not isinstance(m, dict): continue
        role = m.get("role")
        content = (m.get("content") or "").strip()
        if role in ("user", "assistant") and content:
            chat_messages.append({"role": role, "content": content[:2000]})
    # Bootstrap with an opener if the conversation is empty.
    if len(chat_messages) == 1:
        chat_messages.append({"role": "user", "content": "Start onboarding."})
    # Helper: turn a possibly-malformed/truncated JSON string into a dict.
    # The previous version errored out when the model truncated mid-object
    # (which the schedule-heavy user case hit reliably). Now we strip code
    # fences, extract the first {...} block, and as a last resort try to
    # close any unclosed braces/quotes before parsing.
    def _tolerant_json(raw):
        if not raw:
            return None
        s = re.sub(r"```json\n?", "", raw)
        s = re.sub(r"```\n?", "", s).strip()
        try:
            return json.loads(s)
        except json.JSONDecodeError:
            pass
        m = re.search(r"\{[\s\S]*\}", s)
        if m:
            try:
                return json.loads(m.group(0))
            except json.JSONDecodeError:
                # Truncated — try repairing by closing braces/brackets.
                frag = m.group(0)
                # Strip trailing comma, then close any unbalanced braces.
                frag = re.sub(r",\s*$", "", frag)
                opens = frag.count("{") - frag.count("}")
                opensq = frag.count("[") - frag.count("]")
                quotes = frag.count('"')
                if quotes % 2 == 1:
                    frag += '"'
                frag += "]" * max(0, opensq)
                frag += "}" * max(0, opens)
                try:
                    return json.loads(frag)
                except json.JSONDecodeError:
                    return None
        return None

    parsed = None
    try:
        result = ai_chat(
            chat_messages, tier="standard", temperature=0.5,
            # Bumped from 600 → 1500. With response_format=json_object the
            # model emits the schedule/availability dict + the reply in
            # one go; 600 truncated on long schedules and the fallback
            # message ("I lost the thread") fired every time.
            max_tokens=1500,
            response_format={"type": "json_object"},
        )
        parsed = _tolerant_json(result)
        if parsed is None:
            raise ValueError("AI returned unparseable JSON after repair")
    except Exception as _e:
        print(f"[onboarding-chat] AI failed: {_e}")
        # When we still couldn't parse anything, fall through with an
        # empty reply so the client just shows nothing rather than the
        # demoralizing "I lost the thread" message after a real answer.
        # The user's message is preserved in the chat history server-side,
        # so a retry sends a fresh, smaller prompt next time.
        last_user = next((m["content"] for m in reversed(msgs) if m.get("role") == "user"), "")
        return jsonify({
            "reply": (
                "Got it. Could you break that into two messages? Long details are tripping up "
                "my structured-data extractor. (Your earlier answers are saved either way.)"
                if last_user and len(last_user) > 250 else
                "Hmm, the AI hiccuped. Try sending that again, or skip to Quick onboarding."
            ),
            "extracted": {}, "complete": False,
        })
    reply = str(parsed.get("reply") or "").strip()[:1200]
    extracted = parsed.get("extracted") or {}
    if not isinstance(extracted, dict):
        extracted = {}
    # Auto-persist any extracted fields immediately so refresh-mid-flow is safe.
    try:
        if extracted.get("grade_level"):
            identity.grade_level = str(extracted["grade_level"])[:32]
        if isinstance(extracted.get("focus_areas"), list) and extracted["focus_areas"]:
            cur = identity.focus_list()
            merged = list(dict.fromkeys(cur + [str(x)[:48] for x in extracted["focus_areas"] if str(x).strip()]))[:12]
            identity.focus_areas = json.dumps(merged)
        if extracted.get("goals"):
            identity.goals = str(extracted["goals"])[:1000]
        if extracted.get("weekly_commitments"):
            identity.weekly_commitments = str(extracted["weekly_commitments"])[:500]
        if isinstance(extracted.get("availability"), dict) and extracted["availability"]:
            identity.availability = json.dumps(extracted["availability"])
        db.session.commit()
    except Exception as _pe:
        print(f"[onboarding-chat] persist failed: {_pe}")
        try: db.session.rollback()
        except Exception: pass
    return jsonify({
        "reply": reply,
        "extracted": extracted,
        "complete": bool(parsed.get("complete")),
        "identity": identity.to_dict(),
    })

_RETURN_TO_ALLOWLIST = (
    ".web.app", ".firebaseapp.com", ".replit.dev", "intelliplan.tech", "localhost",
)

def _safe_return_to(url):
    """Return url if its host is on the allowlist, else None.

    Allowlist entries are matched on a dot/exact boundary so a lookalike
    domain (e.g. evilintelliplan.tech, notlocalhost) can't pass: an entry
    "intelliplan.tech" matches only that host or a "*.intelliplan.tech"
    subdomain, never a domain that merely ends with that string. Only http(s)
    URLs are accepted."""
    if not url:
        return None
    import urllib.parse as _up
    try:
        parsed = _up.urlparse(url)
        if parsed.scheme and parsed.scheme.lower() not in ("http", "https"):
            return None
        host = (parsed.hostname or "").lower()
        if not host:
            return None
        for a in _RETURN_TO_ALLOWLIST:
            domain = a.lstrip(".").lower()
            if host == domain or host.endswith("." + domain):
                return url
    except Exception:
        pass
    return None


@app.route("/login/google")
def login_google():
    if not GCAL_AVAILABLE:
        return redirect(url_for("login"))
    # Store a validated return_to URL so the callback can send the user back
    # to an embedding partner (e.g. Lotus) after sign-in completes.
    return_to = _safe_return_to(request.args.get("return_to", "").strip())
    if return_to:
        session["oauth_return_to"] = return_to
    # A desktop-initiated sign-in arrives here in the system browser, not in
    # the app, carrying the PKCE challenge the app generated. Remember it so
    # the callback knows to hand the result back over intelliplan:// instead
    # of just landing this browser tab on the command centre. Anything
    # malformed is dropped rather than trusted — see desktop_auth.py.
    desktop_challenge = request.args.get("desktop", "").strip()
    session.pop("desktop_auth_challenge", None)
    if desktop_challenge:
        if desktop_auth.is_valid_challenge(desktop_challenge):
            session["desktop_auth_challenge"] = desktop_challenge
        else:
            print("[DESKTOP AUTH] rejected malformed challenge on /login/google")
    state = secrets_module.token_urlsafe(32)
    session["oauth_state"] = state
    session["oauth_purpose"] = "login"
    session.permanent = True
    session.modified = True
    auth_url, code_verifier = get_auth_url(state, purpose="login")
    session["oauth_code_verifier"] = code_verifier
    session.modified = True
    print(f"[GOOGLE LOGIN] state={state[:8]}..., return_to={return_to!r}")
    return redirect(auth_url)

@app.route("/api/desktop/auth/exchange", methods=["POST"])
@limiter.limit("10 per minute;40 per hour")
def desktop_auth_exchange():
    """Spend a one-time code for a real session in the desktop app.

    The app calls this itself, so the ``Set-Cookie`` on the response lands
    in the app's own cookie jar — which is the entire point of the dance:
    the browser that actually did the Google sign-in has a session the app
    could never read.

    Every failure returns the same flat 400. Distinguishing "no such code"
    from "expired" from "wrong verifier" would tell someone grinding at the
    endpoint which half of their guess was right.
    """
    payload = request.get_json(silent=True) or {}
    code = (payload.get("code") or "").strip()
    verifier = (payload.get("verifier") or "").strip()

    def refuse():
        return jsonify({"status": "error", "error": "invalid_code"}), 400

    if not code or not desktop_auth.is_valid_verifier(verifier):
        return refuse()

    row = DesktopAuthCode.query.filter_by(
        code_hash=desktop_auth.hash_code(code)
    ).first()
    if row is None or row.used_at is not None:
        return refuse()
    if desktop_auth.is_expired(row.created_at, datetime.utcnow()):
        return refuse()
    if not desktop_auth.verifier_matches(verifier, row.code_challenge):
        # A valid code with the wrong verifier is the signature of a stolen
        # deep link, so burn it rather than letting them try again.
        row.used_at = datetime.utcnow()
        db.session.commit()
        print("[DESKTOP AUTH] verifier mismatch — code burned")
        return refuse()

    user = User.query.get(row.user_id)
    if user is None:
        return refuse()

    # Burn before issuing. If the commit below fails the student retries a
    # sign-in; if it succeeded and we crashed after, the code must not still
    # be live.
    row.used_at = datetime.utcnow()
    db.session.commit()

    login_user(user, remember=True)
    session.permanent = True
    session.modified = True
    print(f"[DESKTOP AUTH] exchanged code for session, user id={user.id}")
    return jsonify({"status": "ok"})


@app.route("/login/account", methods=["GET", "POST"])
@limiter.limit("10 per minute;60 per hour", methods=["POST"])
def login_account():
    if current_user.is_authenticated:
        return redirect("/command-center")
    error = None
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "").strip()
        try:
            user = User.query.filter_by(email=email).first()
        except Exception as _e:
            # If the SELECT itself blows up (e.g. a new column hasn't
            # been migrated yet on production), retry the migration and
            # try ONCE more before surfacing a friendly error.
            print(f"[login] User lookup failed: {_e}")
            try:
                db.session.rollback()
            except Exception:
                pass
            try:
                _run_boot_migration_once()
                user = User.query.filter_by(email=email).first()
            except Exception as _e2:
                print(f"[login] retry also failed: {_e2}")
                user = None
                error = "Login is briefly unavailable. Please try again in a moment."
        if not error:
            if user and user.password_hash and bcrypt.check_password_hash(user.password_hash, password):
                # COPPA: block sign-in for under-13 accounts that haven't been approved yet.
                if hasattr(user, "parent_consent_granted") and user.parent_consent_granted is False and (user.parent_email or ""):
                    error = ("This account is waiting for parental consent. We emailed "
                             f"{user.parent_email} a consent link — once they click it, "
                             "you'll be able to sign in.")
                    return render_template("login_account.html", active_page="login", error=error)
                login_user(user, remember=True)
                # Auto-join any group whose invite link was clicked pre-login.
                joined_gid = _apply_pending_group_join()
                if joined_gid:
                    return redirect(url_for("groups_page") + f"?open={joined_gid}")
                try:
                    acct = LinkedAccount.query.filter_by(user_id=user.id, is_active=True).first()
                except Exception:
                    acct = None
                if not acct:
                    return redirect(url_for("connect_account"))
                return redirect("/command-center")
            else:
                error = "Invalid email or password."
    return render_template("login_account.html", active_page="login", error=error)

@app.route("/connect", methods=["GET"])
def connect_account():
    return render_template("connect.html", active_page="login")

@app.route("/login/canvas", methods=["GET", "POST"])
def login_canvas():
    error = None
    if request.method == "POST":
        token = request.form.get("canvas_token", "").strip()
        canvas_url = request.form.get("canvas_url", "").strip().rstrip("/")
        profile_name = request.form.get("profile_name", "").strip() or "Canvas Account"
        if canvas_url and not canvas_url.startswith(("http://", "https://")):
            canvas_url = "https://" + canvas_url
        if not token or not canvas_url:
            error = "Please fill in both fields."
        else:
            test = requests.get(f"{canvas_url}/api/v1/courses", headers={"Authorization": f"Bearer {token}"}, timeout=20)
            if test.status_code == 200:
                creds = {"canvas_token": token, "canvas_url": canvas_url}
                if current_user.is_authenticated:
                    LinkedAccount.query.filter_by(user_id=current_user.id).update({"is_active": False})
                    acct = LinkedAccount(user_id=current_user.id, name=profile_name, login_type="canvas", is_active=True)
                    acct.set_credentials(creds)
                    db.session.add(acct)
                    db.session.commit()
                else:
                    session.permanent = True
                    session["canvas_token"] = token
                    session["canvas_url"] = canvas_url
                    session["login_type"] = "canvas"
                return redirect("/command-center")
            else:
                error = "Invalid token or Canvas URL."
    return render_template(
        "login_canvas.html",
        active_page="login",
        error=error,
        canvas_oauth_available=(CANVAS_OAUTH_AVAILABLE and canvas_oauth_configured()),
    )

# ── CANVAS OAUTH ──────────────────────────────────────────────
# Lets students connect Canvas with one click instead of finding and
# pasting a personal access token. Backed by canvas_oauth.py.
@app.route("/oauth/canvas/check")
def oauth_canvas_check():
    """Lightweight probe used by the login page to decide whether to show
    the "Continue with Canvas" button for a given canvas_base."""
    base = (request.args.get("canvas_base") or DEFAULT_CANVAS_BASE).strip()
    if not CANVAS_OAUTH_AVAILABLE:
        return flask.jsonify({"available": False, "reason": "library_missing"})
    return flask.jsonify({"available": bool(canvas_oauth_configured(base)), "canvas_base": base})


@app.route("/oauth/canvas")
def oauth_canvas_start():
    canvas_base = request.args.get("canvas_base") or DEFAULT_CANVAS_BASE
    if not CANVAS_OAUTH_AVAILABLE or not canvas_oauth_configured(canvas_base):
        return redirect(url_for("login_canvas") + "?reason=oauth_unavailable&canvas_base=" + urllib.parse.quote(canvas_base))
    profile_name = request.args.get("profile_name") or "Canvas Account"
    state = secrets_module.token_urlsafe(24)
    session["canvas_oauth_state"] = state
    session["canvas_oauth_base"] = canvas_base
    session["canvas_oauth_profile_name"] = profile_name
    session.modified = True
    redirect_uri = os.getenv("CANVAS_REDIRECT_URI") or (
        APP_BASE_URL.rstrip("/") + "/oauth/canvas/callback"
    )
    try:
        url = get_canvas_auth_url(state, canvas_base=canvas_base, redirect_uri=redirect_uri)
    except Exception as e:
        return render_template("error.html", active_page="error", error_code=500,
                               error_id=f"CANVAS-OAUTH-{make_error_id()}",
                               message=str(e)), 500
    return redirect(url)


@app.route("/oauth/canvas/callback")
def oauth_canvas_callback():
    if not CANVAS_OAUTH_AVAILABLE:
        return redirect(url_for("login_canvas"))
    code = request.args.get("code")
    state = request.args.get("state")
    err = request.args.get("error")
    err_desc = request.args.get("error_description") or ""
    if err:
        # Canvas returns ?error=invalid_client when the institution hasn't
        # registered IntelliPlan as a Developer Key. Redirect to the token
        # paste page with a clear explanation instead of a generic error.
        if "invalid_client" in err or "unauthorized" in err:
            return redirect(
                url_for("login_canvas")
                + "?reason=oauth_not_registered"
                + ("&desc=" + urllib.parse.quote(err_desc) if err_desc else "")
            )
        return redirect(url_for("login_canvas") + f"?error={err}")
    expected_state = session.pop("canvas_oauth_state", None)
    canvas_base = session.pop("canvas_oauth_base", None) or DEFAULT_CANVAS_BASE
    profile_name = session.pop("canvas_oauth_profile_name", None) or "Canvas Account"
    if not code or not state or state != expected_state:
        return render_template("error.html", active_page="error", error_code=400,
                               error_id="CANVAS-OAUTH-STATE",
                               message="Canvas OAuth state did not match. Try again."), 400
    redirect_uri = os.getenv("CANVAS_REDIRECT_URI") or (
        APP_BASE_URL.rstrip("/") + "/oauth/canvas/callback"
    )
    try:
        tokens = exchange_canvas_code(code, canvas_base, redirect_uri=redirect_uri)
    except Exception as e:
        msg = str(e)
        if "invalid_client" in msg or "unauthorized_client" in msg or "redirect_uri" in msg:
            return redirect(url_for("login_canvas") + "?reason=oauth_not_registered")
        return render_template("error.html", active_page="error", error_code=500,
                               error_id=f"CANVAS-OAUTH-{make_error_id()}",
                               message=msg), 500

    access_token = tokens.get("access_token")
    refresh_token = tokens.get("refresh_token")
    expires_in = tokens.get("expires_in")
    expires_at = (utcnow() + timedelta(seconds=int(expires_in))) if expires_in else None
    user_info = tokens.get("user") or {}
    canvas_user_id = str(user_info.get("id") or "") or None
    canvas_user_name = user_info.get("name") or profile_name

    creds = {
        "canvas_token": access_token,
        "canvas_url": tokens.get("canvas_base") or canvas_base,
        "canvas_refresh_token": refresh_token,
        "canvas_token_expires_at": expires_at.isoformat() if expires_at else None,
        "canvas_oauth": True,
    }

    if current_user.is_authenticated:
        # Persist long-lived OAuth tokens in their own table for refresh,
        # AND mirror short-lived access_token onto LinkedAccount so the
        # existing /live, /grades, /classes paths keep working unchanged.
        existing_oauth = CanvasIntegration.query.filter_by(user_id=current_user.id).first()
        if existing_oauth:
            existing_oauth.canvas_base = tokens.get("canvas_base") or canvas_base
            existing_oauth.access_token = access_token
            existing_oauth.refresh_token = refresh_token or existing_oauth.refresh_token
            existing_oauth.token_expires_at = expires_at
            existing_oauth.canvas_user_id = canvas_user_id or existing_oauth.canvas_user_id
            existing_oauth.canvas_user_name = canvas_user_name or existing_oauth.canvas_user_name
            existing_oauth.connected_at = utcnow()
        else:
            db.session.add(CanvasIntegration(
                user_id=current_user.id,
                canvas_base=tokens.get("canvas_base") or canvas_base,
                access_token=access_token,
                refresh_token=refresh_token,
                token_expires_at=expires_at,
                canvas_user_id=canvas_user_id,
                canvas_user_name=canvas_user_name,
            ))
        LinkedAccount.query.filter_by(user_id=current_user.id).update({"is_active": False})
        acct = LinkedAccount(
            user_id=current_user.id,
            name=canvas_user_name or profile_name,
            login_type="canvas",
            is_active=True,
        )
        acct.set_credentials(creds)
        db.session.add(acct)
        db.session.commit()
    else:
        session.permanent = True
        session["canvas_token"] = access_token
        session["canvas_url"] = tokens.get("canvas_base") or canvas_base
        session["canvas_refresh_token"] = refresh_token
        session["canvas_oauth"] = True
        session["login_type"] = "canvas"
    return redirect("/command-center")


@app.route("/oauth/canvas/disconnect", methods=["POST"])
def oauth_canvas_disconnect():
    if current_user.is_authenticated:
        ci = CanvasIntegration.query.filter_by(user_id=current_user.id).first()
        if ci and ci.access_token:
            try:
                revoke_canvas_token(ci.access_token, ci.canvas_base)
            except Exception:
                pass
            db.session.delete(ci)
            db.session.commit()
    session.pop("canvas_token", None)
    session.pop("canvas_url", None)
    session.pop("canvas_refresh_token", None)
    session.pop("canvas_oauth", None)
    return flask.jsonify({"status": "ok"})


@app.route("/login/studentvue", methods=["GET", "POST"])
def login_studentvue():
    error = None
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        district_url = normalize_district_url(request.form.get("district_url", ""))
        profile_name = request.form.get("profile_name", "").strip() or "StudentVue Account"
        if not username or not password or not district_url:
            error = "Please fill in all fields."
        else:
            if test_login(district_url, username, password):
                creds = {"sv_username": username, "sv_password": password, "sv_district_url": district_url}
                if current_user.is_authenticated:
                    LinkedAccount.query.filter_by(user_id=current_user.id).update({"is_active": False})
                    acct = LinkedAccount(user_id=current_user.id, name=profile_name, login_type="studentvue", is_active=True)
                    acct.set_credentials(creds)
                    db.session.add(acct)
                    db.session.commit()
                else:
                    session.permanent = True
                    session["sv_username"] = username
                    session["sv_password"] = password
                    session["sv_district_url"] = district_url
                    session["login_type"] = "studentvue"
                return redirect("/command-center")
            else:
                error = "Invalid StudentVUE credentials or district URL. Try your district's StudentVUE web address, like https://district-psv.edupoint.com."
    return render_template("login_studentvue.html", active_page="login", error=error)


@app.route("/oauth/studentvue")
def oauth_studentvue_start():
    return redirect(url_for("login_studentvue") + "?oauth=studentvue")

@app.route("/login/schoology", methods=["GET", "POST"])
def login_schoology():
    error = None
    if request.method == "POST":
        key = request.form.get("api_key", "").strip()
        secret = request.form.get("api_secret", "").strip()
        profile_name = request.form.get("profile_name", "").strip() or "Schoology Account"
        if not key or not secret:
            error = "Please fill in both fields."
        else:
            try:
                from schoology_helper import test_schoology_login
                if test_schoology_login(key, secret):
                    creds = {"schoology_key": key, "schoology_secret": secret}
                    if current_user.is_authenticated:
                        LinkedAccount.query.filter_by(user_id=current_user.id).update({"is_active": False})
                        acct = LinkedAccount(user_id=current_user.id, name=profile_name, login_type="schoology", is_active=True)
                        acct.set_credentials(creds)
                        db.session.add(acct)
                        db.session.commit()
                    else:
                        session["schoology_key"] = key
                        session["schoology_secret"] = secret
                        session["login_type"] = "schoology"
                    return redirect("/command-center")
                else:
                    error = "Invalid Schoology credentials."
            except Exception as e:
                error = f"Schoology error: {str(e)}"
    return render_template("login_schoology.html", active_page="login", error=error)

@app.route("/logout", methods=["POST", "GET"])
def logout():
    logout_user()
    session.clear()
    response = redirect(url_for("login"))
    response.delete_cookie(app.config.get("SESSION_COOKIE_NAME", "session"))
    response.delete_cookie("remember_token")
    return response

# ── GOOGLE OAUTH ──────────────────────────────────────────────
# ── FIX: Single consolidated OAuth callback that handles both
#         "login" (create/find account + log in) and "calendar"
#         (link calendar to existing logged-in user).
#
#   Both /oauth/google/callback and /oauth2callback point here so
#   either redirect URI registered in Google Cloud Console works.
# ─────────────────────────────────────────────────────────────

def _handle_google_callback():
    import traceback

    print(f"[GOOGLE CALLBACK] args={dict(request.args)}")
    print(f"[GOOGLE CALLBACK] session_keys={list(session.keys())}")

    # Google returned an error (user denied, etc.)
    error_msg = request.args.get("error")
    if error_msg:
        print(f"[GOOGLE CALLBACK] Google error param: {error_msg}")
        return redirect(url_for("login"))

    returned_state = request.args.get("state")
    stored_state = session.pop("oauth_state", None)
    purpose = session.pop("oauth_purpose", "calendar")

    # ── State check — primary guard against CSRF and session loss ──
    if not stored_state:
        # This is the exact cause of IPE-XXXXXXXX errors:
        # session was empty because a different container handled the callback.
        # Switching to sqlalchemy sessions (above) prevents this permanently.
        print("[GOOGLE CALLBACK] FATAL: oauth_state missing from session (session was lost)")
        return redirect(url_for("login"))

    if returned_state != stored_state:
        print(f"[GOOGLE CALLBACK] FATAL: state mismatch returned={returned_state!r} stored={stored_state!r}")
        return redirect(url_for("login"))

    code = request.args.get("code")
    if not code:
        print("[GOOGLE CALLBACK] FATAL: no code param")
        return redirect(url_for("login"))

    # ── Exchange code for tokens ──
    try:
        code_verifier = session.pop("oauth_code_verifier", None)
        token_dict = exchange_code_for_token(code, code_verifier=code_verifier)
        print(f"[GOOGLE CALLBACK] token_dict keys={list(token_dict.keys())}")
    except Exception:
        print(f"[GOOGLE CALLBACK] token exchange failed:\n{traceback.format_exc()}")
        return redirect(url_for("login"))

    # exchange_code_for_token returns {"token": ..., "refresh_token": ...}
    # guard against either key name
    access_token = token_dict.get("access_token") or token_dict.get("token")
    if not access_token:
        print(f"[GOOGLE CALLBACK] FATAL: no access_token in token_dict keys={list(token_dict.keys())}")
        return redirect(url_for("login"))

    # ── Fetch Google user profile ──
    try:
        ui_resp = requests.get(
            "https://www.googleapis.com/oauth2/v3/userinfo",
            headers={"Authorization": f"Bearer {access_token}"},
            timeout=10,
        )
        ui_resp.raise_for_status()
        userinfo = ui_resp.json()
    except Exception:
        print(f"[GOOGLE CALLBACK] userinfo fetch failed:\n{traceback.format_exc()}")
        return redirect(url_for("login"))

    google_id = userinfo.get("sub")
    email = (userinfo.get("email") or "").lower().strip()
    name = userinfo.get("name") or email.split("@")[0]
    print(f"[GOOGLE CALLBACK] google_id={google_id} email={email} purpose={purpose}")

    if not google_id or not email:
        print("[GOOGLE CALLBACK] FATAL: missing sub or email in userinfo")
        return redirect(url_for("login"))

    if purpose == "calendar":
        if not current_user.is_authenticated:
            print("[GOOGLE CALLBACK] calendar link attempted without an authenticated app user")
            return redirect(url_for("login"))
        user = current_user
        if user.email.lower() == email and not user.google_id:
            user.google_id = google_id
        if not user.name:
            user.name = name
        db.session.commit()
        if has_calendar_scope(token_dict):
            # Multi-account: match the row by account_email so a second
            # Google sign-in adds a new row instead of overwriting the first.
            gi = (
                GoogleIntegration.query
                .filter_by(user_id=user.id, account_email=email)
                .first()
            )
            existing_token = json.loads(gi.token_data) if gi else {}
            token_dict = merge_token_data(existing_token, token_dict)
            # Newly-linked account becomes the active one.
            GoogleIntegration.query.filter_by(user_id=user.id).update({"is_active": False})
            if gi:
                gi.token_data = json.dumps(token_dict)
                gi.account_name = name or gi.account_name
                gi.is_active = True
            else:
                db.session.add(GoogleIntegration(
                    user_id=user.id,
                    token_data=json.dumps(token_dict),
                    account_email=email,
                    account_name=name,
                    is_active=True,
                ))
            db.session.commit()
            session["google_token"] = token_dict
            session.permanent = True
            session.modified = True
        return redirect(url_for("settings") if session.pop("oauth_return_to_settings", False) else "/command-center")

    # ── Find or create User (defensive — see login_account for rationale) ──
    try:
        user = User.query.filter_by(google_id=google_id).first()
    except Exception as _e:
        print(f"[GOOGLE CALLBACK] User.query failed, retrying after migration: {_e}")
        try: db.session.rollback()
        except Exception: pass
        try:
            _run_boot_migration_once()
            user = User.query.filter_by(google_id=google_id).first()
        except Exception as _e2:
            print(f"[GOOGLE CALLBACK] User.query still failing after migration: {_e2}")
            return redirect(url_for("login") + "?err=oauth_db")
    if not user:
        try:
            user = User.query.filter_by(email=email).first()
        except Exception:
            user = None
        if user:
            # Existing email-based account — link the Google ID
            user.google_id = google_id
            if not user.name:
                user.name = name
            print(f"[GOOGLE CALLBACK] linked google_id to existing user id={user.id}")
        else:
            # Brand-new user — create the account
            user = User(
                email=email,
                google_id=google_id,
                name=name,
                password_hash="",
            )
            db.session.add(user)
            print(f"[GOOGLE CALLBACK] created new user email={email}")
    else:
        print(f"[GOOGLE CALLBACK] found existing google user id={user.id}")

    db.session.commit()

    # ── Persist the Google token for calendar use ──
    if has_calendar_scope(token_dict):
        gi = (
            GoogleIntegration.query
            .filter_by(user_id=user.id, account_email=email)
            .first()
        )
        existing_token = json.loads(gi.token_data) if gi else {}
        token_dict = merge_token_data(existing_token, token_dict)
        GoogleIntegration.query.filter_by(user_id=user.id).update({"is_active": False})
        if gi:
            gi.token_data = json.dumps(token_dict)
            gi.account_name = name or gi.account_name
            gi.is_active = True
        else:
            db.session.add(GoogleIntegration(
                user_id=user.id,
                token_data=json.dumps(token_dict),
                account_email=email,
                account_name=name,
                is_active=True,
            ))
        db.session.commit()
        session["google_token"] = token_dict
        session.permanent = True
        session.modified = True

    # ── Log the user in ──
    login_user(user, remember=True)
    print(f"[GOOGLE CALLBACK] logged in user id={user.id}")

    # ── Redirect ──
    if purpose == "login":
        # A desktop-initiated sign-in finishes in the system browser, which
        # is the wrong place for the session to stop: the app has its own
        # cookie jar and cannot see this one. Mint a one-time code and bounce
        # through the protocol handler so the app can claim it.
        desktop_challenge = session.pop("desktop_auth_challenge", None)
        if desktop_challenge and desktop_auth.is_valid_challenge(desktop_challenge):
            code = desktop_auth.new_code()
            db.session.add(DesktopAuthCode(
                code_hash=desktop_auth.hash_code(code),
                code_challenge=desktop_challenge,
                user_id=user.id,
            ))
            db.session.commit()
            # The code itself is deliberately not logged.
            print(f"[DESKTOP AUTH] minted code for user id={user.id}")
            return redirect(desktop_auth.deep_link_for(code))
        return_to = _safe_return_to(session.pop("oauth_return_to", None))
        if return_to:
            print(f"[GOOGLE CALLBACK] redirecting to partner return_to={return_to!r}")
            return redirect(return_to)
        return redirect("/command-center")
    return redirect("/command-center")


@app.route("/oauth/google")
def google_oauth_start():
    """Start the OAuth flow from the calendar-linking UI.
    Accepts ?add=1 to force the Google account chooser (so the user can
    add a second account) and ?return=settings to come back to /settings."""
    if not is_logged_in():
        return redirect(url_for("login"))
    if not GCAL_AVAILABLE:
        return "Google Calendar not configured", 500
    state = secrets_module.token_urlsafe(32)
    session["oauth_state"] = state
    session["oauth_purpose"] = "calendar"
    if request.args.get("return") == "settings":
        session["oauth_return_to_settings"] = True
    session.permanent = True
    session.modified = True
    auth_url, code_verifier = get_auth_url(state, purpose="calendar")
    # When the user clicks "Add another account", force Google to show
    # the account picker even if they're already signed in.
    if request.args.get("add") == "1":
        sep = "&" if "?" in auth_url else "?"
        auth_url = f"{auth_url}{sep}prompt=select_account"
    session["oauth_code_verifier"] = code_verifier
    session.modified = True
    return redirect(auth_url)


@app.route("/gcal/status")
def gcal_status():
    """List the user's connected Google accounts and which one is active."""
    if not current_user.is_authenticated:
        return flask.jsonify({"connected": False, "accounts": [], "active_id": None})
    rows = (
        GoogleIntegration.query
        .filter_by(user_id=current_user.id)
        .order_by(GoogleIntegration.id.asc())
        .all()
    )
    accounts = [{
        "id": r.id,
        "email": r.account_email or "",
        "name": r.account_name or "",
        "is_active": bool(r.is_active),
    } for r in rows]
    active = next((a for a in accounts if a["is_active"]), accounts[0] if accounts else None)
    return flask.jsonify({
        "connected": bool(accounts),
        "accounts": accounts,
        "active_id": active["id"] if active else None,
        "active_email": active["email"] if active else "",
    })


@app.route("/gcal/activate", methods=["POST"])
def gcal_activate():
    """Mark a specific GoogleIntegration row as the active account."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    body = request.get_json(silent=True) or {}
    row_id = body.get("id")
    if not row_id:
        return flask.jsonify({"status": "error", "message": "id required"}), 400
    target = GoogleIntegration.query.filter_by(user_id=current_user.id, id=row_id).first()
    if not target:
        return flask.jsonify({"status": "error", "message": "not found"}), 404
    GoogleIntegration.query.filter_by(user_id=current_user.id).update({"is_active": False})
    target.is_active = True
    db.session.commit()
    try:
        session["google_token"] = json.loads(target.token_data)
        session.modified = True
    except Exception:
        pass
    return flask.jsonify({"status": "ok"})


@app.route("/gcal/remove", methods=["POST"])
def gcal_remove():
    """Disconnect a single Google account by row id (multi-account aware)."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    body = request.get_json(silent=True) or {}
    row_id = body.get("id")
    if not row_id:
        return flask.jsonify({"status": "error", "message": "id required"}), 400
    target = GoogleIntegration.query.filter_by(user_id=current_user.id, id=row_id).first()
    if not target:
        return flask.jsonify({"status": "error", "message": "not found"}), 404
    was_active = bool(target.is_active)
    db.session.delete(target)
    db.session.commit()
    if was_active:
        # Promote the next-most-recent account to active.
        nxt = GoogleIntegration.query.filter_by(user_id=current_user.id).order_by(GoogleIntegration.id.desc()).first()
        if nxt:
            nxt.is_active = True
            db.session.commit()
            try: session["google_token"] = json.loads(nxt.token_data)
            except Exception: pass
        else:
            session.pop("google_token", None)
        session.modified = True
    return flask.jsonify({"status": "ok"})


# ── FIX: Both route paths registered so either redirect URI works.
#   Set GOOGLE_REDIRECT_URI=https://intelliplan.tech/oauth2callback
#   in Railway and add that URI in Google Cloud Console.
#   Keep /oauth/google/callback as a fallback alias.
@app.route("/oauth2callback")
@app.route("/oauth/google/callback")
def google_oauth_callback():
    return _handle_google_callback()


@app.route("/oauth/google/disconnect", methods=["POST"])
def google_disconnect():
    if current_user.is_authenticated:
        GoogleIntegration.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
    else:
        session.pop("google_token", None)
    return flask.jsonify({"status": "ok"})

@app.route("/calendar/events")
def calendar_events():
    if not GCAL_AVAILABLE:
        return flask.jsonify({"connected": False, "events": []})
    token = get_google_token()
    if not token:
        return flask.jsonify({"connected": False, "events": []})
    try:
        events = get_upcoming_events(token)
        session["google_token"] = token
        session.modified = True
        return flask.jsonify({"connected": True, "events": events})
    except Exception as e:
        print(f"Calendar events error: {e}")
        session.pop("google_token", None)
        if current_user.is_authenticated:
            GoogleIntegration.query.filter_by(user_id=current_user.id).delete()
            db.session.commit()
        return flask.jsonify({"connected": False, "error": safe_error_message(e), "events": []})

@app.route("/calendar/free-slot")
def calendar_free_slot():
    if not GCAL_AVAILABLE:
        return flask.jsonify({"slot": "7:00 PM", "connected": False})
    token = get_google_token()
    date_str = request.args.get("date", datetime.now().strftime("%Y-%m-%d"))
    if not token:
        return flask.jsonify({"slot": "7:00 PM", "connected": False})
    try:
        slot = find_free_slots(token, date_str)
        free_hours = compute_free_hours(token, date_str)
        return flask.jsonify({"slot": slot, "connected": True, "free_hours": free_hours})
    except Exception as e:
        return flask.jsonify({"slot": "7:00 PM", "connected": False, "error": safe_error_message(e)})

@app.route("/calendar/export", methods=["POST"])
def calendar_export():
    if not GCAL_AVAILABLE:
        return flask.jsonify({"status": "error", "message": "Google Calendar not configured"})
    token = get_google_token()
    if not token:
        return flask.jsonify({"status": "error", "message": "Google Calendar not connected"})
    data = request.get_json(silent=True) or {}
    schedule_data = data.get("schedule_data")
    if not schedule_data:
        return flask.jsonify({"status": "error", "message": "No schedule data supplied"}), 400
    skip_overlaps = data.get("skip_overlaps", False)
    try:
        existing_events = []
        if skip_overlaps:
            try:
                existing_events = get_upcoming_events(token)
            except Exception:
                existing_events = []
        ids, new_token, skipped = add_schedule_to_calendar(token, schedule_data, existing_events if skip_overlaps else [])
        if new_token:
            session["google_token"] = {**token, "token": new_token}
            session.modified = True
            if current_user.is_authenticated:
                gi = GoogleIntegration.query.filter_by(user_id=current_user.id).first()
                if gi:
                    td = json.loads(gi.token_data)
                    td["token"] = new_token
                    gi.token_data = json.dumps(td)
                    db.session.commit()
        return flask.jsonify({"status": "ok", "created": len(ids), "skipped": skipped})
    except Exception as e:
        print(f"Calendar export error: {e}")
        return flask.jsonify({"status": "error", "message": "Google Calendar export failed. Please try again."})

# ── PROFILE MANAGEMENT ────────────────────────────────────────
@app.route("/profiles/list")
def profiles_list():
    if not current_user.is_authenticated:
        login_type = session.get("login_type")
        if login_type:
            return flask.jsonify({"is_guest": True, "profiles": [{"id": "guest", "name": "Guest Session", "login_type": login_type, "is_active": True}], "active": "guest"})
        return flask.jsonify({"is_guest": True, "profiles": [], "active": None})
    accounts = LinkedAccount.query.filter_by(user_id=current_user.id).all()
    active = next((a for a in accounts if a.is_active), None)
    return flask.jsonify({
        "is_guest": False,
        "email": current_user.email,
        "profiles": [{"id": a.profile_id, "name": a.name, "login_type": a.login_type, "is_active": a.is_active} for a in accounts],
        "active": active.profile_id if active else None
    })

@app.route("/profiles/switch", methods=["POST"])
def profiles_switch():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"})
    profile_id = (request.json or {}).get("id")
    acct = LinkedAccount.query.filter_by(user_id=current_user.id, profile_id=profile_id).first()
    if not acct:
        return flask.jsonify({"status": "error"})
    LinkedAccount.query.filter_by(user_id=current_user.id).update({"is_active": False})
    acct.is_active = True
    db.session.commit()
    return flask.jsonify({"status": "ok"})

@app.route("/profiles/delete", methods=["POST"])
def profiles_delete():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"})
    profile_id = (request.json or {}).get("id")
    acct = LinkedAccount.query.filter_by(user_id=current_user.id, profile_id=profile_id).first()
    if acct:
        db.session.delete(acct)
        db.session.commit()
    return flask.jsonify({"status": "ok"})

@app.route("/profiles/rename", methods=["POST"])
def profiles_rename():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"})
    profile_id = (request.json or {}).get("id")
    name = (request.json or {}).get("name", "").strip()
    acct = LinkedAccount.query.filter_by(user_id=current_user.id, profile_id=profile_id).first()
    if acct and name:
        acct.name = name
        db.session.commit()
    return flask.jsonify({"status": "ok"})

@app.route("/account/delete", methods=["POST"])
def account_delete():  # noqa: C901
    # NOTE: this whole body is wrapped in a top-level try/except so any
    # unexpected failure returns a clean JSON 500 instead of letting the
    # global Exception handler render the HTML error.html page (which
    # confuses every JS caller that expects JSON).
    try:
        return _account_delete_impl()
    except Exception as _outer_e:
        import traceback as _tb
        tb = _tb.format_exc()
        print(f"[account_delete] OUTER FAIL: {_outer_e}\n{tb}")
        return flask.jsonify({
            "status": "error",
            "message": "Account deletion failed. Please try again.",
            "debug": str(_outer_e)[:300],
        }), 500


def _account_delete_impl():
    """Delete the current user's account and every row that references them.

    The previous implementation just did `db.session.delete(user)`, which:
      - on PostgreSQL hits the first FK constraint that *isn't* on a cascading
        relationship (UserIdentity, ImportedGrade, ManualTask, every
        per-LMS integrations table, etc.) and raises IntegrityError →
        transaction rolls back → user row stays → "delete didn't work".
      - on SQLite without FK pragma enforcement, leaves orphan rows behind
        but at least removes the User row.

    The fix: explicitly nuke every table that holds a user_id FK first,
    inside a single transaction. We use bulk deletes so it stays fast
    even for users with thousands of TaskFeedback / StudySession rows.
    """
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "message": "Not signed in"}), 401
    from sqlalchemy import text as _t
    try:
        user_id = current_user.id
    except Exception as _ce:
        print(f"[account_delete] could not read current_user.id: {_ce}")
        return flask.jsonify({"status": "error", "message": "Session error"}), 500
    # We capture the id first, then logout so Flask-Login doesn't try to
    # re-load the about-to-be-deleted row at request teardown.
    try:
        logout_user()
    except Exception as _e:
        print(f"[account_delete] logout warning: {_e}")
    # Every table that holds a user_id FK. Keep this list in sync with
    # any new model that adds `user_id = db.Column(... ForeignKey("users.id"))`.
    user_owned_tables = [
        # Integrations
        "google_integrations", "notion_integrations", "canvas_integrations",
        "classroom_integrations", "blackboard_integrations", "moodle_integrations",
        # Identity / profile data
        "user_identities",
        # Tasks + scheduling
        "manual_tasks", "saved_schedules", "day_archive", "task_feedback",
        # Imports (new)
        "imported_grades",
        # Notes / lessons / study
        "course_notes", "lessons", "study_sessions", "study_points",
        "study_mastery", "session_messages", "syllabus_records",
        "saved_meetings", "live_sessions",
        # Push / reminders / groups / extension
        "push_subscriptions", "reminders_sent", "extension_tokens",
        "study_group_members", "test_marks", "custom_descriptions",
        "dismissed_assignments", "linked_accounts",
    ]
    # Per-statement connections so one failing table (missing on older DBs,
    # or a constraint we haven't enumerated) can't poison the whole
    # transaction. The previous attempt batched everything into one
    # SQLAlchemy session and the very first missing-table error aborted
    # the rest under Postgres semantics.
    try:
        db.session.close()  # release any session-bound transaction
    except Exception:
        pass
    deleted_any = False
    for table in user_owned_tables:
        try:
            with db.engine.connect() as conn:
                conn.execute(_t(f"DELETE FROM {table} WHERE user_id = :uid"), {"uid": user_id})
                try: conn.commit()
                except Exception: pass
            deleted_any = True
        except Exception as _de:
            # Most often: table doesn't exist on this DB. Safe to skip.
            print(f"[account_delete] skip {table}: {_de}")
    # Null out outgoing referrals so we don't cascade-delete other users.
    try:
        with db.engine.connect() as conn:
            conn.execute(_t("UPDATE users SET referred_by_id = NULL WHERE referred_by_id = :uid"), {"uid": user_id})
            try: conn.commit()
            except Exception: pass
    except Exception as _re:
        print(f"[account_delete] referral null-out failed: {_re}")
    # Finally drop the user row itself. This is the one delete we
    # actually require to succeed.
    try:
        with db.engine.connect() as conn:
            conn.execute(_t("DELETE FROM users WHERE id = :uid"), {"uid": user_id})
            try: conn.commit()
            except Exception: pass
    except Exception as e:
        print(f"[account_delete] FAILED for user {user_id}: {e}")
        return flask.jsonify({
            "status": "error",
            "message": "Could not delete your account right now. Try again or email support@intelliplan.tech.",
        }), 500
    session.clear()
    return flask.jsonify({"status": "ok"})

# ── DATA ROUTES ───────────────────────────────────────────────
@app.route("/live")
@limiter.limit("30 per minute")
def get_live_schedule():
    acct = get_active_account()
    if not acct:
        return flask.jsonify([])
    dismissed = get_dismissed_titles()
    test_titles = get_test_titles()
    excluded = dismissed | test_titles
    login_type = acct["login_type"]
    if login_type == "studentvue":
        try:
            print(f"Fetching StudentVue data for {acct['sv_username']}...")
            result = get_sv_assignments(acct["sv_district_url"], acct["sv_username"], acct["sv_password"])
            if not isinstance(result, list):
                print("StudentVue returned non-list data")
                return flask.jsonify([])
            filtered = [a for a in result if isinstance(a, dict) and a.get("title") not in excluded]
            return flask.jsonify(filtered)
        except Exception as e:
            print(f"StudentVue Live Error: {str(e)}")
            return flask.jsonify([]), 500
    if login_type == "schoology":
        try:
            from schoology_helper import get_schoology_assignments
            result = get_schoology_assignments(acct["schoology_key"], acct["schoology_secret"])
            return flask.jsonify([a for a in result if a.get("title", "") not in excluded])
        except Exception as e:
            print(f"Schoology Error: {e}")
            return flask.jsonify([])
    try:
        token = acct["canvas_token"]
        canvas_url = acct.get("canvas_url", "https://canvas.instructure.com")
        base = f"{canvas_url}/api/v1"
        headers = {"Authorization": f"Bearer {token}"}
        course_response = requests.get(f"{base}/courses", headers=headers, timeout=20)
        courses = course_response.json()
        course_map = {}
        for c in courses:
            if isinstance(c, dict) and "id" in c:
                course_map[c["id"]] = c.get("name", "Unknown")
        assignments = []
        for course_id in course_map:
            response = requests.get(f"{base}/courses/{course_id}/assignments?per_page=100", headers=headers, timeout=20)
            data = response.json()
            if isinstance(data, list):
                assignments += data
        schedule = []
        today = datetime.now(timezone.utc)
        for a in assignments:
            if not isinstance(a, dict): continue
            if a.get("due_at") is None: continue
            if a.get("points_possible") is None:
                a["points_possible"] = 60
            due_str = a["due_at"]
            due_date = datetime.fromisoformat(due_str.replace("Z", "+00:00"))
            days = (due_date - today).days
            if days < -14: continue
            priority = compute_priority(days, a["points_possible"], a.get("name", ""))
            rounded_minutes, description = _lms_row_sizing(a, a["points_possible"])
            difficulty = infer_task_difficulty(a["points_possible"], priority, due_str[:10])
            title = a["name"]
            if title in excluded: continue
            schedule.append({
                "id": str(a["id"]),
                "course_id": str(a["course_id"]),
                "title": title,
                "course": course_map.get(a["course_id"], "Unknown Course"),
                "due_date": due_str[:10],
                "points_possible": a["points_possible"],
                "priority": priority,
                "difficulty": difficulty,
                "estimated_time": rounded_minutes,
                "description": description,
                "color": PRIORITY_COLORS.get(priority, "#60a5fa"),
            })
        return flask.jsonify(sorted(schedule, key=lambda x: x["due_date"]))
    except Exception as e:
        print(f"Canvas live error: {e}")
        return flask.jsonify([])

@app.route("/courses")
def get_courses():
    acct = get_active_account()
    if not acct:
        return flask.jsonify([])
    login_type = acct["login_type"]
    if login_type == "studentvue":
        from studentvue_helper import get_courses as get_sv_courses
        return flask.jsonify(get_sv_courses(acct["sv_district_url"], acct["sv_username"], acct["sv_password"]))
    if login_type == "schoology":
        try:
            from schoology_helper import get_schoology_courses
            return flask.jsonify(get_schoology_courses(acct["schoology_key"], acct["schoology_secret"]))
        except Exception:
            return flask.jsonify([])
    token = acct["canvas_token"]
    canvas_url = acct.get("canvas_url", "https://canvas.instructure.com")
    headers = {"Authorization": f"Bearer {token}"}
    course_response = requests.get(f"{canvas_url}/api/v1/courses", headers=headers, timeout=20)
    courses = course_response.json()
    return flask.jsonify([{"name": c.get("name", "Unknown")} for c in courses if isinstance(c, dict) and "id" in c])

def _imported_grades_payload():
    """Read ImportedGrade rows (CSV / smart-paste / extension scraper) and
    return them in the same shape /grades/data normally produces."""
    q = ImportedGrade.query
    if current_user.is_authenticated:
        q = q.filter_by(user_id=current_user.id)
    else:
        q = q.filter_by(guest_session_id=session.get("guest_id"))
    rows = q.order_by(ImportedGrade.last_synced.desc()).all()
    if not rows:
        return []
    # De-dupe by course — keep the most-recently-synced row per course.
    seen, out = set(), []
    for r in rows:
        key = (r.course or "").lower()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append({
            "course": r.course, "class_name": r.course,
            "percentage": r.percentage, "grade": r.letter or "",
            "letter": r.letter or "", "teacher": r.teacher or "",
            "period": r.period or "",
            "source": r.source_label or r.source,
        })
    return out


@app.route("/grades/data")
def grades_data():
    acct = get_grade_account()
    if not acct:
        # No connected LMS — fall back to imported grades from CSV / paste /
        # extension scraper. This is how unsupported-LMS users see grades.
        return flask.jsonify(_imported_grades_payload())
    login_type = acct["login_type"]
    if login_type == "studentvue":
        from studentvue_helper import get_grades as get_sv_grades
        return flask.jsonify(get_sv_grades(acct["sv_district_url"], acct["sv_username"], acct["sv_password"]))
    if login_type == "schoology":
        try:
            from schoology_helper import get_schoology_grades
            return flask.jsonify(get_schoology_grades(acct["schoology_key"], acct["schoology_secret"]))
        except Exception:
            return flask.jsonify(_imported_grades_payload())
    if login_type == "canvas":
        try:
            from canvas_helper import get_grades as get_canvas_grades
            return flask.jsonify(get_canvas_grades(
                acct.get("canvas_url", "https://canvas.instructure.com"),
                acct["canvas_token"]
            ))
        except Exception as e:
            print(f"Canvas grades error: {e}")
            return flask.jsonify(_imported_grades_payload())
    return flask.jsonify(_imported_grades_payload())

@app.route("/gradebook/detail")
def gradebook_detail():
    acct = get_grade_account()
    if not acct:
        return flask.jsonify([])
    if acct["login_type"] == "studentvue":
        from studentvue_helper import get_gradebook_detail
        return flask.jsonify(get_gradebook_detail(acct["sv_district_url"], acct["sv_username"], acct["sv_password"]))
    if acct["login_type"] == "canvas":
        try:
            from canvas_helper import get_gradebook_detail as get_canvas_gradebook
            return flask.jsonify(get_canvas_gradebook(
                acct.get("canvas_url", "https://canvas.instructure.com"),
                acct["canvas_token"]
            ))
        except Exception as e:
            print(f"Canvas gradebook error: {e}")
            return flask.jsonify([])
    return flask.jsonify([])

@app.route("/dismissed/data")
def dismissed_data():
    rows = get_dismissed_rows()
    result = []
    for r in rows:
        try:
            result.append(json.loads(r.data))
        except Exception:
            result.append({"title": r.title})
    return flask.jsonify(result)

@app.route("/notes/list")
def notes_list():
    course_name = request.args.get("course_name", "").strip()
    course_id = request.args.get("course_id", "").strip()
    course_source = request.args.get("course_source", "").strip()
    q = get_notes_owner_query()
    if course_name:
        q = q.filter(db.func.lower(CourseNote.course_name) == course_name.lower())
    if course_id:
        q = q.filter(CourseNote.course_id == course_id)
    if course_source:
        q = q.filter(CourseNote.course_source == course_source)
    q = q.order_by(CourseNote.note_date.desc(), CourseNote.created_at.desc())
    payload = paginate_query(q, course_note_payload, default_size=30)
    payload["notes"] = payload["items"]
    return flask.jsonify(payload)

@app.route("/notes/upload", methods=["POST"])
def upload_note():
    course_name = request.form.get("course_name", "").strip()
    course_id = request.form.get("course_id", "").strip()
    course_source = request.form.get("course_source", "").strip()
    note_date = request.form.get("note_date", "").strip() or datetime.now().strftime("%Y-%m-%d")
    title = request.form.get("title", "").strip() or f"{course_name} Notes"
    if not course_name:
        return flask.jsonify({"status": "error", "message": "Course name is required"}), 400
    file = request.files.get("file")
    text_content = request.form.get("text_content", "").strip()
    original_filename = None
    stored_filename = None
    if file and file.filename:
        original_filename = file.filename
        ext = os.path.splitext(original_filename)[1].lower()
        if ext not in NOTE_ALLOWED_EXTENSIONS:
            return flask.jsonify({"status": "error", "message": "Only TXT, MD, CSV, PDF, and DOCX files are supported."}), 400
        owner_folder = get_notes_owner_folder()
        owner_dir = os.path.join(app.config["NOTES_UPLOAD_FOLDER"], owner_folder)
        os.makedirs(owner_dir, exist_ok=True)
        stored_filename = f"{uuid.uuid4().hex}{ext}"
        file_path = os.path.join(owner_dir, stored_filename)
        file.save(file_path)
        extracted = extract_text_from_note_file(file_path)
        if extracted:
            text_content = extracted
    if not text_content and not stored_filename:
        return flask.jsonify({"status": "error", "message": "Upload a note file or paste note text."}), 400
    note = CourseNote(
        user_id=current_user.id if current_user.is_authenticated else None,
        guest_session_id=None if current_user.is_authenticated else get_guest_session_id(),
        course_name=course_name,
        course_id=course_id or None,
        course_source=course_source or None,
        note_date=note_date,
        title=title,
        original_filename=original_filename,
        stored_filename=stored_filename,
        text_content=text_content or "",
    )
    db.session.add(note)
    db.session.commit()
    return flask.jsonify({"status": "ok", "note": course_note_payload(note)})

@app.route("/notes/<int:note_id>")
def get_note(note_id):
    note = db.session.get(CourseNote, note_id)
    if not note or not note_belongs_to_current_user(note):
        return flask.jsonify({"status": "error", "message": "Note not found"}), 404
    return flask.jsonify({"status": "ok", "note": course_note_payload(note, include_text=True)})

@app.route("/notes/<int:note_id>/download")
def download_note(note_id):
    note = db.session.get(CourseNote, note_id)
    if not note or not note.stored_filename or not note_belongs_to_current_user(note):
        return flask.jsonify({"status": "error", "message": "File not found"}), 404
    owner_dir = os.path.join(app.config["NOTES_UPLOAD_FOLDER"], f"user_{note.user_id}" if note.user_id else f"guest_{note.guest_session_id}")
    return flask.send_from_directory(owner_dir, note.stored_filename, as_attachment=True, download_name=note.original_filename or note.stored_filename)

@app.route("/notes/<int:note_id>/summarize", methods=["POST"])
def summarize_note(note_id):
    note = db.session.get(CourseNote, note_id)
    if not note or not note_belongs_to_current_user(note):
        return flask.jsonify({"status": "error", "message": "Note not found"}), 404
    if not (note.text_content or "").strip():
        return flask.jsonify({"status": "error", "message": "No extracted text is available for this note."}), 400
    if not ai_available():
        return flask.jsonify({"status": "error", "message": "AI summarization is not configured."}), 500
    text = (note.text_content or "")[:12000]
    prompt = f"""Summarize these class notes for a student.\n\nReturn:\n- 5 to 8 bullet points\n- a short "Key takeaways" section\n- keep it clear, practical, and concise\n\nNotes:\n{text}"""
    try:
        summary = ai_chat([{"role": "user", "content": prompt}], tier="standard", temperature=0.2, max_tokens=900)
        note.summary_cache = summary
        db.session.commit()
        return flask.jsonify({"status": "ok", "summary": summary})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": "AI summarization is temporarily unavailable. Please try again."}), 500

@app.route("/notes/<int:note_id>/study", methods=["POST"])
def study_note_route(note_id):
    note = db.session.get(CourseNote, note_id)
    if not note or not note_belongs_to_current_user(note):
        return flask.jsonify({"status": "error", "message": "Note not found"}), 404
    if not (note.text_content or "").strip():
        return flask.jsonify({"status": "error", "message": "No extracted text is available for this note."}), 400
    text = (note.text_content or "")[:12000]
    prompt = f"""Turn these notes into study material.\n\nReturn ONLY valid JSON:\n{{\n  "title": "Study Guide",\n  "summary": "short summary",\n  "cards": [{{"question": "Q1", "answer": "A1"}}],\n  "quiz": [{{"question": "Q1", "answer": "A1"}}]\n}}\n\nNotes:\n{text}"""
    try:
        raw = ai_chat([{"role": "user", "content": prompt}], tier="standard", temperature=0.2, max_tokens=1200)
        raw = re.sub(r"```json\s*", "", raw)
        raw = re.sub(r"```", "", raw).strip()
        try:
            study_data = json.loads(raw)
        except Exception:
            study_data = {"title": "Study Guide", "summary": raw, "cards": [], "quiz": []}
        return flask.jsonify({"status": "ok", "study": study_data})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": "Study generation is temporarily unavailable."}), 500

@app.route("/notes/<int:note_id>", methods=["DELETE"])
def delete_note(note_id):
    note = db.session.get(CourseNote, note_id)
    if not note or not note_belongs_to_current_user(note):
        return flask.jsonify({"status": "error", "message": "Note not found"}), 404
    if note.stored_filename:
        owner_dir = os.path.join(app.config["NOTES_UPLOAD_FOLDER"], f"user_{note.user_id}" if note.user_id else f"guest_{note.guest_session_id}")
        file_path = os.path.join(owner_dir, note.stored_filename)
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Could not remove note file: {e}")
    db.session.delete(note)
    db.session.commit()
    return flask.jsonify({"status": "ok"})

@app.route("/notes/<int:note_id>/quiz", methods=["POST"])
def notes_quiz(note_id):
    note = None
    if current_user.is_authenticated:
        note = CourseNote.query.filter_by(id=note_id, user_id=current_user.id).first()
    else:
        note = CourseNote.query.filter_by(id=note_id, guest_session_id=get_guest_session_id()).first()
    if not note:
        return flask.jsonify({"status": "error", "message": "Note not found"}), 404
    note_text = (note.text_content or "").strip()
    if not note_text:
        return flask.jsonify({"status": "error", "message": "No note text available"}), 400
    history = (request.json or {}).get("history", []) if request.is_json else []
    history_text = json.dumps(history[-8:], ensure_ascii=False)
    prompt = f"""Generate one study question from the note below.\n\nPrior questions:\n{history_text}\n\nNote:\n{note_text[:12000]}\n\nReturn JSON:\n{{\n  "question": "one question",\n  "answer": "one correct answer",\n  "key_points": ["point 1", "point 2"]\n}}"""
    try:
        raw = ai_chat([{"role": "user", "content": prompt}], tier="standard", temperature=0.5, max_tokens=900)
        raw = re.sub(r"```json\s*", "", raw)
        raw = re.sub(r"```", "", raw)
        quiz = json.loads(raw)
        return flask.jsonify({"status": "ok", "quiz": quiz})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

@app.route("/notes/<int:note_id>/file", methods=["GET"])
def notes_file(note_id):
    note = None
    if current_user.is_authenticated:
        note = CourseNote.query.filter_by(id=note_id, user_id=current_user.id).first()
    else:
        note = CourseNote.query.filter_by(id=note_id, guest_session_id=get_guest_session_id()).first()
    if not note:
        return flask.jsonify({"status": "error", "message": "Note not found"}), 404
    return flask.jsonify({"status": "ok", "view_url": getattr(note, "download_url", None), "filename": getattr(note, "original_filename", None), "text_content": getattr(note, "text_content", "")})

@app.route("/dismiss", methods=["POST"])
def dismiss():
    data = request.get_json(silent=True) or {}
    title = (data.get("title") or "").strip()
    if not title:
        return flask.jsonify({"status": "error", "message": "Missing title"}), 400
    try:
        save_dismissed(title, data)
        # Streak: dismissing a task = completing it
        if current_user.is_authenticated:
            try:
                _record_streak_qualifying_action(current_user.id, data.get("timezone"))
            except Exception as e:
                print(f"[streak] error on dismiss: {e}")
            try:
                _award_pet_xp(current_user.id, "task_completed", browser_tz=data.get("timezone"))
            except Exception as e:
                print(f"[pet] error on dismiss: {e}")
            # Bust the LMS cache so the dismissed task vanishes immediately
            try:
                invalidate_lms_cache_for_user(current_user.id)
            except Exception:
                pass
        return flask.jsonify({"status": "ok"})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

@app.route("/restore", methods=["POST"])
def restore():
    data = request.get_json(silent=True) or {}
    title = (data.get("title") or "").strip()
    if not title:
        return flask.jsonify({"status": "error", "message": "Missing title"}), 400
    try:
        delete_dismissed(title)
        return flask.jsonify({"status": "ok"})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

@app.route("/tests")
def tests_page():
    return flask.render_template("tests.html", active_page="tests",
                                 logged_in=current_user.is_authenticated)

@app.route("/api/tests", methods=["GET"])
def api_get_tests():
    rows = get_test_marks()
    result = []
    for r in rows:
        try:
            d = json.loads(r.data) if r.data else {}
        except Exception:
            d = {}
        d["title"] = r.title
        d["marked_at"] = r.created_at.isoformat() if r.created_at else ""
        result.append(d)
    return flask.jsonify(result)

@app.route("/test/mark", methods=["POST"])
def mark_as_test():
    data = request.get_json(silent=True) or {}
    title = (data.get("title") or "").strip()
    if not title:
        return flask.jsonify({"status": "error", "message": "Missing title"}), 400
    try:
        existing_titles = get_test_titles()
        if title in existing_titles:
            return flask.jsonify({"status": "ok", "already_marked": True})
        save_test_mark(title, data)
        return flask.jsonify({"status": "ok", "already_marked": False})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

@app.route("/test/unmark", methods=["POST"])
def unmark_as_test():
    data = request.get_json(silent=True) or {}
    title = (data.get("title") or "").strip()
    if not title:
        return flask.jsonify({"status": "error", "message": "Missing title"}), 400
    try:
        delete_test_mark(title)
        return flask.jsonify({"status": "ok"})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

@app.route("/assignment/description", methods=["GET"])
def get_description():
    assignment_id = request.args.get("id")
    course_id = request.args.get("course_id")
    title = request.args.get("title", "")
    custom = get_custom_description(title)
    if custom:
        return flask.jsonify({"description": custom, "source": "custom"})
    acct = get_active_account()
    if acct and acct["login_type"] == "canvas" and assignment_id and course_id:
        token = acct["canvas_token"]
        canvas_url = acct.get("canvas_url", "https://canvas.instructure.com")
        resp = requests.get(f"{canvas_url}/api/v1/courses/{course_id}/assignments/{assignment_id}", headers={"Authorization": f"Bearer {token}"}, timeout=20)
        if resp.status_code == 200:
            raw = resp.json().get("description") or ""
            clean = re.sub(r"<[^>]+>", " ", raw).strip()
            clean = re.sub(r"\s+", " ", clean)
            if clean:
                return flask.jsonify({"description": clean, "source": "canvas"})
    return flask.jsonify({"description": "", "source": "none"})

@app.route("/assignment/description", methods=["POST"])
def save_description():
    data = request.json or {}
    title = data.get("title")
    description = data.get("description", "").strip()
    if title and description:
        save_custom_description(title, description)
    return flask.jsonify({"status": "ok"})


@app.route("/api/assignment/notes", methods=["GET"])
def api_get_assignment_notes():
    title = request.args.get("title", "").strip()
    if not title:
        return flask.jsonify({"notes": "", "due_date": ""})
    raw = get_custom_description(title)
    if raw:
        try:
            data = json.loads(raw)
            return flask.jsonify({"notes": data.get("notes", ""), "due_date": data.get("due_date", "")})
        except Exception:
            return flask.jsonify({"notes": raw, "due_date": ""})
    return flask.jsonify({"notes": "", "due_date": ""})


@app.route("/api/assignment/notes", methods=["POST"])
def api_save_assignment_notes():
    body = request.get_json(silent=True) or {}
    title = (body.get("title") or "").strip()
    notes = (body.get("notes") or "").strip()
    if not title:
        return flask.jsonify({"status": "error", "message": "title required"}), 400
    raw = get_custom_description(title)
    try:
        existing = json.loads(raw) if raw else {}
    except Exception:
        existing = {"notes": raw} if raw else {}
    existing["notes"] = notes
    save_custom_description(title, json.dumps(existing))
    return flask.jsonify({"status": "ok"})


@app.route("/api/assignment/due-date", methods=["POST"])
def api_save_assignment_due_date():
    body = request.get_json(silent=True) or {}
    title = (body.get("title") or "").strip()
    due_date = (body.get("due_date") or "").strip()
    if not title:
        return flask.jsonify({"status": "error", "message": "title required"}), 400
    raw = get_custom_description(title)
    try:
        existing = json.loads(raw) if raw else {}
    except Exception:
        existing = {"notes": raw} if raw else {}
    existing["due_date"] = due_date
    save_custom_description(title, json.dumps(existing))
    return flask.jsonify({"status": "ok"})

def _planner_busy_by_date(horizon_days=14):
    """Dated committed time from the student's Google Calendar.

    Weekly commitments typed into settings recur; a dentist appointment does
    not. Until this was wired up the scheduler knew only about the recurring
    kind, so it would book an hour of chemistry directly on top of an event
    sitting right there in the calendar it already had permission to read.

    Every failure path returns ``{}``. A calendar we cannot reach means we
    know less about the student's week, not that they get no plan — and a
    scheduler that hard-fails on a third-party outage is worse than one that
    occasionally suggests a busy hour.
    """
    if not current_user.is_authenticated:
        return {}
    try:
        token = get_google_token()
        if not token or not has_calendar_scope(token):
            return {}
        from google_calendar_helper import busy_minutes_by_date

        offset = getattr(current_user, "utc_offset_minutes", 0) or 0
        return busy_minutes_by_date(
            token, date.today(), days=horizon_days, utc_offset_minutes=offset
        )
    except Exception as e:
        print(f"[planner] calendar busy lookup failed: {e}")
        return {}


def _lms_row_sizing(raw, points_possible, kind=""):
    """``(minutes, description)`` for one raw LMS payload.

    Every connector used to carry its own copy of ``points_possible × 1.5``
    rounded to the half hour. That number cannot distinguish two assignments
    inside one course, so this routes all of them through the shared sizing
    module and keeps the old heuristic only as the fallback it always should
    have been.

    ``description`` is returned as well so the caller can put it on the task
    dict — the planner reads it again later, and re-fetching it per assignment
    would be a network call per row.
    """
    description = ""
    if isinstance(raw, dict):
        for key in ("description", "intro", "instructions", "body", "summary"):
            value = raw.get(key)
            if value:
                description = str(value)
                break
    try:
        points = float(points_possible or 0)
    except (TypeError, ValueError):
        points = 0.0
    try:
        from intelliplan.intelligence.sizing import size_from_metadata, strip_markup

        clean = strip_markup(description)
        sized = size_from_metadata(
            title=str((raw or {}).get("name") or (raw or {}).get("title") or "")
            if isinstance(raw, dict) else "",
            kind=kind,
            description=clean,
            points_possible=points or None,
            submission_types=(raw or {}).get("submission_types") if isinstance(raw, dict) else None,
            rubric_rows=len(raw["rubric"]) if isinstance(raw, dict) and isinstance(raw.get("rubric"), list) else 0,
        )
        if sized.is_measured:
            return sized.minutes, clean[:4000]
        return max(30, round((points or 60) * 1.5 / 30) * 30), clean[:4000]
    except Exception as e:
        print(f"[sizing] fell back to the points heuristic: {e}")
        return max(30, round((points or 60) * 1.5 / 30) * 30), ""


def _sized_estimate(assignment, kind, saved_description=None):
    """Base minutes for one assignment, before the student's own bias model.

    Order of trust, highest first:

    1. **What the student told us.** A duration they typed into the clarify
       prompt is not a guess we should overrule.
    2. **What the work describes.** Word counts, page ranges, problem counts,
       rubrics, quiz time limits — see
       :mod:`intelliplan.intelligence.sizing`. This is the layer that can tell
       "problems 12–18" from "write a 2,000-word essay" when both are worth
       fifty points.
    3. **What the upstream helper guessed.** Historically
       ``points_possible × 1.5``, which is better than nothing and worse than
       either of the above.

    Returns ``(minutes, subtask_count, signal_details)``.
    """
    from intelliplan.intelligence.sizing import size_assignment

    try:
        declared = int(assignment.get("estimated_time") or 0)
    except (TypeError, ValueError):
        declared = 0

    if declared > 0 and str(assignment.get("estimate_source") or "").lower() == "student":
        return declared, 0, []

    payload = {
        **assignment,
        "kind": kind,
        "description": (
            assignment.get("description")
            or assignment.get("instructions")
            or saved_description
            or ""
        ),
    }
    sized = size_assignment(payload)
    details = [s.detail for s in sized.signals]
    if sized.is_measured:
        return sized.minutes, sized.subtask_count, details
    # Nothing in the metadata described the work, so keep whatever the source
    # already guessed rather than replacing one prior with another.
    return (declared or sized.minutes), sized.subtask_count, []


def _saved_descriptions_for(assignments):
    """``{title: description}`` from what the student has typed themselves.

    Students paste the real instructions into IntelliPlan for assignments
    whose LMS description is empty — that text is the best sizing signal in
    the system and until now it was only ever rendered, never read. Loaded in
    one query rather than one per assignment.
    """
    titles = [
        (a.get("title") or "").strip()
        for a in assignments
        if isinstance(a, dict) and (a.get("title") or "").strip()
    ]
    if not titles:
        return {}
    try:
        query = CustomDescription.query.filter(
            CustomDescription.assignment_title.in_(titles[:300])
        )
        if current_user.is_authenticated:
            query = query.filter(CustomDescription.user_id == current_user.id)
        else:
            query = query.filter(
                CustomDescription.guest_session_id == get_guest_session_id()
            )
        out = {}
        for row in query.limit(300).all():
            text = row.description or ""
            # The notes endpoint stores a JSON blob in the same column.
            if text.lstrip().startswith("{"):
                try:
                    text = (json.loads(text) or {}).get("notes") or ""
                except Exception:
                    pass
            if text:
                out[row.assignment_title] = text
        return out
    except Exception as e:
        print(f"[planner] saved description load failed: {e}")
        return {}


def _planner_task_rows(normalized_assignments, custom_tasks, descriptions=None):
    """Flatten assignments and free-text custom tasks into planner input.

    Custom tasks carry no metadata at all — they are a title someone typed
    — so they get neutral defaults and let the estimation model's kind-level
    prior do the sizing. Inventing a due date for them would be worse than
    having none: the planner treats a missing deadline as "no deadline
    pressure", which is the truth, rather than as urgency we made up.
    """
    _KIND_WORDS = (
        ("exam", ("exam", "final", "midterm")),
        ("test", ("test", "quiz")),
        ("project", ("project", "essay", "paper", "presentation", "report")),
        ("lab", ("lab",)),
        ("classwork", ("classwork", "worksheet", "notes")),
    )

    def kind_of(title, declared=""):
        declared = (declared or "").strip().lower()
        if declared in ("exam", "test", "project", "lab", "homework", "classwork"):
            return declared
        text = (title or "").lower()
        for kind, words in _KIND_WORDS:
            if any(w in text for w in words):
                return kind
        return "homework"

    # One priority engine for the whole product. The scheduler used to run a
    # three-bucket lookup while the Command Center ran the real five-component
    # model — so the surface that *acts* on priority, deciding what gets pulled
    # earlier and what gets dropped when the week is over capacity, was using
    # the weaker of the two.
    scores = {}
    try:
        from intelliplan.services.prioritisation import score_rows

        scores = score_rows(normalized_assignments, date.today())
    except Exception as e:
        print(f"[planner] priority scoring failed, using labels: {e}")

    rows = []
    for a in normalized_assignments:
        title = a.get("title") or "Task"
        kind = kind_of(title, a.get("kind") or a.get("type"))
        est_minutes, subtasks, size_signals = _sized_estimate(
            a, kind, (descriptions or {}).get(title)
        )
        row_id = str(a.get("id") or a.get("source_ref") or title)
        scored = scores.get(row_id) or scores.get(title)
        rows.append({
            "id": row_id,
            "title": title,
            "course": a.get("course") or "",
            "kind": kind,
            "due_date": a.get("due_date") or "",
            "est_minutes": est_minutes,
            "difficulty": (a.get("difficulty") or "Medium").lower(),
            "priority": scored.score if scored else _priority_score_for(a),
            "priority_reasons": list(scored.reasons) if scored else [],
            "points_possible": a.get("points_possible"),
            "subtask_count": max(len(a.get("subtasks") or []), subtasks),
            "size_signals": size_signals,
            "description": (a.get("description") or (descriptions or {}).get(title) or ""),
        })
    for title in custom_tasks or []:
        title = (title or "").strip()
        if not title:
            continue
        rows.append({
            "id": f"custom:{title}",
            "title": title,
            "course": "",
            "kind": kind_of(title),
            "due_date": "",
            "difficulty": "medium",
            "priority": 50,
        })
    return rows


def _priority_score_for(assignment):
    """Map the UI's three-level priority onto the planner's 0..100 scale.

    Deliberately not a bare 30/50/80 lookup: an assignment worth a large
    share of the grade matters more than one that is not, whatever bucket
    the label puts it in.
    """
    base = {"high": 80, "medium": 50, "low": 30}.get(
        str(assignment.get("priority") or "Medium").strip().lower(), 50
    )
    try:
        points = float(assignment.get("points_possible") or 0)
    except (TypeError, ValueError):
        points = 0
    if points >= 100:
        base += 10
    elif points >= 50:
        base += 5
    return max(0, min(100, base))


def _build_planner_schedule(normalized_assignments, custom_tasks, uid, gid,
                            availability, commitments, preferred_time, hours_per_day,
                            dna=None):
    """Run the v2 planner and return ``schedule_data``, or None to fall back.

    Returning None rather than raising is deliberate: if anything here is
    wrong, the student gets the old AI path and a schedule, instead of an
    error page. The failure is logged loudly so it does not stay hidden.
    """
    try:
        from intelliplan.intelligence.planner import PlannerConfig
        from intelliplan.services.scheduling import SchedulingService, StudentContext

        rows = _planner_task_rows(
            normalized_assignments,
            custom_tasks,
            descriptions=_saved_descriptions_for(normalized_assignments),
        )
        if not rows:
            return None

        try:
            comfort_minutes = int(round(float(hours_per_day) * 60)) or None
        except (TypeError, ValueError):
            comfort_minutes = None

        context = StudentContext(
            availability=availability,
            commitments=commitments,
            preferred_time=preferred_time,
            feedback_rows=_planner_feedback_rows(uid, gid),
            session_rows=_planner_session_rows(uid, gid),
            concept_mastery=_planner_concept_mastery(uid),
            # Measured under-delivery days. The planner prices them rather
            # than banning them; without this the day_quality weight has
            # nothing to act on and every day looks equally good.
            weak_days=tuple(getattr(dna, "weak_days", ()) or ()),
            # The student's stated hours-per-day is a ceiling on comfort, not
            # on capacity: their availability windows are the hard limit, and
            # this sets how full those windows get before the optimizer starts
            # charging for it.
            daily_target_minutes=comfort_minutes,
            # Real calendar events. Study time booked over a dentist
            # appointment is a plan the student cannot follow.
            busy_by_date=_planner_busy_by_date(),
        )
        config = PlannerConfig()
        service = SchedulingService(context, config)
        plan = service.plan(rows)
        data = service.to_schedule_data(plan)

        # Reuse the existing enrichment so the Interactive View, checklists,
        # and redirects keep working exactly as they do on the AI path.
        try:
            data = enrich_schedule_data(
                data, normalized_assignments, preferred_time, hours_per_day
            )
        except Exception as ee:
            print(f"[planner] enrich_schedule_data failed (non-fatal): {ee}")
        return data
    except Exception as pe:
        import traceback
        print(f"[planner] v2 planning failed, falling back to AI path: {pe}")
        traceback.print_exc()
        return None


def _planner_feedback_rows(uid, gid):
    """``TaskFeedback`` rows the estimation model fits on."""
    try:
        query = TaskFeedback.query
        if uid:
            query = query.filter(TaskFeedback.user_id == uid)
        elif gid:
            query = query.filter(
                TaskFeedback.user_id.is_(None),
                TaskFeedback.guest_session_id == gid,
            )
        else:
            return []
        rows = query.order_by(TaskFeedback.completed_at.desc()).limit(300).all()
        return [{
            "estimated_time": r.estimated_time,
            "actual_time": r.actual_time,
            "course": r.course or "",
            "completed_at": r.completed_at,
        } for r in rows]
    except Exception as e:
        print(f"[planner] feedback load failed: {e}")
        return []


def _planner_session_rows(uid, gid):
    """Active-study sittings — richer than feedback, and the source of truth
    for how long this student actually sits still."""
    try:
        from intelliplan.repositories.active_sessions import ActiveSessionRepository

        repo = ActiveSessionRepository(ActiveSession, ActiveFocusSample, db.session)
        return repo.observations(user_id=uid, guest_id=gid)
    except Exception as e:
        print(f"[planner] session load failed: {e}")
        return []


def _planner_concept_mastery(uid):
    """``{concept: mastery}`` so shaky material gets more room."""
    if not uid:
        return {}
    try:
        rows = ConceptMastery.query.filter(ConceptMastery.user_id == uid).limit(500).all()
        out = {}
        for r in rows:
            name = (getattr(r, "concept", "") or "").strip().lower()
            score = getattr(r, "mastery_score", None)
            if name and score is not None:
                out[name] = max(0.0, min(1.0, float(score)))
        return out
    except Exception as e:
        print(f"[planner] concept mastery load failed: {e}")
        return {}


@app.route("/generate_schedule", methods=["GET", "POST"])
# 10/hour was sized for a route whose every call cost an AI round trip.
# Planning is now deterministic and free, and re-planning is something a
# student should be able to do as often as their week changes — which, in
# the week before finals, is many times a day. The limit still exists
# because the AI fallback path is still reachable.
@limiter.limit("60 per hour", methods=["POST"])
def generate_schedule():
    if request.method != "POST":
        return flask.jsonify({"status": "error", "message": "POST required"}), 405
    data = request.json or {}
    assignments = data.get("assignments", [])
    hours_per_day = data.get("hours_per_day", 2)
    preferred_time = data.get("preferred_time", "evening")
    custom_tasks = data.get("custom_tasks", [])
    if not assignments and not custom_tasks:
        return flask.jsonify({"status": "error", "message": "No assignments to schedule."})

    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()

    # ── Clarification gate ────────────────────────────────────────
    # A task called "Study" with no subject and no goal can only ever produce
    # a block that says "Study" back at the student. Rather than dress that up
    # with better prompt wording, ask for the missing pieces — once — and
    # remember the answers.
    clarifications = data.get("clarifications") or {}
    presets = load_scheduler_presets(uid, gid)
    # Free-text custom tasks are just titles; give them the same shape so one
    # code path assesses everything.
    task_views = [*normalized_custom_task_views(custom_tasks), *assignments]
    preset_answers = scheduler_clarify.answers_from_presets(task_views, presets)
    effective_answers = {**preset_answers, **clarifications}

    # Persist new answers immediately, before the gate re-checks. Clarifying
    # can take more than one round (we only ask three things at a time), and
    # saving after the gate meant round one's answers were thrown away every
    # time round two was needed.
    if clarifications:
        save_scheduler_presets(
            scheduler_clarify.preset_payload_from_answers(
                clarifications, scheduler_clarify.labels_for(task_views)),
            uid, gid)

    if not data.get("skip_clarify"):
        pending = [
            q for q in scheduler_clarify.assess_tasks(
                scheduler_clarify.apply_answers(task_views, effective_answers),
                known_courses=[a.get("course", "") for a in assignments if a.get("course")],
            )
        ]
        if pending:
            offered = {
                k: v for k, v in presets.items()
                if any(scheduler_clarify.preset_key(q.task) == k for q in pending)
            }
            return flask.jsonify({
                "status": "needs_clarification",
                "questions": [q.to_dict() for q in pending],
                "presets": offered,
                "message": "A couple of quick questions so this plan is actually about your work.",
            })

    used_presets = []
    if effective_answers:
        applied = scheduler_clarify.apply_answers(task_views, effective_answers)
        # Split back out: custom tasks stay strings, assignments stay dicts.
        n_custom = len(custom_tasks)
        custom_tasks = [t.get("title", "") for t in applied[:n_custom]]
        assignments = applied[n_custom:]
        matched = {k for k in presets if any(
            scheduler_clarify.preset_key(t.get("title") or "") == k for t in task_views)}
        _mark_presets_used(matched, uid, gid)
        # Tell the student a saved answer was used. Silently reusing something
        # they typed weeks ago and never mentioning it is how a planner starts
        # feeling like it's ignoring them.
        used_presets = [
            {"task_key": k, "label": presets[k].get("label") or k,
             "answers": presets[k].get("answers") or {}}
            for k in sorted(matched)
        ]

    normalized_assignments = []
    for assignment in assignments:
        difficulty = assignment.get("difficulty") or infer_task_difficulty(
            assignment.get("points_possible"),
            assignment.get("priority", "Medium"),
            assignment.get("due_date"),
        )
        normalized_assignments.append({
            **assignment,
            "difficulty": difficulty,
            "color": assignment.get("color") or PRIORITY_COLORS.get(assignment.get("priority", "Medium"), "#60a5fa"),
        })
    # Measured habits + real weekly availability. These drive deterministic
    # block placement after the model returns; the prompt blocks built from
    # them only help the model propose durations that survive placement.
    dna, availability, commitments = build_scheduler_personalization(
        user_id=current_user.id if current_user.is_authenticated else None,
        guest_id=None if current_user.is_authenticated else get_guest_session_id(),
    )
    # ── Deterministic planning ────────────────────────────────────
    # Day allocation — how much work exists, how it splits into sittings,
    # which day each lands on, how much buffer sits before each deadline —
    # is decided by the planner, not by a language model. The model was
    # never able to reason about the student's measured pace, their real
    # free hours, or the interaction between six deadlines in one week; it
    # produced plausible-looking days that placement then had to repair.
    #
    # This also takes an AI round trip off the critical path, so the
    # scheduler stays fast and keeps working when the provider is down.
    #
    # Flag-gated so the previous behaviour is one toggle away.
    if feature_enabled("planner_v2"):
        planned = _build_planner_schedule(
            normalized_assignments, custom_tasks, uid, gid,
            availability, commitments, preferred_time, hours_per_day,
            dna=dna,
        )
        if planned is not None:
            return flask.jsonify({
                "status": "ok",
                "data": planned,
                "used_presets": used_presets,
            })

    # Re-baseline estimates against how long work actually takes THIS student,
    # before they reach the prompt — a plan built on "60 min" for someone who
    # reliably needs 90 is a plan they will fall behind on by lunchtime.
    #
    # This applies to the AI path *only*. The planner has its own estimation
    # model, fit on the same TaskFeedback rows, and it applies its own learned
    # ratio inside build_plan(). Correcting here as well multiplied the two
    # together: a student who reliably runs 1.5x over was sized at 2.25x, and
    # the plan they got was a third longer than the week they have.
    for a in normalized_assignments:
        try:
            raw_est = int(a.get("estimated_time") or 60)
        except (TypeError, ValueError):
            raw_est = 60
        a["estimated_time"] = dna.adjust_estimate(raw_est, a.get("course", ""))

    today_str = datetime.now().strftime("%Y-%m-%d")
    overdue = [a for a in normalized_assignments if a.get("due_date", "9999") < today_str]
    upcoming = [a for a in normalized_assignments if a.get("due_date", "9999") >= today_str]
    upcoming.sort(key=lambda x: x.get("due_date", "9999"))
    if not ai_available():
        return flask.jsonify({"status": "error", "message": API_ERROR_MESSAGES["ai"], "retryable": True}), 503
    overdue_text = ""
    if overdue:
        overdue_text = f"\nOVERDUE — MUST BE SCHEDULED TODAY ({len(overdue)} assignments):\n" + "\n".join([
            f"  ⚠ {a['title']} ({a['course']}) — was due {a['due_date']}, Priority: HIGH, Est: {a['estimated_time']}min"
            for a in overdue
        ])
    upcoming_text = ""
    if upcoming:
        upcoming_text = f"\nUPCOMING ({len(upcoming)} assignments):\n" + "\n".join([
            f"  - {a['title']} ({a['course']}) — Due: {a['due_date']}, Priority: {a['priority']}, Difficulty: {a['difficulty']}, Est: {a['estimated_time']}min"
            for a in upcoming
        ])
    custom_text = ""
    if custom_tasks:
        custom_text = f"\nCUSTOM TASKS ADDED BY STUDENT — use EXACT names as written ({len(custom_tasks)}):\n" + "\n".join([f"  - {t}" for t in custom_tasks])
    today = datetime.now().strftime("%Y-%m-%d")
    total = len(normalized_assignments) + len(custom_tasks)
    # Build personalized student context. When the user has opted in
    # (Settings → Privacy), we layer in grade signals so the AI can weight
    # weaker subjects more carefully and play to strengths. When opted-out,
    # build_student_context() returns "" and the scheduler behaves
    # exactly as it did before personalization existed.
    grades_summary = None
    if current_user.is_authenticated and _ai_personalization_enabled():
        try:
            grades_summary = _summarize_grade_signals(_fetch_grades_for_personalization())
        except Exception as _ge:
            print(f"[scheduler] grade summary failed: {_ge}")
    profile_context = build_student_context(
        user_id=current_user.id if current_user.is_authenticated else None,
        grades_summary=grades_summary,
        depth="full",
    )
    week_context = scheduler_engine.describe_week(availability, commitments)
    habits_context = dna.to_prompt() if _ai_personalization_enabled() else ""
    prompt = f"""You are IntelliPlan — an adaptive academic study-planning system. Today is {today}.

You must schedule ALL {total} items below. Every single one must appear in the schedule.
{profile_context}{week_context}{habits_context}{overdue_text}
{upcoming_text}
{custom_text}

Student availability: {hours_per_day} hours/day, prefers {preferred_time}.
The estimates above are already corrected for how long work actually takes
this student — treat them as accurate and do not pad them further.

RULES:
1. ALL {total} items must appear in the schedule — no exceptions
2. Overdue items go on Day 1 as first priority blocks
3. Custom task names must be copied EXACTLY as written — do not rename them
4. Spread assignments across multiple days — max 3 assignments per day unless unavoidable
5. Split long assignments (>90min) across multiple days
6. Add a 10min break after every 45min work block
7. Never put the same assignment twice in one day
8. Schedule must end before the latest due date
9. If REAL WEEK is present above, put NO work on days marked "no study time
   available", and keep each day's total within that day's listed hours.
   Blocks that don't fit get pushed to the next day automatically, so an
   overstuffed day silently loses work — size days honestly.

PERSONALIZATION (only if STUDENT CONTEXT is present above):
- Allocate ~30% more study time to assignments in subjects listed under "needing more time".
- For weak-subject blocks, use the notes field to suggest a concrete starting tactic
  (e.g. "review the last quiz error first", "redo example problems before starting").
- For strong-subject blocks, keep blocks shorter and frame notes around polish/depth,
  not foundations.
- Never reference specific past grades back to the student in the notes — the context
  is for YOUR planning only, the student does not need to read their own GPA back.

WRITE FOR THIS STUDENT, NOT A GENERIC ONE:
- "notes" must name the actual next physical action for THAT assignment
  ("outline the three body paragraphs", "redo problems 12–18"), never filler
  like "focus on this task", "work steadily", or "review the material".
- "daily_tip" must reference something concrete about THAT day — the specific
  assignment that matters most, or the shape of that day's free time. A tip
  that would read identically on any other day is a failed tip.

Return ONLY valid JSON:
{{
  "schedule": [
    {{
      "date": "YYYY-MM-DD",
      "day_name": "Monday",
      "total_hours": {hours_per_day},
      "blocks": [
        {{
          "assignment": "Exact title here",
          "course": "Course name",
          "duration_minutes": 45,
          "time_slot": "7:00 PM - 7:45 PM",
          "notes": "What to focus on",
          "is_break": false
        }}
      ],
      "daily_tip": "Actionable tip"
    }}
  ],
  "overview": "Plan covering all {total} items",
  "total_study_time": "X hours Y minutes"
}}"""
    # Scale the ceiling to the size of the plan. A 30-item schedule simply
    # needs more room to write than a 3-item one, and running out mid-object
    # is what surfaces to the student as "the AI returned an invalid schedule".
    max_tokens = max(8000, min(32000, 2000 + total * 900))
    try:
        schedule_data = None
        last_err = None
        result = ""
        for attempt in range(2):
            try:
                result = ai_chat(
                    [{"role": "user", "content": prompt}],
                    tier="standard",
                    temperature=0.3,
                    max_tokens=max_tokens,
                    response_format={"type": "json_object"},
                    # Second pass drops reasoning entirely — if the first
                    # attempt failed, a plainer plan beats another error page.
                    thinking_budget=0 if attempt else None,
                )
                schedule_data = _parse_schedule_json(result)
                break
            except (json.JSONDecodeError, ValueError) as pe:
                last_err = pe
                print(f"[scheduler] schedule JSON parse failed (attempt {attempt + 1}): {pe}")
                print(f"[scheduler] raw head: {str(result)[:300]!r}")
                print(f"[scheduler] raw tail: {str(result)[-300:]!r}")
        if schedule_data is None:
            raise last_err or json.JSONDecodeError("no schedule", "", 0)
        schedule_data = enrich_schedule_data(schedule_data, normalized_assignments, preferred_time, hours_per_day)
        # Adaptive humanization pass — fix spacing, anti-cluster, attach
        # checklist + redirect data the Interactive View needs.
        try:
            schedule_data = humanize_schedule(
                schedule_data, preferred_time, hours_per_day,
                availability=availability, commitments=commitments, dna=dna,
            )
        except Exception as he:
            print(f"[scheduler] humanize_schedule failed (non-fatal): {he}")
        return flask.jsonify({"status": "ok", "data": schedule_data,
                              "used_presets": used_presets})
    except json.JSONDecodeError:
        return flask.jsonify({"status": "error", "message": "The AI returned an invalid schedule. Please try again.", "retryable": True})
    except Exception as e:
        err_str = str(e).lower()
        if "rate" in err_str or "429" in err_str:
            return flask.jsonify({"status": "error", "message": "AI usage limit reached. Please wait a minute and try again.", "retryable": True}), 429
        if "timeout" in err_str:
            return flask.jsonify({"status": "error", "message": "The AI took too long to respond. Please try again.", "retryable": True}), 504
        print(f"Schedule generation error: {e}")
        return flask.jsonify({"status": "error", "message": API_ERROR_MESSAGES["ai"], "retryable": True}), 503

@app.route("/static/sw.js")
def service_worker():
    response = flask.make_response(flask.send_from_directory("static", "sw.js"))
    response.headers["Content-Type"] = "application/javascript"
    response.headers["Service-Worker-Allowed"] = "/"
    return response

# ── NOTION OAUTH ──────────────────────────────────────────────
@app.route("/oauth/notion")
def oauth_notion_start():
    """Begin the Notion OAuth flow. Requires NOTION_CLIENT_ID."""
    if not NOTION_AVAILABLE:
        return redirect(url_for("settings"))
    if not os.getenv("NOTION_CLIENT_ID"):
        return render_template(
            "error.html",
            active_page="error",
            error_code=503,
            error_id="NOTION-OAUTH-NOT-CONFIGURED",
            message="Notion OAuth isn't set up on this deployment yet. "
                    "Use the integration token form instead.",
        ), 503
    state = secrets_module.token_urlsafe(24)
    session["notion_oauth_state"] = state
    session.modified = True
    redirect_uri = os.getenv("NOTION_REDIRECT_URI") or (
        APP_BASE_URL.rstrip("/") + "/oauth/notion/callback"
    )
    try:
        url = get_notion_auth_url(state, redirect_uri=redirect_uri)
    except Exception as e:
        return render_template("error.html", active_page="error", error_code=500,
                               error_id=f"NOTION-OAUTH-{make_error_id()}",
                               message=str(e)), 500
    return redirect(url)


@app.route("/oauth/notion/callback")
def oauth_notion_callback():
    if not NOTION_AVAILABLE:
        return redirect(url_for("settings"))
    code = request.args.get("code")
    state = request.args.get("state")
    expected_state = session.pop("notion_oauth_state", None)
    if not code or not state or state != expected_state:
        return render_template("error.html", active_page="error", error_code=400,
                               error_id="NOTION-OAUTH-STATE",
                               message="Notion OAuth state did not match. Try again."), 400
    redirect_uri = os.getenv("NOTION_REDIRECT_URI") or (
        APP_BASE_URL.rstrip("/") + "/oauth/notion/callback"
    )
    try:
        tokens = exchange_notion_code(code, redirect_uri=redirect_uri)
    except Exception as e:
        return render_template("error.html", active_page="error", error_code=500,
                               error_id=f"NOTION-OAUTH-{make_error_id()}",
                               message=str(e)), 500

    access_token = tokens.get("access_token")
    if not access_token:
        return render_template("error.html", active_page="error", error_code=500,
                               error_id="NOTION-OAUTH-NO-TOKEN",
                               message="Notion did not return an access token."), 500

    # Persist on the user (or in session for guests).
    session["notion_token"] = access_token
    session.modified = True
    if current_user.is_authenticated:
        existing = NotionIntegration.query.filter_by(user_id=current_user.id).first()
        if existing:
            existing.token = access_token
            existing.auth_type = "oauth"
            existing.workspace_id = tokens.get("workspace_id")
            existing.workspace_name = tokens.get("workspace_name")
            existing.workspace_icon = tokens.get("workspace_icon")
            existing.bot_id = tokens.get("bot_id")
            existing.connected_at = utcnow()
            # Don't clobber database_id if the user already chose one.
        else:
            db.session.add(NotionIntegration(
                user_id=current_user.id,
                token=access_token,
                auth_type="oauth",
                workspace_id=tokens.get("workspace_id"),
                workspace_name=tokens.get("workspace_name"),
                workspace_icon=tokens.get("workspace_icon"),
                bot_id=tokens.get("bot_id"),
            ))
        db.session.commit()
    return redirect(url_for("settings") + "?notion=connected")


@app.route("/notion/status")
def notion_status():
    """Mirror of /calendar/events — does the dashboard widget have data?"""
    if not NOTION_AVAILABLE:
        return flask.jsonify({"connected": False, "configured": False})
    token, db_id = get_notion_token_and_db()
    if not token:
        return flask.jsonify({
            "connected": False,
            "configured": True,
            "oauth_available": bool(os.getenv("NOTION_CLIENT_ID")),
        })
    workspace = None
    if current_user.is_authenticated:
        ni = NotionIntegration.query.filter_by(user_id=current_user.id).first()
        if ni and ni.workspace_name:
            workspace = {
                "id": ni.workspace_id,
                "name": ni.workspace_name,
                "icon": ni.workspace_icon,
                "auth_type": ni.auth_type or "manual",
                "connected_at": ni.connected_at.isoformat() if ni.connected_at else None,
            }
    return flask.jsonify({
        "connected": True,
        "configured": True,
        "has_database": bool(db_id),
        "workspace": workspace,
        "oauth_available": bool(os.getenv("NOTION_CLIENT_ID")),
    })


@app.route("/notion/upcoming")
def notion_upcoming():
    """Notion analog of /calendar/events — used by the dashboard widget."""
    if not NOTION_AVAILABLE:
        return flask.jsonify({"connected": False, "tasks": []})
    token, db_id = get_notion_token_and_db()
    if not token or not db_id:
        return flask.jsonify({"connected": False, "tasks": []})
    try:
        days = int(request.args.get("days", 7))
        tasks = get_upcoming_notion_tasks(token, db_id, days=days)
        return flask.jsonify({"connected": True, "tasks": tasks})
    except Exception as e:
        return flask.jsonify({"connected": False, "error": safe_error_message(e), "tasks": []})


@app.route("/notion/export", methods=["POST"])
def notion_export_schedule():
    """Push a generated schedule to Notion — mirror of /calendar/export."""
    if not NOTION_AVAILABLE:
        return flask.jsonify({"status": "error", "message": "Notion not configured"}), 503
    token, db_id = get_notion_token_and_db()
    if not token or not db_id:
        return flask.jsonify({"status": "error", "message": "Notion not connected"}), 400
    schedule = (request.get_json(silent=True) or {}).get("schedule") or {}
    try:
        created, skipped = add_schedule_to_notion(token, db_id, schedule)
        return flask.jsonify({"status": "ok", "created": len(created), "skipped": skipped, "ids": created})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


# ── NOTION ────────────────────────────────────────────────────
@app.route("/notion/connect", methods=["POST"])
def notion_connect():
    if not NOTION_AVAILABLE:
        return flask.jsonify({
            "status": "error",
            "message": "Notion library is not installed on this server. The admin needs to add notion-client to requirements.txt and redeploy."
        }), 503
    token = (request.json or {}).get("token", "").strip()
    if not token:
        return flask.jsonify({"status": "error", "message": "No token provided"}), 400
    ok, info = test_notion_token_detail(token)
    if not ok:
        return flask.jsonify({"status": "error", "message": f"Notion rejected the token: {info}"}), 400
    # Populate workspace metadata from the bot's user record so the UI can
    # show "Connected to <workspace>" even for manual-token connections.
    bot_info = info if isinstance(info, dict) else {}
    bot = (bot_info.get("bot") or {}) if bot_info.get("type") == "bot" else {}
    workspace_name = bot.get("workspace_name") or bot_info.get("name") or "Notion workspace"
    bot_id = bot_info.get("id")
    session["notion_token"] = token
    session.modified = True
    if current_user.is_authenticated:
        existing = NotionIntegration.query.filter_by(user_id=current_user.id).first()
        if existing:
            existing.token = token
            # Keep the previously-selected database so re-pasting the token to
            # refresh the connection doesn't silently stop the task sync. The
            # user can still change it from "Pick a database".
            existing.auth_type = "manual"
            existing.workspace_name = workspace_name
            existing.bot_id = bot_id
            existing.connected_at = utcnow()
        else:
            db.session.add(NotionIntegration(
                user_id=current_user.id,
                token=token,
                auth_type="manual",
                workspace_name=workspace_name,
                bot_id=bot_id,
            ))
        db.session.commit()
    try:
        dbs = get_notion_databases(token)
    except Exception as e:
        return flask.jsonify({
            "status": "ok",
            "databases": [],
            "warning": f"Connected, but could not list databases yet: {e}. Share a database with your integration in Notion, then refresh."
        })
    return flask.jsonify({"status": "ok", "databases": dbs})

@app.route("/notion/disconnect", methods=["POST"])
def notion_disconnect():
    if current_user.is_authenticated:
        NotionIntegration.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
    else:
        session.pop("notion_token", None)
        session.pop("notion_database_id", None)
    return flask.jsonify({"status": "ok"})

@app.route("/notion/set-database", methods=["POST"])
def notion_set_database():
    db_id = (request.json or {}).get("database_id")
    if not db_id:
        return flask.jsonify({"status": "error"})
    if current_user.is_authenticated:
        ni = NotionIntegration.query.filter_by(user_id=current_user.id).first()
        if ni:
            ni.database_id = db_id
            db.session.commit()
        else:
            token = session.get("notion_token")
            if token:
                db.session.add(NotionIntegration(user_id=current_user.id, token=token, database_id=db_id))
                db.session.commit()
    session["notion_database_id"] = db_id
    session.modified = True
    return flask.jsonify({"status": "ok"})

@app.route("/notion/databases")
def notion_databases_route():
    if not NOTION_AVAILABLE:
        return flask.jsonify({"status": "error", "databases": []}), 503
    token, db_id = get_notion_token_and_db()
    if not token:
        return flask.jsonify({"status": "error", "message": "Notion not connected", "databases": []}), 400
    try:
        return flask.jsonify({
            "status": "ok",
            "selected": db_id,
            "databases": get_notion_databases(token),
        })
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e), "databases": []}), 500


@app.route("/notion/pages")
def notion_pages_route():
    """Pages the integration is allowed to see — used as parent options
    when creating a new IntelliPlan database in Notion."""
    if not NOTION_AVAILABLE:
        return flask.jsonify({"status": "error", "pages": []}), 503
    token, _ = get_notion_token_and_db()
    if not token:
        return flask.jsonify({"status": "error", "message": "Notion not connected", "pages": []}), 400
    try:
        return flask.jsonify({"status": "ok", "pages": get_shared_pages(token)})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e), "pages": []}), 500


@app.route("/notion/create-database", methods=["POST"])
def notion_create_database_route():
    """One-click "set up Notion for me" — create a tasks database with
    the right schema under a page the user has shared with us."""
    if not NOTION_AVAILABLE:
        return flask.jsonify({"status": "error", "message": "Notion library not installed"}), 503
    token, _ = get_notion_token_and_db()
    if not token:
        return flask.jsonify({"status": "error", "message": "Notion not connected"}), 400
    body = request.json or {}
    parent_id = body.get("parent_page_id")
    name = (body.get("name") or "IntelliPlan Tasks").strip()
    if not parent_id:
        return flask.jsonify({"status": "error", "message": "Pick a parent page first"}), 400
    try:
        new_id = create_intelliplan_database(token, parent_id, name=name)
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500
    # Auto-select the new database as the active target.
    if current_user.is_authenticated:
        ni = NotionIntegration.query.filter_by(user_id=current_user.id).first()
        if ni:
            ni.database_id = new_id
            db.session.commit()
    session["notion_database_id"] = new_id
    session.modified = True
    return flask.jsonify({"status": "ok", "database_id": new_id, "name": name})

@app.route("/notion/tasks")
def notion_tasks_route():
    if not NOTION_AVAILABLE:
        return flask.jsonify({"connected": False, "tasks": []})
    token, db_id = get_notion_token_and_db()
    if not token or not db_id:
        return flask.jsonify({"connected": False, "tasks": []})
    try:
        tasks = get_notion_tasks(token, db_id)
        return flask.jsonify({"connected": True, "tasks": tasks})
    except Exception as e:
        return flask.jsonify({"connected": False, "error": safe_error_message(e), "tasks": []})

@app.route("/notion/tasks/create", methods=["POST"])
def notion_create_task():
    if not NOTION_AVAILABLE:
        return flask.jsonify({"status": "error"})
    token, db_id = get_notion_token_and_db()
    if not token or not db_id:
        return flask.jsonify({"status": "error", "message": "Notion not connected"})
    data = request.json or {}
    try:
        page_id = create_notion_task(token, db_id, data.get("title", ""), data.get("due_date"), data.get("priority", "Medium"))
        return flask.jsonify({"status": "ok", "page_id": page_id})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/notion/tasks/update", methods=["POST"])
def notion_update_task():
    if not NOTION_AVAILABLE:
        return flask.jsonify({"status": "error"})
    token, _ = get_notion_token_and_db()
    if not token:
        return flask.jsonify({"status": "error"})
    data = request.json or {}
    try:
        update_notion_task(token, data["page_id"], data.get("updates", {}))
        return flask.jsonify({"status": "ok"})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/notion/tasks/complete", methods=["POST"])
def notion_complete_task():
    if not NOTION_AVAILABLE:
        return flask.jsonify({"status": "error"})
    token, _ = get_notion_token_and_db()
    if not token:
        return flask.jsonify({"status": "error"})
    page_id = (request.json or {}).get("page_id")
    try:
        complete_notion_task(token, page_id)
        return flask.jsonify({"status": "ok"})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

# ── UNIFIED TASKS ─────────────────────────────────────────────
# ── LMS aggregation cache ────────────────────────────────────────────
# The command center renders on every page load. Without this cache,
# each render does N+1 HTTP calls against Canvas (1 for courses + 1 per
# course) plus full StudentVue/Schoology pulls — easily 8–30s wall-clock.
# 5-minute TTL keeps the surface fresh while making the page feel instant.
#
# In-memory only on purpose: a single Railway/Gunicorn worker pinned to
# one user serves them most of the time. For multi-worker fan-out we can
# move this to Redis later.

_LMS_CACHE: dict[int, tuple[float, list[dict]]] = {}
_LMS_CACHE_TTL_SECONDS = 300  # 5 minutes


def _lms_cache_get(user_id: int) -> list[dict] | None:
    """Return cached LMS data if it's still fresh, else None."""
    import time as _time
    entry = _LMS_CACHE.get(user_id)
    if not entry:
        return None
    ts, data = entry
    if _time.time() - ts > _LMS_CACHE_TTL_SECONDS:
        _LMS_CACHE.pop(user_id, None)
        return None
    return data


def _lms_cache_put(user_id: int, data: list[dict]) -> None:
    import time as _time
    _LMS_CACHE[user_id] = (_time.time(), data)


def invalidate_lms_cache_for_user(user_id: int) -> None:
    """Force the next render to re-fetch from upstream. Called when the
    user creates/dismisses/edits a manual task or finishes an LMS connect."""
    _LMS_CACHE.pop(user_id, None)


def collect_lms_assignments_for_user(user_id: int, *, use_cache: bool = True) -> list[dict]:
    """Return canonical LMS-source task dicts for the given user.

    Pulls live data from Canvas / StudentVue / Schoology / Classroom /
    Blackboard / Moodle / Notion based on the user's active account and
    LMS preferences. Used by:
      - /tasks/unified (legacy route)
      - the Command Center repository (real data, not mocks)
      - the agentic Plani chatbot

    Cached per user for ``_LMS_CACHE_TTL_SECONDS`` to make repeat
    page loads instant. Pass ``use_cache=False`` to bypass.

    Failures from any single source are isolated — one broken provider
    must not blank out the entire feed.
    """
    from datetime import date as date_type

    if not user_id:
        return []
    if use_cache:
        cached = _lms_cache_get(user_id)
        if cached is not None:
            return cached

    tasks: list[dict] = []
    today = date_type.today()
    try:
        dismissed = _DismissedSet(d.title for d in DismissedAssignment.query.filter_by(user_id=user_id).all())
    except Exception:
        dismissed = _DismissedSet([])

    # Active LMS account (resolves through the LinkedAccount table)
    acct_dict = None
    try:
        acct = LinkedAccount.query.filter_by(user_id=user_id, is_active=True).first()
        if acct:
            creds = acct.get_credentials() or {}
            creds["login_type"] = acct.login_type
            if not creds.get("canvas_url"):
                creds["canvas_url"] = "https://canvas.instructure.com"
            acct_dict = creds
    except Exception:
        acct_dict = None

    if acct_dict:
        login_type = acct_dict["login_type"]
        if login_type == "studentvue":
            try:
                raw = get_sv_assignments(acct_dict["sv_district_url"], acct_dict["sv_username"], acct_dict["sv_password"])
                if isinstance(raw, list):
                    for a in raw:
                        if isinstance(a, dict) and a.get("title") not in dismissed:
                            a["source"] = "studentvue"
                            a.setdefault("priority", "Medium")
                            tasks.append(a)
            except Exception as e:
                print(f"[lms-collect] SV err: {e}")
            try:
                missing_raw = get_missing_assignments(acct_dict["sv_district_url"], acct_dict["sv_username"], acct_dict["sv_password"])
                if isinstance(missing_raw, list):
                    for a in missing_raw:
                        if isinstance(a, dict) and a.get("title") not in dismissed:
                            a["source"] = "studentvue_missing"
                            a.setdefault("priority", "High")
                            tasks.append(a)
            except Exception as e:
                print(f"[lms-collect] SV missing err: {e}")
        elif login_type == "canvas":
            try:
                from concurrent.futures import ThreadPoolExecutor, as_completed
                token = acct_dict["canvas_token"]
                base = f"{acct_dict['canvas_url']}/api/v1"
                headers = {"Authorization": f"Bearer {token}"}
                # Single 6s budget for the courses list — anything slower
                # means Canvas is degraded and we'd rather show stale data
                course_response = requests.get(f"{base}/courses", headers=headers, timeout=6)
                courses = course_response.json()
                course_map = {c["id"]: c.get("name", "Unknown") for c in courses if isinstance(c, dict) and "id" in c}

                # Fan out per-course requests in parallel — 8 concurrent
                # cuts a 10-course pull from ~50s sequential to ~6s.
                def _fetch_course(course_id):
                    try:
                        r = requests.get(
                            f"{base}/courses/{course_id}/assignments?per_page=100",
                            headers=headers, timeout=6,
                        )
                        return course_id, r.json()
                    except Exception:
                        return course_id, []

                with ThreadPoolExecutor(max_workers=8) as pool:
                    futures = [pool.submit(_fetch_course, cid) for cid in course_map]
                    for fut in as_completed(futures, timeout=15):
                        try:
                            course_id, data = fut.result()
                        except Exception:
                            continue
                        if not isinstance(data, list):
                            continue
                        for a in data:
                            if not isinstance(a, dict) or not a.get("due_at"):
                                continue
                            due_str = a["due_at"][:10]
                            try:
                                due = datetime.strptime(due_str, "%Y-%m-%d").date()
                            except Exception:
                                continue
                            days = (due - today).days
                            if days < -14:
                                continue
                            title = a.get("name") or ""
                            if not title or title in dismissed:
                                continue
                            priority = compute_priority(days, a.get("points_possible") or 0, title)
                            est_minutes, description = _lms_row_sizing(
                                a, a.get("points_possible")
                            )
                            tasks.append({
                                "id": str(a["id"]),
                                "title": title,
                                "course": course_map.get(a.get("course_id"), course_map.get(course_id, "Unknown")),
                                "due_date": due_str,
                                "priority": priority,
                                "source": "canvas",
                                "estimated_time": est_minutes,
                                "description": description,
                                "difficulty": "Medium",
                            })
            except Exception as e:
                print(f"[lms-collect] canvas err: {e}")
        elif login_type == "schoology":
            try:
                from schoology_helper import get_schoology_assignments
                raw = get_schoology_assignments(acct_dict["schoology_key"], acct_dict["schoology_secret"])
                if isinstance(raw, list):
                    for a in raw:
                        if isinstance(a, dict) and a.get("title") not in dismissed:
                            a["source"] = "schoology"
                            a.setdefault("priority", "Medium")
                            tasks.append(a)
            except Exception as e:
                print(f"[lms-collect] schoology err: {e}")

    # Google Classroom (independent of active account)
    try:
        ctok, _crow = _classroom_access_token_for(user_id)
        if ctok:
            for a in _classroom_fetch_assignments(ctok):
                if a.get("title") and a["title"] not in dismissed:
                    tasks.append(a)
    except Exception as e:
        print(f"[lms-collect] classroom err: {e}")

    # Blackboard
    try:
        btok, brow = _blackboard_access_token_for(user_id)
        if btok and brow and brow.institution_url:
            for a in _blackboard_fetch_assignments(brow.institution_url, btok, brow.bb_user_id):
                if a.get("title") and a["title"] not in dismissed:
                    tasks.append(a)
    except Exception as e:
        print(f"[lms-collect] blackboard err: {e}")

    # Moodle
    try:
        mrow = MoodleIntegration.query.filter_by(user_id=user_id).order_by(MoodleIntegration.id.desc()).first()
        if mrow and mrow.moodle_url and mrow.ws_token:
            for a in _moodle_fetch_assignments(mrow.moodle_url, mrow.ws_token, mrow.moodle_user_id):
                if a.get("title") and a["title"] not in dismissed:
                    tasks.append(a)
    except Exception as e:
        print(f"[lms-collect] moodle err: {e}")

    # Deduplicate (title + course + due_date)
    seen = set()
    deduped = []
    for t in tasks:
        title = (t.get("title") or "").strip().lower()
        course = (t.get("course") or "").strip().lower()
        due_date = (t.get("due_date") or "").strip()
        key = (title, course, due_date)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(t)
    # Cache so the next render is instant
    if use_cache:
        _lms_cache_put(user_id, deduped)
    return deduped


@app.route("/tasks/unified")
def unified_tasks():
    from datetime import date as date_type

    def dedupe_tasks(task_list):
        seen = set()
        unique = []
        for t in task_list:
            title = (t.get("title") or "").strip().lower()
            course = (t.get("course") or "").strip().lower()
            due_date = (t.get("due_date") or "").strip()
            key = (title, course, due_date)
            if key in seen:
                continue
            seen.add(key)
            unique.append(t)
        return unique

    tasks = []
    dismissed = get_dismissed_titles()
    today = date_type.today()
    priority_order = {"High": 0, "Medium": 1, "Low": 2}
    acct = get_active_account()
    if acct:
        login_type = acct["login_type"]
        if login_type == "studentvue":
            try:
                raw = get_sv_assignments(acct["sv_district_url"], acct["sv_username"], acct["sv_password"])
                if isinstance(raw, list):
                    for a in raw:
                        if isinstance(a, dict) and a.get("title") not in dismissed:
                            a["source"] = "studentvue"
                            if "priority" not in a:
                                a["priority"] = "Medium"
                            a.setdefault("color", PRIORITY_COLORS.get(a.get("priority", "Medium"), "#f59e0b"))
                            tasks.append(a)
            except Exception as e:
                print(f"SV assignments error: {e}")
            try:
                missing_raw = get_missing_assignments(acct["sv_district_url"], acct["sv_username"], acct["sv_password"])
                if isinstance(missing_raw, list):
                    for a in missing_raw:
                        if isinstance(a, dict) and a.get("title") not in dismissed:
                            a["source"] = "studentvue_missing"
                            if "priority" not in a:
                                a["priority"] = "High"
                            a.setdefault("color", PRIORITY_COLORS.get(a.get("priority", "High"), "#ef4444"))
                            tasks.append(a)
            except Exception as e:
                print(f"Missing assignments error: {e}")
        elif login_type == "canvas":
            try:
                token = acct["canvas_token"]
                canvas_url = acct.get("canvas_url", "https://canvas.instructure.com")
                base = f"{canvas_url}/api/v1"
                headers = {"Authorization": f"Bearer {token}"}
                course_response = requests.get(f"{base}/courses", headers=headers, timeout=20)
                courses = course_response.json()
                course_map = {c["id"]: c.get("name", "Unknown") for c in courses if isinstance(c, dict) and "id" in c}
                for course_id in course_map:
                    resp = requests.get(f"{base}/courses/{course_id}/assignments?per_page=100", headers=headers, timeout=20)
                    data = resp.json()
                    if not isinstance(data, list):
                        continue
                    for a in data:
                        if not isinstance(a, dict) or not a.get("due_at"):
                            continue
                        due_str = a["due_at"][:10]
                        try:
                            due = datetime.strptime(due_str, "%Y-%m-%d").date()
                        except Exception:
                            continue
                        days = (due - today).days
                        if days < -14:
                            continue
                        title = a["name"]
                        priority = compute_priority(days, a.get("points_possible") or 0, title)
                        if title in dismissed:
                            continue
                        est_minutes, description = _lms_row_sizing(
                            a, a.get("points_possible")
                        )
                        tasks.append({
                            "id": str(a["id"]),
                            "course_id": str(a["course_id"]),
                            "title": title,
                            "course": course_map.get(a["course_id"], "Unknown"),
                            "due_date": due_str,
                            "priority": priority,
                            "source": "canvas",
                            "estimated_time": est_minutes,
                            "description": description,
                            "difficulty": "Medium",
                            "color": PRIORITY_COLORS.get(priority, "#f59e0b")
                        })
            except Exception as e:
                print(f"Canvas unified error: {e}")
        elif login_type == "schoology":
            try:
                from schoology_helper import get_schoology_assignments
                raw = get_schoology_assignments(acct["schoology_key"], acct["schoology_secret"])
                if isinstance(raw, list):
                    for a in raw:
                        if isinstance(a, dict) and a.get("title") not in dismissed:
                            a["source"] = "schoology"
                            a.setdefault("priority", "Medium")
                            a.setdefault("color", PRIORITY_COLORS.get(a.get("priority", "Medium"), "#f59e0b"))
                            tasks.append(a)
            except Exception as e:
                print(f"Schoology unified error: {e}")

    lms_prefs = get_user_lms_prefs()
    assignment_sources = set(lms_prefs.get("assignment_sources") or _default_lms_prefs()["assignment_sources"])

    # ── Google Classroom (independent of active LMS account) ──
    # A user may have Classroom connected alongside Canvas or StudentVue, so
    # we always check for stored Classroom tokens, not just the active account.
    try:
        if current_user.is_authenticated and "google_classroom" in assignment_sources:
            ctok, _crow = _classroom_access_token_for(current_user.id)
            if ctok:
                for a in _classroom_fetch_assignments(ctok):
                    if a.get("title") in dismissed:
                        continue
                    tasks.append(a)
    except Exception as e:
        print(f"Classroom unified error: {e}")

    # ── Blackboard Learn (independent of active LMS account) ──
    try:
        if current_user.is_authenticated and "blackboard" in assignment_sources:
            btok, brow = _blackboard_access_token_for(current_user.id)
            if btok and brow and brow.institution_url:
                for a in _blackboard_fetch_assignments(brow.institution_url, btok, brow.bb_user_id):
                    if a.get("title") in dismissed:
                        continue
                    tasks.append(a)
    except Exception as e:
        print(f"Blackboard unified error: {e}")

    # ── Moodle (independent of active LMS account) ──
    try:
        if current_user.is_authenticated and "moodle" in assignment_sources:
            mrow = MoodleIntegration.query.filter_by(user_id=current_user.id).order_by(MoodleIntegration.id.desc()).first()
            if mrow and mrow.moodle_url and mrow.ws_token:
                for a in _moodle_fetch_assignments(mrow.moodle_url, mrow.ws_token, mrow.moodle_user_id):
                    if a.get("title") in dismissed:
                        continue
                    tasks.append(a)
    except Exception as e:
        print(f"Moodle unified error: {e}")

    if NOTION_AVAILABLE and "notion" in assignment_sources:
        try:
            notion_token, notion_db_id = get_notion_token_and_db()
            if notion_token and notion_db_id:
                notion_raw = get_notion_tasks(notion_token, notion_db_id)
                for t in notion_raw:
                    if t.get("title") not in dismissed:
                        tasks.append(t)
        except Exception as e:
            print(f"Notion tasks error: {e}")
    try:
        if "manual" in assignment_sources:
            if current_user.is_authenticated:
                manual = ManualTask.query.filter_by(user_id=current_user.id, done=False).all()
            else:
                gid = get_guest_session_id()
                manual = ManualTask.query.filter_by(guest_session_id=gid, done=False).all()
        else:
            manual = []
        for t in manual:
            if t.title not in dismissed:
                tasks.append({
                    "id": t.id,
                    "title": t.title,
                    "due_date": t.due_date or "",
                    "priority": t.priority,
                    "course": t.course,
                    "estimated_time": t.estimated_time,
                    "notes": t.notes,
                    "source": "manual",
                    "notion_page_id": t.notion_page_id,
                    "color": PRIORITY_COLORS.get(t.priority, "#f59e0b")
                })
    except Exception as e:
        print(f"Manual tasks error: {e}")
    tasks = dedupe_tasks(tasks)
    result = {"today": [], "upcoming": [], "overdue": []}
    for t in tasks:
        due = t.get("due_date", "")
        if not due:
            result["upcoming"].append(t)
            continue
        try:
            due_date = datetime.strptime(due, "%Y-%m-%d").date()
            if due_date < today:
                result["overdue"].append(t)
            elif due_date == today:
                result["today"].append(t)
            else:
                result["upcoming"].append(t)
        except Exception:
            result["upcoming"].append(t)
    for key in result:
        result[key].sort(key=lambda x: (x.get("due_date", "9999-12-31"), priority_order.get(x.get("priority", "Low"), 2)))
    return flask.jsonify(result)

@app.route("/missing/data")
def missing_data():
    acct = get_active_account()
    if not acct:
        return flask.jsonify([])
    login_type = acct["login_type"]
    try:
        if login_type == "studentvue":
            return flask.jsonify(get_missing_assignments(
                acct["sv_district_url"], acct["sv_username"], acct["sv_password"]
            ))
        if login_type == "canvas":
            from canvas_helper import get_missing_assignments as get_canvas_missing
            return flask.jsonify(get_canvas_missing(
                acct.get("canvas_url", "https://canvas.instructure.com"),
                acct["canvas_token"]
            ))
    except Exception as e:
        print(f"Missing data error ({login_type}): {e}")
        return flask.jsonify([])
    return flask.jsonify([])

# ── MANUAL TASKS ──────────────────────────────────────────────
@app.route("/tasks/manual/create", methods=["POST"])
def manual_create_task():
    data = request.json or {}
    title = data.get("title", "").strip()
    if not title:
        return flask.jsonify({"status": "error", "message": "Title required"})
    task = ManualTask(
        user_id=current_user.id if current_user.is_authenticated else None,
        guest_session_id=None if current_user.is_authenticated else get_guest_session_id(),
        title=title,
        due_date=data.get("due_date", ""),
        priority=data.get("priority", "Medium"),
        course=data.get("course", "Personal"),
        estimated_time=int(data.get("estimated_time", 60)),
        notes=data.get("notes", "")
    )
    db.session.add(task)
    db.session.commit()
    if NOTION_AVAILABLE and data.get("sync_notion"):
        notion_token, notion_db_id = get_notion_token_and_db()
        if notion_token and notion_db_id:
            try:
                page_id = create_notion_task(notion_token, notion_db_id, title, data.get("due_date"), data.get("priority", "Medium"))
                task.notion_page_id = page_id
                db.session.commit()
            except Exception:
                pass
    if current_user.is_authenticated:
        try: invalidate_lms_cache_for_user(current_user.id)
        except Exception: pass
    return flask.jsonify({"status": "ok", "id": task.id})

def _owns_manual_task(task) -> bool:
    """True if the current session owns this ManualTask.

    Works for both authenticated users (user_id) and guests
    (guest_session_id) so we close the IDOR without breaking guest task
    editing. Prevents mutating another user's task by guessing its id.
    """
    if task is None:
        return False
    if current_user.is_authenticated:
        return task.user_id == current_user.id
    gid = get_guest_session_id()
    return bool(task.guest_session_id) and task.guest_session_id == gid


@app.route("/tasks/manual/update", methods=["POST"])
def manual_update_task():
    data = request.json or {}
    task_id = data.get("id")
    task = db.session.get(ManualTask, task_id)
    if not _owns_manual_task(task):
        return flask.jsonify({"status": "error", "message": "Not found"}), 404
    if "title" in data: task.title = data["title"]
    if "due_date" in data: task.due_date = data["due_date"]
    if "priority" in data: task.priority = data["priority"]
    if "course" in data: task.course = data["course"]
    if "estimated_time" in data: task.estimated_time = int(data["estimated_time"])
    if "notes" in data: task.notes = data["notes"]
    was_done_before = task.done
    if "done" in data: task.done = data["done"]
    db.session.commit()
    if NOTION_AVAILABLE and task.notion_page_id:
        notion_token, _ = get_notion_token_and_db()
        if notion_token:
            try:
                update_notion_task(notion_token, task.notion_page_id, data)
            except Exception:
                pass
    # Streak: trigger on task completion (done flipped False → True)
    if data.get("done") and not was_done_before and current_user.is_authenticated:
        try:
            _record_streak_qualifying_action(current_user.id, data.get("timezone"))
        except Exception as e:
            print(f"[streak] error on task complete: {e}")
        try:
            _award_pet_xp(current_user.id, "task_completed", browser_tz=data.get("timezone"))
        except Exception as e:
            print(f"[pet] error on task complete: {e}")
    if current_user.is_authenticated:
        try: invalidate_lms_cache_for_user(current_user.id)
        except Exception: pass
    return flask.jsonify({"status": "ok"})

@app.route("/tasks/manual/delete", methods=["POST"])
def manual_delete_task():
    task_id = (request.json or {}).get("id")
    task = db.session.get(ManualTask, task_id)
    if task and not _owns_manual_task(task):
        return flask.jsonify({"status": "error", "message": "Not found"}), 404
    if task:
        db.session.delete(task)
        db.session.commit()
    if current_user.is_authenticated:
        try: invalidate_lms_cache_for_user(current_user.id)
        except Exception: pass
    return flask.jsonify({"status": "ok"})


# ── Task-completion streak engine ────────────────────────────────────

def _is_streak_enabled_for_user(user_id: int) -> bool:
    """Check if streak_v1 flag is enabled for this user."""
    try:
        flag = FeatureFlag.query.filter_by(key="streak_v1").first()
        if not flag:
            return False
        if not flag.enabled:
            return False
        pct = flag.rollout_percentage if flag.rollout_percentage is not None else 100
        return streak_engine.is_user_in_rollout(user_id, "streak_v1", pct)
    except Exception:
        return False


def _get_or_create_streak(user_id: int) -> UserStreak:
    row = UserStreak.query.filter_by(user_id=user_id).first()
    if not row:
        row = UserStreak(user_id=user_id)
        db.session.add(row)
        db.session.commit()
    return row


def _record_streak_qualifying_action(user_id: int, browser_tz: str | None = None) -> dict | None:
    """Core streak update path. Called from task-complete and plan-review."""
    if not _is_streak_enabled_for_user(user_id):
        return None

    row = _get_or_create_streak(user_id)

    # Resolve timezone: prefer stored, fall back to browser-sent, then UTC
    tz_name = row.timezone or browser_tz or "UTC"
    if not row.timezone and browser_tz:
        row.timezone = browser_tz

    last_date = None
    if row.last_qualifying_local_date:
        try:
            last_date = date.fromisoformat(row.last_qualifying_local_date)
        except (ValueError, TypeError):
            last_date = None

    result = streak_engine.compute_streak_update(
        current_streak=row.current_streak,
        longest_streak=row.longest_streak,
        last_qualifying_local_date=last_date,
        freezes_available=row.freezes_available,
        freezes_used_total=row.freezes_used_total,
        user_tz=tz_name,
    )

    today_local = streak_engine.resolve_local_date(tz_name)

    # Persist
    row.current_streak = result.current_streak
    row.longest_streak = result.longest_streak
    row.freezes_available = result.freezes_available
    row.last_qualifying_action_at = datetime.now(timezone.utc)
    row.last_qualifying_local_date = today_local.isoformat()

    # Track qualified dates for the week dots
    try:
        qualified = set(json.loads(row.qualified_dates_json or "[]"))
    except (json.JSONDecodeError, TypeError):
        qualified = set()
    qualified.add(today_local.isoformat())
    # Keep only last 30 days to avoid unbounded growth
    cutoff = (today_local - timedelta(days=30)).isoformat()
    qualified = {d for d in qualified if d >= cutoff}
    row.qualified_dates_json = json.dumps(sorted(qualified))

    if result.event == "streak_freeze_consumed":
        row.freezes_used_total += result.event_props.get("days_covered", 0)

    db.session.commit()

    # Analytics
    if result.event:
        app_analytics.track(user_id, result.event, result.event_props)
    if result.freeze_earned:
        app_analytics.track(user_id, "streak_freeze_earned", {
            "milestone_day": result.current_streak,
        })

    # ── Plani Pet: feed XP for streak milestones (3, 7, 14, 30, 60, 100)
    try:
        milestone_event = pet_engine.streak_milestone_event(result.current_streak)
        if milestone_event:
            _award_pet_xp(user_id, milestone_event, browser_tz=browser_tz)
    except Exception as e:
        print(f"[pet] streak-milestone error: {e}")

    return {
        "current_streak": result.current_streak,
        "longest_streak": result.longest_streak,
        "freezes_available": result.freezes_available,
        "toast": result.toast_message,
        "event": result.event,
    }


@app.route("/api/streak/status")
def streak_status():
    """Return the task-completion streak state for the current user."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "message": "Not logged in"}), 401
    if not _is_streak_enabled_for_user(current_user.id):
        return flask.jsonify({"status": "ok", "enabled": False})
    row = _get_or_create_streak(current_user.id)
    tz_name = row.timezone or "UTC"
    # Visiting the app counts as a qualifying action — but only the first
    # status fetch of each local day actually advances the streak, so the
    # nav badge polling this endpoint doesn't write on every call.
    try:
        _today_iso = streak_engine.resolve_local_date(tz_name).isoformat()
        if row.last_qualifying_local_date != _today_iso:
            _record_streak_qualifying_action(current_user.id)
            row = _get_or_create_streak(current_user.id)
    except Exception as _streak_e:
        print(f"[streak] visit-record skipped: {_streak_e}")
    try:
        qualified = set()
        for d in json.loads(row.qualified_dates_json or "[]"):
            try:
                qualified.add(date.fromisoformat(d))
            except (ValueError, TypeError):
                pass
    except (json.JSONDecodeError, TypeError):
        qualified = set()
    dots = streak_engine.week_dots(tz_name, qualified)
    last_date = None
    if row.last_qualifying_local_date:
        try:
            last_date = date.fromisoformat(row.last_qualifying_local_date)
        except (ValueError, TypeError):
            pass
    show_nudge = streak_engine.should_show_nudge(
        current_streak=row.current_streak,
        last_qualifying_local_date=last_date,
        user_tz=tz_name,
        nudge_shown_today=(row.nudge_shown_date == streak_engine.resolve_local_date(tz_name).isoformat()),
    )
    cohort = streak_engine.user_cohort(current_user.id, "streak_v1",
        getattr(FeatureFlag.query.filter_by(key="streak_v1").first(), "rollout_percentage", 100) or 100)
    app_analytics.identify(current_user.id, {"streak_v1_cohort": cohort})
    return flask.jsonify({
        "status": "ok",
        "enabled": True,
        "current_streak": row.current_streak,
        "longest_streak": row.longest_streak,
        "freezes_available": row.freezes_available,
        "freezes_max": 3,
        "week_dots": dots,
        "show_nudge": show_nudge,
        "timezone": tz_name,
        "cohort": cohort,
    })


@app.route("/api/streak/nudge-shown", methods=["POST"])
def streak_nudge_shown():
    """Mark the nudge as shown for today so it doesn't repeat."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    row = _get_or_create_streak(current_user.id)
    tz_name = row.timezone or "UTC"
    row.nudge_shown_date = streak_engine.resolve_local_date(tz_name).isoformat()
    db.session.commit()
    app_analytics.track(current_user.id, "nudge_shown")
    return flask.jsonify({"status": "ok"})


@app.route("/api/streak/nudge-tapped", methods=["POST"])
def streak_nudge_tapped():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    app_analytics.track(current_user.id, "nudge_tapped")
    return flask.jsonify({"status": "ok"})


@app.route("/api/streak/pill-tapped", methods=["POST"])
def streak_pill_tapped():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    row = _get_or_create_streak(current_user.id)
    app_analytics.track(current_user.id, "streak_pill_tapped", {
        "current_streak": row.current_streak,
    })
    return flask.jsonify({"status": "ok"})


@app.route("/api/streak/plan-review", methods=["POST"])
def streak_plan_review():
    """Triggered when the user views the Today/dashboard page.
    Counts as a qualifying streak action (plan review)."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    browser_tz = (request.json or {}).get("timezone")
    result = _record_streak_qualifying_action(current_user.id, browser_tz)
    # Also grant the once-per-day pet visit XP
    try:
        pet_award = _maybe_award_daily_visit(current_user.id, browser_tz)
    except Exception as e:
        print(f"[pet] error on plan-review visit: {e}")
        pet_award = None
    if result is None:
        return flask.jsonify({"status": "ok", "enabled": False, "pet": pet_award})
    return flask.jsonify({"status": "ok", "pet": pet_award, **result})


@app.route("/api/streak/set-timezone", methods=["POST"])
def streak_set_timezone():
    """Persist the user's IANA timezone from the browser."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    tz = (request.json or {}).get("timezone", "")
    if not tz:
        return flask.jsonify({"status": "error", "message": "timezone required"}), 400
    row = _get_or_create_streak(current_user.id)
    row.timezone = tz
    db.session.commit()
    return flask.jsonify({"status": "ok"})


# ── Plani Pet (creature that grows with site usage) ─────────────────

def _get_or_create_pet(user_id: int) -> PlaniPet:
    row = PlaniPet.query.filter_by(user_id=user_id).first()
    if not row:
        row = PlaniPet(user_id=user_id, name="Plani", xp=0)
        db.session.add(row)
        db.session.commit()
    return row


def _user_tz_for_pet(user_id: int) -> str:
    """Use the same timezone the streak engine has stored, fall back to UTC."""
    sr = UserStreak.query.filter_by(user_id=user_id).first()
    return (sr.timezone if sr and sr.timezone else "UTC")


def _award_pet_xp(user_id: int, event_key: str, *, browser_tz: str | None = None) -> dict | None:
    """Award XP for a pet event. Returns delta info for client feedback.

    `event_key` must be a key in pet_engine.XP_REWARDS. Safe to call from
    anywhere — silently no-ops for unauthenticated callers or unknown events.
    """
    if not user_id:
        return None
    delta = pet_engine.XP_REWARDS.get(event_key)
    if not delta:
        return None
    pet = _get_or_create_pet(user_id)
    tz_name = _user_tz_for_pet(user_id) or browser_tz or "UTC"

    before_stage = pet_engine.stage_for_xp(pet.xp)["id"]
    pet.xp = (pet.xp or 0) + delta
    after_stage = pet_engine.stage_for_xp(pet.xp)["id"]

    # daily_visit also stamps the visit date so we don't double-grant
    if event_key == "daily_visit":
        today_local = pet_engine.datetime.now(pet_engine.ZoneInfo(tz_name)).date()
        pet.last_visit_local_date = today_local.isoformat()

    db.session.commit()

    evolved = after_stage != before_stage
    try:
        app_analytics.track(user_id, "pet_xp_earned", {
            "event": event_key,
            "delta": delta,
            "total_xp": pet.xp,
            "evolved": evolved,
        })
        if evolved:
            app_analytics.track(user_id, "pet_evolved", {
                "from": before_stage,
                "to": after_stage,
                "xp": pet.xp,
            })
    except Exception:
        pass

    return {
        "delta": delta,
        "event": event_key,
        "evolved": evolved,
        "before_stage": before_stage,
        "after_stage": after_stage,
        "xp": pet.xp,
    }


def _maybe_award_daily_visit(user_id: int, browser_tz: str | None = None) -> dict | None:
    """Grant the per-day visit XP exactly once per local day."""
    if not user_id:
        return None
    pet = _get_or_create_pet(user_id)
    tz_name = _user_tz_for_pet(user_id) or browser_tz or "UTC"
    last_date = None
    if pet.last_visit_local_date:
        try:
            last_date = date.fromisoformat(pet.last_visit_local_date)
        except (ValueError, TypeError):
            last_date = None
    if not pet_engine.should_grant_daily_visit(last_date, tz_name):
        return None
    return _award_pet_xp(user_id, "daily_visit", browser_tz=browser_tz)


@app.route("/api/pet/status")
def pet_status():
    """Return the user's pet state (xp, stage, level, mood, progress)."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "message": "Not logged in"}), 401
    pet = _get_or_create_pet(current_user.id)
    tz_name = _user_tz_for_pet(current_user.id) or "UTC"
    last_date = None
    if pet.last_visit_local_date:
        try:
            last_date = date.fromisoformat(pet.last_visit_local_date)
        except (ValueError, TypeError):
            last_date = None
    state = pet_engine.resolve_pet_state(
        xp=pet.xp or 0,
        last_visit_local=last_date,
        user_tz=tz_name,
    )
    # Also surface streak info so the pet UI can show the combo
    streak_row = UserStreak.query.filter_by(user_id=current_user.id).first()

    # Care action availability (drives the feed/play/pet/study buttons)
    last_at_map = {
        "feed": pet.last_fed_at,
        "play": pet.last_played_at,
        "pet": pet.last_petted_at,
        "study_with": pet.last_studied_at,
    }
    care = {}
    for action, cfg in pet_engine.CARE_ACTIONS.items():
        allowed, wait_s = pet_engine.can_perform_care_action(action, last_at_map.get(action))
        care[action] = {
            "label": cfg["label"],
            "emoji": cfg["emoji"],
            "xp": cfg["xp"],
            "ready": allowed,
            "wait_seconds": wait_s,
            "cooldown_hours": cfg["cooldown_hours"],
        }

    # Chest availability
    today_iso = pet_engine.datetime.now(pet_engine.ZoneInfo(tz_name)).date().isoformat()
    chest_ready = pet.last_chest_local_date != today_iso
    chest_reward = pet_engine.daily_chest_reward(
        (pet.chest_streak_days or 0) + (1 if chest_ready else 0)
    )

    return flask.jsonify({
        "status": "ok",
        "name": pet.name or "Plani",
        "xp": state.xp,
        "level": state.level,
        "stage_id": state.stage_id,
        "stage_name": state.stage_name,
        "stage_title": state.stage_title,
        "stage_color": state.stage_color,
        "progress_to_next": state.progress_to_next,
        "xp_into_stage": state.xp_into_stage,
        "xp_for_next_stage": state.xp_for_next_stage,
        "next_stage_id": state.next_stage_id,
        "next_stage_name": state.next_stage_name,
        "mood": state.mood,
        "days_since_visit": state.days_since_visit,
        "stages": pet_engine.STAGES,
        "hatched_at": (pet.hatched_at.isoformat() if pet.hatched_at else None),
        "streak": (streak_row.current_streak if streak_row else 0),
        "longest_streak": (streak_row.longest_streak if streak_row else 0),
        "care": care,
        "chest": {
            "ready": chest_ready,
            "current_streak_days": pet.chest_streak_days or 0,
            "next_reward": chest_reward,
        },
        "accessories": pet_engine.accessories_for(state.stage_id),
    })


@app.route("/api/pet/visit", methods=["POST"])
def pet_visit():
    """Mark today's visit. Grants daily_visit XP once per local day."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    browser_tz = (request.json or {}).get("timezone")
    result = _maybe_award_daily_visit(current_user.id, browser_tz)
    return flask.jsonify({"status": "ok", "awarded": result})


@app.route("/api/pet/rename", methods=["POST"])
def pet_rename():
    """Let the user name their pet."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    name = ((request.json or {}).get("name") or "").strip()[:40]
    if not name:
        return flask.jsonify({"status": "error", "message": "Name required"}), 400
    pet = _get_or_create_pet(current_user.id)
    pet.name = name
    db.session.commit()
    return flask.jsonify({"status": "ok", "name": name})


@app.route("/pet")
def pet_page():
    """Full pet page — like a Tamagotchi dashboard."""
    if not current_user.is_authenticated:
        return flask.redirect("/login")
    return flask.render_template(
        "pet.html",
        active_page="pet",
        logged_in=True,
    )


@app.route("/api/pet/care", methods=["POST"])
def pet_care():
    """Perform a care action (feed/play/pet/study_with) on the user's Plani.

    Cooldown-gated by pet_engine.CARE_ACTIONS. Awards XP and bumps mood.
    """
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    action = ((request.json or {}).get("action") or "").strip()
    if action not in pet_engine.CARE_ACTIONS:
        return flask.jsonify({"status": "error", "message": "Unknown action"}), 400

    pet = _get_or_create_pet(current_user.id)
    last_at_map = {
        "feed": pet.last_fed_at,
        "play": pet.last_played_at,
        "pet": pet.last_petted_at,
        "study_with": pet.last_studied_at,
    }
    allowed, wait_s = pet_engine.can_perform_care_action(action, last_at_map[action])
    if not allowed:
        return flask.jsonify({
            "status": "cooldown",
            "message": f"{pet_engine.CARE_ACTIONS[action]['label']} on cooldown",
            "wait_seconds": wait_s,
        }), 429

    cfg = pet_engine.CARE_ACTIONS[action]
    before_stage = pet_engine.stage_for_xp(pet.xp or 0)["id"]
    pet.xp = (pet.xp or 0) + cfg["xp"]
    after_stage = pet_engine.stage_for_xp(pet.xp)["id"]
    now = utcnow()
    if action == "feed":
        pet.last_fed_at = now
    elif action == "play":
        pet.last_played_at = now
    elif action == "pet":
        pet.last_petted_at = now
    elif action == "study_with":
        pet.last_studied_at = now
    db.session.commit()

    try:
        app_analytics.track(current_user.id, "pet_care_action", {
            "action": action, "xp": cfg["xp"], "total_xp": pet.xp,
        })
    except Exception:
        pass

    return flask.jsonify({
        "status": "ok",
        "action": action,
        "xp_awarded": cfg["xp"],
        "total_xp": pet.xp,
        "copy": cfg["copy"],
        "evolution": pet_engine.evolution_payload(before_stage, after_stage),
        "cooldown_hours": cfg["cooldown_hours"],
    })


@app.route("/api/pet/chest", methods=["POST"])
def pet_chest():
    """Open the daily check-in chest. Idempotent per local day."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    pet = _get_or_create_pet(current_user.id)
    tz_name = _user_tz_for_pet(current_user.id) or "UTC"
    today = pet_engine.datetime.now(pet_engine.ZoneInfo(tz_name)).date()

    # Already claimed today?
    if pet.last_chest_local_date == today.isoformat():
        reward = pet_engine.daily_chest_reward(pet.chest_streak_days or 1)
        return flask.jsonify({
            "status": "already_claimed",
            "chest": reward,
            "chest_streak_days": pet.chest_streak_days,
        })

    # Continuation check
    yesterday = (today - timedelta(days=1)).isoformat()
    if pet.last_chest_local_date == yesterday:
        pet.chest_streak_days = (pet.chest_streak_days or 0) + 1
    else:
        pet.chest_streak_days = 1

    reward = pet_engine.daily_chest_reward(pet.chest_streak_days)
    before_stage = pet_engine.stage_for_xp(pet.xp or 0)["id"]
    pet.xp = (pet.xp or 0) + reward["xp"]
    after_stage = pet_engine.stage_for_xp(pet.xp)["id"]
    pet.last_chest_local_date = today.isoformat()
    db.session.commit()

    try:
        app_analytics.track(current_user.id, "pet_chest_opened", {
            "xp": reward["xp"], "streak_days": pet.chest_streak_days, "tier": reward["tier"],
        })
    except Exception:
        pass

    return flask.jsonify({
        "status": "ok",
        "chest": reward,
        "chest_streak_days": pet.chest_streak_days,
        "total_xp": pet.xp,
        "evolution": pet_engine.evolution_payload(before_stage, after_stage),
    })


@app.route("/api/streak/risk")
def streak_risk():
    """Real-time risk assessment for the at-risk banner / push toast."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    row = _get_or_create_streak(current_user.id)
    tz_name = row.timezone or "UTC"
    last_date = None
    if row.last_qualifying_local_date:
        try:
            last_date = date.fromisoformat(row.last_qualifying_local_date)
        except (ValueError, TypeError):
            last_date = None
    risk = streak_engine.assess_streak_risk(
        current_streak=row.current_streak,
        last_qualifying_local_date=last_date,
        user_tz=tz_name,
    )
    # Compute the week dots for the perfect-week badge
    try:
        qualified = set()
        for d in json.loads(row.qualified_dates_json or "[]"):
            try:
                qualified.add(date.fromisoformat(d))
            except (ValueError, TypeError):
                pass
    except (json.JSONDecodeError, TypeError):
        qualified = set()
    dots = streak_engine.week_dots(tz_name, qualified)
    perfect = streak_engine.perfect_week_bonus(dots)

    # Pay out the perfect-week bonus exactly once per ISO week
    iso_year, iso_week, _ = pet_engine.datetime.now(pet_engine.ZoneInfo(tz_name)).date().isocalendar()
    week_key = f"{iso_year}-W{iso_week:02d}"
    paid_bonus = None
    if perfect.get("perfect"):
        pet = _get_or_create_pet(current_user.id)
        if pet.perfect_week_paid != week_key:
            pet.xp = (pet.xp or 0) + perfect["bonus_xp"]
            pet.perfect_week_paid = week_key
            db.session.commit()
            paid_bonus = perfect["bonus_xp"]

    return flask.jsonify({
        "status": "ok",
        "level": risk.level,
        "hours_until_break": risk.hours_until_break,
        "message": risk.message,
        "urgency_score": risk.urgency_score,
        "perfect_week": perfect,
        "perfect_week_paid": paid_bonus,
        "freeze_offer": streak_engine.streak_freeze_offer(
            current_streak=row.current_streak,
            last_qualifying_local_date=last_date,
            freezes_available=row.freezes_available,
            user_tz=tz_name,
        ),
    })


# ── Pagination ──────────────────────────────────────────────────

# Ceiling on `per_page` so a hand-edited query string cannot ask the DB for
# the whole table. Every paginated endpoint clamps to this.
PAGE_SIZE_DEFAULT = 25
PAGE_SIZE_MAX = 100


def _page_args(default_size: int = PAGE_SIZE_DEFAULT) -> tuple[int, int]:
    """Read and clamp `page` / `per_page` from the query string."""
    try:
        page = max(1, int(request.args.get("page", 1)))
    except (TypeError, ValueError):
        page = 1
    try:
        per_page = int(request.args.get("per_page", default_size))
    except (TypeError, ValueError):
        per_page = default_size
    per_page = max(1, min(per_page, PAGE_SIZE_MAX))
    return page, per_page


def paginate_query(query, serialize, default_size: int = PAGE_SIZE_DEFAULT,
                   count_total: bool = True) -> dict:
    """Run `query` for one page and shape the response the client expects.

    Fetches `per_page + 1` rows and discards the extra. That one spare row is
    what tells us whether another page exists without paying for a COUNT on
    every request — `count_total=False` skips the COUNT entirely for tables
    where the exact total is not worth the scan.

    Returns {items, page, per_page, has_more, total?} — the shape
    IP.paginate() in ip-async.js consumes.
    """
    page, per_page = _page_args(default_size)
    offset = (page - 1) * per_page
    rows = query.limit(per_page + 1).offset(offset).all()
    has_more = len(rows) > per_page
    rows = rows[:per_page]

    payload = {
        "status": "ok",
        "items": [serialize(r) for r in rows],
        "page": page,
        "per_page": per_page,
        "has_more": has_more,
    }
    if count_total:
        try:
            payload["total"] = query.order_by(None).count()
        except Exception:
            # A COUNT is a nicety, not a requirement — never fail the page
            # over it (some queries carry DISTINCT/GROUP BY that upset it).
            pass
    return payload


# ── Feedback collector ──────────────────────────────────────────

@app.route("/api/feedback/submit", methods=["POST"])
@limiter.limit("10 per hour")
def feedback_submit():
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "Login required"}), 401
    data = request.get_json(silent=True) or {}
    message = (data.get("message") or "").strip()
    if not message or len(message) < 3:
        return jsonify({"status": "error", "message": "Message too short"}), 400
    if len(message) > 5000:
        return jsonify({"status": "error", "message": "Message too long"}), 400
    category = data.get("category", "general")
    if category not in ("bug", "feature", "praise", "general"):
        category = "general"
    mood = data.get("mood")
    if isinstance(mood, int) and 1 <= mood <= 5:
        pass
    else:
        mood = None
    try:
        fb = SiteFeedback(
            user_id=current_user.id,
            category=category,
            mood=mood,
            message=message[:5000],
            page_url=(data.get("page_url") or "")[:512],
        )
        db.session.add(fb)
        db.session.commit()
        return jsonify({"status": "ok", "id": fb.id})
    except Exception as e:
        db.session.rollback()
        print(f"[feedback] submit failed: {e}")
        return jsonify({"status": "error", "message": "Could not save"}), 500


def _serialize_feedback(r: "SiteFeedback") -> dict:
    return {
        "id": r.id,
        "category": r.category,
        "mood": r.mood,
        "message": r.message,
        "page_url": r.page_url,
        "status": r.status,
        "admin_note": r.admin_note or "",
        "created_at": r.created_at.isoformat() if r.created_at else None,
    }


@app.route("/api/feedback/mine")
def feedback_mine():
    if not current_user.is_authenticated:
        return jsonify({"status": "error"}), 401
    q = SiteFeedback.query.filter_by(user_id=current_user.id)\
        .order_by(SiteFeedback.created_at.desc())
    return jsonify(paginate_query(q, _serialize_feedback, default_size=20))


# ── Client error reporting ──────────────────────────────────────

def _error_fingerprint(kind: str, message: str, source: str, line: int) -> str:
    """Stable id for "the same bug", so repeats increment instead of piling up.

    The message is truncated before hashing because many browser errors embed
    a varying id or URL in the tail; keeping the first 180 characters groups
    those together while still separating genuinely different failures.
    """
    raw = f"{kind}|{message[:180]}|{source[-120:]}|{line}"
    return hashlib.sha256(raw.encode("utf-8", "replace")).hexdigest()[:32]


@app.route("/api/client-error", methods=["POST"])
@limiter.limit("60 per hour")
def client_error_report():
    """Record a JavaScript failure reported by the browser.

    Deliberately permissive: anonymous reports are accepted (unauthenticated
    pages break too), the response is always 204, and every failure path is
    swallowed. A reporting endpoint that can itself error, 401, or block is
    worse than no reporting at all.
    """
    data = request.get_json(silent=True, force=True) or {}
    message = (data.get("message") or "").strip()[:512]
    if not message:
        return ("", 204)

    kind = (data.get("kind") or "error")[:32]
    source = (data.get("source") or "")[:512]
    try:
        line = int(data.get("line") or 0)
    except (TypeError, ValueError):
        line = 0

    fp = _error_fingerprint(kind, message, source, line)
    uid = current_user.id if current_user.is_authenticated else None

    try:
        row = ClientErrorLog.query.filter_by(fingerprint=fp).first()
        if row:
            row.count = (row.count or 0) + 1
            row.last_seen = utcnow()
            # A recurrence after someone marked it fixed means it is not fixed.
            if row.resolved:
                row.resolved = False
            if uid and not row.user_id:
                row.user_id = uid
        else:
            db.session.add(ClientErrorLog(
                user_id=uid,
                fingerprint=fp,
                kind=kind,
                message=message,
                stack=(data.get("stack") or "")[:4000],
                source=source,
                line=line,
                page_url=(data.get("url") or "")[:512],
                user_agent=(data.get("ua") or "")[:300],
                viewport=(data.get("viewport") or "")[:24],
                context=(data.get("context") or "")[:1000],
            ))
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"[client-error] could not record: {e}")
    return ("", 204)


@app.route("/api/bug-report", methods=["POST"])
@limiter.limit("20 per hour")
def bug_report_submit():
    """User-written bug report plus the diagnostics the dialog collected."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "Login required"}), 401

    data = request.get_json(silent=True) or {}
    message = (data.get("message") or "").strip()
    if len(message) < 8:
        return jsonify({"status": "error",
                        "message": "Add a little more detail so we can find the problem."}), 400
    if len(message) > 5000:
        return jsonify({"status": "error", "message": "Message too long"}), 400

    category = data.get("category", "bug")
    if category not in ("bug", "feature", "praise", "general"):
        category = "bug"

    diagnostics = data.get("diagnostics") or {}
    try:
        diag_json = json.dumps(diagnostics)[:8000]
    except (TypeError, ValueError):
        diag_json = ""

    try:
        fb = SiteFeedback(
            user_id=current_user.id,
            category=category,
            message=message[:5000],
            page_url=(data.get("page_url") or "")[:512],
            diagnostics=diag_json,
        )
        db.session.add(fb)
        db.session.commit()
        return jsonify({"status": "ok", "id": fb.id})
    except Exception as e:
        db.session.rollback()
        print(f"[bug-report] submit failed: {e}")
        return jsonify({"status": "error", "message": "Could not save your report"}), 500


@app.route("/api/admin/client-errors")
def admin_client_errors():
    """Paginated JS error log, newest occurrence first. Admin only."""
    if not is_admin(current_user):
        return jsonify({"status": "error", "message": "Not permitted"}), 403

    q = ClientErrorLog.query
    if request.args.get("unresolved") == "1":
        q = q.filter(ClientErrorLog.resolved.is_(False))
    q = q.order_by(ClientErrorLog.last_seen.desc())

    def _row(r):
        return {
            "id": r.id,
            "fingerprint": r.fingerprint,
            "kind": r.kind,
            "message": r.message,
            "source": r.source,
            "line": r.line,
            "page_url": r.page_url,
            "viewport": r.viewport,
            "user_agent": r.user_agent,
            "count": r.count,
            "resolved": bool(r.resolved),
            "user_id": r.user_id,
            "stack": (r.stack or "")[:1500],
            "first_seen": r.first_seen.isoformat() if r.first_seen else None,
            "last_seen": r.last_seen.isoformat() if r.last_seen else None,
        }

    return jsonify(paginate_query(q, _row, default_size=30))


@app.route("/api/admin/client-errors/<int:error_id>/resolve", methods=["POST"])
def admin_resolve_client_error(error_id: int):
    if not is_admin(current_user):
        return jsonify({"status": "error", "message": "Not permitted"}), 403
    row = db.session.get(ClientErrorLog, error_id)
    if not row:
        return jsonify({"status": "error", "message": "Not found"}), 404
    row.resolved = not row.resolved
    db.session.commit()
    return jsonify({"status": "ok", "resolved": row.resolved})


@app.route("/api/admin/bug-reports")
def admin_bug_reports():
    """Paginated user-submitted reports, with the attached diagnostics."""
    if not is_admin(current_user):
        return jsonify({"status": "error", "message": "Not permitted"}), 403

    q = SiteFeedback.query
    category = request.args.get("category")
    if category in ("bug", "feature", "praise", "general"):
        q = q.filter(SiteFeedback.category == category)
    q = q.order_by(SiteFeedback.created_at.desc())

    def _row(r):
        out = _serialize_feedback(r)
        out["user_id"] = r.user_id
        try:
            out["diagnostics"] = json.loads(r.diagnostics) if r.diagnostics else None
        except (TypeError, ValueError):
            out["diagnostics"] = None
        return out

    return jsonify(paginate_query(q, _row, default_size=30))


# ── My Stats page ──────────────────────────────────────────────

@app.route("/my-stats")
def my_stats_page():
    if not current_user.is_authenticated:
        return flask.redirect("/login")
    return flask.render_template("my_stats.html", active_page="my_stats", logged_in=True)


@app.route("/api/my-stats")
def api_my_stats():
    if not current_user.is_authenticated:
        return jsonify({"status": "error"}), 401
    uid = current_user.id

    # Tasks completed
    tasks_done = DismissedAssignment.query.filter_by(user_id=uid).count()
    manual_done = ManualTask.query.filter_by(user_id=uid, done=True).count()

    # Task feedback stats
    # These are pure aggregates — compute them in SQL. Hydrating every
    # TaskFeedback / StudySession / StudyMastery row into ORM objects just to
    # sum them made this endpoint scale with the user's lifetime history.
    total_actual_min, total_estimated_min = db.session.query(
        db.func.coalesce(db.func.sum(TaskFeedback.actual_time), 0),
        db.func.coalesce(db.func.sum(TaskFeedback.estimated_time), 0),
    ).filter(TaskFeedback.user_id == uid).one()

    diff_map = {"Easy": 1, "Medium": 2, "Hard": 3}
    diff_rows = (
        db.session.query(TaskFeedback.difficulty, db.func.count(TaskFeedback.id))
        .filter(TaskFeedback.user_id == uid)
        .group_by(TaskFeedback.difficulty)
        .all()
    )
    avg_difficulty = None
    fb_total = sum(n for _, n in diff_rows)
    if fb_total:
        avg_difficulty = round(
            sum(diff_map.get(d, 2) * n for d, n in diff_rows) / fb_total, 1
        )

    courses_worked = [
        c for (c,) in db.session.query(TaskFeedback.course)
        .filter(TaskFeedback.user_id == uid, TaskFeedback.course != "")
        .distinct().all() if c
    ]

    # Study sessions
    total_secs, total_questions, total_correct, session_count = db.session.query(
        db.func.coalesce(db.func.sum(StudySession.duration_seconds), 0),
        db.func.coalesce(db.func.sum(StudySession.questions_total), 0),
        db.func.coalesce(db.func.sum(StudySession.questions_correct), 0),
        db.func.count(StudySession.id),
    ).filter(StudySession.user_id == uid, StudySession.completed == True).one()
    total_study_min = int(total_secs or 0) // 60

    # Streak
    streak_row = UserStreak.query.filter_by(user_id=uid).first()
    current_streak = streak_row.current_streak if streak_row else 0
    longest_streak = streak_row.longest_streak if streak_row else 0

    # Pet
    pet = PlaniPet.query.filter_by(user_id=uid).first()
    pet_data = None
    if pet:
        pet_data = {
            "name": pet.name,
            "xp": pet.xp or 0,
            "level": pet_engine.level_for_xp(pet.xp or 0),
            "stage": pet_engine.stage_for_xp(pet.xp or 0),
        }

    # Grades
    grades = ImportedGrade.query.filter_by(user_id=uid).all()
    grade_list = [{"course": g.course, "percentage": g.percentage, "letter": g.letter}
                  for g in grades if g.percentage is not None]
    gpa = None
    if grade_list:
        gpa = round(sum(g["percentage"] for g in grade_list) / len(grade_list), 1)

    # Mastery — two counts, no need to hydrate the rows.
    mastery_count = StudyMastery.query.filter_by(user_id=uid).count()
    mastered = StudyMastery.query.filter(
        StudyMastery.user_id == uid, StudyMastery.mastery_level >= 4
    ).count()

    # Study points
    sp = StudyPoints.query.filter_by(user_id=uid).first()
    points_data = None
    if sp:
        points_data = {
            "total_points": sp.total_points,
            "level": sp.level,
            "sparks": sp.spark_balance,
            "badges": json.loads(sp.badges or "[]"),
        }

    # Daily activity (last 30 days)
    thirty_days_ago = utcnow() - timedelta(days=30)
    recent_feedback = TaskFeedback.query.filter(
        TaskFeedback.user_id == uid,
        TaskFeedback.completed_at >= thirty_days_ago,
    ).all()
    daily_completions = {}
    for r in recent_feedback:
        day = r.completed_at.strftime("%Y-%m-%d") if r.completed_at else None
        if day:
            daily_completions[day] = daily_completions.get(day, 0) + 1

    recent_sessions = StudySession.query.filter(
        StudySession.user_id == uid,
        StudySession.completed == True,
        StudySession.created_at >= thirty_days_ago,
    ).all()
    daily_study_min = {}
    for s in recent_sessions:
        day = s.created_at.strftime("%Y-%m-%d") if s.created_at else None
        if day:
            daily_study_min[day] = daily_study_min.get(day, 0) + (s.duration_seconds or 0) // 60

    # Account age
    account_days = (utcnow() - current_user.created_at).days if current_user.created_at else 0

    return jsonify({
        "status": "ok",
        "account_days": account_days,
        "tasks_completed": tasks_done + manual_done,
        "tasks_lms": tasks_done,
        "tasks_manual": manual_done,
        "feedback_count": fb_total,
        "total_actual_minutes": total_actual_min,
        "total_estimated_minutes": total_estimated_min,
        "avg_difficulty": avg_difficulty,
        "courses_worked": courses_worked,
        "study_sessions": session_count,
        "study_minutes": total_study_min,
        "questions_answered": total_questions,
        "questions_correct": total_correct,
        "accuracy": round(total_correct / total_questions * 100, 1) if total_questions else None,
        "current_streak": current_streak,
        "longest_streak": longest_streak,
        "pet": pet_data,
        "grades": grade_list,
        "gpa_avg": gpa,
        "mastery_concepts": mastery_count,
        "mastery_mastered": mastered,
        "points": points_data,
        "daily_completions": daily_completions,
        "daily_study_minutes": daily_study_min,
    })


@app.route("/tasks/manual/list")
def manual_list_tasks():
    if current_user.is_authenticated:
        tasks = ManualTask.query.filter_by(user_id=current_user.id, done=False).all()
    else:
        gid = get_guest_session_id()
        tasks = ManualTask.query.filter_by(guest_session_id=gid, done=False).all()
    return flask.jsonify([{
        "id": t.id, "title": t.title, "due_date": t.due_date,
        "priority": t.priority, "course": t.course,
        "estimated_time": t.estimated_time, "notes": t.notes,
        "source": "manual", "color": PRIORITY_COLORS.get(t.priority, "#f59e0b")
    } for t in tasks])


# ── UNSUPPORTED-LMS IMPORT (CSV / smart-paste / extension scraper) ──
# These endpoints exist so students whose districts aren't supported
# (or whose LMS we don't have a connector for yet) can still get
# their assignments + grades into IntelliPlan. Three tiers:
#   1. /api/import/csv          — generic CSV template
#   2. /api/import/smart_paste  — paste a gradebook table, AI parses it
#   3. /api/import/scraper      — extension posts structured JSON

CSV_TEMPLATE_ROWS = [
    ["type", "title", "course", "due_date", "priority", "estimated_minutes", "grade_percent", "letter_grade", "teacher", "notes"],
    ["assignment", "Chapter 4 Reading", "Honors Bio", "2026-06-12", "Medium", "45", "", "", "Ms. Patel", "Pages 88-104"],
    ["assignment", "Lab Report 3", "Chemistry", "2026-06-15", "High", "120", "", "", "Mr. Chen", "Include data table"],
    ["grade", "", "Honors Bio", "", "", "", "92", "A-", "Ms. Patel", "Current overall"],
    ["grade", "", "Chemistry", "", "", "", "78", "C+", "Mr. Chen", "Current overall"],
]


def _import_owner_filter(query, model):
    """Apply the right ownership filter to an Imported* query depending on auth."""
    if current_user.is_authenticated:
        return query.filter(model.user_id == current_user.id)
    return query.filter(model.guest_session_id == get_guest_session_id())


def _import_owner_kwargs():
    if current_user.is_authenticated:
        return {"user_id": current_user.id, "guest_session_id": None}
    return {"user_id": None, "guest_session_id": get_guest_session_id()}


def _parse_csv_rows(text):
    """Parse the IntelliPlan CSV format. Returns (assignments, grades, errors)."""
    import csv as _csv
    from io import StringIO as _SIO
    assignments, grades, errors = [], [], []
    reader = _csv.reader(_SIO(text))
    rows = list(reader)
    if not rows:
        return [], [], ["Empty file"]
    header = [h.strip().lower() for h in rows[0]]
    expected = {"type", "title", "course", "due_date", "priority",
                "estimated_minutes", "grade_percent", "letter_grade", "teacher", "notes"}
    if not expected.issubset(set(header)):
        missing = expected - set(header)
        errors.append(f"CSV is missing required columns: {', '.join(sorted(missing))}. "
                      f"Download the template and fill it in.")
        return [], [], errors
    idx = {col: header.index(col) for col in header}

    def _g(row, col):
        i = idx.get(col)
        return (row[i].strip() if i is not None and i < len(row) else "")

    for line_no, row in enumerate(rows[1:], start=2):
        if not any(cell.strip() for cell in row):
            continue
        kind = _g(row, "type").lower()
        title = _g(row, "title")
        course = _g(row, "course") or "Imported"
        if kind == "assignment":
            if not title:
                errors.append(f"Line {line_no}: assignment missing title")
                continue
            try:
                est = int(_g(row, "estimated_minutes") or "60")
            except ValueError:
                est = 60
            assignments.append({
                "title": title[:500],
                "course": course[:255],
                "due_date": _g(row, "due_date"),
                "priority": (_g(row, "priority") or "Medium").title(),
                "estimated_time": max(15, min(est, 600)),
                "notes": _g(row, "notes")[:800],
            })
        elif kind == "grade":
            try:
                pct = float(_g(row, "grade_percent").rstrip("%") or "nan")
                if pct != pct:  # NaN
                    pct = None
            except ValueError:
                pct = None
            grades.append({
                "course": course[:255],
                "percentage": pct,
                "letter": _g(row, "letter_grade")[:4],
                "teacher": _g(row, "teacher")[:255],
                "notes": _g(row, "notes")[:800],
            })
        else:
            errors.append(f"Line {line_no}: unknown type '{kind}' (expected 'assignment' or 'grade')")
    return assignments, grades, errors


def _persist_import(assignments, grades, source="csv", source_label="", batch_id=None, replace_batch=True):
    """Write parsed assignments/grades to the DB, optionally replacing a prior
    batch with the same source (for auto-sync from the extension)."""
    import uuid as _uuid
    batch_id = batch_id or str(_uuid.uuid4())
    own = _import_owner_kwargs()
    # Refresh-replace: when the same scraper pushes a new sync, drop the old
    # rows for that source so we don't accumulate duplicates.
    if replace_batch and source.startswith("scraper:"):
        q = ManualTask.query.filter_by(**own).filter(ManualTask.import_source == source)
        q.delete(synchronize_session=False)
        gq = ImportedGrade.query.filter_by(**own).filter(ImportedGrade.source == source)
        gq.delete(synchronize_session=False)
    created_assignments = 0
    for a in assignments:
        db.session.add(ManualTask(
            **own,
            title=a["title"], course=a.get("course", "Imported"),
            due_date=a.get("due_date", ""), priority=a.get("priority", "Medium"),
            estimated_time=a.get("estimated_time", 60), notes=a.get("notes", ""),
            import_source=source, import_batch_id=batch_id,
            external_id=a.get("external_id", ""),
        ))
        created_assignments += 1
    created_grades = 0
    for g in grades:
        db.session.add(ImportedGrade(
            **own,
            course=g["course"], percentage=g.get("percentage"),
            letter=g.get("letter", ""), teacher=g.get("teacher", ""),
            period=g.get("period", ""), source=source,
            source_label=source_label or source,
        ))
        created_grades += 1
    db.session.commit()
    # Learning Graph: grade import event
    try:
        uid = None
        if current_user.is_authenticated:
            uid = current_user.id
        if uid and grades:
            from learning_graph_glue import _learning_graph_on_grade_changed
            for g in grades:
                course = g.get("course", "")
                new_pct = g.get("percentage")
                if course and new_pct is not None:
                    _learning_graph_on_grade_changed(uid, course, None, float(new_pct))
    except Exception:
        pass
    return batch_id, created_assignments, created_grades


@app.route("/api/import/csv/template")
def api_import_csv_template():
    """Download the blank IntelliPlan import template."""
    import csv as _csv
    from io import StringIO as _SIO
    buf = _SIO()
    writer = _csv.writer(buf)
    for row in CSV_TEMPLATE_ROWS:
        writer.writerow(row)
    response = flask.make_response(buf.getvalue())
    response.headers["Content-Type"] = "text/csv; charset=utf-8"
    response.headers["Content-Disposition"] = 'attachment; filename="intelliplan-import-template.csv"'
    return response


@app.route("/api/import/csv", methods=["POST"])
@limiter.limit("20 per hour")
def api_import_csv():
    """Accept an uploaded CSV (or raw `text` field) and import its rows."""
    text = ""
    if "file" in request.files:
        try:
            text = request.files["file"].read().decode("utf-8", errors="replace")
        except Exception as e:
            return flask.jsonify({"status": "error", "message": f"Could not read file: {e}"}), 400
    else:
        body = request.get_json(silent=True) or {}
        text = body.get("text", "") or ""
    if not text.strip():
        return flask.jsonify({"status": "error", "message": "No CSV content provided."}), 400
    assignments, grades, errors = _parse_csv_rows(text)
    if errors and not assignments and not grades:
        return flask.jsonify({"status": "error", "message": errors[0], "errors": errors}), 400
    source_label = (request.form.get("source_label") if "file" in request.files
                    else (request.get_json(silent=True) or {}).get("source_label", "")) or "CSV import"
    batch_id, a_count, g_count = _persist_import(
        assignments, grades, source="csv", source_label=source_label, replace_batch=False
    )
    return flask.jsonify({
        "status": "ok", "batch_id": batch_id,
        "assignments_imported": a_count, "grades_imported": g_count,
        "warnings": errors,
    })


@app.route("/api/import/smart_paste", methods=["POST"])
@limiter.limit("15 per hour")
def api_import_smart_paste():
    """Paste a gradebook or assignment list as raw text; AI parses it.

    Body: {"text": "...pasted content...", "hint": "optional context"}.
    Returns the parsed assignments+grades AND persists them. The user can
    later delete any row that came out wrong from the unified task list.
    """
    if not ai_available():
        return flask.jsonify({"status": "error",
                              "message": "Smart paste needs the AI service, which is unavailable right now."}), 503
    body = request.get_json(silent=True) or {}
    text = (body.get("text") or "").strip()
    hint = (body.get("hint") or "").strip()
    if not text:
        return flask.jsonify({"status": "error", "message": "Paste some text first."}), 400
    if len(text) > 24000:
        text = text[:24000]
    today = datetime.now().strftime("%Y-%m-%d")
    prompt = f"""You convert pasted gradebook/assignment text into IntelliPlan JSON. Today is {today}.

The student pasted this from an unsupported LMS{(' — ' + hint) if hint else ''}:
---
{text}
---

Extract every assignment and every course grade you can identify. Return ONLY valid JSON in this exact shape:

{{
  "assignments": [
    {{"title": "...", "course": "...", "due_date": "YYYY-MM-DD or empty", "priority": "High|Medium|Low",
      "estimated_time": 45, "notes": "..."}}
  ],
  "grades": [
    {{"course": "...", "percentage": 92.5, "letter": "A-", "teacher": ""}}
  ]
}}

Rules:
- If a row is clearly a course summary (no due date, has a percent or letter grade), put it in `grades`.
- If a row is an individual task with a due date or work to do, put it in `assignments`.
- When the year is missing from a date, assume the current academic year.
- Priority: assignments due within 3 days = High, within 14 days = Medium, beyond = Low.
- estimated_time should be a reasonable minute estimate based on the assignment type (reading 30, problem set 60, project 120, essay 90).
- If you cannot determine a field, leave it as an empty string. Never invent course names.
- Strip HTML, table pipes, and column letters. Return clean text only."""
    try:
        result = ai_chat(
            [{"role": "user", "content": prompt}],
            tier="standard", temperature=0.2, max_tokens=4000,
            response_format={"type": "json_object"},
        )
        result = re.sub(r"```json\n?", "", result)
        result = re.sub(r"```\n?", "", result).strip()
        try:
            parsed = json.loads(result)
        except json.JSONDecodeError:
            m = re.search(r"\{[\s\S]*\}", result)
            if not m:
                raise
            parsed = json.loads(m.group(0))
    except Exception as e:
        print(f"[smart_paste] AI parse failed: {e}")
        return flask.jsonify({"status": "error",
                              "message": "Couldn't read that paste. Try shorter text or use the CSV template."}), 502

    assignments = parsed.get("assignments") or []
    grades = parsed.get("grades") or []
    # Sanitize before persisting — the LLM occasionally returns numbers as strings.
    clean_assignments = []
    for a in assignments[:200]:
        if not isinstance(a, dict) or not a.get("title"):
            continue
        try:
            est = int(a.get("estimated_time") or 60)
        except (TypeError, ValueError):
            est = 60
        clean_assignments.append({
            "title": str(a["title"])[:500],
            "course": str(a.get("course") or "Imported")[:255],
            "due_date": str(a.get("due_date") or ""),
            "priority": str(a.get("priority") or "Medium").title(),
            "estimated_time": max(15, min(est, 600)),
            "notes": str(a.get("notes") or "")[:800],
        })
    clean_grades = []
    for g in grades[:50]:
        if not isinstance(g, dict) or not g.get("course"):
            continue
        try:
            pct = float(str(g.get("percentage")).rstrip("%")) if g.get("percentage") is not None else None
        except (TypeError, ValueError):
            pct = None
        clean_grades.append({
            "course": str(g["course"])[:255],
            "percentage": pct,
            "letter": str(g.get("letter") or "")[:4],
            "teacher": str(g.get("teacher") or "")[:255],
        })
    batch_id, a_count, g_count = _persist_import(
        clean_assignments, clean_grades, source="paste",
        source_label=hint[:60] or "Smart paste", replace_batch=False,
    )
    return flask.jsonify({
        "status": "ok", "batch_id": batch_id,
        "assignments_imported": a_count, "grades_imported": g_count,
        "preview": {"assignments": clean_assignments[:5], "grades": clean_grades[:5]},
    })


@app.route("/api/import/scraper", methods=["POST"])
@limiter.limit("60 per hour")
def api_import_scraper():
    """Endpoint the IntelliPlan browser extension posts structured data to.

    Body shape:
      {"lms": "powerschool|aeries|infinitecampus|other",
       "label": "PowerSchool (West HS)",
       "assignments": [{"title", "course", "due_date", "priority",
                        "estimated_time", "notes", "external_id"}, ...],
       "grades": [{"course", "percentage", "letter", "teacher", "period"}, ...]}

    Each subsequent post for the same `lms` replaces the prior batch so
    auto-sync stays idempotent.
    """
    if not (current_user.is_authenticated or session.get("guest_id")):
        return flask.jsonify({"status": "error", "message": "Sign in first"}), 401
    body = request.get_json(silent=True) or {}
    lms = (body.get("lms") or "").strip().lower() or "other"
    label = (body.get("label") or lms.title())[:60]
    assignments = body.get("assignments") or []
    grades = body.get("grades") or []
    if not isinstance(assignments, list) or not isinstance(grades, list):
        return flask.jsonify({"status": "error", "message": "Bad payload"}), 400
    if not assignments and not grades:
        return flask.jsonify({"status": "ok", "assignments_imported": 0, "grades_imported": 0})
    # Same sanitization as smart_paste — never trust the extension blindly.
    clean_a = []
    for a in assignments[:500]:
        if not isinstance(a, dict) or not a.get("title"):
            continue
        try:
            est = int(a.get("estimated_time") or 60)
        except (TypeError, ValueError):
            est = 60
        clean_a.append({
            "title": str(a["title"])[:500],
            "course": str(a.get("course") or "Imported")[:255],
            "due_date": str(a.get("due_date") or ""),
            "priority": str(a.get("priority") or "Medium").title(),
            "estimated_time": max(15, min(est, 600)),
            "notes": str(a.get("notes") or "")[:800],
            "external_id": str(a.get("external_id") or "")[:128],
        })
    clean_g = []
    for g in grades[:100]:
        if not isinstance(g, dict) or not g.get("course"):
            continue
        try:
            pct = float(str(g.get("percentage")).rstrip("%")) if g.get("percentage") is not None else None
        except (TypeError, ValueError):
            pct = None
        clean_g.append({
            "course": str(g["course"])[:255], "percentage": pct,
            "letter": str(g.get("letter") or "")[:4],
            "teacher": str(g.get("teacher") or "")[:255],
            "period": str(g.get("period") or "")[:64],
        })
    batch_id, a_count, g_count = _persist_import(
        clean_a, clean_g, source=f"scraper:{lms}",
        source_label=label, replace_batch=True,
    )
    return flask.jsonify({
        "status": "ok", "batch_id": batch_id,
        "assignments_imported": a_count, "grades_imported": g_count,
    })


@app.route("/api/import/status")
def api_import_status():
    """Summary of what's been imported (used by the connect page + extension)."""
    own_q_a = _import_owner_filter(ManualTask.query, ManualTask).filter(
        ManualTask.import_source != ""
    )
    own_q_g = _import_owner_filter(ImportedGrade.query, ImportedGrade)
    by_source = {}
    for t in own_q_a.all():
        by_source.setdefault(t.import_source, {"assignments": 0, "grades": 0})["assignments"] += 1
    for g in own_q_g.all():
        by_source.setdefault(g.source, {"assignments": 0, "grades": 0})["grades"] += 1
    last_grade = own_q_g.order_by(ImportedGrade.last_synced.desc()).first()
    return flask.jsonify({
        "status": "ok",
        "sources": by_source,
        "last_synced": last_grade.last_synced.isoformat() if last_grade else None,
    })


# ── SAVED SCHEDULE ────────────────────────────────────────────
@app.route("/schedule/save", methods=["POST"])
def save_schedule():
    data = request.json or {}
    schedule_data = data.get("schedule_data")
    name = data.get("name", f"Schedule {datetime.now().strftime('%b %d')}")
    if not schedule_data:
        return flask.jsonify({"status": "error", "message": "No schedule data"})
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if current_user.is_authenticated else get_guest_session_id()
    if uid:
        SavedSchedule.query.filter_by(user_id=uid).update({"is_active": False})
    else:
        SavedSchedule.query.filter_by(guest_session_id=gid).update({"is_active": False})
    s = SavedSchedule(user_id=uid, guest_session_id=gid, name=name, schedule_data=json.dumps(schedule_data), is_active=True)
    db.session.add(s)
    db.session.commit()
    return flask.jsonify({"status": "ok", "id": s.id})

@app.route("/schedule/saved")
def get_saved_schedule():
    if current_user.is_authenticated:
        s = SavedSchedule.query.filter_by(user_id=current_user.id, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    else:
        gid = get_guest_session_id()
        s = SavedSchedule.query.filter_by(guest_session_id=gid, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    if not s:
        return flask.jsonify({"status": "none"})
    data = json.loads(s.schedule_data)
    # Backfill block_id / kind / checklist / redirect on schedules saved before
    # the Interactive View shipped, so the new UI works without re-generating.
    try:
        needs_backfill = False
        for d in (data.get("schedule") or []):
            for b in (d.get("blocks") or []):
                if not b.get("block_id") or not b.get("checklist"):
                    needs_backfill = True
                    break
            if needs_backfill: break
        if needs_backfill:
            data = humanize_schedule(data, "evening", 2)
    except Exception as e:
        print(f"[scheduler] backfill on saved schedule failed: {e}")
    try:
        progress = json.loads(s.progress_json) if s.progress_json else {}
        if not isinstance(progress, dict):
            progress = {}
    except Exception:
        progress = {}
    return flask.jsonify({"status": "ok", "name": s.name, "created_at": s.created_at.strftime("%b %d, %Y"), "data": data, "progress": progress})


@app.route("/schedule/progress", methods=["POST"])
def save_schedule_progress():
    """Persist Interactive View progress for the active saved schedule.

    The client sends {progress: {block_id: {done, checked, ...}}} debounced
    after each interaction, so checked-off blocks follow the student to any
    device instead of living only in one browser's localStorage."""
    data = request.json or {}
    progress = data.get("progress")
    if not isinstance(progress, dict):
        return flask.jsonify({"status": "error", "message": "progress object required"}), 400
    raw = json.dumps(progress)
    if len(raw) > 200_000:
        return flask.jsonify({"status": "error", "message": "progress payload too large"}), 413
    if current_user.is_authenticated:
        s = SavedSchedule.query.filter_by(user_id=current_user.id, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    else:
        s = SavedSchedule.query.filter_by(guest_session_id=get_guest_session_id(), is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    if not s:
        return flask.jsonify({"status": "none", "message": "No active saved schedule"})
    s.progress_json = raw
    db.session.commit()
    return flask.jsonify({"status": "ok"})


def _active_saved_schedule():
    """The student's live plan row, or None. Owner-scoped."""
    query = SavedSchedule.query
    if current_user.is_authenticated:
        query = query.filter_by(user_id=current_user.id, is_active=True)
    else:
        query = query.filter_by(guest_session_id=get_guest_session_id(), is_active=True)
    return query.order_by(SavedSchedule.created_at.desc()).first()


def _session_reality(task_ids, uid, gid):
    """What Active-study sittings say about these tasks.

    Checkboxes say whether a block was finished; sittings say how many minutes
    were actually spent, including on attempts the student abandoned. Both are
    needed — a plan that ignores abandoned time re-schedules work the student
    has already done.
    """
    try:
        from intelliplan.repositories.active_sessions import ActiveSessionRepository

        repo = ActiveSessionRepository(ActiveSession, ActiveFocusSample, db.session)
        return repo.reality_for(task_ids, user_id=uid, guest_id=gid)
    except Exception as e:
        print(f"[recover] session reality load failed: {e}")
        return {"completed": {}, "abandoned": {}, "finished": set()}


@app.route("/schedule/recover", methods=["POST"])
@limiter.limit("30 per hour")
def recover_schedule():
    """Re-solve the plan from where the student actually is.

    The naive recovery — push what slipped onto tomorrow — is worse than none,
    because it builds a day nobody can do and then a second missed day. This
    credits what was done, drops what is finished, and re-optimises the whole
    remaining horizon under the same cost function, so balance, spacing, and
    deadline buffer all still apply. When the remaining work genuinely does not
    fit, the response says so instead of inventing hours.

    Returns ``changed: false`` when nothing slipped. Re-solving a plan that is
    on track is not free — it moves blocks the student has already made peace
    with, for no gain.
    """
    from intelliplan.intelligence.planner import Reality
    from intelliplan.services.recovery import build_reality, summarise_changes

    data = request.json or {}
    assignments = data.get("assignments", [])
    custom_tasks = data.get("custom_tasks", [])
    hours_per_day = data.get("hours_per_day", 2)
    preferred_time = data.get("preferred_time", "evening")

    row = _active_saved_schedule()
    if not row:
        return flask.jsonify({"status": "none", "message": "No active saved schedule."})

    try:
        before = json.loads(row.schedule_data) if row.schedule_data else {}
    except Exception:
        return flask.jsonify({"status": "error", "message": "Saved plan is unreadable."}), 500
    try:
        progress = json.loads(row.progress_json) if row.progress_json else {}
        if not isinstance(progress, dict):
            progress = {}
    except Exception:
        progress = {}

    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()

    task_ids = {
        str(b.get("task_id"))
        for d in (before.get("schedule") or [])
        for b in (d.get("blocks") or [])
        if isinstance(b, dict) and b.get("task_id")
    }
    sessions = _session_reality(task_ids, uid, gid)

    facts = build_reality(
        before, progress, date.today(), abandoned_minutes=sessions.get("abandoned") or {}
    )
    if not facts.needs_replan:
        return flask.jsonify({
            "status": "ok",
            "changed": False,
            "message": "Your plan is still on track — nothing needed moving.",
            "completed_minutes": facts.completed_minutes,
        })

    if not assignments and not custom_tasks:
        return flask.jsonify({
            "status": "error",
            "message": "Send your current assignments so the plan can be rebuilt.",
        }), 400

    # Two sources of truth about minutes spent, and they overlap: a block the
    # student checked off may also have an Active-study sitting. Taking the
    # larger of the two credits the work once instead of twice.
    completed = dict(facts.reality.completed)
    for task_id, minutes in (sessions.get("completed") or {}).items():
        completed[str(task_id)] = max(completed.get(str(task_id), 0), int(minutes or 0))
    reality = Reality(
        completed=completed,
        abandoned=facts.reality.abandoned,
        missed_task_ids=facts.reality.missed_task_ids,
        finished_task_ids=facts.reality.finished_task_ids
        | frozenset(str(t) for t in (sessions.get("finished") or ())),
    )

    normalized = []
    for a in assignments:
        if not isinstance(a, dict):
            continue
        normalized.append({
            **a,
            "difficulty": a.get("difficulty") or infer_task_difficulty(
                a.get("points_possible"), a.get("priority", "Medium"), a.get("due_date")
            ),
        })

    try:
        from intelliplan.intelligence.planner import PlannerConfig
        from intelliplan.services.scheduling import SchedulingService, StudentContext

        dna, availability, commitments = build_scheduler_personalization(
            user_id=uid, guest_id=gid
        )
        rows = _planner_task_rows(
            normalized, custom_tasks,
            descriptions=_saved_descriptions_for(normalized),
        )
        try:
            comfort = int(round(float(hours_per_day) * 60)) or None
        except (TypeError, ValueError):
            comfort = None
        service = SchedulingService(
            StudentContext(
                availability=availability,
                commitments=commitments,
                preferred_time=preferred_time,
                feedback_rows=_planner_feedback_rows(uid, gid),
                session_rows=_planner_session_rows(uid, gid),
                concept_mastery=_planner_concept_mastery(uid),
                weak_days=tuple(getattr(dna, "weak_days", ()) or ()),
                daily_target_minutes=comfort,
                busy_by_date=_planner_busy_by_date(),
            ),
            PlannerConfig(),
        )
        after = service.to_schedule_data(service.replan(rows, reality))
        try:
            after = enrich_schedule_data(after, normalized, preferred_time, hours_per_day)
        except Exception as ee:
            print(f"[recover] enrich failed (non-fatal): {ee}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        return flask.jsonify({
            "status": "error", "message": safe_error_message(e), "retryable": True,
        }), 500

    changes = summarise_changes(before, after, facts)

    # Auto-save in place. A recovery the student has to remember to save is a
    # recovery that silently does not happen.
    try:
        row.schedule_data = json.dumps(after)
        db.session.commit()
        invalidate_schedule_cache(user_id=uid, guest_id=gid)
    except Exception as e:
        db.session.rollback()
        print(f"[recover] save failed: {e}")

    try:
        notifications_glue.on_plan_rescheduled(
            uid, len([c for c in changes if c.kind == "moved"]), "missed sessions"
        ) if uid else None
    except Exception as e:
        print(f"[recover] notification failed: {e}")

    return flask.jsonify({
        "status": "ok",
        "changed": True,
        "data": after,
        "missed_minutes": facts.missed_minutes,
        "completed_minutes": facts.completed_minutes,
        "lost_days": [d.isoformat() for d in facts.lost_days],
        "missed": [
            {"title": m.title, "day": m.day.isoformat(), "minutes": m.minutes}
            for m in facts.missed
        ],
        "changes": [c.to_dict() for c in changes],
        "overloaded": bool(after.get("overloaded")),
    })


@app.route("/schedule/update", methods=["POST"])
def update_saved_schedule():
    """Replace the active saved schedule's data in place.

    Used when the student reschedules blocks in the Interactive View.
    Unlike /schedule/save this never inserts a new row, so the plan's
    name, created date, and synced progress all survive the edit."""
    data = request.json or {}
    schedule_data = data.get("schedule_data")
    if not isinstance(schedule_data, dict) or not schedule_data.get("schedule"):
        return flask.jsonify({"status": "error", "message": "schedule_data required"}), 400
    raw = json.dumps(schedule_data)
    if len(raw) > 500_000:
        return flask.jsonify({"status": "error", "message": "schedule too large"}), 413
    if current_user.is_authenticated:
        s = SavedSchedule.query.filter_by(user_id=current_user.id, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    else:
        s = SavedSchedule.query.filter_by(guest_session_id=get_guest_session_id(), is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    if not s:
        return flask.jsonify({"status": "none", "message": "No active saved schedule"})
    s.schedule_data = raw
    db.session.commit()
    return flask.jsonify({"status": "ok"})

@app.route("/schedule/reflow", methods=["POST"])
def reflow_saved_schedule():
    """Re-time a hand-rearranged schedule against the student's real free time.

    The Interactive View calls this after a drag-and-drop. The client has
    already applied an optimistic local re-time so the UI feels instant; this
    is the authoritative pass that knows the student's availability and
    commitments, so it's what catches "you just dragged this into your soccer
    practice". Block order is preserved exactly — only clock times move.
    """
    data = request.json or {}
    schedule_data = data.get("schedule_data")
    if not isinstance(schedule_data, dict) or not schedule_data.get("schedule"):
        return flask.jsonify({"status": "error", "message": "schedule_data required"}), 400
    if len(json.dumps(schedule_data)) > 500_000:
        return flask.jsonify({"status": "error", "message": "schedule too large"}), 413
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    dna, availability, commitments = build_scheduler_personalization(user_id=uid, guest_id=gid)
    try:
        schedule_data = reflow_schedule(
            schedule_data, availability=availability, commitments=commitments, dna=dna,
            preferred_time=data.get("preferred_time") or "evening",
            hours_per_day=data.get("hours_per_day") or 2,
        )
    except Exception as e:
        print(f"[reflow] failed: {e}")
        return flask.jsonify({"status": "error", "message": "Could not re-time the plan."}), 500
    # Persist in place so the rearrangement survives a reload, exactly like
    # /schedule/update. Missing row is not an error — an unsaved preview can
    # still be reflowed, the client just keeps it client-side.
    if uid:
        s = SavedSchedule.query.filter_by(user_id=uid, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    else:
        s = SavedSchedule.query.filter_by(guest_session_id=gid, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    if s:
        s.schedule_data = json.dumps(schedule_data)
        db.session.commit()
    return flask.jsonify({"status": "ok", "data": schedule_data, "saved": bool(s)})


@app.route("/schedule/delete", methods=["POST"])
def delete_saved_schedule():
    if current_user.is_authenticated:
        SavedSchedule.query.filter_by(user_id=current_user.id).delete()
    else:
        SavedSchedule.query.filter_by(guest_session_id=get_guest_session_id()).delete()
    db.session.commit()
    return flask.jsonify({"status": "ok"})


# ── UNIVERSAL CALENDAR EXPORT (.ics) ──────────────────────────
def _ics_escape(text):
    """Escape text per RFC 5545 (backslash, semicolon, comma, newline)."""
    return (str(text or "")
            .replace("\\", "\\\\")
            .replace(";", "\\;")
            .replace(",", "\\,")
            .replace("\r\n", "\\n")
            .replace("\n", "\\n"))


def _parse_12h_time(raw):
    """Parse '7:00 PM' / '7 PM' into a time object, or None."""
    cleaned = str(raw or "").strip().upper().replace(".", "")
    for fmt in ("%I:%M %p", "%I %p"):
        try:
            return datetime.strptime(cleaned, fmt).time()
        except ValueError:
            continue
    return None


def _schedule_to_ics(schedule_data, name="IntelliPlan Study Plan"):
    """Convert schedule JSON into an RFC 5545 iCalendar string.

    Events use floating local times (no TZID) so blocks land at the same
    wall-clock hour in any calendar app — Apple, Outlook, Google, Proton.
    Block times come from start_iso/end_iso when present, but are always
    re-anchored onto the day's own date: humanize_schedule() stamps every
    day's ISO times with *today's* date, so only the clock part is trusted.
    """
    now_stamp = utcnow().strftime("%Y%m%dT%H%M%SZ")
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//IntelliPlan//Study Scheduler//EN",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
        f"X-WR-CALNAME:{_ics_escape(name)}",
    ]
    seq = 0
    for day in (schedule_data.get("schedule") or []):
        date_str = str(day.get("date") or "")[:10]
        try:
            day_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        except ValueError:
            continue
        for block in (day.get("blocks") or []):
            if block.get("is_break"):
                continue
            start_t = end_t = None
            for key, target in (("start_iso", "start"), ("end_iso", "end")):
                if block.get(key):
                    try:
                        parsed = datetime.fromisoformat(str(block[key])).time()
                        if target == "start":
                            start_t = parsed
                        else:
                            end_t = parsed
                    except ValueError:
                        pass
            if start_t is None and block.get("time_slot"):
                parts = str(block["time_slot"]).split("-")
                start_t = _parse_12h_time(parts[0]) if parts else None
                if len(parts) > 1:
                    end_t = _parse_12h_time(parts[1])
            if start_t is None:
                continue
            start_dt = datetime.combine(day_date, start_t)
            if end_t is not None:
                end_dt = datetime.combine(day_date, end_t)
                if end_dt <= start_dt:
                    end_dt += timedelta(days=1)
            else:
                end_dt = start_dt + timedelta(minutes=int(block.get("duration_minutes") or 30))
            seq += 1
            summary = block.get("assignment") or "Study block"
            if block.get("course"):
                summary = f"{summary} ({block['course']})"
            desc_bits = []
            if block.get("notes"):
                desc_bits.append(str(block["notes"]))
            desc_bits.append("Planned with IntelliPlan.")
            lines += [
                "BEGIN:VEVENT",
                f"UID:ipd-{now_stamp}-{seq}@intelliplan.tech",
                f"DTSTAMP:{now_stamp}",
                f"DTSTART:{start_dt.strftime('%Y%m%dT%H%M%S')}",
                f"DTEND:{end_dt.strftime('%Y%m%dT%H%M%S')}",
                f"SUMMARY:{_ics_escape(summary)}",
                f"DESCRIPTION:{_ics_escape(' '.join(desc_bits))}",
                "END:VEVENT",
            ]
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines) + "\r\n"


@app.route("/schedule/export.ics", methods=["POST"])
def export_schedule_ics():
    """Download the current schedule as a universal .ics calendar file.

    Accepts schedule_data in the POST body; falls back to the active
    saved schedule so the button also works after a page reload."""
    data = request.json or {}
    schedule_data = data.get("schedule_data")
    name = (data.get("name") or "IntelliPlan Study Plan").strip() or "IntelliPlan Study Plan"
    if not schedule_data:
        if current_user.is_authenticated:
            s = SavedSchedule.query.filter_by(user_id=current_user.id, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
        else:
            s = SavedSchedule.query.filter_by(guest_session_id=get_guest_session_id(), is_active=True).order_by(SavedSchedule.created_at.desc()).first()
        if s:
            try:
                schedule_data = json.loads(s.schedule_data)
                name = s.name or name
            except Exception:
                schedule_data = None
    if not schedule_data:
        return flask.jsonify({"status": "error", "message": "No schedule to export"}), 400
    ics = _schedule_to_ics(schedule_data, name)
    resp = flask.Response(ics, mimetype="text/calendar")
    fname = re.sub(r"[^A-Za-z0-9._-]+", "-", name).strip("-") or "study-plan"
    resp.headers["Content-Disposition"] = f'attachment; filename="{fname}.ics"'
    return resp

# ── DAY ARCHIVE / MEMORIES ────────────────────────────────────
ARCHIVE_TYPE_LABELS = {
    "schedule": "Full schedule",
    "schedule_day": "Single day",
    "resources": "Resources",
    "notes": "Notes",
    "assignments": "Assignments",
    "custom_tasks": "Custom tasks",
    "identity": "Profile & availability",
    "snapshot": "Full day snapshot",
    "manual_tasks": "Manual tasks",
    "library": "Library items",
}

def _archive_owner():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if current_user.is_authenticated else get_guest_session_id()
    return uid, gid

def _archive_query():
    uid, gid = _archive_owner()
    if uid:
        return DayArchive.query.filter_by(user_id=uid)
    return DayArchive.query.filter_by(guest_session_id=gid)

def _parse_archive_date(value):
    if not value:
        return utcnow().date()
    try:
        return datetime.strptime(str(value)[:10], "%Y-%m-%d").date()
    except ValueError:
        return utcnow().date()

def _archive_row_dict(row):
    meta = {}
    if row.meta_json:
        try:
            meta = json.loads(row.meta_json)
        except Exception:
            meta = {}
    return {
        "id": row.id,
        "date": row.archive_date.isoformat(),
        "item_type": row.item_type,
        "type_label": ARCHIVE_TYPE_LABELS.get(row.item_type, row.item_type.replace("_", " ").title()),
        "title": row.title or ARCHIVE_TYPE_LABELS.get(row.item_type, "Saved item"),
        "meta": meta,
        "created_at": row.created_at.strftime("%b %d, %Y %I:%M %p"),
    }

def _apply_schedule_payload(payload, name=None):
    uid, gid = _archive_owner()
    if uid:
        SavedSchedule.query.filter_by(user_id=uid).update({"is_active": False})
    else:
        SavedSchedule.query.filter_by(guest_session_id=gid).update({"is_active": False})
    s = SavedSchedule(
        user_id=uid,
        guest_session_id=gid,
        name=name or "Restored schedule",
        schedule_data=json.dumps(payload),
        is_active=True,
    )
    db.session.add(s)
    return s

def _apply_archive_item(row):
    """Restore one archived item into the live app. Returns a short label."""
    try:
        payload = json.loads(row.payload)
    except Exception:
        return None
    if row.item_type in ("schedule", "snapshot"):
        sched = payload.get("schedule_data") if row.item_type == "snapshot" else payload
        if sched:
            _apply_schedule_payload(sched, row.title or "Restored schedule")
            return "schedule"
    elif row.item_type == "schedule_day":
        day_data = payload
        uid, gid = _archive_owner()
        active = None
        if uid:
            active = SavedSchedule.query.filter_by(user_id=uid, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
        else:
            active = SavedSchedule.query.filter_by(guest_session_id=gid, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
        if active:
            data = json.loads(active.schedule_data)
            days = data.get("schedule") or []
            target_date = day_data.get("date") or row.archive_date.isoformat()
            replaced = False
            for i, d in enumerate(days):
                if d.get("date") == target_date:
                    days[i] = day_data
                    replaced = True
                    break
            if not replaced:
                days.append(day_data)
            data["schedule"] = sorted(days, key=lambda x: x.get("date") or "")
            active.schedule_data = json.dumps(data)
        else:
            _apply_schedule_payload({"schedule": [day_data]}, row.title or "Restored day")
        return "schedule_day"
    elif row.item_type == "custom_tasks":
        session["archived_custom_tasks"] = payload
        return "custom_tasks"
    elif row.item_type == "manual_tasks":
        if current_user.is_authenticated:
            for t in payload if isinstance(payload, list) else []:
                if not isinstance(t, dict):
                    continue
                db.session.add(ManualTask(
                    user_id=current_user.id,
                    title=t.get("title") or "Task",
                    due_date=t.get("due_date") or "",
                    priority=t.get("priority") or "medium",
                    course=t.get("course") or "",
                    estimated_time=t.get("estimated_time") or 60,
                    notes=t.get("notes") or "",
                ))
        return "manual_tasks"
    elif row.item_type == "identity" and current_user.is_authenticated:
        ident = _get_or_create_identity(current_user.id)
        for key in ("grade_level", "focus_areas", "goals", "availability", "weekly_commitments", "class_schedule"):
            if key in payload:
                val = payload[key]
                if key == "focus_areas" and isinstance(val, list):
                    ident.focus_areas = json.dumps(val)
                elif key == "class_schedule" and isinstance(val, list):
                    ident.class_schedule = json.dumps(val)
                elif key == "availability" and isinstance(val, dict):
                    ident.availability = json.dumps(val)
                elif hasattr(ident, key):
                    setattr(ident, key, val)
        return "identity"
    elif row.item_type == "resources":
        session["archived_resources"] = payload
        return "resources"
    elif row.item_type == "notes":
        session["archived_notes"] = payload
        return "notes"
    return row.item_type

@app.route("/archive/context")
def archive_context():
    """Return one-shot restored context after loading a memory (custom tasks, resources)."""
    if not is_logged_in():
        return flask.jsonify({})
    ctx = {}
    for key in ("archived_custom_tasks", "archived_resources", "archived_notes"):
        val = session.pop(key, None)
        if val is not None:
            ctx[key.replace("archived_", "")] = val
    return flask.jsonify(ctx)

@app.route("/memories")
def memories_page():
    return render_template("memories.html", active_page="memories")

@app.route("/archive/save", methods=["POST"])
def archive_save():
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "Login required"}), 401
    data = request.json or {}
    payload = data.get("payload")
    if payload is None:
        return flask.jsonify({"status": "error", "message": "No payload"})
    uid, gid = _archive_owner()
    archive_date = _parse_archive_date(data.get("date"))
    item_type = (data.get("item_type") or "snapshot").strip()[:64]
    title = (data.get("title") or ARCHIVE_TYPE_LABELS.get(item_type, "Saved item"))[:256]
    meta = data.get("meta")
    row = DayArchive(
        user_id=uid,
        guest_session_id=gid,
        archive_date=archive_date,
        item_type=item_type,
        title=title,
        payload=json.dumps(payload),
        meta_json=json.dumps(meta) if meta is not None else None,
    )
    db.session.add(row)
    db.session.commit()
    return flask.jsonify({"status": "ok", "id": row.id, "date": archive_date.isoformat()})

@app.route("/archive/snapshot", methods=["POST"])
def archive_snapshot():
    """Save everything available right now for a given day."""
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "Login required"}), 401
    data = request.json or {}
    archive_date = _parse_archive_date(data.get("date"))
    uid, gid = _archive_owner()
    saved_ids = []
    client_payload = data.get("client") or {}

    if client_payload.get("schedule_data"):
        row = DayArchive(
            user_id=uid, guest_session_id=gid, archive_date=archive_date,
            item_type="schedule", title=client_payload.get("schedule_name") or f"Schedule {archive_date.strftime('%b %d')}",
            payload=json.dumps(client_payload["schedule_data"]),
        )
        db.session.add(row)
        db.session.flush()
        saved_ids.append(row.id)

    if client_payload.get("custom_tasks"):
        row = DayArchive(
            user_id=uid, guest_session_id=gid, archive_date=archive_date,
            item_type="custom_tasks", title="Custom tasks",
            payload=json.dumps(client_payload["custom_tasks"]),
        )
        db.session.add(row)
        db.session.flush()
        saved_ids.append(row.id)

    if client_payload.get("resources"):
        row = DayArchive(
            user_id=uid, guest_session_id=gid, archive_date=archive_date,
            item_type="resources", title=client_payload.get("resources_title") or "Resources",
            payload=json.dumps(client_payload["resources"]),
        )
        db.session.add(row)
        db.session.flush()
        saved_ids.append(row.id)

    active = None
    if uid:
        active = SavedSchedule.query.filter_by(user_id=uid, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    else:
        active = SavedSchedule.query.filter_by(guest_session_id=gid, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
    if active and not client_payload.get("schedule_data"):
        row = DayArchive(
            user_id=uid, guest_session_id=gid, archive_date=archive_date,
            item_type="schedule", title=active.name or "Active schedule",
            payload=active.schedule_data,
        )
        db.session.add(row)
        db.session.flush()
        saved_ids.append(row.id)

    if uid:
        manual = ManualTask.query.filter_by(user_id=uid, done=False).all()
        if manual:
            row = DayArchive(
                user_id=uid, guest_session_id=gid, archive_date=archive_date,
                item_type="manual_tasks", title="Manual tasks",
                payload=json.dumps([{
                    "title": t.title, "due_date": t.due_date, "priority": t.priority,
                    "course": t.course, "estimated_time": t.estimated_time, "notes": t.notes,
                } for t in manual]),
            )
            db.session.add(row)
            db.session.flush()
            saved_ids.append(row.id)
        try:
            ident = _get_or_create_identity(uid)
            ident_payload = ident.to_dict()
            row = DayArchive(
                user_id=uid, guest_session_id=gid, archive_date=archive_date,
                item_type="identity", title="Profile & availability",
                payload=json.dumps(ident_payload),
            )
            db.session.add(row)
            db.session.flush()
            saved_ids.append(row.id)
        except Exception:
            pass

    if not saved_ids:
        return flask.jsonify({"status": "error", "message": "Nothing to save"})
    db.session.commit()
    return flask.jsonify({"status": "ok", "ids": saved_ids, "date": archive_date.isoformat(), "count": len(saved_ids)})

@app.route("/archive/days")
def archive_days():
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "Login required"}), 401
    # Paged by *day*, not by row, so a day's items never split across two
    # pages. We select the page of distinct dates first, then fetch only the
    # rows belonging to them.
    page, per_page = _page_args(14)
    date_rows = (_archive_query()
                 .with_entities(DayArchive.archive_date)
                 .distinct()
                 .order_by(DayArchive.archive_date.desc())
                 .limit(per_page + 1)
                 .offset((page - 1) * per_page)
                 .all())
    has_more = len(date_rows) > per_page
    dates = [d[0] for d in date_rows[:per_page]]

    days = []
    if dates:
        rows = (_archive_query()
                .filter(DayArchive.archive_date.in_(dates))
                .order_by(DayArchive.archive_date.desc(), DayArchive.created_at.desc())
                .all())
        by_date = {}
        for row in rows:
            key = row.archive_date.isoformat()
            if key not in by_date:
                by_date[key] = {"date": key,
                                "label": row.archive_date.strftime("%A, %B %d, %Y"),
                                "items": [], "count": 0}
            by_date[key]["items"].append(_archive_row_dict(row))
            by_date[key]["count"] += 1
        days = sorted(by_date.values(), key=lambda d: d["date"], reverse=True)

    # `total` counts distinct archived dates, not the rows on this page —
    # the page unit is a day, so summing this page's items produced
    # nonsense like "showing 4 of 2" in the paginator.
    try:
        total_days = (_archive_query()
                      .with_entities(DayArchive.archive_date)
                      .distinct()
                      .count())
    except Exception:
        total_days = None

    payload = {
        "status": "ok",
        "days": days,
        "items": days,
        "page": page,
        "per_page": per_page,
        "has_more": has_more,
        "items_on_page": sum(d["count"] for d in days),
    }
    if total_days is not None:
        payload["total"] = total_days
    return flask.jsonify(payload)

@app.route("/archive/day/<date_str>")
def archive_day(date_str):
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "Login required"}), 401
    try:
        target = datetime.strptime(date_str[:10], "%Y-%m-%d").date()
    except ValueError:
        return flask.jsonify({"status": "error", "message": "Invalid date"}), 400
    rows = _archive_query().filter_by(archive_date=target).order_by(DayArchive.created_at.desc()).all()
    return flask.jsonify({
        "status": "ok",
        "date": target.isoformat(),
        "label": target.strftime("%A, %B %d, %Y"),
        "items": [_archive_row_dict(r) for r in rows],
    })

@app.route("/archive/load/<int:item_id>", methods=["POST"])
def archive_load_item(item_id):
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "Login required"}), 401
    row = _archive_query().filter_by(id=item_id).first()
    if not row:
        return flask.jsonify({"status": "error", "message": "Not found"}), 404
    loaded = _apply_archive_item(row)
    db.session.commit()
    redirect_to = "/scheduler/saved" if loaded in ("schedule", "schedule_day", "snapshot") else "/dashboard"
    return flask.jsonify({"status": "ok", "loaded": loaded, "redirect": redirect_to})

@app.route("/archive/load-day/<date_str>", methods=["POST"])
def archive_load_day(date_str):
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "Login required"}), 401
    try:
        target = datetime.strptime(date_str[:10], "%Y-%m-%d").date()
    except ValueError:
        return flask.jsonify({"status": "error", "message": "Invalid date"}), 400
    rows = _archive_query().filter_by(archive_date=target).order_by(DayArchive.created_at.asc()).all()
    if not rows:
        return flask.jsonify({"status": "error", "message": "No items for this day"}), 404
    loaded = []
    for row in rows:
        label = _apply_archive_item(row)
        if label:
            loaded.append(label)
    db.session.commit()
    redirect_to = "/scheduler/saved" if any(x in loaded for x in ("schedule", "schedule_day", "snapshot")) else "/dashboard"
    return flask.jsonify({"status": "ok", "loaded": loaded, "redirect": redirect_to, "count": len(loaded)})

@app.route("/archive/delete/<int:item_id>", methods=["POST"])
def archive_delete_item(item_id):
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "Login required"}), 401
    row = _archive_query().filter_by(id=item_id).first()
    if not row:
        return flask.jsonify({"status": "error", "message": "Not found"}), 404
    db.session.delete(row)
    db.session.commit()
    return flask.jsonify({"status": "ok"})

def _archive_csv_rows(rows):
    import csv
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Date", "Type", "Title", "Created", "Summary"])
    for row in rows:
        summary = ""
        try:
            payload = json.loads(row.payload)
            if row.item_type in ("schedule", "schedule_day"):
                blocks = payload.get("blocks") if row.item_type == "schedule_day" else None
                if blocks is None and isinstance(payload.get("schedule"), list):
                    blocks = []
                    for d in payload["schedule"]:
                        blocks.extend(d.get("blocks") or [])
                summary = f"{len(blocks or [])} blocks"
            elif isinstance(payload, list):
                summary = f"{len(payload)} items"
            elif isinstance(payload, dict):
                summary = ", ".join(list(payload.keys())[:5])
        except Exception:
            summary = ""
        writer.writerow([
            row.archive_date.isoformat(),
            ARCHIVE_TYPE_LABELS.get(row.item_type, row.item_type),
            row.title,
            row.created_at.strftime("%Y-%m-%d %H:%M"),
            summary,
        ])
    output.seek(0)
    return output.getvalue()

@app.route("/archive/export")
def archive_export_all():
    if not is_logged_in():
        return redirect(url_for("login"))
    rows = _archive_query().order_by(DayArchive.archive_date.desc(), DayArchive.created_at.desc()).all()
    csv_data = _archive_csv_rows(rows)
    return flask.Response(
        csv_data,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment;filename=intelliplan_memories.csv"},
    )

@app.route("/archive/export/<date_str>")
def archive_export_day(date_str):
    if not is_logged_in():
        return redirect(url_for("login"))
    try:
        target = datetime.strptime(date_str[:10], "%Y-%m-%d").date()
    except ValueError:
        return flask.jsonify({"status": "error", "message": "Invalid date"}), 400
    rows = _archive_query().filter_by(archive_date=target).order_by(DayArchive.created_at.asc()).all()
    csv_data = _archive_csv_rows(rows)
    fname = f"intelliplan_memories_{date_str[:10]}.csv"
    return flask.Response(
        csv_data,
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment;filename={fname}"},
    )

# ── FEEDBACK ──────────────────────────────────────────────────
@app.route("/feedback/predict-time", methods=["GET"])
def feedback_predict_time():
    """Predict how long this task will take based on the user's past
    TaskFeedback records. Falls back to the original estimate if there's
    no signal yet.

    Lookup strategy (most → least specific):
      1. Median actual_time of past completions of THIS exact title.
      2. Weighted blend of:
         - course average actual_time
         - priority-bucket average actual_time
         (weights = 0.6 / 0.4 when both exist; else whichever is present)
      3. Global median actual_time across all the user's completions.
      4. The estimate the caller already had (or 60 min).
    """
    if not is_logged_in():
        return flask.jsonify({"status": "ok", "predicted_minutes": None, "source": "anon"})
    title = (request.args.get("title") or "").strip()
    course = (request.args.get("course") or "").strip()
    priority = (request.args.get("priority") or "Medium").strip()
    fallback = request.args.get("fallback")
    try:
        fallback = int(fallback) if fallback else 60
    except (TypeError, ValueError):
        fallback = 60

    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    q = TaskFeedback.query
    if uid:
        q = q.filter_by(user_id=uid)
    else:
        q = q.filter_by(guest_session_id=gid)
    rows = q.filter(TaskFeedback.actual_time.isnot(None)).order_by(TaskFeedback.id.desc()).limit(500).all()

    def median(xs):
        xs = sorted(int(x) for x in xs if x is not None)
        if not xs:
            return None
        n = len(xs)
        return xs[n // 2] if n % 2 else (xs[n // 2 - 1] + xs[n // 2]) // 2

    same_title = [r.actual_time for r in rows if r.title.lower() == title.lower()] if title else []
    if same_title:
        med = median(same_title)
        if med:
            return flask.jsonify({"status": "ok", "predicted_minutes": med, "source": "same_title", "sample_size": len(same_title)})

    same_course = [r.actual_time for r in rows if course and (r.course or "").lower() == course.lower()]
    same_priority = [r.actual_time for r in rows if (r.priority or "").lower() == priority.lower()]

    course_med = median(same_course)
    prio_med = median(same_priority)

    if course_med and prio_med:
        blended = int(round(course_med * 0.6 + prio_med * 0.4))
        return flask.jsonify({"status": "ok", "predicted_minutes": blended, "source": "course+priority",
                              "course_sample": len(same_course), "priority_sample": len(same_priority)})
    if course_med:
        return flask.jsonify({"status": "ok", "predicted_minutes": course_med, "source": "course",
                              "sample_size": len(same_course)})
    if prio_med:
        return flask.jsonify({"status": "ok", "predicted_minutes": prio_med, "source": "priority",
                              "sample_size": len(same_priority)})

    global_med = median([r.actual_time for r in rows])
    if global_med:
        return flask.jsonify({"status": "ok", "predicted_minutes": global_med, "source": "global",
                              "sample_size": len(rows)})

    return flask.jsonify({"status": "ok", "predicted_minutes": fallback, "source": "fallback"})


@app.route("/feedback/complete", methods=["POST"])
def feedback_complete():
    data = request.json or {}
    title = data.get("title", "").strip()
    actual_time = data.get("actual_time")
    if not title:
        return flask.jsonify({"status": "error"})
    now = datetime.now()
    hour = now.hour
    time_of_day = "morning" if hour < 12 else "afternoon" if hour < 17 else "evening"
    feedback = TaskFeedback(
        user_id=current_user.id if current_user.is_authenticated else None,
        guest_session_id=None if current_user.is_authenticated else get_guest_session_id(),
        title=title,
        course=data.get("course", ""),
        estimated_time=data.get("estimated_time", 60),
        actual_time=int(actual_time) if actual_time else None,
        difficulty=data.get("difficulty", "Medium"),
        priority=data.get("priority", "Medium"),
        day_of_week=now.strftime("%a"),
        time_of_day=time_of_day
    )
    db.session.add(feedback)
    if data.get("dismiss"):
        save_dismissed(title, data)
    db.session.commit()
    # Streak: feedback-complete = task done
    if current_user.is_authenticated:
        try:
            _record_streak_qualifying_action(current_user.id, data.get("timezone"))
        except Exception as e:
            print(f"[streak] error on feedback complete: {e}")
        # Learning Graph: record task completion event
        try:
            from learning_graph_glue import _learning_graph_on_task_completed
            _learning_graph_on_task_completed(current_user.id, {
                "title": title, "course": data.get("course", ""),
                "estimated_time": data.get("estimated_time", 60),
                "actual_time": int(actual_time) if actual_time else None,
                "difficulty": data.get("difficulty", "Medium"),
            })
        except Exception:
            pass
    return flask.jsonify({"status": "ok"})

@app.route("/feedback/export")
def feedback_export():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"})
    rows = TaskFeedback.query.filter_by(user_id=current_user.id).all()
    import csv
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Subject", "Estimate", "Actual", "Difficulty", "Priority", "DayOfWeek", "TimeOfDay"])
    for r in rows:
        writer.writerow([r.course, r.estimated_time, r.actual_time or "", r.difficulty, r.priority, r.day_of_week, r.time_of_day])
    output.seek(0)
    return flask.Response(output.getvalue(), mimetype="text/csv", headers={"Content-Disposition": "attachment;filename=intelliplan_data.csv"})

# ── PUSH NOTIFICATIONS ────────────────────────────────────────

#: Seconds a push service should hold an undelivered message for.
#:
#: This is not a nicety. pywebpush defaults to TTL 0, meaning "deliver only
#: if the device is connected right now, otherwise discard". Microsoft's
#: WNS — which is where every Edge subscription points — rejects TTL 0
#: outright with `400 Ttl value conflicts with X-WNS-Cache-Policy`, so
#: every push to an Edge user failed, silently, on the server side. Even
#: where TTL 0 is accepted it is the wrong contract: a phone in a bag
#: misses the notification entirely rather than getting it on wake.
#:
#: Twelve hours is chosen against what these messages are: a reminder that
#: something is due is still worth reading later today, and is worth
#: nothing tomorrow.
PUSH_DEFAULT_TTL = 12 * 60 * 60

#: Never send TTL 0 even when a caller asks for it — see above.
PUSH_MIN_TTL = 60

#: Browsers a student can hold subscriptions for at once. Generous enough
#: that phone + laptop + tablet + a spare browser all work, low enough that
#: a loop re-subscribing on every page load cannot grow the table forever.
MAX_PUSH_SUBSCRIPTIONS = 10


def _prune_push_subscriptions(uid, gid):
    """Drop the oldest rows so a new one stays inside the per-owner cap.

    Oldest-first because a subscription the student has not refreshed in
    months is the one most likely to be a browser they no longer open —
    and push services expire those anyway.
    """
    try:
        rows = (PushSubscription.query
                .filter_by(user_id=uid, guest_session_id=gid)
                .order_by(PushSubscription.created_at.asc())
                .all())
        for row in rows[:max(0, len(rows) - (MAX_PUSH_SUBSCRIPTIONS - 1))]:
            db.session.delete(row)
    except Exception as e:
        print(f"[push] prune failed: {e}")


@app.route("/push/subscribe", methods=["POST"])
def push_subscribe():
    data = request.get_json(silent=True) or {}
    sub = data.get("subscription") or {}
    # Validate shape so we don't store garbage payloads. A real Web Push
    # subscription always has an https endpoint + p256dh + auth keys.
    endpoint = (sub.get("endpoint") if isinstance(sub, dict) else None) or ""
    keys = sub.get("keys") if isinstance(sub, dict) else None
    if (not isinstance(endpoint, str) or not endpoint.startswith("https://")
            or not isinstance(keys, dict) or "p256dh" not in keys or "auth" not in keys):
        return flask.jsonify({"status": "error", "message": "invalid subscription"}), 400
    sub_json = json.dumps(sub)
    # Cap size to avoid abusive payloads filling the DB.
    if len(sub_json) > 4000:
        return flask.jsonify({"status": "error", "message": "subscription too large"}), 400
    if len(endpoint) > 512:
        return flask.jsonify({"status": "error", "message": "endpoint too long"}), 400
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if current_user.is_authenticated else get_guest_session_id()
    try:
        # Match on the endpoint, which identifies the browser. Matching on
        # the owner instead let a second device overwrite the first, so a
        # student who enabled notifications on a laptop stopped getting
        # them on their phone.
        existing = PushSubscription.query.filter_by(
            user_id=uid, guest_session_id=gid, endpoint=endpoint).first()
        if existing:
            existing.subscription_json = sub_json
        else:
            _prune_push_subscriptions(uid, gid)
            db.session.add(PushSubscription(
                user_id=uid, guest_session_id=gid,
                endpoint=endpoint, subscription_json=sub_json))
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500
    return flask.jsonify({"status": "ok"})

@app.route("/push/test", methods=["POST"])
def push_test():
    """Fire a test notification at every browser the caller has registered.

    Every one, not the first: the point of the button is to answer "will I
    actually get reminders on this device?", and testing only one row
    answers it for the wrong device as often as the right one.
    """
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if current_user.is_authenticated else get_guest_session_id()
    subs = PushSubscription.query.filter_by(user_id=uid, guest_session_id=gid).all()
    if not subs:
        return flask.jsonify({
            "status": "error",
            "message": "No subscription on file for this device. Turn notifications on first.",
        })
    if not os.getenv("VAPID_PRIVATE_KEY"):
        return flask.jsonify({
            "status": "error",
            "message": "Push is not configured on the server (no VAPID key).",
        })

    payload = json.dumps({
        "title": "IntelliPlan",
        "body": "Notifications are working!",
        "url": "/dashboard",
    })
    claims = {"sub": f"mailto:{os.getenv('VAPID_EMAIL', 'hello@intelliplan.tech')}"}
    sent, expired, last_error = 0, 0, None
    try:
        from pywebpush import webpush
    except ImportError:
        return flask.jsonify({"status": "error", "message": "Push library unavailable."})

    for sub in list(subs):
        try:
            webpush(
                subscription_info=json.loads(sub.subscription_json),
                data=payload,
                vapid_private_key=os.getenv("VAPID_PRIVATE_KEY"),
                vapid_claims=dict(claims),
                # Short, but not zero: a test the student asked for a minute
                # ago is worth delivering when their phone wakes up, and
                # zero is rejected outright by WNS. See PUSH_DEFAULT_TTL.
                ttl=15 * 60,
            )
            sent += 1
        except Exception as e:
            msg = str(e)
            # 404/410 mean the browser dropped the subscription. Keeping the
            # row makes every later send fail on a device that no longer
            # exists, so bin it here as well as in _send_push_to_user.
            if "410" in msg or "404" in msg:
                expired += 1
                try:
                    db.session.delete(sub)
                    db.session.commit()
                except Exception:
                    db.session.rollback()
            else:
                last_error = safe_error_message(e)
    if sent:
        return flask.jsonify({"status": "ok", "sent": sent, "expired": expired})
    if expired:
        return flask.jsonify({
            "status": "error",
            "message": "Your saved subscription had expired. Turn notifications off and on again.",
        })
    return flask.jsonify({"status": "error", "message": last_error or "Push delivery failed."})

@app.route("/push/vapid-public")
def vapid_public():
    return flask.jsonify({"key": os.getenv("VAPID_PUBLIC_KEY", "")})

@app.route("/notifications/silence", methods=["POST"])
def silence_notifications():
    data = request.json or {}
    minutes = int(data.get("minutes", 0))
    if minutes <= 0:
        return jsonify({"status": "error", "message": "Invalid duration"})
    silenced_until = utcnow() + timedelta(minutes=minutes)
    session["notifications_silenced_until"] = silenced_until.isoformat()
    return jsonify({"status": "ok", "silenced_until": silenced_until.isoformat()})

@app.route("/notifications/status")
def notification_status():
    return jsonify({"silenced_until": session.get("notifications_silenced_until")})

# ── DEBUG ─────────────────────────────────────────────────────
@app.route("/debug/auth")
def debug_auth():
    return flask.jsonify({
        "is_authenticated": current_user.is_authenticated,
        "user_id": current_user.id if current_user.is_authenticated else None,
        "session_keys": list(session.keys()),
        "has_google_session": "google_token" in session,
        "has_notion_session": "notion_token" in session,
        "google_db_row": GoogleIntegration.query.filter_by(user_id=current_user.id).first() is not None if current_user.is_authenticated else False,
        "notion_db_row": NotionIntegration.query.filter_by(user_id=current_user.id).first() is not None if current_user.is_authenticated else False,
    })

# ── NOTES HELPERS ─────────────────────────────────────────────
NOTE_ALLOWED_EXTENSIONS = {".txt", ".md", ".csv", ".pdf", ".docx"}

def get_notes_owner_folder():
    if current_user.is_authenticated:
        return f"user_{current_user.id}"
    return f"guest_{get_guest_session_id()}"

def get_notes_owner_query():
    if current_user.is_authenticated:
        return CourseNote.query.filter_by(user_id=current_user.id)
    return CourseNote.query.filter_by(guest_session_id=get_guest_session_id())

def note_belongs_to_current_user(note):
    if current_user.is_authenticated:
        return note.user_id == current_user.id
    return note.guest_session_id == get_guest_session_id()

def extract_text_from_note_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext in {".txt", ".md", ".csv"}:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        if ext == ".pdf":
            try:
                from pypdf import PdfReader
            except Exception:
                from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            pages = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")
            return "\n".join(pages).strip()
        if ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception as e:
        print(f"Note extraction error: {e}")
    return ""

def course_note_payload(note, include_text=False):
    payload = {
        "id": note.id,
        "course_name": note.course_name,
        "course_id": note.course_id,
        "course_source": note.course_source,
        "note_date": note.note_date,
        "title": note.title,
        "original_filename": note.original_filename,
        "has_file": bool(note.stored_filename),
        "download_url": f"/notes/{note.id}/download" if note.stored_filename else None,
        "summary_available": bool((note.summary_cache or "").strip()),
        "created_at": note.created_at.strftime("%b %d, %Y %I:%M %p"),
        "preview": (note.text_content or "")[:240],
    }
    if include_text:
        payload["text_content"] = note.text_content or ""
        payload["summary_cache"] = note.summary_cache or ""
    return payload

# ── EXTENSION ROUTES ──────────────────────────────────────────
@app.route("/extension/login", methods=["POST", "OPTIONS"])
def extension_login():
    if request.method == "OPTIONS":
        response = flask.make_response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Extension-Token"
        response.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        return response
    try:
        data = request.get_json(force=True, silent=True) or {}
        email = data.get("email", "").strip().lower()
        password = data.get("password", "").strip()
        if not email or not password:
            return flask.jsonify({"status": "error", "message": "Email and password required"}), 400
        try:
            user = User.query.filter_by(email=email).first()
        except Exception as lookup_error:
            # A failing User SELECT here is almost always a schema that is
            # behind the code — a column the model knows about and the
            # database does not. Retry the migration once and look again,
            # rather than returning 500 to every extension user until
            # somebody notices.
            print(f"[extension-login] user lookup failed, retrying migration: {lookup_error}")
            db.session.rollback()
            _run_boot_migration_once()
            user = User.query.filter_by(email=email).first()

        if not user or not user.password_hash or not bcrypt.check_password_hash(user.password_hash, password):
            # 401, not 200. A failed login is not a successful request, and
            # a client that only checks the status code was being told it
            # had signed in.
            return flask.jsonify({"status": "error", "message": "Invalid email or password"}), 401
        token = secrets_module.token_hex(32)
        db.session.add(ExtensionToken(user_id=user.id, token=token))
        db.session.commit()
        resp = flask.jsonify({"status": "ok", "token": token, "email": user.email})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp
    except Exception as e:
        # Never return the exception text. This endpoint is unauthenticated,
        # and str(e) on a SQLAlchemy error contains the failing SQL and the
        # table's column list — a free schema dump for anyone who asks.
        print(f"Extension login error: {e}")
        return flask.jsonify({"status": "error", "message": "Login is temporarily unavailable."}), 500

@app.route("/extension/register", methods=["POST", "OPTIONS"])
def extension_register():
    if request.method == "OPTIONS":
        response = flask.make_response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Extension-Token"
        response.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        return response
    try:
        data = request.get_json(force=True, silent=True) or {}
        email = data.get("email", "").strip().lower()
        password = data.get("password", "").strip()
        if not email or not password:
            return flask.jsonify({"status": "error", "message": "Email and password required"})
        if len(password) < 8:
            return flask.jsonify({"status": "error", "message": "Password must be at least 8 characters"})
        if User.query.filter_by(email=email).first():
            return flask.jsonify({"status": "error", "message": "Account already exists"})
        pw_hash = bcrypt.generate_password_hash(password).decode("utf-8")
        user = User(email=email, password_hash=pw_hash)
        db.session.add(user)
        db.session.commit()
        token = secrets_module.token_hex(32)
        db.session.add(ExtensionToken(user_id=user.id, token=token))
        db.session.commit()
        resp = flask.jsonify({"status": "ok", "token": token, "email": user.email})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp
    except Exception as e:
        # Same reasoning as extension_login: no exception text to an
        # unauthenticated caller.
        print(f"Extension register error: {e}")
        return flask.jsonify({"status": "error", "message": "Sign-up is temporarily unavailable."}), 500

@app.route("/extension/logout", methods=["POST", "OPTIONS"])
def extension_logout():
    if request.method == "OPTIONS":
        response = flask.make_response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Extension-Token"
        return response
    token = request.headers.get("X-Extension-Token")
    if token:
        ExtensionToken.query.filter_by(token=token).delete()
        db.session.commit()
    resp = flask.jsonify({"status": "ok"})
    resp.headers["Access-Control-Allow-Origin"] = "*"
    return resp

def get_extension_user(token):
    if not token:
        return None
    try:
        row = ExtensionToken.query.filter_by(token=token).first()
        if not row:
            return None
        return db.session.get(User, row.user_id)
    except Exception:
        return None

def ext_response(data, status=200):
    resp = flask.jsonify(data)
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Extension-Token"
    resp.status_code = status
    return resp

@app.route("/extension/tasks")
def extension_tasks():
    token = request.headers.get("X-Extension-Token")
    user = get_extension_user(token)
    if not user:
        return ext_response({"status": "error", "message": "Not authenticated"}, 401)
    try:
        from datetime import date as date_type
        today = date_type.today()
        priority_order = {"High": 0, "Medium": 1, "Low": 2}
        tasks = []
        dismissed_rows = DismissedAssignment.query.filter_by(user_id=user.id).all()
        dismissed = _DismissedSet(r.title for r in dismissed_rows)
        acct = LinkedAccount.query.filter_by(user_id=user.id, is_active=True).first()
        if acct:
            creds = acct.get_credentials()
            login_type = acct.login_type
            if login_type == "studentvue":
                try:
                    raw = get_sv_assignments(creds["sv_district_url"], creds["sv_username"], creds["sv_password"])
                    missing = get_missing_assignments(creds["sv_district_url"], creds["sv_username"], creds["sv_password"])
                    for a in raw + missing:
                        if a["title"] not in dismissed:
                            a.setdefault("source", "studentvue")
                            tasks.append(a)
                except Exception as e:
                    print(f"Ext SV error: {e}")
            elif login_type == "canvas":
                try:
                    canvas_token = creds["canvas_token"]
                    canvas_url = creds.get("canvas_url", "https://canvas.instructure.com")
                    headers = {"Authorization": f"Bearer {canvas_token}"}
                    courses = requests.get(f"{canvas_url}/api/v1/courses", headers=headers, timeout=10).json()
                    course_map = {c["id"]: c.get("name", "Unknown") for c in courses if isinstance(c, dict) and "id" in c}
                    for course_id in course_map:
                        resp = requests.get(f"{canvas_url}/api/v1/courses/{course_id}/assignments?per_page=100", headers=headers, timeout=10).json()
                        if not isinstance(resp, list):
                            continue
                        for a in resp:
                            if not isinstance(a, dict) or not a.get("due_at"):
                                continue
                            due_str = a["due_at"][:10]
                            try:
                                due = datetime.strptime(due_str, "%Y-%m-%d").date()
                            except Exception:
                                continue
                            days = (due - today).days
                            if days < -14:
                                continue
                            title = a["name"]
                            priority = compute_priority(days, a.get("points_possible") or 0, title)
                            if title in dismissed:
                                continue
                            est_minutes, description = _lms_row_sizing(
                                a, a.get("points_possible")
                            )
                            tasks.append({
                                "title": title,
                                "course": course_map.get(a["course_id"], "Unknown"),
                                "due_date": due_str,
                                "priority": priority,
                                "source": "canvas",
                                "estimated_time": est_minutes,
                                "description": description,
                                "color": PRIORITY_COLORS.get(priority, "#f59e0b")
                            })
                except Exception as e:
                    print(f"Ext Canvas error: {e}")
        manual = ManualTask.query.filter_by(user_id=user.id, done=False).all()
        for t in manual:
            if t.title not in dismissed:
                tasks.append({"id": t.id, "title": t.title, "due_date": t.due_date or "", "priority": t.priority, "course": t.course, "estimated_time": t.estimated_time, "source": "manual", "color": PRIORITY_COLORS.get(t.priority, "#f59e0b")})
        result = {"today": [], "upcoming": [], "overdue": []}
        for t in tasks:
            due = t.get("due_date", "")
            if not due:
                result["upcoming"].append(t)
                continue
            try:
                due_date = datetime.strptime(due, "%Y-%m-%d").date()
                if due_date < today:
                    result["overdue"].append(t)
                elif due_date == today:
                    result["today"].append(t)
                else:
                    result["upcoming"].append(t)
            except Exception:
                result["upcoming"].append(t)
        for key in result:
            result[key].sort(key=lambda x: (x.get("due_date", "9999"), priority_order.get(x.get("priority", "Low"), 2)))
        return ext_response(result)
    except Exception as e:
        print(f"Extension tasks error: {e}")
        return ext_response({"status": "error", "message": safe_error_message(e)}, 500)

@app.route("/extension/schedule")
def extension_schedule():
    token = request.headers.get("X-Extension-Token")
    user = get_extension_user(token)
    if not user:
        return ext_response({"status": "error"}, 401)
    try:
        s = SavedSchedule.query.filter_by(user_id=user.id, is_active=True).order_by(SavedSchedule.created_at.desc()).first()
        if not s:
            return ext_response({"status": "none"})
        return ext_response({"status": "ok", "name": s.name, "created_at": s.created_at.strftime("%b %d, %Y"), "data": json.loads(s.schedule_data)})
    except Exception as e:
        print(f"Extension schedule error: {e}")
        return ext_response({"status": "error"}, 500)

@app.route("/extension/grades")
def extension_grades():
    token = request.headers.get("X-Extension-Token")
    user = get_extension_user(token)
    if not user:
        return ext_response([], 401)
    try:
        acct = LinkedAccount.query.filter_by(user_id=user.id, is_active=True).first()
        if not acct:
            return ext_response([])
        creds = acct.get_credentials()
        if acct.login_type == "studentvue":
            from studentvue_helper import get_grades as get_sv_grades
            grades = get_sv_grades(creds["sv_district_url"], creds["sv_username"], creds["sv_password"])
            return ext_response(grades)
        return ext_response([])
    except Exception as e:
        print(f"Extension grades error: {e}")
        return ext_response([], 500)

@app.route("/extension/dismiss", methods=["POST", "OPTIONS"])
def extension_dismiss():
    if request.method == "OPTIONS":
        response = flask.make_response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type, X-Extension-Token"
        return response
    token = request.headers.get("X-Extension-Token")
    user = get_extension_user(token)
    if not user:
        return ext_response({"status": "error"}, 401)
    try:
        data = request.get_json(force=True, silent=True) or {}
        title = data.get("title", "")
        if title:
            existing = DismissedAssignment.query.filter_by(user_id=user.id, title=title).first()
            if not existing:
                db.session.add(DismissedAssignment(user_id=user.id, title=title, data=json.dumps(data)))
                db.session.commit()
        return ext_response({"status": "ok"})
    except Exception as e:
        return ext_response({"status": "error"}, 500)


@app.route("/extension/session-token", methods=["GET", "OPTIONS"])
def extension_session_token():
    """Return an extension token for the user's active web session.
    Called by the Chrome extension with credentials:'include' so the
    browser's session cookie is sent. If the user is already logged in
    on intelliplan.tech, this auto-logs them into the extension too."""
    origin = request.headers.get("Origin", "")

    def _cors(resp):
        allowed = (
            origin.startswith("chrome-extension://") or
            origin.startswith("http://localhost") or
            origin.startswith("http://127.0.0.1") or
            "intelliplan.tech" in origin
        )
        if allowed and origin:
            resp.headers["Access-Control-Allow-Origin"] = origin
            resp.headers["Access-Control-Allow-Credentials"] = "true"
            resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
            resp.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
            resp.headers["Vary"] = "Origin"
        return resp

    if request.method == "OPTIONS":
        return _cors(flask.make_response("", 204))

    if not current_user.is_authenticated:
        r = flask.jsonify({"status": "error", "message": "Not logged in"})
        return _cors(r), 401

    import secrets as _secrets
    existing = ExtensionToken.query.filter_by(user_id=current_user.id).first()
    if existing:
        token = existing.token
    else:
        token = _secrets.token_hex(32)
        db.session.add(ExtensionToken(user_id=current_user.id, token=token))
        db.session.commit()

    r = flask.jsonify({"status": "ok", "token": token, "email": current_user.email})
    return _cors(r), 200


# ── STUDY ROUTES ──────────────────────────────────────────────
@app.route("/study/evaluate", methods=["POST"])
def study_evaluate():
    data = request.json or {}
    question = data.get("question", "").strip()
    correct_answer = data.get("correct_answer", "").strip()
    user_answer = data.get("user_answer", "").strip()
    confidence = data.get("confidence", "medium")
    if not user_answer:
        return flask.jsonify({"status": "error", "message": "No answer provided"})
    prompt = f'''Evaluate this student answer against the correct answer SEMANTICALLY and LENIENTLY.

QUESTION: {question}
CORRECT ANSWER: {correct_answer}
STUDENT'S ANSWER: {user_answer}

Guidelines:
- Focus on meaning, not exact wording.
- If the student captures the main idea, mark at least "partial".
- Minor wording differences or missing detail should NOT be marked incorrect.
- Only mark "incorrect" if the core concept is wrong or missing.
- Reward approximate understanding.

Return ONLY valid JSON:
{{
  "verdict": "correct" | "partial" | "incorrect",
  "score": 0-100,
  "what_was_right": "Encouraging feedback on what was correct",
  "what_was_missing": "Precise gaps or misconceptions",
  "critique": "2-3 sentence constructive critique",
  "memory_anchor": "One vivid way to remember this concept",
  "better_answer": "Ideal concise response"
}}

Scoring guide:
- correct: 70-100 (main idea + mostly accurate)
- partial: 40-69 (some understanding present)
- incorrect: 0-39 (core idea missing or wrong)
'''
    try:
        raw = ai_chat(
            [{"role": "user", "content": prompt}],
            tier="standard",
            temperature=0.4,
            max_tokens=800,
        )
        raw = re.sub(r"```json\n?", "", raw)
        raw = re.sub(r"```\n?", "", raw)
        result = json.loads(raw)
        ua = user_answer.lower()
        ca = correct_answer.lower()
        keywords = [w for w in re.findall(r'\w+', ca) if len(w) > 4]
        keyword_hits = sum(1 for w in keywords if w in ua)
        from difflib import SequenceMatcher
        similarity = SequenceMatcher(None, ua, ca).ratio()
        if result["verdict"] == "incorrect":
            if keyword_hits >= 2 or similarity > 0.5:
                result["verdict"] = "partial"
                result["score"] = max(result.get("score", 0), 45)
        if result["verdict"] == "partial" and similarity > 0.75:
            result["verdict"] = "correct"
            result["score"] = max(result.get("score", 0), 75)
        base = {"correct": 10, "partial": 7, "incorrect": 3}.get(result["verdict"], 3)
        conf_mult = {"high": 1.5, "medium": 1.0, "low": 0.7}.get(confidence, 1.0)
        if result["verdict"] == "incorrect" and confidence == "high":
            conf_mult = 0.6
        result["points_earned"] = max(1, round(base * conf_mult))
        return flask.jsonify({"status": "ok", "evaluation": result})
    except Exception as e:
        print(f"Study evaluate error: {e}")
        return flask.jsonify({
            "status": "error",
            "message": "Evaluation temporarily unavailable. Please try again."
        })

@app.route("/study/analyze-image", methods=["POST"])
def study_analyze_image():
    if "image" not in request.files:
        return flask.jsonify({"status": "error", "message": "No image provided"})
    img_file = request.files["image"]
    allowed_types = {"image/jpeg", "image/png", "image/webp", "image/gif"}
    if img_file.content_type not in allowed_types:
        return flask.jsonify({"status": "error", "message": "Only JPEG, PNG, WebP, or GIF images are supported"})
    try:
        raw = img_file.read()
        if len(raw) > 10 * 1024 * 1024:
            return flask.jsonify({"status": "error", "message": "Image too large. Max 10MB."})
        b64 = base64.b64encode(raw).decode("utf-8")
        media_type = img_file.content_type
        extracted = ai_vision(
            system_prompt="You extract educational content from images for students.",
            user_text="Extract ALL text, formulas, diagrams, tables, and key information from this educational image. Format as clean, readable study material. Preserve all text exactly, describe visual elements, preserve mathematical formulas, and note labels and captions. Output the extracted content directly without any preamble.",
            image_b64=b64,
            image_mime=media_type,
            temperature=0.1,
            max_tokens=2000,
        )
        if not extracted:
            return flask.jsonify({"status": "error", "message": "No content could be extracted from this image"})
        return flask.jsonify({"status": "ok", "text": extracted, "char_count": len(extracted)})
    except Exception as e:
        print(f"Image analysis error: {e}")
        return flask.jsonify({"status": "error", "message": "Image analysis is temporarily unavailable. Try pasting the text manually."})

@app.route("/analyze-image", methods=["POST"])
def analyze_image_general():
    if "image" not in request.files:
        return flask.jsonify({"status": "error", "message": "No image provided"})
    img_file = request.files["image"]
    question = request.form.get("question", "Describe what you see in this image in detail.")
    try:
        raw = img_file.read()
        if len(raw) > 10 * 1024 * 1024:
            return flask.jsonify({"status": "error", "message": "Image too large. Max 10MB."})
        b64 = base64.b64encode(raw).decode("utf-8")
        media_type = img_file.content_type or "image/jpeg"
        answer = ai_vision(
            system_prompt="You are a helpful educational assistant that analyzes images for students.",
            user_text=question,
            image_b64=b64,
            image_mime=media_type,
            temperature=0.3,
            max_tokens=1500,
        )
        return flask.jsonify({"status": "ok", "response": answer})
    except Exception as e:
        print(f"General image analysis error: {e}")
        return flask.jsonify({"status": "error", "message": "Service temporarily unavailable. Please try again later."})

@app.route("/study/points", methods=["GET"])
@app.route("/study/streak", methods=["GET"])
def study_get_points():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    try:
        p = get_study_profile(uid, gid)
        streak_event = reconcile_missed_streak(p)
        tier = current_study_tier(p.streak_count)
        p.freeze_capacity = tier["freeze_cap"]
        weekly_quests = ensure_weekly_quests(p)
        history = safe_json_load(p.streak_history, [])
        sessions = safe_json_load(p.session_history, [])
        badges = safe_json_load(p.badges, [])
        cosmetics = safe_json_load(p.active_cosmetics, {})
        booster = safe_json_load(p.active_booster, None)
        level_info = level_for_sparks(p.sparks_earned_total)
        now = datetime.now()
        today_str = now.strftime("%Y-%m-%d")
        at_risk = bool(p.last_active_date != today_str and now.hour >= 18 and int(p.streak_count or 0) > 0)
        repair_until = p.repair_eligible_until.isoformat() if p.repair_eligible_until else None
        repair_available = bool(p.repair_eligible_until and p.repair_eligible_until > utcnow())
        weekly_item_ids = sorted(SHOP_ITEMS.keys())
        weekly_deal_id = random.Random(active_week_id()).choice(weekly_item_ids)
        db.session.commit()
        return flask.jsonify({
            "status": "ok",
            "total_points": p.total_points,
            "spark_balance": p.spark_balance or 0,
            "sparks_earned_total": p.sparks_earned_total or p.total_points or 0,
            "level": p.level or level_info["level"],
            "level_title": level_info["title"],
            "next_level": level_info["next"],
            "streak_count": p.streak_count,
            "streak_freeze_count": p.streak_freeze_count,
            "freeze_capacity": p.freeze_capacity or tier["freeze_cap"],
            "last_active_date": p.last_active_date,
            "streak_tier": tier,
            "at_risk": at_risk,
            "repair_eligible_until": repair_until,
            "repair_available": repair_available,
            "repair_cost": repair_cost_for(p.longest_streak or p.streak_count),
            "streak_event": streak_event,
            "streak_history": history,
            "session_history": sessions[-20:],
            "longest_streak": p.longest_streak or p.streak_count,
            "total_sessions": p.total_sessions or 0,
            "badges": [{"id": bid, **BADGE_CATALOG.get(bid, {"name": bid.replace("_", " ").title(), "kind": "earned"})} for bid in badges],
            "active_booster": booster,
            "active_cosmetics": cosmetics,
            "weekly_quests": weekly_quests,
            "shop": {
                "items": SHOP_ITEMS,
                "weekly_deal": {
                    "item_id": weekly_deal_id,
                    "discount_percent": 30,
                    "price": int(round(SHOP_ITEMS[weekly_deal_id]["price"] * 0.7))
                }
            }
        })
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/study/sparks/update", methods=["POST"])
@app.route("/study/points/update", methods=["POST"])
def study_update_points():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    data = request.json or {}
    delta = int(data.get("delta", data.get("sparks", 0)) or 0)
    try:
        p = get_study_profile(uid, gid)
        result = grant_sparks(p, delta, data.get("reason", "study"), apply_booster=bool(data.get("apply_booster", True)))
        p.updated_at = utcnow()
        db.session.commit()
        return flask.jsonify({
            "status": "ok",
            "awarded": result["awarded"],
            "base": result["base"],
            "multiplier": result["multiplier"],
            "total_points": p.total_points,
            "spark_balance": p.spark_balance,
            "sparks_earned_total": p.sparks_earned_total,
            "level": p.level,
            "level_title": level_for_sparks(p.sparks_earned_total)["title"],
            "level_up": result["level_up"],
            "active_booster": safe_json_load(p.active_booster, None)
        })
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/study/streak/update", methods=["POST"])
def study_update_streak():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    data = request.json or {}
    today_str = (data.get("local_date") or datetime.now().strftime("%Y-%m-%d"))[:10]
    questions_answered = int(data.get("questions_answered", data.get("questions_total", 0)) or 0)
    if questions_answered < 5:
        return flask.jsonify({"status": "error", "message": "Complete at least 5 questions to count toward your streak."})
    try:
        p = get_study_profile(uid, gid)
        streak_event = reconcile_missed_streak(p)
        history = safe_json_load(p.streak_history, [])
        last = p.last_active_date
        bonus_points = 0
        milestone = None
        new_badges = []
        freeze_awarded = 0
        if last == today_str:
            return flask.jsonify({"status": "ok", "streak_count": p.streak_count, "bonus_points": 0, "total_points": p.total_points, "spark_balance": p.spark_balance, "streak_history": history, "longest_streak": p.longest_streak or p.streak_count, "streak_event": streak_event})
        yesterday = (datetime.strptime(today_str, "%Y-%m-%d") - timedelta(days=1)).strftime("%Y-%m-%d")
        if last == yesterday:
            p.streak_count += 1
        elif last == "":
            p.streak_count = 1
        else:
            if p.streak_freeze_count > 0:
                p.streak_freeze_count -= 1
                p.streak_count += 1
                streak_event = {"type": "freeze_consumed", "message": "Streak Freeze used - your streak is safe. Back tomorrow."}
            else:
                p.streak_count = 1
        p.last_active_date = today_str
        if today_str not in history:
            history.append(today_str)
        history = sorted(history)[-90:]
        p.streak_history = json.dumps(history)
        if p.streak_count > (p.longest_streak or 0):
            p.longest_streak = p.streak_count
        tier = current_study_tier(p.streak_count)
        p.freeze_capacity = tier["freeze_cap"]
        daily = grant_sparks(p, tier["bonus"], "daily_streak", apply_booster=False)
        bonus_points += daily["awarded"]
        if p.streak_count in STREAK_MILESTONES:
            reward = STREAK_MILESTONES[p.streak_count]
            milestone_grant = grant_sparks(p, reward["sparks"], f"milestone:{p.streak_count}", apply_booster=False)
            freeze_awarded += reward["freezes"]
            p.streak_freeze_count = min(int(p.freeze_capacity or tier["freeze_cap"]), int(p.streak_freeze_count or 0) + freeze_awarded)
            new_badges.extend(add_badges(p, [reward.get("badge")]))
            bonus_points += milestone_grant["awarded"]
            milestone = {"day": p.streak_count, "sparks": reward["sparks"], "freezes": reward["freezes"], "badge": reward.get("badge"), "title": reward.get("title")}
        passive = passive_freezes_due(p)
        if passive:
            p.streak_freeze_count = min(int(p.freeze_capacity or tier["freeze_cap"]), int(p.streak_freeze_count or 0) + passive)
            freeze_awarded += passive
        p.repair_eligible_until = None
        db.session.commit()
        return flask.jsonify({
            "status": "ok",
            "streak_count": p.streak_count,
            "streak_freeze_count": p.streak_freeze_count,
            "freeze_capacity": p.freeze_capacity,
            "bonus_points": bonus_points,
            "daily_bonus": tier["bonus"],
            "total_points": p.total_points,
            "spark_balance": p.spark_balance,
            "streak_history": history,
            "longest_streak": p.longest_streak or p.streak_count,
            "streak_tier": tier,
            "milestone": milestone,
            "badges_unlocked": new_badges,
            "freezes_awarded": freeze_awarded,
            "streak_event": streak_event
        })
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/api/activity", methods=["POST"])
def api_activity():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    data = request.json or {}
    today_str = (data.get("local_date") or datetime.now().strftime("%Y-%m-%d"))[:10]
    try:
        p = get_study_profile(uid, gid)
        reconcile_missed_streak(p)
        history = safe_json_load(p.streak_history, [])
        last = p.last_active_date
        if last == today_str:
            return flask.jsonify({"status": "ok", "already_counted": True, "streak_count": p.streak_count, "spark_balance": p.spark_balance, "streak_history": history})
        yesterday = (datetime.strptime(today_str, "%Y-%m-%d") - timedelta(days=1)).strftime("%Y-%m-%d")
        if last == yesterday:
            p.streak_count += 1
        elif last == "":
            p.streak_count = 1
        else:
            if p.streak_freeze_count > 0:
                p.streak_freeze_count -= 1
                p.streak_count += 1
            else:
                p.streak_count = 1
        p.last_active_date = today_str
        if today_str not in history:
            history.append(today_str)
        history = sorted(history)[-90:]
        p.streak_history = json.dumps(history)
        if p.streak_count > (p.longest_streak or 0):
            p.longest_streak = p.streak_count
        grant_sparks(p, 5, "daily_activity")
        db.session.commit()
        return flask.jsonify({"status": "ok", "streak_count": p.streak_count, "spark_balance": p.spark_balance, "streak_history": history})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/study/mastery/update", methods=["POST"])
def study_mastery_update():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    data = request.json or {}
    question_key = data.get("question_key", "")[:512]
    verdict = data.get("verdict", "incorrect")
    score = int(data.get("score", 0))
    if not question_key:
        return flask.jsonify({"status": "error"})
    today = datetime.now().strftime("%Y-%m-%d")
    try:
        q = StudyMastery.query.filter_by(user_id=uid, guest_session_id=gid, question_key=question_key).first()
        if not q:
            q = StudyMastery(user_id=uid, guest_session_id=gid, question_key=question_key, question_text=data.get("question_text", "")[:1000], answer_text=data.get("answer_text", "")[:1000], topic=data.get("topic", "")[:256])
            db.session.add(q)
        q.times_seen += 1
        q.last_seen = today
        if verdict == "correct":
            q.times_correct += 1
            q.easiness_factor = max(1.3, q.easiness_factor + 0.1 - (5 - min(5, score // 20)) * (0.08 + (5 - min(5, score // 20)) * 0.02))
            if q.times_correct == 1:
                q.interval_days = 1
            elif q.times_correct == 2:
                q.interval_days = 6
            else:
                q.interval_days = round(q.interval_days * q.easiness_factor)
            q.mastery_level = min(3, q.mastery_level + 1)
        elif verdict == "partial":
            q.times_partial += 1
            q.interval_days = max(1, q.interval_days // 2)
            q.easiness_factor = max(1.3, q.easiness_factor - 0.15)
        else:
            q.interval_days = 1
            q.easiness_factor = max(1.3, q.easiness_factor - 0.2)
            q.mastery_level = max(0, q.mastery_level - 1)
        q.next_review = (datetime.now() + timedelta(days=q.interval_days)).strftime("%Y-%m-%d")
        db.session.commit()
        # Learning Graph: concept reviewed
        if uid:
            try:
                from learning_graph_glue import _learning_graph_on_concept_reviewed
                _learning_graph_on_concept_reviewed(uid, {
                    "question_key": question_key, "verdict": verdict,
                    "topic": data.get("topic", ""),
                })
            except Exception:
                pass
        mastery_labels = ["Not Learned", "Learning", "Familiar", "Mastered"]
        return flask.jsonify({"status": "ok", "mastery_level": q.mastery_level, "mastery_label": mastery_labels[q.mastery_level], "next_review": q.next_review, "interval_days": q.interval_days})
    except Exception as e:
        print(f"Mastery update error: {e}")
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/study/mastery/due", methods=["GET"])
def study_mastery_due():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    today = datetime.now().strftime("%Y-%m-%d")
    try:
        if uid:
            items = StudyMastery.query.filter(StudyMastery.user_id == uid, StudyMastery.next_review <= today, StudyMastery.mastery_level < 3).order_by(StudyMastery.next_review.asc()).limit(20).all()
        else:
            items = StudyMastery.query.filter(StudyMastery.guest_session_id == gid, StudyMastery.next_review <= today, StudyMastery.mastery_level < 3).order_by(StudyMastery.next_review.asc()).limit(20).all()
        mastery_labels = ["Not Learned", "Learning", "Familiar", "Mastered"]
        return flask.jsonify([{"question_key": m.question_key, "question_text": m.question_text, "answer_text": m.answer_text, "topic": m.topic, "mastery_level": m.mastery_level, "mastery_label": mastery_labels[m.mastery_level], "times_seen": m.times_seen, "times_correct": m.times_correct, "next_review": m.next_review} for m in items])
    except Exception as e:
        return flask.jsonify([])

@app.route("/study/mastery/all", methods=["GET"])
def study_mastery_all():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    try:
        if uid:
            items = StudyMastery.query.filter_by(user_id=uid).order_by(StudyMastery.mastery_level.asc()).limit(100).all()
        else:
            items = StudyMastery.query.filter_by(guest_session_id=gid).order_by(StudyMastery.mastery_level.asc()).limit(100).all()
        mastery_labels = ["Not Learned", "Learning", "Familiar", "Mastered"]
        return flask.jsonify([{"question_key": m.question_key, "question_text": m.question_text, "topic": m.topic, "mastery_level": m.mastery_level, "mastery_label": mastery_labels[m.mastery_level], "times_seen": m.times_seen, "times_correct": m.times_correct, "accuracy": round(m.times_correct / m.times_seen * 100) if m.times_seen else 0, "next_review": m.next_review} for m in items])
    except Exception as e:
        return flask.jsonify([])

@app.route("/study/session/complete", methods=["POST"])
def study_session_complete():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    data = request.json or {}
    try:
        p = get_study_profile(uid, gid)
        sessions = safe_json_load(p.session_history, [])
        questions_total = int(data.get("questions_total", 0) or 0)
        questions_correct = int(data.get("questions_correct", 0) or 0)
        questions_partial = int(data.get("questions_partial", 0) or 0)
        points_earned = int(data.get("points_earned", data.get("sparks_earned", 0)) or 0)
        duration_seconds = int(data.get("duration_seconds", 0) or 0)
        rec = {
            "date": (data.get("local_date") or datetime.now().strftime("%Y-%m-%d"))[:10],
            "mode": data.get("mode", "casual"),
            "questions": questions_total,
            "correct": questions_correct,
            "partial": questions_partial,
            "points": points_earned,
            "duration": duration_seconds
        }
        sessions.append(rec)
        p.session_history = json.dumps(sessions[-50:])
        if not data.get("already_awarded"):
            grant_sparks(p, points_earned, "session_complete_import", apply_booster=False)
        p.total_sessions = (p.total_sessions or 0) + 1
        badges_to_add = []
        if p.total_sessions >= 1:
            badges_to_add.append("first_session")
        if p.total_sessions >= 10:
            badges_to_add.append("getting_serious")
        if p.total_sessions >= 25:
            badges_to_add.append("dedicated")
        if p.total_sessions >= 50:
            badges_to_add.append("committed")
        if p.total_sessions >= 100:
            badges_to_add.append("unstoppable")
        total_questions = sum(int(s.get("questions", 0) or 0) for s in sessions)
        total_correct = sum(int(s.get("correct", 0) or 0) for s in sessions)
        accuracy = round(total_correct / total_questions * 100) if total_questions else 0
        if accuracy >= 75 and total_questions >= 20:
            badges_to_add.append("sharp")
        if accuracy >= 85 and total_questions >= 20:
            badges_to_add.append("precise")
        if accuracy >= 95 and len(sessions) >= 20:
            badges_to_add.append("flawless")
        if questions_total > 0 and questions_correct >= questions_total:
            badges_to_add.append("perfect_week")
        if duration_seconds and duration_seconds < 300:
            badges_to_add.append("speed_demon")
        local_hour = int(data.get("local_hour", datetime.now().hour) or 0)
        if local_hour < 7:
            badges_to_add.append("early_bird")
        if local_hour == 0:
            badges_to_add.append("night_owl")
        new_badges = add_badges(p, badges_to_add)
        quest_rewards = update_quest_progress(p, {
            "date": rec["date"],
            "mode": rec["mode"],
            "questions": questions_total,
            "correct": questions_correct,
            "sparks": points_earned,
            "duration": duration_seconds,
            "mastered_concepts": int(data.get("mastered_concepts", 0) or 0)
        })
        db.session.commit()
        # Learning Graph: study session ended
        if uid:
            try:
                from learning_graph_glue import _learning_graph_on_study_session_ended
                _learning_graph_on_study_session_ended(uid, {
                    "questions_total": questions_total,
                    "questions_correct": questions_correct,
                    "duration_seconds": duration_seconds,
                    "mode": data.get("mode", "casual"),
                })
            except Exception:
                pass
        return flask.jsonify({
            "status": "ok",
            "total_points": p.total_points,
            "spark_balance": p.spark_balance,
            "total_sessions": p.total_sessions,
            "badges_unlocked": new_badges,
            "quest_rewards": quest_rewards,
            "weekly_quests": safe_json_load(p.weekly_quests, {})
        })
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/study/quests", methods=["GET"])
def study_quests():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    try:
        p = get_study_profile(uid, gid)
        quests = ensure_weekly_quests(p)
        db.session.commit()
        return flask.jsonify({"status": "ok", "weekly_quests": quests})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/study/quests/update", methods=["POST"])
def study_quests_update():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    data = request.json or {}
    try:
        p = get_study_profile(uid, gid)
        rewards = update_quest_progress(p, data)
        db.session.commit()
        return flask.jsonify({"status": "ok", "quest_rewards": rewards, "weekly_quests": safe_json_load(p.weekly_quests, {}), "spark_balance": p.spark_balance})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/study/shop/buy", methods=["POST"])
def study_shop_buy():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    data = request.json or {}
    item_id = data.get("item_id")
    if item_id not in SHOP_ITEMS:
        return flask.jsonify({"status": "error", "message": "Unknown shop item."}), 400
    try:
        p = get_study_profile(uid, gid)
        item = SHOP_ITEMS[item_id]
        weekly_deal_id = random.Random(active_week_id()).choice(sorted(SHOP_ITEMS.keys()))
        price = int(round(item["price"] * 0.7)) if item_id == weekly_deal_id else int(item["price"])
        if int(p.spark_balance or 0) < price:
            return flask.jsonify({"status": "error", "message": "Not enough Sparks.", "spark_balance": p.spark_balance}), 400
        p.spark_balance = int(p.spark_balance or 0) - price
        cosmetics = safe_json_load(p.active_cosmetics, {})
        badges_to_add = ["shopper"]
        if item["kind"] == "protection":
            tier = current_study_tier(p.streak_count)
            p.freeze_capacity = tier["freeze_cap"]
            p.streak_freeze_count = min(int(p.freeze_capacity or 2), int(p.streak_freeze_count or 0) + int(item["value"]))
            badges_to_add.append("freeze_ready")
        elif item["kind"] == "booster":
            booster = {"type": item_id, "multiplier": item["multiplier"], "created_at": utcnow().isoformat()}
            if item.get("uses"):
                booster["uses"] = item["uses"]
            if item.get("hours"):
                booster["expires_at"] = (utcnow() + timedelta(hours=item["hours"])).isoformat()
            p.active_booster = json.dumps(booster)
            badges_to_add.append("booster_pilot")
        elif item["kind"] == "inventory":
            field = item["field"]
            cosmetics[field] = int(cosmetics.get(field, 0) or 0) + int(item["value"])
            p.active_cosmetics = json.dumps(cosmetics)
        elif item["kind"] == "cosmetic":
            owned = set(cosmetics.get("owned", []))
            owned.add(item_id)
            cosmetics["owned"] = sorted(owned)
            cosmetics[item["slot"]] = item["value"]
            p.active_cosmetics = json.dumps(cosmetics)
            badges_to_add.append("style_setter")
        if item_id == weekly_deal_id:
            badges_to_add.append("deal_hunter")
        if int(p.spark_balance or 0) >= 1000:
            badges_to_add.append("spark_saver")
        new_badges = add_badges(p, badges_to_add)
        purchases = safe_json_load(p.shop_purchases, [])
        purchases.append({"id": str(uuid.uuid4()), "item_id": item_id, "price": price, "created_at": utcnow().isoformat()})
        p.shop_purchases = json.dumps(purchases[-200:])
        db.session.commit()
        return flask.jsonify({
            "status": "ok",
            "item": item,
            "item_id": item_id,
            "price": price,
            "spark_balance": p.spark_balance,
            "streak_freeze_count": p.streak_freeze_count,
            "freeze_capacity": p.freeze_capacity,
            "active_booster": safe_json_load(p.active_booster, None),
            "active_cosmetics": safe_json_load(p.active_cosmetics, {}),
            "badges_unlocked": new_badges
        })
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

@app.route("/study/repair", methods=["POST"])
def study_repair():
    uid = current_user.id if current_user.is_authenticated else None
    gid = None if uid else get_guest_session_id()
    try:
        p = get_study_profile(uid, gid)
        cosmetics = safe_json_load(p.active_cosmetics, {})
        broken_streak = int(cosmetics.get("broken_streak_count", p.longest_streak or 0) or 0)
        if not p.repair_eligible_until or p.repair_eligible_until < utcnow():
            return flask.jsonify({"status": "error", "message": "Repair window expired."}), 400
        if p.repair_last_used:
            try:
                last_repair = datetime.strptime(p.repair_last_used[:10], "%Y-%m-%d")
                if last_repair > datetime.now() - timedelta(days=30):
                    return flask.jsonify({"status": "error", "message": "Streak Repair can only be used once every 30 days."}), 400
            except Exception:
                pass
        cost = repair_cost_for(broken_streak)
        used_repair_credit = False
        if int(cosmetics.get("repair_credits", 0) or 0) > 0:
            cost = max(1, int(round(cost * 0.5)))
            cosmetics["repair_credits"] = max(0, int(cosmetics.get("repair_credits", 0) or 0) - 1)
            used_repair_credit = True
        if int(p.spark_balance or 0) < cost:
            return flask.jsonify({"status": "error", "message": "Not enough Sparks to repair this streak.", "cost": cost, "spark_balance": p.spark_balance}), 400
        p.spark_balance = int(p.spark_balance or 0) - cost
        p.streak_count = max(int(p.streak_count or 0), broken_streak)
        p.longest_streak = max(int(p.longest_streak or 0), p.streak_count)
        p.last_active_date = cosmetics.get("broken_last_active_date") or p.last_active_date
        p.repair_last_used = datetime.now().strftime("%Y-%m-%d")
        p.repair_eligible_until = None
        cosmetics.pop("broken_streak_count", None)
        cosmetics.pop("broken_last_active_date", None)
        p.active_cosmetics = json.dumps(cosmetics)
        new_badges = add_badges(p, ["comeback_kid"])
        db.session.commit()
        return flask.jsonify({"status": "ok", "streak_count": p.streak_count, "spark_balance": p.spark_balance, "badges_unlocked": new_badges, "message": "Streak repaired.", "repair_cost": cost, "used_repair_credit": used_repair_credit})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)})

# ── STUDY ACCESS LIMITS ───────────────────────────────────────
GUEST_STUDY_LIMITS = {
    "uploads": 1,
    "generations": 1,
    "max_chars": 6000,
    "max_questions": 5
}

def _get_guest_usage():
    if "guest_study_usage" not in session:
        session["guest_study_usage"] = {"uploads": 0, "generations": 0}
    return session["guest_study_usage"]

def _save_guest_usage(usage):
    session["guest_study_usage"] = usage
    session.modified = True

def _guest_limit_response():
    return flask.jsonify({
        "status": "error",
        "code": "login_required",
        "message": "Create an account to continue using Study & Learn.",
    }), 403

def _is_guest():
    return not current_user.is_authenticated

@app.route("/study/access", methods=["GET"])
def study_access():
    if current_user.is_authenticated:
        return flask.jsonify({"status": "ok", "logged_in": True, "limits": None})
    usage = _get_guest_usage()
    remaining_uploads = max(0, GUEST_STUDY_LIMITS["uploads"] - usage["uploads"])
    remaining_generations = max(0, GUEST_STUDY_LIMITS["generations"] - usage["generations"])
    return flask.jsonify({
        "status": "ok",
        "logged_in": False,
        "limits": {
            "remaining_uploads": remaining_uploads,
            "remaining_generations": remaining_generations,
            "max_questions": GUEST_STUDY_LIMITS["max_questions"]
        }
    })

@app.route("/study/extract-pdf", methods=["POST"])
def study_extract_pdf():
    if "file" not in request.files:
        return flask.jsonify({"status": "error", "message": "No file"}), 400
    f = request.files["file"]
    if not f.filename.lower().endswith(".pdf"):
        return flask.jsonify({"status": "error", "message": "Only PDF files"}), 400
    if _is_guest():
        usage = _get_guest_usage()
        if usage["uploads"] >= GUEST_STUDY_LIMITS["uploads"]:
            return _guest_limit_response()
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(f.read()))
        text = " ".join(page.extract_text() or "" for page in reader.pages)
        if _is_guest():
            usage = _get_guest_usage()
            usage["uploads"] += 1
            _save_guest_usage(usage)
        return flask.jsonify({"status": "ok", "text": text[:15000]})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

@app.route("/study/youtube", methods=["POST"])
def study_youtube():
    """Turn a YouTube link into study text by fetching its caption transcript.
    The existing /study/generate pipeline then turns that text into notes,
    flashcards, and a quiz. Mirrors /study/extract-pdf (returns {status, text})
    and counts against the same guest upload limit."""
    data = request.json or {}
    url = (data.get("url") or "").strip()
    if not url:
        return flask.jsonify({"status": "error", "message": "No URL provided"}), 400
    # Pull the 11-char video id out of any common YouTube URL shape.
    vid = None
    m = re.search(r"(?:v=|youtu\.be/|/shorts/|/embed/|/live/)([A-Za-z0-9_-]{11})", url)
    if m:
        vid = m.group(1)
    elif re.fullmatch(r"[A-Za-z0-9_-]{11}", url):
        vid = url
    if not vid:
        return flask.jsonify({"status": "error", "message": "That doesn't look like a valid YouTube link."}), 400
    if _is_guest():
        usage = _get_guest_usage()
        if usage["uploads"] >= GUEST_STUDY_LIMITS["uploads"]:
            return _guest_limit_response()
    # YouTube blocks transcript requests from datacenter IPs (Railway etc.).
    # Route through a proxy if YOUTUBE_PROXY is set (e.g. a Webshare residential
    # endpoint: http://user:pass@host:port). Works with any HTTP/HTTPS proxy.
    proxy = (os.getenv("YOUTUBE_PROXY") or "").strip()
    proxies = {"http": proxy, "https": proxy} if proxy else None
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        langs = ["en", "en-US", "en-GB"]
        # Support both the <1.0 classmethod API and the 1.x instance API.
        try:
            if proxies:
                segments = YouTubeTranscriptApi.get_transcript(vid, languages=langs, proxies=proxies)
            else:
                segments = YouTubeTranscriptApi.get_transcript(vid, languages=langs)
            text = " ".join((s.get("text") or "") for s in segments)
        except AttributeError:
            if proxies:
                from youtube_transcript_api.proxies import GenericProxyConfig
                api = YouTubeTranscriptApi(proxy_config=GenericProxyConfig(http_url=proxy, https_url=proxy))
            else:
                api = YouTubeTranscriptApi()
            fetched = api.fetch(vid, languages=langs)
            text = " ".join(getattr(snip, "text", "") for snip in fetched)
        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            raise ValueError("empty transcript")
        if _is_guest():
            usage = _get_guest_usage()
            usage["uploads"] += 1
            _save_guest_usage(usage)
        return flask.jsonify({"status": "ok", "text": text[:15000], "video_id": vid})
    except Exception as e:
        print(f"YouTube transcript error for {vid}: {e}")
        return flask.jsonify({
            "status": "error",
            "message": "Couldn't fetch this video's transcript. Make sure the video has captions, or paste the transcript into the Paste Text tab."
        }), 502

_AUDIO_EXTS = (".mp3", ".m4a", ".wav", ".webm", ".ogg", ".oga", ".flac", ".mp4", ".mpeg", ".mpga")
MAX_AUDIO_BYTES = 25 * 1024 * 1024  # Audio upload cap (Gemini / Groq Whisper)

@app.route("/study/transcribe", methods=["POST"])
def study_transcribe():
    """Transcribe a recorded or uploaded lecture/audio file with Gemini (Groq fallback).
    The existing /study/generate pipeline then turns the transcript into notes,
    flashcards, and a quiz. Mirrors /study/extract-pdf (returns {status, text})
    and counts against the same guest upload limit."""
    if "file" not in request.files:
        return flask.jsonify({"status": "error", "message": "No audio file"}), 400
    f = request.files["file"]
    fname = (f.filename or "audio.webm").lower()
    if not fname.endswith(_AUDIO_EXTS):
        return flask.jsonify({"status": "error", "message": "Unsupported audio format. Use mp3, m4a, wav, webm, ogg, or mp4."}), 400
    if _is_guest():
        usage = _get_guest_usage()
        if usage["uploads"] >= GUEST_STUDY_LIMITS["uploads"]:
            return _guest_limit_response()
    try:
        data = f.read()
        if not data:
            return flask.jsonify({"status": "error", "message": "Empty audio file."}), 400
        if len(data) > MAX_AUDIO_BYTES:
            return flask.jsonify({"status": "error", "message": "Audio is too large (max 25 MB). Record a shorter clip or upload a compressed file."}), 413
        text = transcribe_audio(fname, data)
        if not text:
            return flask.jsonify({"status": "error", "message": "Couldn't hear any speech in that audio."}), 422
        if _is_guest():
            usage = _get_guest_usage()
            usage["uploads"] += 1
            _save_guest_usage(usage)
        return flask.jsonify({"status": "ok", "text": text[:15000]})
    except Exception as e:
        print(f"Transcription error: {e}")
        return flask.jsonify({"status": "error", "message": "Transcription temporarily unavailable. Please try again."}), 500

@app.route("/study/generate", methods=["POST"])
def study_generate():
    data = request.json or {}
    content = data.get("content", "").strip()
    mode = data.get("mode", "casual")
    num_questions = int(data.get("num_questions", 8))
    if not content:
        return flask.jsonify({"status": "error", "message": "No content provided"}), 400
    if _is_guest():
        usage = _get_guest_usage()
        if usage["generations"] >= GUEST_STUDY_LIMITS["generations"]:
            return _guest_limit_response()
        mode = "casual"
        num_questions = min(num_questions, GUEST_STUDY_LIMITS["max_questions"])
        content = content[:GUEST_STUDY_LIMITS["max_chars"]]
    else:
        if len(content) > 20000:
            content = content[:20000]
    prompt = f'''You are an expert study assistant. Analyze the following study material and generate exactly {num_questions} study questions.

STUDY MATERIAL:
{content}

Generate a mix of:
- 3-4 recall/definition questions (straightforward facts)
- 2-3 conceptual questions (understanding why/how)
- 2-3 short-answer questions (application or explanation)

Also extract 5-8 key concepts from the material.

Respond ONLY with valid JSON in this exact format:
{{
  "title": "Brief topic title (5 words max)",
  "key_concepts": [
    {{"term": "Term name", "definition": "Clear definition in 1-2 sentences"}}
  ],
  "questions": [
    {{
      "id": 1,
      "type": "recall",
      "question": "Question text here?",
      "answer": "Complete, detailed answer here. Be thorough.",
      "hint": "Optional one-word hint"
    }}
  ]
}}

Question types: "recall", "conceptual", "short-answer"
Make answers comprehensive (2-4 sentences). Make questions specific to the content.
Be accurate, but keep the tone supportive and student-friendly.'''
    try:
        raw = ai_chat(
            [{"role": "user", "content": prompt}],
            tier="standard",
            temperature=0.5,
            max_tokens=1200 if _is_guest() else 3000,
        )
        raw = re.sub(r"```json\s*", "", raw)
        raw = re.sub(r"```", "", raw).strip()
        result = json.loads(raw)
        if _is_guest():
            usage = _get_guest_usage()
            usage["generations"] += 1
            _save_guest_usage(usage)
        return flask.jsonify({"status": "ok", "data": result})
    except Exception as e:
        print(f"Study generate error: {e}")
        return flask.jsonify({
            "status": "error",
            "message": "Study generation temporarily unavailable. Please try again later."
        }), 500

# ── ERROR HANDLERS ────────────────────────────────────────────
def _referrer_is_internal():
    """True iff the request's Referer header points at our own host.
    Used by the 404 handler to decide between rendering the 404 page
    (internal navigation hit a dead link — show it so we can fix it)
    and a soft redirect to /dashboard (someone arrived from outside
    on a URL that doesn't exist — be friendly).
    """
    try:
        ref = (request.referrer or "").strip()
    except Exception:
        return False
    if not ref:
        return False
    try:
        from urllib.parse import urlparse
        ref_host = (urlparse(ref).netloc or "").lower().split(":")[0]
        our_host = (request.host or "").lower().split(":")[0]
    except Exception:
        return False
    if not ref_host or not our_host:
        return False
    # Accept the apex + any subdomain of the same registrable domain so
    # www. → apex, embedded iframe origins, etc. all count as internal.
    if ref_host == our_host:
        return True
    if ref_host.endswith("." + our_host) or our_host.endswith("." + ref_host):
        return True
    return False


@app.errorhandler(404)
def error_404(e):
    # API + extension callers always get JSON, never an HTML redirect —
    # a redirect would corrupt the JSON they're parsing.
    if request.path.startswith("/extension/") or request.path.startswith("/api/"):
        return flask.jsonify({"status": "error", "message": "Not found"}), 404
    # Don't redirect missing static assets (favicon, images, manifests,
    # service-worker probes, sourcemaps) — those need a real 404 so the
    # browser stops asking, and an HTML redirect would just confuse it.
    _path_l = (request.path or "").lower()
    _asset_exts = (".ico", ".png", ".jpg", ".jpeg", ".gif", ".svg",
                   ".webp", ".css", ".js", ".map", ".woff", ".woff2",
                   ".ttf", ".json", ".xml", ".txt", ".webmanifest")
    if (_path_l.startswith("/static/") or _path_l.endswith(_asset_exts)
            or _path_l in ("/favicon.ico", "/robots.txt", "/sitemap.xml")):
        return flask.Response("Not Found", status=404, mimetype="text/plain")
    # If the visitor came from outside intelliplan.tech (typed-in URL,
    # stale Google result, social link, blank Referer), bounce them to
    # the dashboard instead of showing a dead-end 404. Internal broken
    # links still render the 404 page so we can spot and fix them.
    if not _referrer_is_internal():
        return redirect(url_for("dashboard"), code=302)
    try:
        return render_template("error.html", active_page="error", error_code=404, error_id=make_error_id()), 404
    except Exception:
        return flask.Response("<h1>404 Not Found</h1><a href='/'>Home</a>", status=404, mimetype="text/html")

@app.errorhandler(403)
def error_403(e):
    if request.path.startswith("/extension/") or request.path.startswith("/api/"):
        return flask.jsonify({"status": "error", "message": "Forbidden"}), 403
    try:
        return render_template("error.html", active_page="error", error_code=403, error_id=make_error_id()), 403
    except Exception:
        return flask.Response("<h1>403 Forbidden</h1><a href='/'>Home</a>", status=403, mimetype="text/html")

# ═════════════════════════════════════════════════════════════════════
# NEW FEATURE MODULES
# (Grade simulator, Lesson Recorder, Study Groups, Writing Assistant,
#  Math Explainer, Task Extractor)
# Each block is self-contained: routes through ai_provider, light DB use, JSON
# responses. The web pages each new module needs sit in /Main_Project/
# templates/ — see lessons.html, groups.html, writing.html, math.html,
# extractor.html. The upgraded grade modeller lives in grademodel.html.
# ═════════════════════════════════════════════════════════════════════

# ── ADMIN / FEATURE FLAGS ───────────────────────────────────────────
# Hidden admin surface for the project owner. Email-gated so even if
# someone discovers the URL, only the configured account can use it.
ADMIN_EMAILS = {
    e.strip().lower()
    for e in (os.getenv("ADMIN_EMAILS") or "uanirudh0811@gmail.com").split(",")
    if e.strip()
}
ADMIN_PATH = os.getenv("ADMIN_PATH", "/admin-x9k2p7")  # not linked anywhere


def is_admin(user):
    try:
        return bool(user and getattr(user, "is_authenticated", False)
                    and (user.email or "").lower() in ADMIN_EMAILS)
    except Exception:
        return False


def require_admin(fn):
    from functools import wraps as _wraps
    @_wraps(fn)
    def w(*a, **kw):
        if not is_admin(current_user):
            # 404 to hide existence from non-admins. Admin gets a fresh
            # /login redirect if simply not yet signed in.
            if not current_user.is_authenticated:
                return redirect(url_for("login"))
            return render_template("error.html", error_code=404,
                                   error_id="ADMIN-403"), 404
        return fn(*a, **kw)
    return w


# Default flag descriptions seeded into the DB on first admin visit so
# the admin page is useful out-of-the-box. Adding new keys here keeps
# them in sync.
DEFAULT_FLAGS = {
    "lessons":       "Lesson Library (uploads + Whisper summary)",
    "groups":        "Study Groups",
    "live_sessions": "Live study sessions (Jitsi)",
    "writing":       "Writing Assistant",
    "math":          "Math Explainer",
    "extractor":     "Task Extractor",
    "referral":      "Referral program",
    "onboarding":    "First-run onboarding modal",
    "ai_chat":       "Plani chat assistant",
    "streak_v1":     "Task-completion streak (retention experiment)",
    "command_center": "AI Daily Command Center (kill switch — default page)",
    "active_study":  "Active study sessions (timer, focus check-in, feedback loop)",
    "planner_v2":    "Deterministic scheduling engine (kill switch — falls back to the AI path)",
}


def feature_enabled(key):
    """Cheap per-request check. Default True so an empty / missing
    flag row never disables a feature accidentally."""
    try:
        row = FeatureFlag.query.filter_by(key=key).first()
        return True if row is None else bool(row.enabled)
    except Exception:
        return True


@app.context_processor
def inject_admin():
    try:
        return {
            "is_admin": is_admin(current_user) if current_user.is_authenticated else False,
            "feature_enabled": feature_enabled,
            "admin_path": ADMIN_PATH,
        }
    except Exception:
        return {"is_admin": False, "feature_enabled": lambda k: True, "admin_path": ADMIN_PATH}


_FLAG_OVERRIDES = {
    "streak_v1": {"enabled": False, "rollout_percentage": 0},
}

def _seed_default_flags():
    try:
        for k, desc in DEFAULT_FLAGS.items():
            row = FeatureFlag.query.filter_by(key=k).first()
            if not row:
                overrides = _FLAG_OVERRIDES.get(k, {})
                flag = FeatureFlag(key=k, enabled=overrides.get("enabled", True), description=desc)
                if "rollout_percentage" in overrides:
                    flag.rollout_percentage = overrides["rollout_percentage"]
                db.session.add(flag)
        db.session.commit()
    except Exception:
        try: db.session.rollback()
        except Exception: pass


@app.route(ADMIN_PATH, methods=["GET"])
@require_admin
def admin_panel():
    _seed_default_flags()
    flags = FeatureFlag.query.order_by(FeatureFlag.key.asc()).all()
    user_count = 0
    try:
        user_count = User.query.count()
    except Exception:
        pass
    return render_template(
        "admin.html",
        active_page="admin",
        flags=flags,
        user_count=user_count,
        admin_email=(current_user.email if current_user.is_authenticated else ""),
    )


@app.route("/api/admin/flag", methods=["POST"])
@require_admin
def admin_set_flag():
    body = request.get_json(silent=True) or {}
    key = (body.get("key") or "").strip()[:64]
    enabled = bool(body.get("enabled"))
    if not key:
        return flask.jsonify({"status": "error", "message": "key required"}), 400
    row = FeatureFlag.query.filter_by(key=key).first()
    if not row:
        row = FeatureFlag(key=key, enabled=enabled, description=DEFAULT_FLAGS.get(key, ""))
        db.session.add(row)
    else:
        row.enabled = enabled
    if "rollout_percentage" in body:
        row.rollout_percentage = max(0, min(100, int(body["rollout_percentage"])))
    db.session.commit()
    return flask.jsonify({
        "status": "ok", "key": key, "enabled": enabled,
        "rollout_percentage": row.rollout_percentage,
    })


@app.route("/api/admin/indexnow/status", methods=["GET"])
@require_admin
def admin_indexnow_status():
    return flask.jsonify({
        "status": "ok",
        "key": INDEXNOW_KEY,
        "key_file_url": f"{APP_BASE_URL.rstrip('/')}/{INDEXNOW_KEY}.txt",
        "keyLocation": _indexnow_key_location(),
        "host": _indexnow_host(),
        "endpoint": INDEXNOW_ENDPOINT,
        "sitemap_url_count": len(_indexnow_sitemap_urls()),
    })


@app.route("/api/admin/indexnow/submit", methods=["POST", "GET"])
@require_admin
def admin_indexnow_submit():
    if request.method == "GET":
        url = (request.args.get("url") or "").strip()
        if url:
            result = notify_indexnow([url])
        else:
            result = notify_indexnow(_indexnow_sitemap_urls())
        code = 200 if result.get("status") == "ok" else 400
        return flask.jsonify(result), code

    body = request.get_json(silent=True) or {}
    urls = body.get("urls")
    if not urls:
        urls = _indexnow_sitemap_urls()
    elif isinstance(urls, str):
        urls = [urls]
    result = notify_indexnow(urls)
    code = 200 if result.get("status") == "ok" else 400
    return flask.jsonify(result), code



@app.route("/api/admin/smtp-status", methods=["GET"])
@require_admin
def admin_smtp_status():
    """Show which SMTP env vars are detected — values masked, just presence/absence."""
    host, port, user, pw, sender = _smtp_config()
    return flask.jsonify({
        "host": host or None,
        "port": port,
        "user_set": bool(user),
        "password_set": bool(pw),
        "sender": sender,
        "ready": bool(host and user and pw),
        "vars_checked": {
            "host": ["SMTP_HOST", "MAIL_HOST", "EMAIL_HOST", "(gmail auto-detect)"],
            "user": ["SMTP_USER", "SMTP_USERNAME", "MAIL_USERNAME", "EMAIL_USERNAME", "USERNAME"],
            "password": ["SMTP_PASSWORD", "MAIL_PASSWORD", "EMAIL_PASSWORD", "PASSWORD"],
            "port": ["SMTP_PORT", "MAIL_PORT", "(default 587)"],
        },
    })


@app.route("/api/admin/sms-blast/preview", methods=["POST"])
@require_admin
def admin_sms_blast_preview():
    """Return recipient count and sample emails for the chosen audience — no SMS sent."""
    body = request.get_json(silent=True) or {}
    audience = (body.get("audience") or "sms_opted_in").strip()
    specific_emails = [e.strip().lower() for e in (body.get("emails") or "").split(",") if e.strip()]

    try:
        q = _admin_sms_audience_query(audience, specific_emails)
        users = q.all()
        return flask.jsonify({
            "status": "ok",
            "count": len(users),
            "sample": [u.email for u in users[:10]],
        })
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


@app.route("/api/admin/sms-blast/send", methods=["POST"])
@require_admin
def admin_sms_blast_send():
    """Send an SMS to every user in the selected audience.
    Returns per-recipient results. Runs synchronously — keep audience small
    or use the cron pipeline for large sends."""
    body = request.get_json(silent=True) or {}
    message = (body.get("message") or "").strip()
    audience = (body.get("audience") or "sms_opted_in").strip()
    specific_emails = [e.strip().lower() for e in (body.get("emails") or "").split(",") if e.strip()]

    if not message:
        return flask.jsonify({"status": "error", "message": "message is required"}), 400
    if len(message) > 160:
        return flask.jsonify({"status": "error", "message": "message must be 160 characters or fewer"}), 400

    try:
        q = _admin_sms_audience_query(audience, specific_emails)
        users = q.all()
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

    sent, failed, skipped = 0, 0, 0
    results = []
    for u in users:
        if not u.phone:
            skipped += 1
            results.append({"email": u.email, "result": "skipped_no_phone"})
            continue
        ok, _err = _sms_send_for_user(u, message)
        if ok:
            sent += 1
            results.append({"email": u.email, "result": "sent"})
        else:
            failed += 1
            results.append({"email": u.email, "result": "failed"})

    print(f"[admin-sms-blast] audience={audience} total={len(users)} sent={sent} failed={failed} skipped={skipped}")
    return flask.jsonify({
        "status": "ok",
        "total": len(users),
        "sent": sent,
        "failed": failed,
        "skipped": skipped,
        "results": results,
    })


def _admin_sms_audience_query(audience, specific_emails=None):
    """Return a SQLAlchemy query for users matching the given audience key."""
    base = User.query.filter(User.phone.isnot(None), User.phone != "")

    if audience == "specific":
        if not specific_emails:
            raise ValueError("No emails provided for specific audience")
        return base.filter(User.email.in_(specific_emails))
    elif audience == "all_with_phone":
        return base
    else:
        return base.filter(User.sms_reminders_opt_in.is_(True))


@app.route("/api/admin/newsletter/preview", methods=["POST"])
@require_admin
def admin_newsletter_preview():
    """Dry run: how many people would get this, and what does it look like.

    Sends nothing. This is the step that must happen before the real send —
    a newsletter is not undoable, and "how many" is the one number worth
    checking twice.
    """
    body = request.get_json(silent=True) or {}
    from intelliplan.email import campaigns

    try:
        summary = campaigns.send_newsletter(body, test=False, dry_run=True)
        return flask.jsonify({"status": "ok", "summary": summary})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


@app.route("/api/admin/newsletter/send", methods=["POST"])
@require_admin
def admin_newsletter_send():
    """Send the newsletter for real, or to the admins only.

    ``test: true``  → only ADMIN_EMAILS addresses, and does not write the
                      deduplication ledger so it can be run repeatedly.
    ``confirm: true`` → required for a full-list send. Without it this
                      refuses and returns the dry-run count instead, so the
                      irreversible action is never the default.
    """
    body = request.get_json(silent=True) or {}
    test = bool(body.get("test"))
    confirm = bool(body.get("confirm"))

    from intelliplan.email import campaigns

    if not test and not confirm:
        try:
            summary = campaigns.send_newsletter(body, test=False, dry_run=True)
        except Exception as e:
            return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500
        return flask.jsonify({
            "status": "confirm_required",
            "message": (
                f"This would email {summary['recipients']} people. "
                "Re-send with \"confirm\": true to go ahead."
            ),
            "summary": summary,
        }), 409

    try:
        summary = campaigns.send_newsletter(body, test=test, dry_run=False)
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

    if summary.get("error"):
        return flask.jsonify({"status": "error", "message": summary["error"], "summary": summary}), 400

    print(
        f"[admin-newsletter] key={summary['email_key']} test={test} "
        f"recipients={summary['recipients']} sent={summary['sent']} failed={summary['failed']}"
    )
    return flask.jsonify({"status": "ok", "summary": summary})


# ── REFERRAL TRACKING ──────────────────────────────────────────────


def _ensure_referral_code(user):
    if user.referral_code:
        return user.referral_code
    for _ in range(5):
        code = secrets_module.token_urlsafe(6).replace("-", "").replace("_", "")[:8].lower()
        if not User.query.filter_by(referral_code=code).first():
            user.referral_code = code
            db.session.commit()
            return code
    return None


@app.route("/ref/<code>")
def referral_landing(code):
    code = (code or "").strip().lower()[:16]
    if code:
        ref_user = User.query.filter_by(referral_code=code).first()
        if ref_user:
            session["pending_referral"] = ref_user.id
            session.modified = True
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    return redirect(url_for("register"))


@app.route("/api/referral", methods=["GET"])
def api_referral():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "message": "login required"}), 401
    code = _ensure_referral_code(current_user)
    invited = User.query.filter_by(referred_by_id=current_user.id).count()
    return flask.jsonify({
        "status": "ok",
        "code": code,
        "invited_count": invited,
    })


# ── PHONE + SMS REMINDERS ─────────────────────────────────────────
import re as _re_phone


def _send_email_via_resend(to_addr, subject, body, html=None, headers=None, reply_to=None):
    """Send an email through the Resend HTTP API.
    Returns True on success, False on failure or if not configured.

    ``html`` and ``headers`` are optional and only added to the payload when
    truthy — Resend rejects a null ``html`` field, and an empty ``headers``
    object is noise in the request log. Callers that want a plain-text send
    keep the original three-argument shape and get the original request.
    """
    api_key = os.getenv("RESEND_API_KEY")
    if not api_key or not to_addr:
        return False
    from_addr = os.getenv("RESEND_FROM") or "IntelliPlan <noreply@intelliplan.tech>"
    try:
        import urllib.request, urllib.error, json as _json
        body_payload = {
            "from": from_addr,
            "to": [to_addr],
            "subject": subject or "",
            "text": body or "",
        }
        if html:
            body_payload["html"] = html
        if headers:
            # List-Unsubscribe and friends. Resend passes these through to
            # the outgoing message verbatim.
            body_payload["headers"] = dict(headers)
        if reply_to:
            # Resend's own field rather than a Reply-To in `headers` — it
            # rejects some standard headers supplied that way. Every
            # lifecycle email invites a reply, and the From address is a
            # no-reply, so without this the answer goes nowhere.
            body_payload["reply_to"] = reply_to
        payload = _json.dumps(body_payload).encode()
        req = urllib.request.Request(
            "https://api.resend.com/emails",
            data=payload,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
        )
        with urllib.request.urlopen(req, timeout=10) as resp:
            result = _json.loads(resp.read())
            message_id = result.get("id")
            print(f"[email-resend] sent to {to_addr}, id={message_id}")
            # The provider id, when there is one, so the caller can record it
            # against the send. Falls back to True so this stays a truthy
            # success for every existing `if _send_email_via_resend(...)`.
            return message_id or True
    except urllib.error.HTTPError as e:
        err = e.read().decode(errors="replace") if hasattr(e, "read") else str(e)
        print(f"[email-resend] failed to {to_addr}: {e.code} {err}")
        return False
    except Exception as e:
        print(f"[email-resend] error: {e}")
        return False


def _send_email(to_addr, subject, body, html=None, headers=None, reply_to=None):
    """Send an email. Tries Resend first (RESEND_API_KEY), then falls back to
    SMTP. If neither is configured, logs the message so the parental-consent
    link can still be delivered manually.

    ``body`` is always the plain-text part and is always required — an
    HTML-only email is a deliverability problem and unreadable in a text
    client. ``html``, when given, is sent as an alternative alongside it.

    Returns a truthy value on success: the provider's message id when Resend
    reports one, otherwise True. The three-positional-argument call
    ``_send_email(addr, subject, body)`` behaves exactly as it did before —
    ``notifications_glue`` depends on that.
    """
    if not to_addr:
        return False
    sent = _send_email_via_resend(
        to_addr, subject, body, html=html, headers=headers, reply_to=reply_to
    )
    if sent:
        return sent

    host, port, user, pw, sender = _smtp_config()
    if not host:
        print(f"[email] neither Resend nor SMTP configured — would send to {to_addr}: {subject}\n{body}")
        return False
    try:
        import smtplib
        from email.message import EmailMessage
        msg = EmailMessage()
        msg["From"] = sender
        msg["To"] = to_addr
        msg["Subject"] = subject
        if reply_to:
            msg["Reply-To"] = reply_to
        for key, value in (headers or {}).items():
            # List-Unsubscribe et al. Assigning over an existing key raises,
            # so drop anything already set above rather than duplicating it.
            if key not in msg:
                msg[key] = value
        msg.set_content(body)
        if html:
            # set_content + add_alternative is what turns this into a
            # multipart/alternative: text first, HTML second, and a text-only
            # client renders the part it understands.
            msg.add_alternative(html, subtype="html")
        clean_pw = "".join(pw.split()) if pw else ""
        with smtplib.SMTP(host, port, timeout=10) as s:
            s.starttls()
            if user and clean_pw:
                s.login(user, clean_pw)
            s.send_message(msg)
        print(f"[email-smtp] sent to {to_addr}")
        return True
    except Exception as e:
        print(f"[email-smtp] send failed: {e}")
        return False


def _mini_page(title: str, body_html: str) -> str:
    """A standalone one-off HTML page, on-brand without a template.

    The COPPA consent flow is often a parent's only visual contact with
    IntelliPlan, and it used to render four copies of the same Arial-on-cool-
    grey stylesheet. This is one place, using the same fonts and warm palette
    as the rest of the site, with a dark-mode block so it does not glare at
    night. Deliberately inline: these responses must survive a half-migrated
    database that could break template rendering.
    """
    return (
        "<!doctype html><html lang='en'><head><meta charset='utf-8'>"
        "<meta name='viewport' content='width=device-width, initial-scale=1'>"
        "<meta name='robots' content='noindex'>"
        f"<title>{title} | IntelliPlan</title>"
        "<link rel='icon' type='image/png' sizes='192x192' href='/static/icons/icon-192.png'>"
        "<link href='https://fonts.googleapis.com/css2?family=DM+Serif+Display"
        "&family=DM+Sans:wght@400;500;600&display=swap' rel='stylesheet'>"
        "<style>"
        ":root{--bg:#f5f4f1;--card:#fff;--ink:#1a1a1a;--muted:#4a4a46;"
        "--accent:#1a56db;--line:rgba(0,0,0,0.08)}"
        "@media(prefers-color-scheme:dark){:root{--bg:#101012;--card:#1c1c20;"
        "--ink:#e6e6e2;--muted:#9e9e9a;--accent:#5b93f5;--line:rgba(255,255,255,0.08)}}"
        "*{box-sizing:border-box}"
        "body{margin:0;min-height:100dvh;display:flex;align-items:center;justify-content:center;"
        "padding:2rem 1rem;background:var(--bg);color:var(--ink);line-height:1.6;"
        "font-family:'DM Sans',system-ui,-apple-system,sans-serif;-webkit-font-smoothing:antialiased}"
        "main{max-width:560px;width:100%;padding:2rem;background:var(--card);"
        "border:1px solid var(--line);border-radius:20px;"
        "box-shadow:0 4px 12px rgba(28,25,20,0.06),0 24px 56px rgba(28,25,20,0.10)}"
        "h1{margin:0 0 .75rem;font-family:'DM Serif Display',Georgia,serif;font-weight:400;"
        "font-size:1.75rem;letter-spacing:-0.02em;text-wrap:balance}"
        "p{margin:0 0 1rem;color:var(--muted);max-width:52ch;text-wrap:pretty}"
        "p:last-child{margin-bottom:0}"
        "strong{color:var(--ink)}"
        "a{color:var(--accent);text-decoration:none}a:hover{text-decoration:underline}"
        "a:focus-visible{outline:2px solid var(--accent);outline-offset:2px;border-radius:4px}"
        "</style></head><body><main>"
        f"{body_html}"
        "</main></body></html>"
    )


@app.route("/parent/consent")
def parent_consent():
    """Public landing page for the COPPA parental-consent link."""
    token = request.args.get("token", "").strip()
    if not token:
        return "Missing consent token.", 400
    user = User.query.filter_by(parent_consent_token=token).first()
    if not user:
        # Was rendering error.html with error_title/error_message, which that
        # template does not read — it takes error_code/message. The result was
        # a parent with a stale link seeing "500 · Something went wrong",
        # served with a 200. Now it says what actually happened, with a 404.
        return _mini_page(
            "Link not valid",
            "<h1>Link not valid</h1>"
            "<p>This consent link no longer works. Either the account was already approved, "
            "or the link has expired.</p>"
            "<p>If your child still needs approval, ask them to sign up again and you will "
            "get a fresh consent email.</p>"
            "<p><a href='/'>Back to IntelliPlan</a></p>"
        ), 404
    if not user.parent_consent_granted:
        user.parent_consent_granted = True
        user.parent_consent_token = None  # one-shot
        db.session.commit()
    # Render a tiny inline confirmation — no template needed.
    return (
        _mini_page(
            "Consent granted",
            "<h1>Consent granted</h1>"
            f"<p>You have approved <strong>{user.email}</strong>'s IntelliPlan account. "
            "They can sign in and start using the app now.</p>"
            "<p>You can revoke consent at any time by emailing "
            "<a href='mailto:uanirudh0811@gmail.com'>uanirudh0811@gmail.com</a>. We delete the "
            "account and all associated data within 30 days.</p>"
            "<p><a href='/'>Back to IntelliPlan</a></p>"
        )
    )


@app.route("/parent/deny")
def parent_deny():
    """COPPA deny path. Hard-deletes the pending under-13 account so a
    rejected child can't sign in and we hold no PII on them. One-shot:
    the consent token is the only handle to the row, so once the
    account is deleted the link can't be replayed.
    """
    token = request.args.get("token", "").strip()
    if not token:
        return "Missing consent token.", 400
    user = User.query.filter_by(parent_consent_token=token).first()
    if not user:
        return (
            _mini_page(
                "Link not valid",
                "<h1>Link not valid</h1>"
                "<p>This link no longer works. The account was probably already approved, "
                "or it has been removed.</p>"
                "<p><a href='/'>Back to IntelliPlan</a></p>"
            )
        ), 404
    # Refuse to delete an account that's already been activated — at that
    # point consent has been granted and removal needs to go through the
    # account-deletion flow under the child's logged-in session.
    if user.parent_consent_granted:
        return (
            _mini_page(
                "Already approved",
                "<h1>Already approved</h1>"
                "<p>This account has already been approved, so it cannot be removed from "
                "this link.</p>"
                f"<p>To remove <strong>{user.email}</strong>, email "
                "<a href='mailto:uanirudh0811@gmail.com'>uanirudh0811@gmail.com</a>. We delete "
                "the account and all associated data within 30 days, as COPPA requires.</p>"
            )
        ), 409
    child_email = user.email
    try:
        db.session.delete(user)
        db.session.commit()
        print(f"[coppa] denied + deleted pending account: {child_email}")
    except Exception as _e:
        db.session.rollback()
        print(f"[coppa] deny delete failed for {child_email}: {_e}")
        return "Could not remove the account right now. Please try again later.", 500
    return (
        _mini_page(
            "Account removed",
            "<h1>Account removed</h1>"
            f"<p>The pending IntelliPlan account for <strong>{child_email}</strong> has been "
            "deleted. We keep no data from it.</p>"
            "<p>If that was a mistake, your child can sign up again at any time and you will "
            "get a new consent email.</p>"
            "<p><a href='/'>Back to IntelliPlan</a></p>"
        )
    )


def _normalise_phone(raw):
    """Return E.164-ish format like +15551234567, or empty string on bad input.
    Reject anything that isn't 10-15 digits (E.164 max). Strips junk like
    extensions ('ext 999') so a half-typed number doesn't get stored as garbage.
    """
    if not raw:
        return ""
    s = _re_phone.sub(r"[^0-9+]", "", raw)
    if not s:
        return ""
    # Collapse any stray '+' chars to a single leading one.
    leading_plus = s.startswith("+")
    digits = _re_phone.sub(r"[^0-9]", "", s)
    if not digits:
        return ""
    # Default-to-US: a bare 10-digit number becomes +1XXXXXXXXXX.
    if not leading_plus:
        if len(digits) == 10:
            digits = "1" + digits
        elif len(digits) < 7 or len(digits) > 15:
            return ""  # too short or too long to be a real phone number
    else:
        if len(digits) < 7 or len(digits) > 15:
            return ""
    return "+" + digits


# ── SMS over carrier email-to-SMS gateways ────────────────────────
# IntelliPlan does NOT use Twilio. Instead, we send the reminder body
# as a plain-text email to the recipient's carrier gateway address
# (e.g. 5551234567@tmomail.net for T-Mobile). Each user picks their
# carrier in Settings; T-Mobile is the default.
SMS_CARRIER_GATEWAYS = {
    "tmobile":   "tmomail.net",                # T-Mobile
    "att":       "txt.att.net",                # AT&T
    "verizon":   "vtext.com",                  # Verizon
    "sprint":    "messaging.sprintpcs.com",    # Sprint (legacy)
    "uscellular":"email.uscc.net",             # US Cellular
    "cricket":   "sms.cricketwireless.net",    # Cricket
    "metropcs":  "mymetropcs.com",             # Metro by T-Mobile
    "boost":     "sms.myboostmobile.com",      # Boost Mobile
    "googlefi":  "msg.fi.google.com",          # Google Fi
}


def _digits_only(s):
    """Return the bare digits of a phone string — strips +, spaces, etc."""
    return _re_phone.sub(r"[^0-9]", "", s or "")


def _smtp_config():
    """Resolve SMTP credentials from environment variables.

    Checks several name variants so Railway vars named HOST / USERNAME /
    PASSWORD work alongside the canonical SMTP_* names. Falls back to
    smtp.gmail.com automatically when the username is a Gmail address and
    no explicit host is set.
    """
    host = (os.getenv("SMTP_HOST")
            or os.getenv("MAIL_HOST")
            or os.getenv("EMAIL_HOST"))
    port = int(os.getenv("SMTP_PORT") or os.getenv("MAIL_PORT") or "587")
    user = (os.getenv("SMTP_USER")
            or os.getenv("SMTP_USERNAME")
            or os.getenv("MAIL_USERNAME")
            or os.getenv("EMAIL_USERNAME")
            or os.getenv("USERNAME"))
    pw = (os.getenv("SMTP_PASSWORD")
          or os.getenv("MAIL_PASSWORD")
          or os.getenv("EMAIL_PASSWORD")
          or os.getenv("PASSWORD"))
    sender = os.getenv("SMTP_FROM") or user or "no-reply@intelliplan.tech"

    # Auto-detect Gmail when no host is given but the address is @gmail.com
    if not host and user and "@gmail.com" in user.lower():
        host = "smtp.gmail.com"
        port = 587

    print(f"[smtp-config] host={'set' if host else 'MISSING'} "
          f"user={'set' if user else 'MISSING'} "
          f"pw={'set' if pw else 'MISSING'} port={port}")
    return host, port, user, pw, sender


def _sms_send_email_gateway(to_phone, body, carrier="tmobile"):
    """Send a short SMS by emailing the carrier's gateway address.
    Returns (True, None) on success or (None, error_str) on failure.

    Tries Resend API first (set RESEND_API_KEY in Railway), then falls
    back to SMTP. Resend is recommended — no App Password headaches.
    """
    digits = _digits_only(to_phone)
    if len(digits) == 11 and digits.startswith("1"):
        digits = digits[1:]
    if len(digits) != 10:
        print(f"[sms-email] refusing to send — number must be 10 digits, got {digits!r}")
        return None, f"phone must be 10 digits, got {len(digits)}"

    gateway = SMS_CARRIER_GATEWAYS.get((carrier or "tmobile").lower())
    if not gateway:
        print(f"[sms-email] unknown carrier {carrier!r} — falling back to T-Mobile")
        gateway = SMS_CARRIER_GATEWAYS["tmobile"]

    sms_recipient = f"{digits}@{gateway}"
    sms_message = (body or "")[:300]

    # ── Resend API (preferred) ────────────────────────────────────────
    resend_key = os.getenv("RESEND_API_KEY")
    if resend_key:
        return _sms_via_resend(resend_key, sms_recipient, sms_message)

    # ── SMTP fallback ─────────────────────────────────────────────────
    smtp_host, smtp_port, smtp_username, smtp_password, smtp_from = _smtp_config()
    if not smtp_host:
        print(f"[sms-email] no sender configured — would email {sms_recipient}: {sms_message}")
        return None, "No email sender configured (set RESEND_API_KEY in Railway)"

    try:
        import smtplib
        from email.message import EmailMessage
        msg = EmailMessage()
        msg["From"] = smtp_from
        msg["To"] = sms_recipient
        msg["Subject"] = ""
        msg.set_content(sms_message)
        clean_pw = "".join(smtp_password.split()) if smtp_password else ""
        print(f"[sms-email] SMTP login as {smtp_username}, pw_len={len(clean_pw)}")
        with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as s:
            s.starttls()
            if smtp_username and clean_pw:
                s.login(smtp_username, clean_pw)
            s.send_message(msg)
        print(f"[sms-email] sent via SMTP to {sms_recipient}")
        return True, None
    except Exception as e:
        print(f"[sms-email] SMTP send to {sms_recipient} failed: {e}")
        return None, str(e)


def _sms_via_resend(api_key, to_addr, text):
    """Send a plain-text email via the Resend HTTP API."""
    import urllib.request, json as _json
    from_addr = os.getenv("RESEND_FROM") or "IntelliPlan <noreply@intelliplan.tech>"
    payload = _json.dumps({
        "from": from_addr,
        "to": [to_addr],
        "subject": "",
        "text": text,
    }).encode()
    req = urllib.request.Request(
        "https://api.resend.com/emails",
        data=payload,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            result = _json.loads(resp.read())
            print(f"[sms-resend] sent to {to_addr}, id={result.get('id')}")
            return True, None
    except urllib.error.HTTPError as e:
        err = e.read().decode(errors="replace")
        print(f"[sms-resend] failed to {to_addr}: {e.code} {err}")
        return None, f"Resend {e.code}: {err}"
    except Exception as e:
        print(f"[sms-resend] error: {e}")
        return None, str(e)


def _sms_send_for_user(user, body):
    """Send an SMS using the carrier email-to-SMS gateway.
    IntelliPlan does NOT use a third-party SMS API (no Twilio etc.) — we
    deliver SMS by emailing the carrier's gateway address (the carrier
    converts the email to an SMS at no cost to us).
    Returns (True, None) on success or (None, error_str) on failure."""
    if not user or not user.phone:
        return None, "no phone on account"
    carrier = (getattr(user, "sms_carrier", None) or "tmobile").lower()
    return _sms_send_email_gateway(user.phone, body, carrier=carrier)


def _twilio_send(to_phone, body):
    """Compatibility shim — older call sites used this name. Routes to
    the carrier email-to-SMS gateway (T-Mobile default) since IntelliPlan
    no longer uses Twilio."""
    return _sms_send_email_gateway(to_phone, body, carrier="tmobile")


def _profile_payload(u):
    return {
        "status": "ok",
        "phone": u.phone or "",
        "sms_reminders_opt_in": bool(getattr(u, "sms_reminders_opt_in", False)),
        "push_reminders_opt_in": bool(getattr(u, "push_reminders_opt_in", False)),
        "reminder_lead_minutes": int(getattr(u, "reminder_lead_minutes", 60) or 60),
        "sms_carrier": (getattr(u, "sms_carrier", "tmobile") or "tmobile"),
        "sms_carrier_choices": list(SMS_CARRIER_GATEWAYS.keys()),
        "ai_personalization_opt_in": bool(getattr(u, "ai_personalization_opt_in", False)),
    }


# ── Focus enforcement (Active study) ─────────────────────────────────
#
# The camera check-in already knows when a student has drifted; until now
# the only response was a dismissible line of text, which is easy to
# ignore and therefore does not do the job it exists to do. These endpoints
# store what should actually happen instead.
#
# Three escalations, chosen by the student — never imposed:
#
#   alarm     a sound at full volume, theirs if they upload one
#   takeover  the screen fills with something impossible to read past
#   stakes    the session forfeits the sparks it has earned so far
#
# "off" keeps the old gentle nudge and stays the default.

FOCUS_ENFORCEMENT_MODES = ("off", "alarm", "takeover", "stakes")

#: Grace period bounds, in seconds. Below ~10s the alarm fires when someone
#: reaches for a textbook; past two minutes it is no longer enforcing
#: anything.
FOCUS_GRACE_BOUNDS = (10, 120)

ALARM_UPLOAD_FOLDER = os.path.join(app.root_path, "uploads", "focus_alarms")
os.makedirs(ALARM_UPLOAD_FOLDER, exist_ok=True)

#: Audio only, and only formats a browser will actually play. The extension
#: is checked against this AND the magic bytes are sniffed below — an
#: extension is a claim by the uploader, not evidence.
ALARM_ALLOWED_EXTENSIONS = {".mp3", ".wav", ".ogg", ".m4a", ".aac"}
ALARM_MAX_BYTES = 2 * 1024 * 1024      # 2 MB — an alarm, not an album

#: Leading bytes for the container formats above. Anything that does not
#: start like audio is rejected regardless of what it is named, so a
#: renamed script cannot be parked in the uploads directory.
_ALARM_MAGIC = (
    b"ID3",              # mp3 with an ID3 tag
    b"\xff\xfb", b"\xff\xf3", b"\xff\xf2",   # bare mp3 frame syncs
    b"RIFF",             # wav
    b"OggS",             # ogg
    b"fLaC",
    b"\xff\xf1",         # aac (ADTS)
)


def _looks_like_audio(head: bytes) -> bool:
    if any(head.startswith(sig) for sig in _ALARM_MAGIC):
        return True
    # m4a/mp4: "ftyp" at offset 4.
    return len(head) >= 12 and head[4:8] == b"ftyp"


def _focus_settings_payload(u):
    return {
        "status": "ok",
        "mode": (getattr(u, "focus_enforcement", "off") or "off"),
        "grace_seconds": int(getattr(u, "focus_grace_seconds", 25) or 25),
        "alarm_url": (
            flask.url_for("focus_alarm_file", user_id=u.id)
            if getattr(u, "focus_alarm_file", None) else ""
        ),
        "modes": list(FOCUS_ENFORCEMENT_MODES),
        "grace_bounds": list(FOCUS_GRACE_BOUNDS),
    }


@app.route("/api/focus/enforcement", methods=["GET", "POST"])
def api_focus_enforcement():
    """Read or write how Active study responds to a detected distraction."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "error": "login required"}), 401

    if request.method == "POST":
        body = request.get_json(silent=True) or {}

        mode = str(body.get("mode") or "").strip().lower()
        if mode:
            if mode not in FOCUS_ENFORCEMENT_MODES:
                return flask.jsonify({
                    "status": "error",
                    "message": "Unknown enforcement mode.",
                }), 400
            current_user.focus_enforcement = mode

        if "grace_seconds" in body:
            try:
                lo, hi = FOCUS_GRACE_BOUNDS
                current_user.focus_grace_seconds = max(lo, min(hi, int(body["grace_seconds"])))
            except (TypeError, ValueError):
                return flask.jsonify({
                    "status": "error",
                    "message": "Grace period must be a number of seconds.",
                }), 400

        db.session.commit()

    return flask.jsonify(_focus_settings_payload(current_user))


@app.route("/api/focus/alarm", methods=["POST", "DELETE"])
def api_focus_alarm():
    """Upload or remove the student's own alarm sound."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "error": "login required"}), 401

    stored_name = f"user_{current_user.id}"

    def _existing():
        name = getattr(current_user, "focus_alarm_file", None)
        return os.path.join(ALARM_UPLOAD_FOLDER, name) if name else None

    if request.method == "DELETE":
        path = _existing()
        if path and os.path.isfile(path):
            try:
                os.remove(path)
            except OSError:
                pass
        current_user.focus_alarm_file = None
        db.session.commit()
        return flask.jsonify(_focus_settings_payload(current_user))

    upload = request.files.get("sound")
    if not upload or not upload.filename:
        return flask.jsonify({"status": "error", "message": "No file was sent."}), 400

    ext = os.path.splitext(secure_filename(upload.filename))[1].lower()
    if ext not in ALARM_ALLOWED_EXTENSIONS:
        return flask.jsonify({
            "status": "error",
            "message": "Use an MP3, WAV, OGG, M4A or AAC file.",
        }), 400

    # Read with a hard ceiling rather than trusting Content-Length, which is
    # supplied by the client. One byte over the limit is enough to know.
    blob = upload.read(ALARM_MAX_BYTES + 1)
    if len(blob) > ALARM_MAX_BYTES:
        return flask.jsonify({
            "status": "error",
            "message": "That file is over 2 MB. Pick something shorter.",
        }), 400
    if not _looks_like_audio(blob[:16]):
        return flask.jsonify({
            "status": "error",
            "message": "That file does not look like audio.",
        }), 400

    # One file per user, named by id — an uploaded filename never reaches
    # the filesystem, so there is nothing to traverse with.
    old = _existing()
    if old and os.path.isfile(old) and os.path.basename(old) != stored_name + ext:
        try:
            os.remove(old)
        except OSError:
            pass

    with open(os.path.join(ALARM_UPLOAD_FOLDER, stored_name + ext), "wb") as fh:
        fh.write(blob)
    current_user.focus_alarm_file = stored_name + ext
    db.session.commit()
    return flask.jsonify(_focus_settings_payload(current_user))


@app.route("/uploads/focus-alarm/<int:user_id>")
def focus_alarm_file(user_id):
    """Serve a student their own alarm sound.

    Scoped to the owner deliberately. These are arbitrary user uploads, and
    the id in the URL is guessable, so without this check the route would
    hand any logged-in user any other user's uploaded file.
    """
    if not current_user.is_authenticated or current_user.id != user_id:
        flask.abort(404)
    name = getattr(current_user, "focus_alarm_file", None)
    if not name:
        flask.abort(404)
    return flask.send_from_directory(
        ALARM_UPLOAD_FOLDER, name,
        # Never let an upload be rendered as anything but a download/media
        # source, whatever its bytes turn out to contain.
        mimetype="application/octet-stream",
    )


@app.route("/api/profile/ai_personalization", methods=["POST"])
def api_profile_ai_personalization():
    """Toggle the AI personalization opt-in. Body: {"enabled": true|false}.

    When enabled, the scheduler/tutor/etc inject the student's grade history
    and identity into AI prompts. When disabled (the default), none of that
    leaves the database.
    """
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    body = request.get_json(silent=True) or {}
    current_user.ai_personalization_opt_in = bool(body.get("enabled"))
    db.session.commit()
    return flask.jsonify({
        "status": "ok",
        "ai_personalization_opt_in": current_user.ai_personalization_opt_in,
    })


@app.route("/api/profile/phone", methods=["GET", "POST"])
def api_profile_phone():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    if request.method == "GET":
        return flask.jsonify(_profile_payload(current_user))
    body = request.get_json(silent=True) or {}
    if "phone" in body:
        normalised = _normalise_phone(body.get("phone") or "")
        current_user.phone = normalised or None
    if "sms_reminders_opt_in" in body:
        current_user.sms_reminders_opt_in = bool(body.get("sms_reminders_opt_in"))
        # Clearing the phone means you can't be reminded — also clear the flag.
        if not current_user.phone:
            current_user.sms_reminders_opt_in = False
    if "push_reminders_opt_in" in body:
        current_user.push_reminders_opt_in = bool(body.get("push_reminders_opt_in"))
    if "reminder_lead_minutes" in body:
        try:
            lead = int(body.get("reminder_lead_minutes") or 60)
            # Clamp to 5 minutes .. 7 days
            current_user.reminder_lead_minutes = max(5, min(lead, 10080))
        except (TypeError, ValueError):
            pass
    if "sms_carrier" in body:
        carrier = (body.get("sms_carrier") or "tmobile").strip().lower()
        # Silently fall back to T-Mobile for unknown values.
        if carrier not in SMS_CARRIER_GATEWAYS:
            carrier = "tmobile"
        current_user.sms_carrier = carrier
    db.session.commit()
    return flask.jsonify(_profile_payload(current_user))


@app.route("/api/profile/phone/test", methods=["POST"])
def api_profile_phone_test():
    """Admin / opt-in user helper: send a one-line test SMS."""
    if not current_user.is_authenticated or not current_user.phone:
        return flask.jsonify({"status": "error", "message": "phone not set"}), 400
    ok, err = _sms_send_for_user(current_user, "IntelliPlan reminders are active for this number.")
    return flask.jsonify({"status": "ok" if ok else "error",
                          "message": "Sent! Check your phone in a moment." if ok else (err or "Send failed")})


def _iso_utc(dt):
    """Normalize a datetime-ish value to an ISO-8601 UTC string ending in Z.
    Returns None if conversion fails.
    """
    if not dt:
        return None
    try:
        if isinstance(dt, str):
            s = dt.strip()
            if not s:
                return None
            if s.endswith("Z"):
                return s
            # Date-only ("2026-05-23") → start of day UTC
            if len(s) == 10 and s.count("-") == 2:
                return s + "T00:00:00Z"
            # Add Z if it parses without a timezone
            try:
                parsed = datetime.fromisoformat(s.replace("Z", "+00:00"))
                if parsed.tzinfo is None:
                    parsed = parsed.replace(tzinfo=timezone.utc)
                return parsed.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
            except Exception:
                return s
        if isinstance(dt, datetime):
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            return dt.astimezone(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    except Exception:
        return None
    return None


def _classify_event_kind(title, description=""):
    """Heuristic: pick "class" | "study" | "exam" | "deadline" from text."""
    blob = ((title or "") + " " + (description or "")).lower()
    if any(w in blob for w in ("exam", "test", "midterm", "final", "quiz")):
        return "exam"
    if any(w in blob for w in ("due", "deadline", "submit", "assignment")):
        return "deadline"
    if any(w in blob for w in ("study", "review", "flashcard", "practice")):
        return "study"
    return "class"


def _build_lotus_snapshot(user):
    """Build the rich Lotus-compatible snapshot payload for an authenticated user.

    Shape matches the Lotus iframe contract:
      nextEvent | deadlines[] | reviewQueue | streak | todayGoal
    Each top-level field may be null when data is unavailable.
    """
    base = APP_BASE_URL.rstrip("/")
    now_utc = datetime.now(timezone.utc)
    today_str = now_utc.strftime("%Y-%m-%d")

    payload = {
        "nextEvent": None,
        "deadlines": None,
        "reviewQueue": None,
        "streak": None,
        "todayGoal": None,
    }

    # ── nextEvent ────────────────────────────────────────────────
    # First preference: Google Calendar (real schedule). Falls back
    # to the next-due ManualTask if Calendar isn't linked.
    try:
        if GCAL_AVAILABLE:
            token = get_google_token()
            if token:
                events = get_upcoming_events(token) or []
                for ev in events:
                    start_iso = _iso_utc(ev.get("start"))
                    if not start_iso:
                        continue
                    end_iso = _iso_utc(ev.get("end")) or start_iso
                    payload["nextEvent"] = {
                        "title": ev.get("title") or "Upcoming",
                        "kind": _classify_event_kind(ev.get("title"), ev.get("description")),
                        "startsAt": start_iso,
                        "endsAt": end_iso,
                        "deepLink": f"{base}/calendar",
                    }
                    break
    except Exception as e:
        print(f"[snapshot] calendar lookup failed: {e}")

    if payload["nextEvent"] is None:
        try:
            soonest = (
                ManualTask.query
                .filter(ManualTask.user_id == user.id, ManualTask.done == False)  # noqa: E712
                .filter(ManualTask.due_date >= today_str)
                .order_by(ManualTask.due_date.asc())
                .first()
            )
            if soonest and soonest.due_date:
                start_iso = _iso_utc(soonest.due_date)
                if start_iso:
                    payload["nextEvent"] = {
                        "title": soonest.title,
                        "kind": "deadline",
                        "startsAt": start_iso,
                        "endsAt": start_iso,
                        "deepLink": f"{base}/dashboard",
                    }
        except Exception as e:
            print(f"[snapshot] manual-task next-event lookup failed: {e}")

    # ── deadlines (up to 3, soonest first) ───────────────────────
    try:
        deadlines = (
            ManualTask.query
            .filter(ManualTask.user_id == user.id, ManualTask.done == False)  # noqa: E712
            .filter(ManualTask.due_date >= today_str)
            .order_by(ManualTask.due_date.asc())
            .limit(3)
            .all()
        )
        if deadlines:
            payload["deadlines"] = [
                {
                    "title": t.title,
                    "course": t.course or "Personal",
                    "dueAt": _iso_utc(t.due_date),
                    "deepLink": f"{base}/dashboard",
                }
                for t in deadlines if t.due_date
            ] or None
    except Exception as e:
        print(f"[snapshot] deadlines lookup failed: {e}")

    # ── reviewQueue (spaced-repetition items due today) ──────────
    try:
        due_count = (
            StudyMastery.query
            .filter(StudyMastery.user_id == user.id)
            .filter(StudyMastery.next_review <= today_str)
            .filter(StudyMastery.mastery_level < 3)
            .count()
        )
        payload["reviewQueue"] = {
            "dueToday": int(due_count or 0),
            "deepLink": f"{base}/study",
        }
    except Exception as e:
        print(f"[snapshot] review-queue lookup failed: {e}")

    # ── streak ───────────────────────────────────────────────────
    try:
        p = get_study_profile(user.id, None)
        if p:
            streak_days = int(p.streak_count or 0)
            # at-risk = no qualifying activity today *and* user already has a streak going
            at_risk = (p.last_active_date != today_str) and streak_days > 0
            payload["streak"] = {
                "days": streak_days,
                "atRisk": bool(at_risk),
            }
    except Exception as e:
        print(f"[snapshot] streak lookup failed: {e}")

    # ── todayGoal ────────────────────────────────────────────────
    # Target: 4 focus sessions/day (the same yardstick Quests use).
    # Completed: distinct completed StudySession rows for *today*.
    try:
        day_start = datetime(now_utc.year, now_utc.month, now_utc.day, tzinfo=timezone.utc)
        completed_today = (
            StudySession.query
            .filter(StudySession.user_id == user.id)
            .filter(StudySession.completed == True)  # noqa: E712
            .filter(StudySession.created_at >= day_start.replace(tzinfo=None))
            .count()
        )
        payload["todayGoal"] = {
            "target": 4,
            "completed": int(completed_today or 0),
            "unit": "focus sessions",
        }
    except Exception as e:
        print(f"[snapshot] today-goal lookup failed: {e}")

    return payload


@app.route("/api/snapshot", methods=["GET"])
def api_snapshot():
    """Return a portable JSON snapshot of the current user's IntelliPlan
    state, suitable for export to external tools (e.g. Lotus).

    The base template fetches this on every authenticated page load and
    writes the result to `localStorage['intelliplan:snapshot']`, which any
    embedded surface can then read synchronously without an extra round
    trip. Anonymous callers get a small public stub so the client code
    can run without branching.
    """
    snapshot = {
        "version": 2,
        "source": "intelliplan",
        "captured_at": datetime.now(timezone.utc).isoformat(),
        "authenticated": bool(current_user.is_authenticated),
        "user": None,
        "linked_account": None,
        "assignments": [],
        # Lotus payload — top-level fields the Lotus iframe consumes
        "nextEvent": None,
        "deadlines": None,
        "reviewQueue": None,
        "streak": None,
        "todayGoal": None,
    }

    if not current_user.is_authenticated:
        return flask.jsonify(snapshot)

    u = current_user
    snapshot["user"] = {
        "id": u.id,
        "email": getattr(u, "email", None),
        "name": getattr(u, "name", None) or getattr(u, "display_name", None),
        "created_at": u.created_at.isoformat() if getattr(u, "created_at", None) else None,
    }

    try:
        acct = LinkedAccount.query.filter_by(user_id=u.id, is_active=True).first()
        if acct:
            snapshot["linked_account"] = {
                "provider": acct.provider,
                "display_name": getattr(acct, "display_name", None),
                "district_url": getattr(acct, "district_url", None),
                "linked_at": acct.created_at.isoformat() if getattr(acct, "created_at", None) else None,
            }
    except Exception as e:
        print(f"[snapshot] linked account lookup failed: {e}")

    # Cached assignments if they were stashed on the session by a recent
    # dashboard fetch. We don't re-hit the LMS here — that would be slow
    # on every page load. The dashboard endpoint keeps this fresh.
    try:
        cached = session.get("snapshot_assignments")
        if isinstance(cached, list):
            snapshot["assignments"] = cached[:200]
    except Exception:
        pass

    # Merge the Lotus payload (nextEvent / deadlines / reviewQueue / streak / todayGoal)
    try:
        lotus = _build_lotus_snapshot(u)
        snapshot.update(lotus)
    except Exception as e:
        print(f"[snapshot] lotus payload build failed: {e}")

    return flask.jsonify(snapshot)


def _send_push_to_user(user_id, payload, ttl=None):
    """Send a webpush payload to every subscription tied to user_id.
    Returns the count of successful deliveries."""
    subs = PushSubscription.query.filter_by(user_id=user_id).all()
    if not subs:
        return 0

    # The phone app registers an Expo token in the same table — one row per
    # device install, which is what these rows already mean. Splitting them
    # out here rather than at every call site means every existing reminder
    # reaches the phone without its sender knowing phones exist.
    from intelliplan.notifications import expo_push

    expo_subs = [s for s in subs if expo_push.is_expo_token(s.endpoint)]
    subs = [s for s in subs if not expo_push.is_expo_token(s.endpoint)]

    ok = 0
    if expo_subs:
        result = expo_push.send([s.endpoint for s in expo_subs], payload)
        ok += result.get("sent", 0)
        # A token Expo calls DeviceNotRegistered will never work again —
        # the app was uninstalled. Same treatment as a 410 below.
        for dead in result.get("invalid") or []:
            for s in expo_subs:
                if s.endpoint == dead:
                    try:
                        db.session.delete(s)
                        db.session.commit()
                    except Exception:
                        db.session.rollback()

    if not subs:
        return ok
    try:
        from pywebpush import webpush, WebPushException
    except ImportError:
        return ok
    private_key = os.getenv("VAPID_PRIVATE_KEY")
    if not private_key:
        # No VAPID configured is a browser-push problem. Anything already
        # delivered to a phone still counts — returning 0 here would report
        # a successful send as a failure.
        return ok
    claims = {"sub": f"mailto:{os.getenv('VAPID_EMAIL', 'hello@intelliplan.tech')}"}
    ttl_seconds = PUSH_DEFAULT_TTL if ttl is None else max(PUSH_MIN_TTL, int(ttl))
    # Deliberately not reset: `ok` already carries the Expo deliveries.
    for sub in list(subs):
        try:
            webpush(
                subscription_info=json.loads(sub.subscription_json),
                data=json.dumps(payload),
                vapid_private_key=private_key,
                vapid_claims=dict(claims),
                ttl=ttl_seconds,
            )
            ok += 1
        except Exception as e:
            msg = str(e)
            # 410 Gone / 404 mean the browser unsubscribed — drop the row.
            if "410" in msg or "404" in msg:
                try: db.session.delete(sub); db.session.commit()
                except Exception: pass
            else:
                print(f"[push] webpush failed: {e}")
    return ok


def _upcoming_tasks_for(user, lead_minutes):
    """Return ManualTasks due within the next `lead_minutes`, not yet done."""
    now = utcnow()
    cutoff = now + timedelta(minutes=int(lead_minutes or 60))
    rows = ManualTask.query.filter_by(user_id=user.id, done=False).all()
    upcoming = []
    for row in rows:
        if not row.due_date:
            continue
        try:
            due = datetime.fromisoformat(str(row.due_date)[:19])
        except (TypeError, ValueError):
            try: due = datetime.fromisoformat(str(row.due_date)[:10])
            except (TypeError, ValueError): continue
        if now <= due <= cutoff:
            upcoming.append((row, due))
    upcoming.sort(key=lambda p: p[1])
    return upcoming


def _send_reminders_for_user(user, mark_sent=True, force=False):
    """Find a user's upcoming tasks and send SMS + push for each, skipping
    any task already reminded.  Returns a dict counting what went out."""
    lead = int(getattr(user, "reminder_lead_minutes", 60) or 60)
    upcoming = _upcoming_tasks_for(user, lead)
    sent = {"sms": 0, "push": 0, "skipped": 0, "tasks": 0}
    if not upcoming:
        return sent
    want_sms  = bool(getattr(user, "sms_reminders_opt_in", False) and user.phone)
    want_push = bool(getattr(user, "push_reminders_opt_in", False))
    # Load this user's already-sent dedupe rows once. Previously this was one
    # query per (task, channel), so the reminder sweep cost
    # users x tasks x channels round-trips.
    already_sent = set()
    if not force:
        keys = [f"manual:{t.id}" for t, _ in upcoming]
        rows = ReminderSent.query.filter(
            ReminderSent.user_id == user.id, ReminderSent.task_key.in_(keys)
        ).all()
        already_sent = {(r.task_key, r.channel) for r in rows}
    for task, due in upcoming:
        sent["tasks"] += 1
        key = f"manual:{task.id}"
        for channel, want in (("sms", want_sms), ("push", want_push)):
            if not want:
                continue
            if not force and (key, channel) in already_sent:
                sent["skipped"] += 1
                continue
            mins = max(1, int((due - utcnow()).total_seconds() // 60))
            body = f"⏰ {task.title} is due in {mins} min ({task.course})."
            ok = False
            if channel == "sms":
                ok, _ = _sms_send_for_user(user, body[:300])
            else:
                ok = _send_push_to_user(user.id, {
                    "title": "Assignment due soon",
                    "body": body,
                    "url": "/dashboard",
                }) > 0
            if ok:
                sent[channel] += 1
                if mark_sent:
                    try:
                        db.session.add(ReminderSent(user_id=user.id, task_key=key, channel=channel))
                        db.session.commit()
                    except Exception:
                        db.session.rollback()
    return sent


@app.route("/api/reminders/check", methods=["GET"])
def api_reminders_check():
    """Read-only check used by the in-app reminder ticker.
    Returns tasks due within the user's lead-time window so the
    browser can show a native Notification — no cron required."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "ok", "tasks": []})
    lead = int(getattr(current_user, "reminder_lead_minutes", 60) or 60)
    try:
        upcoming = _upcoming_tasks_for(current_user, lead)
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500
    out = []
    now = utcnow()
    for task, due in upcoming:
        mins = max(1, int((due - now).total_seconds() // 60))
        out.append({
            "id":         task.id,
            "key":        f"manual:{task.id}",
            "title":      task.title or "(untitled task)",
            "course":     task.course or "",
            "due_iso":    due.isoformat() + "Z",
            "minutes_until_due": mins,
        })
    return flask.jsonify({
        "status": "ok",
        "tasks": out,
        "lead_minutes": lead,
        "checked_at": now.isoformat() + "Z",
    })


@app.route("/api/reminders/preview", methods=["POST"])
def api_reminders_preview():
    """User-facing test: ignore the dedupe table and fire any pending
    reminders for the current user right now, so the user can verify
    SMS/push delivery without waiting for the cron to run."""
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    result = _send_reminders_for_user(current_user, mark_sent=False, force=True)
    msg = (
        f"Found {result['tasks']} upcoming task(s). "
        f"Sent {result['sms']} SMS, {result['push']} push notification(s)."
        if result["tasks"] else
        "Nothing due within your reminder window yet."
    )
    return flask.jsonify({"status": "ok", "result": result, "message": msg})


@app.route("/cron/send-reminders", methods=["GET", "POST"])
def cron_send_reminders():
    """Hit by Railway cron / external scheduler. Requires CRON_SECRET in the
    X-Cron-Secret header (preferred) or a `secret` query param to prevent abuse.
    Uses hmac.compare_digest to avoid timing attacks on the secret."""
    expected = os.getenv("CRON_SECRET", "")
    provided = request.headers.get("X-Cron-Secret") or request.args.get("secret") or ""
    if not expected:
        return flask.jsonify({"status": "error", "message": "cron not configured"}), 503
    import hmac as _hmac
    if not _hmac.compare_digest(str(expected), str(provided)):
        return flask.jsonify({"status": "error", "message": "unauthorized"}), 401
    total = {"users": 0, "sms": 0, "push": 0, "tasks": 0}
    candidates = User.query.filter(
        (User.sms_reminders_opt_in.is_(True)) | (User.push_reminders_opt_in.is_(True))
    ).all()
    for u in candidates:
        r = _send_reminders_for_user(u, mark_sent=True, force=False)
        if r["tasks"]:
            total["users"] += 1
            total["sms"]  += r["sms"]
            total["push"] += r["push"]
            total["tasks"] += r["tasks"]
    return flask.jsonify({"status": "ok", "summary": total})


# ── LIFECYCLE EMAIL ───────────────────────────────────────────────
# Welcome, feedback request, newsletter. The gate, the templates and the
# deduplication ledger all live in intelliplan/email/ — these are the HTTP
# edges only.


@app.route("/cron/lifecycle-emails", methods=["GET", "POST"])
def cron_lifecycle_emails():
    """Daily sweep for the welcome and feedback emails.

    Same CRON_SECRET + hmac.compare_digest guard as /cron/send-reminders
    above; deliberately a copy of that pattern rather than a new scheme.
    The newsletter is *not* here — it is admin-triggered only.
    """
    expected = os.getenv("CRON_SECRET", "")
    provided = request.headers.get("X-Cron-Secret") or request.args.get("secret") or ""
    if not expected:
        return flask.jsonify({"status": "error", "message": "cron not configured"}), 503
    import hmac as _hmac
    if not _hmac.compare_digest(str(expected), str(provided)):
        return flask.jsonify({"status": "error", "message": "unauthorized"}), 401

    from intelliplan.email import campaigns

    summary = {}
    for name, sweep in (("welcome", campaigns.sweep_welcome), ("feedback", campaigns.sweep_feedback)):
        try:
            summary[name] = sweep()
        except Exception as exc:
            # One sweep failing must not stop the other. A cron that returns
            # a 500 tells the scheduler to retry the whole thing, which for
            # the sweep that already ran means re-walking work it finished.
            # The full traceback goes to the log; the response gets a
            # sanitised message. Raw exception text has leaked connection
            # strings before — see test_error_sanitising.
            app.logger.exception("lifecycle sweep %s failed: %s", name, exc)
            summary[name] = {"error": safe_error_message(exc)}
    return flask.jsonify({"status": "ok", **summary})


@app.route("/cron/weekly-newsletter", methods=["GET", "POST"])
def cron_weekly_newsletter():
    """Generate and send the weekly newsletter. Unattended, by design.

    Same CRON_SECRET guard as the sweeps above. Point a weekly cron at it.

    This one sends to the full marketing list with no human reviewing the
    content first — a deliberate operator decision. What still protects it:
    every recipient passes the marketing gate, the ISO-week email key makes
    a repeat fire a no-op, MARKETING_POSTAL_ADDRESS is still required, and
    the changelog section is built from an allow-list of commit types and
    scopes so an internal commit message cannot reach a student.

    Pass ?dry_run=1 to see the issue and the recipient count without
    sending.
    """
    expected = os.getenv("CRON_SECRET", "")
    provided = request.headers.get("X-Cron-Secret") or request.args.get("secret") or ""
    if not expected:
        return flask.jsonify({"status": "error", "message": "cron not configured"}), 503
    import hmac as _hmac
    if not _hmac.compare_digest(str(expected), str(provided)):
        return flask.jsonify({"status": "error", "message": "unauthorized"}), 401

    dry_run = request.args.get("dry_run") in {"1", "true", "yes"}

    from intelliplan.email import campaigns

    try:
        summary = campaigns.send_weekly_newsletter(dry_run=dry_run)
    except Exception as exc:
        app.logger.exception("weekly newsletter failed: %s", exc)
        return flask.jsonify({"status": "error", "message": safe_error_message(exc)}), 500
    return flask.jsonify({"status": "ok", "summary": summary})


@app.route("/api/admin/newsletter/weekly-preview", methods=["GET", "POST"])
@require_admin
def admin_weekly_newsletter_preview():
    """What this week's automatic issue will contain, and who would get it.

    Generation is deterministic for a given ISO week, so what this returns
    is what the cron will actually send.
    """
    from intelliplan.email import campaigns

    try:
        issue = campaigns.generate_weekly_issue()
        summary = campaigns.send_weekly_newsletter(dry_run=True)
        return flask.jsonify({"status": "ok", "issue": issue, "summary": summary})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


@app.route("/api/admin/email/preflight", methods=["GET", "POST"])
@require_admin
def admin_email_preflight():
    """Can the email system actually send right now, and will replies land?

    Read-only. Worth running before any blast and after any env change —
    every failure it reports is otherwise invisible until a student does
    not get an email.
    """
    from intelliplan.email import preflight

    try:
        result = preflight.check()
        return flask.jsonify({"status": "ok", **result})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


@app.route("/api/admin/feedback-blast/preview", methods=["POST", "GET"])
@require_admin
def admin_feedback_blast_preview():
    """Dry run for the feedback ask to the whole eligible list. Sends nothing."""
    from intelliplan.email import campaigns

    try:
        return flask.jsonify(
            {"status": "ok", "summary": campaigns.send_feedback_now(dry_run=True)}
        )
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


@app.route("/api/admin/feedback-blast/send", methods=["POST"])
@require_admin
def admin_feedback_blast_send():
    """Send the feedback ask to everyone who passes the marketing gate.

    Ignores the 14-day window and the activity requirement — those are
    quality heuristics. Consent, age and suppression still apply in full.
    Requires "confirm": true, so the irreversible call is never the default.
    """
    body = request.get_json(silent=True) or {}
    from intelliplan.email import campaigns

    if not body.get("confirm"):
        try:
            summary = campaigns.send_feedback_now(dry_run=True)
        except Exception as e:
            return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500
        return flask.jsonify({
            "status": "confirm_required",
            "message": (
                f"This would email {summary['recipients']} people who passed the "
                "consent gate. Re-send with \"confirm\": true to go ahead."
            ),
            "summary": summary,
        }), 409

    try:
        summary = campaigns.send_feedback_now(dry_run=False)
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500

    if summary.get("error"):
        return flask.jsonify({"status": "error", "message": summary["error"], "summary": summary}), 400

    print(f"[admin-feedback-blast] recipients={summary['recipients']} sent={summary['sent']} failed={summary['failed']}")
    return flask.jsonify({"status": "ok", "summary": summary})


@app.route("/email/unsubscribe/<token>", methods=["GET", "POST"])
def email_unsubscribe(token):
    """One-click unsubscribe. No login, no confirmation form.

    Deliberately works logged-out: someone who cannot get into their account
    must still be able to get off the list, and CAN-SPAM does not accept
    "sign in first" as an unsubscribe mechanism. POST is accepted because
    List-Unsubscribe-Post one-click sends one.
    """
    from intelliplan.email.eligibility import suppress
    from intelliplan.email.sender import parse_unsubscribe_token

    payload = parse_unsubscribe_token(token or "")
    if not payload:
        return _mini_page(
            "Link not recognised",
            "<h1>We couldn't read that link</h1>"
            "<p>The unsubscribe link looks incomplete or altered. Forward the "
            'email to <a href="mailto:support@intelliplan.tech">support@intelliplan.tech</a> '
            "and we'll take you off the list by hand.</p>",
        ), 400

    address = (payload.get("email") or "").strip().lower()
    suppress(address, reason="unsubscribe")

    # Also clear the flag on any account with this address, so the settings
    # page agrees with reality. The suppression list is the authority; this
    # keeps the UI honest.
    try:
        for user in User.query.filter(db.func.lower(User.email) == address).all():
            user.marketing_emails_opt_in = False
        db.session.commit()
    except Exception:
        db.session.rollback()

    # One-click clients want a bare 200, not a page.
    if request.method == "POST":
        return flask.jsonify({"status": "ok", "unsubscribed": True})

    # markupsafe, not flask.escape — the latter was removed in Flask 3.
    from markupsafe import escape as _escape

    return _mini_page(
        "Unsubscribed",
        "<h1>You're unsubscribed</h1>"
        f"<p>We won't send marketing email to <strong>{_escape(address)}</strong> again.</p>"
        "<p>Deadline reminders are separate and are unaffected — if you had those "
        "switched on, they'll keep arriving. You can change them any time in "
        '<a href="/settings">Settings</a>.</p>',
    )


# ── DAILY CHECK-IN CHEST (Duolingo-style) ─────────────────────────
@app.route("/api/streak/daily-claim", methods=["GET"])
def api_daily_claim_status():
    """Read-only check used by the dashboard on load. Returns whether the
    chest has been claimed today so we can render it blacked out instead
    of waiting for the user to click and discover it themselves."""
    if not is_logged_in():
        return flask.jsonify({"status": "ok", "claimed": False, "logged_in": False})
    try:
        p = get_study_profile(
            user_id=current_user.id if current_user.is_authenticated else None,
            guest_id=None if current_user.is_authenticated else get_guest_session_id(),
        )
        today = utcnow().date().isoformat()
        last_claim = (p.last_daily_claim or "") if hasattr(p, "last_daily_claim") else ""
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500
    streak = int(getattr(p, "streak_count", 0) or 0)
    reward = 10 + min(40, streak * 2)
    return flask.jsonify({
        "status": "ok",
        "claimed": last_claim == today,
        "logged_in": True,
        "reward_if_claimed_now": reward,
        "last_claim": last_claim,
    })


@app.route("/api/streak/daily-claim", methods=["POST"])
def api_daily_claim():
    """One-shot daily Sparks claim. Returns the reward + whether it was
    already claimed today. Front-end uses this for the chest animation
    on the dashboard."""
    if not is_logged_in():
        return flask.jsonify({"status": "error"}), 401
    try:
        p = get_study_profile(
            user_id=current_user.id if current_user.is_authenticated else None,
            guest_id=None if current_user.is_authenticated else get_guest_session_id(),
        )
        today = utcnow().date().isoformat()
        last_claim = (p.last_daily_claim or "") if hasattr(p, "last_daily_claim") else ""
        # last_daily_claim might not exist yet — try/except handles that.
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500
    already = (last_claim == today)
    if already:
        return flask.jsonify({"status": "ok", "claimed": False, "reward": 0,
                              "message": "Already claimed today. Come back tomorrow!"})
    # Reward scales with streak length to reward consistency.
    streak = int(getattr(p, "streak_count", 0) or 0)
    reward = 10 + min(40, streak * 2)
    try:
        p.spark_balance = int(getattr(p, "spark_balance", 0) or 0) + reward
        if hasattr(p, "last_daily_claim"):
            p.last_daily_claim = today
        db.session.commit()
    except Exception:
        try: db.session.rollback()
        except Exception: pass
    return flask.jsonify({"status": "ok", "claimed": True, "reward": reward,
                          "spark_balance": getattr(p, "spark_balance", 0)})


# ── LIVE STUDY SESSIONS (audio/video + materials) ─────────────────
def _jitsi_embed_url(room_url, audio_only=False):
    """Same minimal embed pattern as study groups — prejoin lets the first
    joiner start the room without the meet.jit.si moderator cold-start."""
    video_muted = "true" if audio_only else "false"
    return (
        f"{room_url}#config.prejoinPageEnabled=true"
        f"&config.startWithAudioMuted=true"
        f"&config.startWithVideoMuted={video_muted}"
    )


def _live_session_to_dict(s):
    room_url = f"https://meet.jit.si/intelliplan-{s.room_slug}"
    return {
        "id": s.id,
        "title": s.title,
        "topic": s.topic or "",
        "room_url": room_url,
        "embed_url": _jitsi_embed_url(room_url, bool(s.audio_only)),
        "room_slug": s.room_slug,
        "audio_only": bool(s.audio_only),
        "video_enabled": bool(s.video_enabled),
        "audio_enabled": bool(s.audio_enabled),
        "materials": s.materials or "",
        "is_open": bool(s.is_open),
        "owner_id": s.owner_id,
        "created_at": s.created_at.isoformat() if s.created_at else None,
        "is_owner": current_user.is_authenticated and s.owner_id == current_user.id,
        "invite_url": (APP_BASE_URL.rstrip("/") if APP_BASE_URL else "") + f"/live/{s.id}",
    }


@app.route("/live/<int:session_id>")
def live_session_page(session_id):
    """Public-by-link landing for a live session. Anyone with the link
    can join. Used for sharing invites among IntelliPlan accounts."""
    if not feature_enabled("live_sessions"):
        return render_template("error.html", error_code=404, error_id="LIVE-DISABLED"), 404
    s = LiveSession.query.get(session_id)
    if not s:
        return render_template("error.html", error_code=404, error_id="LIVE-NOT-FOUND",
                               message="That study room no longer exists."), 404
    return render_template("live_session.html", active_page="study", session=_live_session_to_dict(s))


@app.route("/api/live", methods=["POST"])
def api_create_live_session():
    if not feature_enabled("live_sessions"):
        return flask.jsonify({"status": "error", "message": "feature disabled"}), 503
    body = request.get_json(silent=True) or {}
    title = (body.get("title") or "Study session").strip()[:160]
    topic = (body.get("topic") or "").strip()[:160]
    audio_only = bool(body.get("audio_only"))
    slug = secrets_module.token_urlsafe(8).replace("-", "").replace("_", "")[:12]
    owner_id = current_user.id if current_user.is_authenticated else None
    s = LiveSession(
        owner_id=owner_id, title=title, topic=topic,
        audio_only=audio_only, video_enabled=not audio_only, audio_enabled=True,
        room_slug=slug, materials=(body.get("materials") or "")[:8000],
    )
    db.session.add(s); db.session.commit()
    return flask.jsonify({"status": "ok", "session": _live_session_to_dict(s)})


@app.route("/api/live/<int:session_id>", methods=["GET"])
def api_get_live_session(session_id):
    s = LiveSession.query.get(session_id)
    if not s:
        return flask.jsonify({"status": "error"}), 404
    return flask.jsonify({"status": "ok", "session": _live_session_to_dict(s)})


@app.route("/api/live/<int:session_id>", methods=["PATCH"])
def api_update_live_session(session_id):
    s = LiveSession.query.get(session_id)
    if not s:
        return flask.jsonify({"status": "error"}), 404
    if s.owner_id is not None:
        if not current_user.is_authenticated or s.owner_id != current_user.id:
            return flask.jsonify({"status": "error", "message": "not the owner"}), 403
    elif not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "message": "login required to edit"}), 401
    body = request.get_json(silent=True) or {}
    if "materials" in body:        s.materials = str(body["materials"])[:8000]
    if "title" in body:            s.title = str(body["title"]).strip()[:160] or s.title
    if "topic" in body:            s.topic = str(body["topic"]).strip()[:160]
    if "audio_only" in body:       s.audio_only = bool(body["audio_only"])
    if "video_enabled" in body:    s.video_enabled = bool(body["video_enabled"])
    if "audio_enabled" in body:    s.audio_enabled = bool(body["audio_enabled"])
    if "is_open" in body:          s.is_open = bool(body["is_open"])
    db.session.commit()
    return flask.jsonify({"status": "ok", "session": _live_session_to_dict(s)})


@app.route("/api/live", methods=["GET"])
def api_list_live_sessions():
    """Owner's live rooms (so they can rejoin)."""
    if not current_user.is_authenticated:
        return flask.jsonify({"sessions": []})
    rows = LiveSession.query.filter_by(owner_id=current_user.id, is_open=True) \
        .order_by(LiveSession.created_at.desc()).limit(20).all()
    return flask.jsonify({"sessions": [_live_session_to_dict(s) for s in rows]})


def _grant_referral_bonus(new_user):
    """Record the referral link between inviter and new signup."""
    ref_id = session.pop("pending_referral", None) if session else None
    if not ref_id:
        return
    inviter = User.query.get(ref_id)
    if not inviter or inviter.id == new_user.id:
        return
    new_user.referred_by_id = inviter.id
    db.session.commit()


def _apply_pending_group_join():
    """Run after login_user to consume a pending /groups/invite/<id> click."""
    try:
        gid = session.pop("pending_group_join", None) if session else None
        if not gid:
            return None
        existing = StudyGroupMember.query.filter_by(group_id=gid, user_id=current_user.id).first()
        if not existing:
            db.session.add(StudyGroupMember(group_id=gid, user_id=current_user.id, role="member"))
            db.session.commit()
        return gid
    except Exception:
        return None





def _owner_filter(model):
    """Build a SQLAlchemy filter that scopes a model query to the current
    user or guest session."""
    if current_user.is_authenticated:
        return model.user_id == current_user.id
    gid = get_guest_session_id()
    return model.guest_session_id == gid


# ── GRADE MODELLER UPGRADES ─────────────────────────────────
# Realistic feasibility + Simulate + Reset are implemented in JS for
# instant feedback; this endpoint is the optional server-side check that
# returns the highest-achievable grade when the user asks for one that
# math forbids.

def _gradebook_compute(courses, simulated=None):
    """Given courses [{name, weight_total, categories:[{name,weight,assignments:[...]}]}],
    return overall course grades and the weighted final."""
    simulated = simulated or {}
    out_courses = []
    overall_total = 0.0
    overall_weight = 0.0
    for course in courses or []:
        cw = float(course.get("weight") or 0)
        cat_grade_total = 0.0
        cat_weight_total = 0.0
        for cat in course.get("categories") or []:
            w = float(cat.get("weight") or 0)
            earned = 0.0
            possible = 0.0
            for a in cat.get("assignments") or []:
                key = str(a.get("id") or a.get("title") or "")
                sim = simulated.get(key)
                if sim is not None:
                    pts = float(sim.get("points_earned", 0))
                    poss = float(sim.get("points_possible", a.get("points_possible") or 0))
                    earned += pts
                    possible += poss
                elif a.get("graded"):
                    earned += float(a.get("points_earned") or 0)
                    possible += float(a.get("points_possible") or 0)
            if possible > 0 and w > 0:
                cat_grade_total += (earned / possible) * w
                cat_weight_total += w
        course_pct = (cat_grade_total / cat_weight_total) * 100 if cat_weight_total > 0 else None
        out_courses.append({"name": course.get("name"), "grade": course_pct, "weight": cw})
        if course_pct is not None and cw > 0:
            overall_total += course_pct * cw
            overall_weight += cw
    overall = (overall_total / overall_weight) if overall_weight > 0 else None
    return {"courses": out_courses, "overall": overall}


@app.route("/api/grademodel/feasibility", methods=["POST"])
@limiter.limit("30 per minute")
def grademodel_feasibility():
    """Decide whether a target grade is mathematically reachable.

    Body:
      {
        "courses": [...],            # full gradebook data (same shape as gradebook page)
        "target": 92,                # desired course or overall percent
        "course_name": "AP Calc BC", # optional — limit to one course
        "max_score_pct": 100         # cap on per-assignment performance (e.g. 100% for perfect)
      }
    """
    body = request.get_json(silent=True) or {}
    courses = body.get("courses") or []
    target = float(body.get("target") or 100)
    course_name = body.get("course_name")
    max_pct = float(body.get("max_score_pct") or 100)

    sims = {}
    for course in courses:
        if course_name and course.get("name") != course_name:
            continue
        for cat in course.get("categories") or []:
            for a in cat.get("assignments") or []:
                if a.get("graded"):
                    continue
                key = str(a.get("id") or a.get("title") or "")
                poss = float(a.get("points_possible") or 0)
                sims[key] = {
                    "points_earned": poss * (max_pct / 100.0),
                    "points_possible": poss,
                    "auto": True,
                }
    best = _gradebook_compute(courses, simulated=sims)
    best_value = None
    if course_name:
        match = next((c for c in best["courses"] if c["name"] == course_name), None)
        best_value = match["grade"] if match else None
    else:
        best_value = best["overall"]
    reachable = best_value is not None and best_value + 1e-6 >= target
    return flask.jsonify({
        "status": "ok",
        "reachable": bool(reachable),
        "best_possible": best_value,
        "target": target,
        "simulated_assignments": sims,
    })


@app.route("/api/grademodel/simulate", methods=["POST"])
@limiter.limit("60 per minute")
def grademodel_simulate():
    """Apply user-supplied simulated scores to UPCOMING assignments only.

    Body: {courses, simulated: {assignment_id_or_title: {points_earned, points_possible}}}
    Existing graded assignments are never modified — the helper just ignores
    sim entries for already-graded items.
    """
    body = request.get_json(silent=True) or {}
    courses = body.get("courses") or []
    simulated_in = body.get("simulated") or {}
    safe_sim = {}
    for course in courses:
        for cat in course.get("categories") or []:
            for a in cat.get("assignments") or []:
                if a.get("graded"):
                    continue
                key = str(a.get("id") or a.get("title") or "")
                if key in simulated_in:
                    s = simulated_in[key]
                    safe_sim[key] = {
                        "points_earned": float(s.get("points_earned", 0) or 0),
                        "points_possible": float(s.get("points_possible") or a.get("points_possible") or 0),
                    }
    result = _gradebook_compute(courses, simulated=safe_sim)
    return flask.jsonify({"status": "ok", "result": result, "applied": safe_sim})


# ── AI GRADE PREDICTIONS ────────────────────────────────────
# Predicts where each course grade is heading based on historical
# performance. The NUMBERS are always computed deterministically from
# the student's real graded work (category averages + recent trend);
# the AI layer only adds narrative insight and study recommendations
# on top, so a flaky model can never invent a grade.

def _grade_prediction_stats(course):
    """Compute prediction stats for one course from its graded history.

    Returns None when the course has no usable graded data.
    """
    assignments = course.get("assignments") or []
    categories = course.get("categories") or []

    def _f(v):
        try:
            n = float(v)
            return n if n == n else None  # filter NaN
        except (TypeError, ValueError):
            return None

    graded = []
    for a in assignments:
        earned = _f(a.get("points_earned"))
        possible = _f(a.get("points_possible"))
        label = str(a.get("display_score") or "").strip().lower()
        pending_labels = {"", "not graded", "not due", "missing", "pending", "-"}
        is_graded = a.get("graded") is True or (
            earned is not None and label not in pending_labels
        )
        if is_graded and earned is not None and possible and possible > 0:
            graded.append({
                "pct": earned / possible,
                "earned": earned,
                "possible": possible,
                "category": str(a.get("category") or "").strip().lower(),
                "due_date": a.get("due_date") or "",
            })
    if not graded:
        return None

    # Per-category averages from real graded work.
    cat_totals = {}
    for g in graded:
        t = cat_totals.setdefault(g["category"], {"earned": 0.0, "possible": 0.0})
        t["earned"] += g["earned"]
        t["possible"] += g["possible"]
    cat_avg = {
        k: (t["earned"] / t["possible"]) for k, t in cat_totals.items() if t["possible"] > 0
    }
    overall_avg = (
        sum(t["earned"] for t in cat_totals.values())
        / max(sum(t["possible"] for t in cat_totals.values()), 1e-9)
    )

    # Recent trend: date-sorted scores, recent third vs earlier portion.
    dated = sorted((g for g in graded if g["due_date"]), key=lambda g: g["due_date"])
    series = dated if len(dated) >= 4 else graded
    trend_delta = 0.0
    if len(series) >= 4:
        split = max(len(series) // 3, 2)
        recent = series[-split:]
        earlier = series[:-split]
        recent_avg = sum(g["pct"] for g in recent) / len(recent)
        earlier_avg = sum(g["pct"] for g in earlier) / len(earlier)
        trend_delta = (recent_avg - earlier_avg) * 100  # percentage points

    # Score variance → confidence.
    pcts = [g["pct"] * 100 for g in graded]
    mean_pct = sum(pcts) / len(pcts)
    variance = sum((p - mean_pct) ** 2 for p in pcts) / len(pcts)
    stddev = variance ** 0.5
    if len(graded) >= 10 and stddev < 10:
        confidence = "high"
    elif len(graded) >= 5:
        confidence = "medium"
    else:
        confidence = "low"

    # Project the final grade: pending work scored at the category's
    # historical average (clamped to 100%), weight-normalised. Same math
    # as the in-page Realistic Forecast, plus a trend adjustment.
    def project(bias_pts=0.0):
        weighted_sum, weight_used = 0.0, 0.0
        for cat in categories:
            ctype = str(cat.get("type") or "").strip().lower()
            w = _f(cat.get("weight")) or 0
            if w <= 0:
                continue
            earned = cat_totals.get(ctype, {}).get("earned", 0.0)
            possible = cat_totals.get(ctype, {}).get("possible", 0.0)
            pend_possible = 0.0
            for a in assignments:
                if str(a.get("category") or "").strip().lower() != ctype:
                    continue
                a_possible = _f(a.get("points_possible"))
                a_earned = _f(a.get("points_earned"))
                label = str(a.get("display_score") or "").strip().lower()
                pending_labels = {"", "not graded", "not due", "missing", "pending", "-"}
                is_graded = a.get("graded") is True or (
                    a_earned is not None and label not in pending_labels
                )
                if not is_graded and a_possible and a_possible > 0:
                    pend_possible += a_possible
            base_pct = cat_avg.get(ctype, overall_avg)
            pend_pct = min(max(base_pct + bias_pts / 100.0, 0.0), 1.0)
            total = possible + pend_possible
            if total <= 0:
                continue
            cat_final = (earned + pend_possible * pend_pct) / total
            weighted_sum += cat_final * w
            weight_used += w
        if weight_used <= 0:
            return None
        return round((weighted_sum / weight_used) * 100, 2)

    current = _f(course.get("percentage"))
    predicted = project()
    if predicted is None:
        predicted = round(overall_avg * 100, 2)
    # Trend-adjusted: assume the recent trajectory partially continues.
    trend_adjusted = project(bias_pts=max(min(trend_delta * 0.5, 8), -8)) or predicted
    band = max(2.0, round(stddev / 2, 1))

    return {
        "course": course.get("course") or course.get("name") or "Course",
        "current": current,
        "predicted": predicted,
        "trend_adjusted": trend_adjusted,
        "band_low": round(max(min(predicted, trend_adjusted) - band, 0), 1),
        "band_high": round(min(max(predicted, trend_adjusted) + band, 100), 1),
        "trend_delta": round(trend_delta, 1),
        "trend": "improving" if trend_delta > 1.5 else "declining" if trend_delta < -1.5 else "steady",
        "confidence": confidence,
        "graded_count": len(graded),
        "stddev": round(stddev, 1),
        "category_averages": {k: round(v * 100, 1) for k, v in cat_avg.items()},
    }


@app.route("/api/grades/predict", methods=["POST"])
@limiter.limit("10 per minute")
def api_grades_predict():
    if not is_logged_in():
        return flask.jsonify({"status": "error", "message": "login required"}), 401

    body = request.get_json(silent=True) or {}
    courses = body.get("courses") or []
    if not isinstance(courses, list) or not courses:
        return flask.jsonify({"status": "error", "message": "no gradebook data"}), 400

    predictions = []
    for course in courses[:12]:
        stats = _grade_prediction_stats(course)
        if stats:
            predictions.append(stats)
    if not predictions:
        return flask.jsonify({
            "status": "error",
            "message": "No graded work found yet — predictions need at least one graded assignment.",
        }), 422

    # AI layer: narrative insight per course + overall recommendations.
    # Numbers come from the deterministic stats above; the model is told
    # to explain, not to compute.
    insights = {}
    summary = ""
    recommendations = []
    if ai_available():
        compact = [
            {
                "course": p["course"],
                "current_pct": p["current"],
                "predicted_pct": p["predicted"],
                "trend": p["trend"],
                "trend_delta_pts": p["trend_delta"],
                "confidence": p["confidence"],
                "graded_assignments": p["graded_count"],
                "score_stddev": p["stddev"],
                "category_averages": p["category_averages"],
            }
            for p in predictions
        ]
        prompt = (
            "You are an academic coach inside IntelliPlan, a student study planner. "
            "Below are grade predictions computed from a student's real graded work. "
            "Do NOT change or invent any numbers. For each course write one specific, "
            "encouraging-but-honest sentence explaining the prediction (mention the "
            "weakest category or the trend when relevant). Then write a 1-2 sentence "
            "overall summary and up to 3 concrete study recommendations targeting the "
            "courses or categories with the most room to improve.\n\n"
            f"DATA:\n{json.dumps(compact, indent=1)}\n\n"
            'Reply with JSON only: {"courses": [{"course": "<name>", "insight": "<sentence>"}], '
            '"summary": "<text>", "recommendations": ["<tip>", ...]}'
        )
        try:
            ai = ai_chat_json(
                [{"role": "user", "content": prompt}],
                tier="standard", temperature=0.4, max_tokens=1200,
            )
            for row in ai.get("courses") or []:
                name = str(row.get("course") or "").strip()
                text = str(row.get("insight") or "").strip()
                if name and text:
                    insights[name.lower()] = text[:400]
            summary = str(ai.get("summary") or "").strip()[:600]
            recommendations = [
                str(r).strip()[:300] for r in (ai.get("recommendations") or [])[:3] if str(r).strip()
            ]
        except Exception as e:
            print(f"[grade-predict] AI insight layer failed, using fallback: {e}")

    for p in predictions:
        fallback = (
            f"Your scores are {p['trend']}"
            + (f" ({p['trend_delta']:+.1f} pts recently)" if p["trend"] != "steady" else "")
            + f" — projected to land near {p['predicted']}% if you keep performing at your averages."
        )
        p["insight"] = insights.get(p["course"].lower(), fallback)

    return flask.jsonify({
        "status": "ok",
        "predictions": predictions,
        "summary": summary,
        "recommendations": recommendations,
        "ai_enhanced": bool(insights),
    })


# ── LESSON RECORDER ─────────────────────────────────────────
LESSON_UPLOAD_FOLDER = os.path.join(app.root_path, "uploads", "lessons")
os.makedirs(LESSON_UPLOAD_FOLDER, exist_ok=True)
LESSON_ALLOWED_EXT = {".mp3", ".m4a", ".wav", ".ogg", ".mp4", ".webm", ".mov", ".mkv"}
LESSON_AUDIO_EXT = {".mp3", ".m4a", ".wav", ".ogg"}


def _transcribe_lesson(lesson):
    """Real audio/video → text transcription via Gemini (Groq Whisper fallback)."""
    if not lesson.stored_filename:
        return ""
    path = os.path.join(LESSON_UPLOAD_FOLDER, lesson.stored_filename)
    if not os.path.exists(path):
        return ""
    try:
        with open(path, "rb") as fh:
            text = transcribe_audio(lesson.original_filename or lesson.stored_filename, fh.read())
        return (text or "").strip()
    except Exception as e:
        print(f"[lesson {lesson.id}] transcription failed: {e}")
        return ""


def _summarize_lesson_async(lesson_id):
    """Two-stage pipeline:

    1. If we don't already have a transcript, transcribe the audio and persist.
    2. Ask the chat model for a clear, detailed study-style summary.
    """
    lesson = Lesson.query.get(lesson_id)
    if not lesson:
        return
    transcript = (lesson.transcript or "").strip()
    if not transcript:
        transcript = _transcribe_lesson(lesson)
        if transcript:
            lesson.transcript = transcript
            db.session.commit()
    seed = transcript
    if not seed:
        seed = (
            f"Title: {lesson.title}\nCourse: {lesson.course or 'general'}\n"
            f"Tags: {', '.join(lesson.tag_list())}\n"
            "No transcript was available — produce a study-friendly placeholder "
            "summary describing what a lesson with this title and tags likely "
            "covers, and list 5 likely sub-topics."
        )
    try:
        summary = ai_chat(
            [
                {"role": "system", "content": "You write simple, clear, detailed study summaries for student-uploaded lesson recordings. Output plain markdown with sections: TL;DR (1 sentence), Key Points (bullets), Examples / Vocabulary (bullets), Suggested Practice (bullets)."},
                {"role": "user", "content": seed[:14000]},
            ],
            tier="standard",
            temperature=0.4,
            max_tokens=900,
        )
        lesson.summary = summary
        lesson.summary_status = "ready"
    except Exception as e:
        lesson.summary_status = "failed"
        lesson.summary = f"Summary generation failed: {e}"
    db.session.commit()


@app.route("/lessons")
def lessons_page():
    if not feature_enabled("lessons"):
        return render_template("error.html", error_code=503, error_id="LESSONS-DISABLED",
                               message="Lesson Library is temporarily disabled."), 503
    return render_template("lessons.html", active_page="lessons")


@app.route("/api/lessons", methods=["GET"])
def api_list_lessons():
    """Paginated lesson list, newest first.

    Was unbounded: a user with a term's worth of uploads pulled every row
    (transcripts included) on each page load. `lessons` is kept alongside
    `items` so existing template code keeps working.
    """
    q = Lesson.query.filter(_owner_filter(Lesson)).order_by(Lesson.created_at.desc())
    payload = paginate_query(q, lambda l: l.to_dict(), default_size=20)
    payload["lessons"] = payload["items"]
    return flask.jsonify(payload)




@app.route("/api/lessons", methods=["POST"])
@limiter.limit("12 per minute")
def api_upload_lesson():
    upload = request.files.get("file")
    title = (request.form.get("title") or "").strip()[:255]
    course = (request.form.get("course") or "").strip()[:128]
    tags_raw = (request.form.get("tags") or "").strip()
    transcript = (request.form.get("transcript") or "").strip()
    if not title:
        return flask.jsonify({"status": "error", "message": "title required"}), 400
    if not upload or not upload.filename:
        return flask.jsonify({"status": "error", "message": "file required"}), 400
    ext = os.path.splitext(upload.filename)[1].lower()
    if ext not in LESSON_ALLOWED_EXT:
        return flask.jsonify({"status": "error", "message": f"unsupported type {ext}"}), 400
    safe_root = secure_filename(os.path.splitext(upload.filename)[0])[:120] or "lesson"
    stored = f"{uuid.uuid4().hex}_{safe_root}{ext}"
    path = os.path.join(LESSON_UPLOAD_FOLDER, stored)
    upload.save(path)
    tags = [t.strip() for t in tags_raw.split(",") if t.strip()][:12]
    lesson = Lesson(
        user_id=current_user.id if current_user.is_authenticated else None,
        guest_session_id=None if current_user.is_authenticated else get_guest_session_id(),
        title=title,
        course=course,
        tags=json.dumps(tags),
        media_kind="audio" if ext in LESSON_AUDIO_EXT else "video",
        original_filename=upload.filename[:255],
        stored_filename=stored,
        mime_type=upload.mimetype or "",
        transcript=transcript[:60000],
        summary_status="pending",
    )
    db.session.add(lesson)
    db.session.commit()
    # Kick the summary in-process (best effort; runs synchronously but
    # cheaply because the prompt is small). For long videos a background
    # worker would be the next upgrade.
    _summarize_lesson_async(lesson.id)
    return flask.jsonify({"status": "ok", "lesson": lesson.to_dict()})


@app.route("/lessons/<int:lesson_id>/stream")
def lesson_stream(lesson_id):
    lesson = Lesson.query.filter(Lesson.id == lesson_id, _owner_filter(Lesson)).first()
    if not lesson or not lesson.stored_filename:
        return flask.jsonify({"status": "error", "message": "not found"}), 404
    return send_from_directory(LESSON_UPLOAD_FOLDER, lesson.stored_filename, as_attachment=False)


@app.route("/api/lessons/<int:lesson_id>", methods=["DELETE"])
def api_delete_lesson(lesson_id):
    lesson = Lesson.query.filter(Lesson.id == lesson_id, _owner_filter(Lesson)).first()
    if not lesson:
        return flask.jsonify({"status": "error"}), 404
    try:
        if lesson.stored_filename:
            os.remove(os.path.join(LESSON_UPLOAD_FOLDER, lesson.stored_filename))
    except Exception:
        pass
    db.session.delete(lesson)
    db.session.commit()
    return flask.jsonify({"status": "ok"})


@app.route("/api/lessons/<int:lesson_id>/resummarize", methods=["POST"])
@limiter.limit("6 per minute")
def api_resummarize_lesson(lesson_id):
    lesson = Lesson.query.filter(Lesson.id == lesson_id, _owner_filter(Lesson)).first()
    if not lesson:
        return flask.jsonify({"status": "error"}), 404
    body = request.get_json(silent=True) or {}
    if body.get("transcript"):
        lesson.transcript = str(body.get("transcript"))[:60000]
        db.session.commit()
    lesson.summary_status = "pending"
    db.session.commit()
    _summarize_lesson_async(lesson.id)
    return flask.jsonify({"status": "ok", "lesson": Lesson.query.get(lesson_id).to_dict()})


# ── STUDY GROUPS ────────────────────────────────────────────
def _group_match_score(group, prefs):
    """Tiny matcher: + for matching topic substring, level, style."""
    score = 0
    topic = (prefs.get("topic") or "").lower().strip()
    if topic and topic in (group.topic or "").lower():
        score += 3
    if prefs.get("level") and (prefs["level"] == group.level or group.level == "any"):
        score += 1
    if prefs.get("style") and (prefs["style"] == group.style or group.style == "any"):
        score += 1
    return score


@app.route("/groups")
def groups_page():
    if not feature_enabled("groups"):
        return render_template("error.html", error_code=503, error_id="GROUPS-DISABLED",
                               message="Study Groups is temporarily disabled."), 503
    return render_template("groups.html", active_page="groups")


@app.route("/api/groups", methods=["GET"])
def api_list_groups():
    """List public groups + groups the user belongs to.

    Paged rather than a flat top-50: the old ceiling meant group 51 onward
    was unreachable from the UI entirely, not merely slow to reach.
    """
    _page, _per_page = _page_args(24)
    public_q = (StudyGroup.query
                .filter_by(visibility="public")
                .order_by(StudyGroup.created_at.desc())
                .limit(_per_page + 1)
                .offset((_page - 1) * _per_page)
                .all())
    _has_more = len(public_q) > _per_page
    public_q = public_q[:_per_page]
    mine_ids = set()
    if current_user.is_authenticated:
        mine_ids = {m.group_id for m in StudyGroupMember.query.filter_by(user_id=current_user.id).all()}

    # One grouped COUNT for the whole page instead of one per group — this
    # was 50 extra round-trips on a full listing.
    counts = {}
    if public_q:
        rows = (
            db.session.query(StudyGroupMember.group_id, db.func.count(StudyGroupMember.id))
            .filter(StudyGroupMember.group_id.in_([g.id for g in public_q]))
            .group_by(StudyGroupMember.group_id)
            .all()
        )
        counts = {gid: n for gid, n in rows}

    def _ser(g):
        return {
            "id": g.id,
            "name": g.name,
            "topic": g.topic,
            "level": g.level,
            "style": g.style,
            "visibility": g.visibility,
            "description": g.description or "",
            "meeting_url": g.meeting_url or "",
            "next_meeting_at": g.next_meeting_at.isoformat() if g.next_meeting_at else None,
            "next_meeting_topic": g.next_meeting_topic or "",
            "member_count": counts.get(g.id, 0),
            "is_member": g.id in mine_ids,
        }

    prefs = {
        "topic": request.args.get("topic", ""),
        "level": request.args.get("level", ""),
        "style": request.args.get("style", ""),
    }
    # Match scoring reorders within the fetched page only — a global ranking
    # would need the whole table, which is exactly what pagination avoids.
    scored = [(_group_match_score(g, prefs), g) for g in public_q]
    scored.sort(key=lambda x: (-x[0], -(x[1].id or 0)))
    items = [_ser(g) for _, g in scored]
    return flask.jsonify({
        "status": "ok",
        "groups": items,
        "items": items,
        "my_group_ids": list(mine_ids),
        "page": _page,
        "per_page": _per_page,
        "has_more": _has_more,
    })


@app.route("/api/groups", methods=["POST"])
def api_create_group():
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "message": "login required"}), 401
    body = request.get_json(silent=True) or {}
    name = (body.get("name") or "").strip()[:120]
    if not name:
        return flask.jsonify({"status": "error", "message": "name required"}), 400
    room_slug = secrets_module.token_urlsafe(8).replace("-", "").replace("_", "")
    group = StudyGroup(
        name=name,
        topic=(body.get("topic") or "").strip()[:120],
        level=(body.get("level") or "any")[:32],
        style=(body.get("style") or "any")[:32],
        visibility=(body.get("visibility") or "public")[:16],
        description=(body.get("description") or "").strip()[:2000],
        owner_id=current_user.id,
        meeting_url=f"https://meet.jit.si/intelliplan-{room_slug}",
    )
    db.session.add(group)
    db.session.commit()
    db.session.add(StudyGroupMember(group_id=group.id, user_id=current_user.id, role="owner"))
    db.session.commit()
    # Async-ish AI plan suggestion (synchronous + cheap).
    try:
        prompt = (
            f"You are an academic coach. Suggest a 4-week study plan for a small group studying \"{group.topic or group.name}\".\n"
            f"Level: {group.level}. Study style: {group.style}.\n"
            "Output plain markdown: Week 1 / Week 2 / Week 3 / Week 4 with 3 bullets each."
        )
        group.suggested_plan = ai_chat(
            [{"role": "user", "content": prompt}],
            tier="standard",
            temperature=0.5,
            max_tokens=700,
        )
        db.session.commit()
    except Exception:
        pass
    return flask.jsonify({"status": "ok", "id": group.id})


@app.route("/api/groups/<int:group_id>", methods=["GET"])
def api_get_group(group_id):
    g = StudyGroup.query.get(group_id)
    if not g:
        return flask.jsonify({"status": "error"}), 404
    rows = StudyGroupMember.query.filter_by(group_id=group_id).all()
    # Resolve every member's display name in one IN query rather than one
    # User.get() per member.
    users = {}
    if rows:
        users = {
            u.id: u
            for u in User.query.filter(User.id.in_([m.user_id for m in rows])).all()
        }
    members = []
    for m in rows:
        u = users.get(m.user_id)
        members.append({
            "user_id": m.user_id,
            "role": m.role,
            "name": (u.name if u else None) or (u.email if u else "Member"),
        })
    is_member = current_user.is_authenticated and any(m["user_id"] == current_user.id for m in members)
    invite_url = (APP_BASE_URL.rstrip("/") + url_for("groups_invite", group_id=g.id)) if APP_BASE_URL else url_for("groups_invite", group_id=g.id)
    return flask.jsonify({
        "id": g.id,
        "name": g.name,
        "topic": g.topic,
        "level": g.level,
        "style": g.style,
        "description": g.description or "",
        "shared_notes": g.shared_notes or "",
        "meeting_url": g.meeting_url or "",
        "next_meeting_at": g.next_meeting_at.isoformat() if g.next_meeting_at else None,
        "next_meeting_topic": g.next_meeting_topic or "",
        "suggested_plan": g.suggested_plan or "",
        "members": members,
        "is_member": is_member,
        "is_owner": current_user.is_authenticated and g.owner_id == current_user.id,
        "invite_url": invite_url,
    })


@app.route("/api/groups/<int:group_id>/join", methods=["POST"])
def api_join_group(group_id):
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "message": "login required"}), 401
    g = StudyGroup.query.get(group_id)
    if not g:
        return flask.jsonify({"status": "error"}), 404
    existing = StudyGroupMember.query.filter_by(group_id=group_id, user_id=current_user.id).first()
    if not existing:
        db.session.add(StudyGroupMember(group_id=group_id, user_id=current_user.id, role="member"))
        db.session.commit()
    return flask.jsonify({"status": "ok"})


@app.route("/groups/invite/<int:group_id>")
def groups_invite(group_id):
    """Shareable invite link. Anyone with an IntelliPlan account who hits
    this URL is auto-joined and dropped into the group room. Logged-out
    visitors are bounced to /login and joined right after sign-in via the
    `pending_group_join` session flag."""
    g = StudyGroup.query.get(group_id)
    if not g:
        return render_template("error.html", active_page="error", error_code=404,
                               error_id="GROUP-NOT-FOUND",
                               message="That study group invite link no longer works."), 404
    if not current_user.is_authenticated:
        session["pending_group_join"] = group_id
        session.modified = True
        return redirect(url_for("login"))
    existing = StudyGroupMember.query.filter_by(group_id=group_id, user_id=current_user.id).first()
    if not existing:
        db.session.add(StudyGroupMember(group_id=group_id, user_id=current_user.id, role="member"))
        db.session.commit()
    return redirect(url_for("groups_page") + f"?open={group_id}")


@app.route("/api/groups/<int:group_id>/leave", methods=["POST"])
def api_leave_group(group_id):
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error"}), 401
    StudyGroupMember.query.filter_by(group_id=group_id, user_id=current_user.id).delete()
    db.session.commit()
    return flask.jsonify({"status": "ok"})


@app.route("/api/groups/<int:group_id>/notes", methods=["POST"])
def api_save_group_notes(group_id):
    g = StudyGroup.query.get(group_id)
    if not g:
        return flask.jsonify({"status": "error"}), 404
    if not current_user.is_authenticated or not StudyGroupMember.query.filter_by(group_id=group_id, user_id=current_user.id).first():
        return flask.jsonify({"status": "error", "message": "not a member"}), 403
    body = request.get_json(silent=True) or {}
    g.shared_notes = (body.get("notes") or "")[:60000]
    db.session.commit()
    return flask.jsonify({"status": "ok"})


@app.route("/api/groups/<int:group_id>/start-meeting", methods=["POST"])
def api_start_group_meeting(group_id):
    if not current_user.is_authenticated:
        return flask.jsonify({"status": "error", "message": "login required"}), 401
    g = StudyGroup.query.get(group_id)
    if not g:
        return flask.jsonify({"status": "error", "message": "not found"}), 404
    if not StudyGroupMember.query.filter_by(group_id=group_id, user_id=current_user.id).first():
        return flask.jsonify({"status": "error", "message": "not a member"}), 403
    if g.meeting_url:
        return flask.jsonify({"status": "ok", "meeting_url": g.meeting_url})
    room_slug = secrets_module.token_urlsafe(8).replace("-", "").replace("_", "")
    g.meeting_url = f"https://meet.jit.si/intelliplan-{room_slug}"
    db.session.commit()
    return flask.jsonify({"status": "ok", "meeting_url": g.meeting_url})


@app.route("/api/groups/<int:group_id>/meeting", methods=["POST"])
def api_set_group_meeting(group_id):
    g = StudyGroup.query.get(group_id)
    if not g:
        return flask.jsonify({"status": "error"}), 404
    if not current_user.is_authenticated or not StudyGroupMember.query.filter_by(group_id=group_id, user_id=current_user.id).first():
        return flask.jsonify({"status": "error"}), 403
    body = request.get_json(silent=True) or {}
    when = (body.get("when") or "").strip()
    topic = (body.get("topic") or "").strip()[:255]
    try:
        g.next_meeting_at = datetime.fromisoformat(when) if when else None
    except Exception:
        g.next_meeting_at = None
    g.next_meeting_topic = topic
    db.session.commit()
    return flask.jsonify({"status": "ok"})


# ── WRITING IMPROVEMENT ASSISTANT ───────────────────────────
@app.route("/writing")
def writing_page():
    if not feature_enabled("writing"):
        return render_template("error.html", error_code=503, error_id="WRITING-DISABLED",
                               message="Writing Assistant is temporarily disabled."), 503
    return render_template("writing.html", active_page="writing")


@app.route("/api/writing/analyze", methods=["POST"])
@limiter.limit("20 per minute")
def api_writing_analyze():
    body = request.get_json(silent=True) or {}
    text = (body.get("text") or "").strip()
    tone = (body.get("tone") or "neutral")[:32]
    purpose = (body.get("purpose") or "essay")[:32]
    if not text:
        return flask.jsonify({"status": "error", "message": "text required"}), 400
    if len(text) > 12000:
        text = text[:12000]
    system = (
        "You are an expert writing coach. Review the user's writing and "
        "respond with STRICT JSON only (no preamble). Schema:\n"
        "{\n"
        "  \"overall\": {\"score\": int 0-100, \"summary\": string},\n"
        "  \"suggestions\": [{\"category\": one of [\"grammar\",\"clarity\",\"tone\",\"structure\",\"argument\"],\n"
        "      \"excerpt\": string (≤120 chars from the original), \"suggestion\": string, \"why\": string}],\n"
        "  \"revised\": string (a clean revised draft)\n"
        "}\n"
        "Return at most 8 suggestions, prioritised by impact. Be specific."
    )
    try:
        data = ai_chat_json(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": f"Purpose: {purpose}\nTarget tone: {tone}\n\nText:\n{text}"},
            ],
            tier="standard",
            temperature=0.3,
            max_tokens=1400,
        )
        return flask.jsonify({"status": "ok", **data})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


# ── MATH EXPLAINER ──────────────────────────────────────────
@app.route("/math")
def math_page():
    if not feature_enabled("math"):
        return render_template("error.html", error_code=503, error_id="MATH-DISABLED",
                               message="Math Explainer is temporarily disabled."), 503
    return render_template("math.html", active_page="math")


@app.route("/api/math/explain", methods=["POST"])
@limiter.limit("30 per minute")
def api_math_explain():
    body = request.get_json(silent=True) or {}
    problem = (body.get("problem") or "").strip()
    level = (body.get("level") or "high_school")[:32]
    if not problem:
        return flask.jsonify({"status": "error", "message": "problem required"}), 400
    system = (
        "You are a patient, rigorous math tutor. Respond with STRICT JSON only (no preamble).\n"
        "Schema:\n"
        "{\n"
        "  \"problem\": string,\n"
        "  \"steps\": [{\"step\": int, \"explanation\": string, \"math\": string}],\n"
        "  \"answer\": string,\n"
        "  \"notes\": string (common pitfalls, intuition)\n"
        "}\n"
        f"Adjust depth for level={level}. Keep each step short and self-contained. Use plain text for math (e.g. x^2 not LaTeX)."
    )
    try:
        data = ai_chat_json(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": problem[:8000]},
            ],
            tier="standard",
            temperature=0.2,
            max_tokens=1400,
        )
        return flask.jsonify({"status": "ok", **data})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


@app.route("/api/math/similar", methods=["POST"])
@limiter.limit("30 per minute")
def api_math_similar():
    body = request.get_json(silent=True) or {}
    problem = (body.get("problem") or "").strip()
    if not problem:
        return flask.jsonify({"status": "error", "message": "problem required"}), 400
    system = (
        "Generate ONE similar practice problem at the same difficulty. Respond with STRICT JSON only:\n"
        "{\"problem\": string, \"answer\": string}"
    )
    try:
        data = ai_chat_json(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": problem[:6000]},
            ],
            tier="standard",
            temperature=0.7,
            max_tokens=600,
        )
        return flask.jsonify({"status": "ok", **data})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


# ── TASK EXTRACTOR ──────────────────────────────────────────
@app.route("/extractor")
def extractor_page():
    if not feature_enabled("extractor"):
        return render_template("error.html", error_code=503, error_id="EXTRACTOR-DISABLED",
                               message="Task Extractor is temporarily disabled."), 503
    return render_template("extractor.html", active_page="extractor")


@app.route("/api/tasks/extract", methods=["POST"])
@limiter.limit("30 per minute")
def api_task_extract():
    body = request.get_json(silent=True) or {}
    text = (body.get("text") or "").strip()
    if not text:
        return flask.jsonify({"status": "error", "message": "text required"}), 400
    today_iso = utcnow().date().isoformat()
    system = (
        "Extract actionable tasks from user-supplied notes / messages / emails. Respond with STRICT JSON only:\n"
        "{\n"
        "  \"tasks\": [{\"title\": string, \"due_date\": null | YYYY-MM-DD, \"priority\": \"High\"|\"Medium\"|\"Low\", \"notes\": string}]\n"
        "}\n"
        f"Today is {today_iso}. Resolve relative dates (e.g. 'next Friday') to absolute YYYY-MM-DD where possible. "
        "Use null for due_date when no date is implied. Use High for urgent or graded items, Low for nice-to-have."
    )
    try:
        data = ai_chat_json(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": text[:12000]},
            ],
            tier="standard",
            temperature=0.2,
            max_tokens=1200,
        )
        tasks = data.get("tasks") or []
        # Optionally persist as ManualTask when ?save=1
        saved_ids = []
        if request.args.get("save") == "1" and current_user.is_authenticated:
            for t in tasks[:25]:
                mt = ManualTask(
                    user_id=current_user.id,
                    title=(t.get("title") or "")[:255],
                    due_date=t.get("due_date") or None,
                    priority=(t.get("priority") or "Medium")[:16],
                    course=(t.get("course") or "Personal")[:128],
                    estimated_time=int(t.get("estimated_time") or 60),
                    notes=(t.get("notes") or "")[:2000],
                )
                db.session.add(mt)
                db.session.flush()
                saved_ids.append(mt.id)
            db.session.commit()
        return flask.jsonify({"status": "ok", "tasks": tasks, "saved_ids": saved_ids})
    except Exception as e:
        return flask.jsonify({"status": "error", "message": safe_error_message(e)}), 500


@app.errorhandler(429)
def error_429(e):
    if request.path.startswith("/extension/") or request.path.startswith("/api/") or request.is_json:
        return flask.jsonify({"status": "error", "message": "Rate limit exceeded. Please wait a moment before trying again."}), 429
    try:
        return render_template("error.html", active_page="error", error_code=429, error_id=make_error_id()), 429
    except Exception:
        return flask.Response("<h1>429 Too Many Requests</h1><a href='/'>Home</a>", status=429, mimetype="text/html")

@app.errorhandler(500)
def error_500(e):
    err_id = make_error_id()
    print(f"[{err_id}] Internal Server Error: {e}")
    if os.getenv("SENTRY_DSN"):
        try:
            sentry_sdk.capture_exception(e)
        except Exception:
            pass
    if request.path.startswith("/extension/") or request.path.startswith("/api/") or request.is_json:
        return flask.jsonify({"status": "error", "message": "Internal server error. Please try again.", "error_id": err_id}), 500
    try:
        return render_template("error.html", active_page="error", error_code=500, error_id=err_id), 500
    except Exception:
        return flask.Response(f"<h1>500 Server Error</h1><p>Error ID: {err_id}</p><a href='/'>Home</a>", status=500, mimetype="text/html")

@app.errorhandler(503)
def error_503(e):
    if request.path.startswith("/extension/") or request.path.startswith("/api/") or request.is_json:
        return flask.jsonify({"status": "error", "message": "Service temporarily unavailable. Please try again later."}), 503
    try:
        return render_template("error.html", active_page="error", error_code=503, error_id=make_error_id()), 503
    except Exception:
        return flask.Response("<h1>503 Service Unavailable</h1><a href='/'>Home</a>", status=503, mimetype="text/html")

@app.errorhandler(Exception)
def handle_unhandled_exception(e):
    import traceback
    err_id = make_error_id()
    print(f"[{err_id}] Unhandled exception:\n{traceback.format_exc()}")
    if os.getenv("SENTRY_DSN"):
        try:
            sentry_sdk.capture_exception(e)
        except Exception:
            pass
    if request.path.startswith("/extension/") or request.path.startswith("/api/") or request.is_json:
        return flask.jsonify({"status": "error", "message": "An unexpected error occurred. Please try again.", "error_id": err_id}), 500
    try:
        return render_template("error.html", active_page="error", error_code=500, error_id=err_id), 500
    except Exception:
        return flask.Response(f"<h1>Server Error</h1><p>Error ID: {err_id}</p><a href='/'>Home</a>", status=500, mimetype="text/html")


# ── Public REST API for third-party clients and the IntelliPlan MCP server.
# Registered at the very end so all models and helpers are defined.
# Attach key references to the Flask app so the blueprint can resolve them
# via `current_app` (avoids a double-import / double-SQLAlchemy-instance bug
# when running as `python App.py` — module loaded as both `__main__` and
# `App`).
app.intelliplan_db = db
app.intelliplan_bcrypt = bcrypt
app.intelliplan_user_model = User
app.intelliplan_get_identity = _get_or_create_identity
# Expose new models so intelliplan/api/*.py blueprints can resolve them
# without importing App (avoids the __main__/App double-load issue).
app.intelliplan_lms_token_model = LMSToken
app.intelliplan_student_link_model = StudentLink
app.intelliplan_study_group_model = StudyGroup
app.intelliplan_study_group_member_model = StudyGroupMember
app.intelliplan_study_group_task_model = StudyGroupTask
app.intelliplan_voice_seat_model = VoiceSeat
app.intelliplan_manual_preset_model = ManualPlanPreset
app.intelliplan_saved_schedule_model = SavedSchedule
app.intelliplan_api_key_model = ApiKey
# Needed by the v1 push-registration endpoints, which store a phone's
# Expo token as another subscription row rather than inventing a
# second table for the same idea.
app.intelliplan_push_subscription_model = PushSubscription
from intelliplan_api import api_bp as intelliplan_api_bp, api_rate_limit_key, api_rate_limit_value
app.register_blueprint(intelliplan_api_bp)

# Developer portal: apply for a key, revoke, roll, and the admin review queue.
from api_keys import developer_api_bp
app.register_blueprint(developer_api_bp)

# Rate-limit the public API per credential rather than per IP. The limit
# string is resolved per request so an approved high-volume integration gets
# the ceiling it was granted instead of the shared default. Applied to the
# whole blueprint so a new endpoint can't be added without a limit.
limiter.limit(api_rate_limit_value, key_func=api_rate_limit_key)(intelliplan_api_bp)
# Applying for a key is cheap to submit and expensive to review, so it gets
# its own much tighter budget.
limiter.limit("10 per hour")(app.view_functions["developer_api_bp.apply_for_key"])

# ── New feature blueprints (grade prediction, LMS sync, roles, group tasks).
# All are self-contained under intelliplan/ and resolve models via
# the current_app.intelliplan_* references attached above.
from intelliplan.api.grade_prediction import bp as grade_prediction_bp
from intelliplan.api.lms_sync import bp as lms_sync_bp
from intelliplan.api.roles import bp as roles_bp
from intelliplan.api.group_tasks import bp as group_tasks_bp
from intelliplan.api.group_voice import bp as group_voice_bp
from intelliplan.api.manual_schedule import bp as manual_schedule_bp
app.register_blueprint(grade_prediction_bp)
app.register_blueprint(lms_sync_bp)
app.register_blueprint(roles_bp)
app.register_blueprint(group_tasks_bp)
app.register_blueprint(group_voice_bp)
app.register_blueprint(manual_schedule_bp)
# The voice heartbeat runs every 12 seconds per person in a room, so the
# default per-IP budget would throttle a household with two students in
# the same call. Its own limit is sized to the poll, not to page loads.
limiter.limit("30 per minute")(app.view_functions["group_voice_bp.voice_heartbeat"])

# ── AI Daily Command Center (docs/command-center/). Registered last so
# the glue module's lazy `from App import ...` calls always resolve.
# The `command_center` feature flag is a KILL SWITCH (default on).
from command_center_glue import command_center_bp
app.register_blueprint(command_center_bp)
from learning_graph_glue import learning_graph_bp
app.register_blueprint(learning_graph_bp)
# ── Active study. Registered here for the same reason: its glue module
# resolves App lazily. The `active_study` flag is a kill switch (default on).
from active_glue import active_bp
app.register_blueprint(active_bp)
# ── Notifications. Outbox-backed: events are queued by the sweep and
# delivered on a timer, so no student request ever waits on an SMS
# gateway or an SMTP handshake.
from notifications_glue import notifications_bp, start_ticker as _start_notification_ticker
app.register_blueprint(notifications_bp)
limiter.exempt(app.view_functions["notifications.cron_notifications"])
# The sweep/flush cycle needs something to drive it. /cron/notifications
# has always been able to, but nothing ever called it — no cron entry, no
# scheduled job, nothing in the Procfile but the web process — so the
# outbox was never swept and no reminder was ever delivered. The app now
# runs its own timer; the endpoint remains for a real external scheduler.
# Set NOTIFICATIONS_INPROCESS_CRON=0 to hand the job back to one.
_start_notification_ticker(app)
# ── Offline write safety. Installs before/after-request hooks that make any
# mutating endpoint replay-safe when the client sends an X-IP-Op-Id, plus the
# one endpoint the offline queue uses to ask "did these ops land?".
from intelliplan.sync import setup as _setup_sync

SyncOp = _setup_sync(
    app, db,
    lambda: current_user.id if current_user.is_authenticated else 0,
)
# The queue flushes every op it holds in one burst on reconnect, which is a
# legitimate spike the per-IP page-load budget would mistake for abuse.
limiter.limit("120 per minute")(app.view_functions["ip_sync.check_ops"])

# ── Keep the Today plan honest after a write ─────────────────────────
# /api/today caches its payload for 90 seconds. That is right for repeated
# page loads and wrong the moment a student adds an assignment: the plan
# they are then shown was computed before the thing they just typed
# existed, and nothing on screen says so. These are the paths that change
# what the plan is made of.
_PLAN_MUTATING_PATHS = (
    "/tasks/manual/create",
    "/tasks/manual/update",
    "/tasks/manual/delete",
    "/dismiss",
    "/restore",
    "/test/mark",
    "/test/unmark",
    "/api/import/csv",
    "/api/syllabus/import",
)


@app.after_request
def _invalidate_today_cache(response):
    """Evict the cached plan after any successful plan-changing write.

    A hook rather than a call in each handler for the same reason the sync
    ledger is one: the set of writes that affect the plan grows, and the
    failure mode of forgetting one is a stale plan nobody can reproduce.
    """
    if request.method == "POST" and response.status_code < 400:
        path = request.path
        if any(path.startswith(p) for p in _PLAN_MUTATING_PATHS):
            try:
                from intelliplan.api.command_center import invalidate_today

                invalidate_today(current_user.id if current_user.is_authenticated else None)
            except Exception as e:
                print(f"[today] cache invalidation failed: {e}")
    return response
limiter.limit("6 per hour")(app.view_functions["notifications.send_test_notification"])
# Heartbeats arrive roughly every 15 seconds while a session runs, so this
# route is sized to the poll rather than to page loads — the default per-IP
# budget would throttle a single student mid-session.
limiter.limit("60 per minute")(app.view_functions["active.heartbeat"])
limiter.limit("30 per hour")(app.view_functions["active.start_session"])
limiter.limit("30 per minute")(app.view_functions["command_center.api_today"])
limiter.limit("6 per hour")(app.view_functions["command_center.api_today_refresh"])
limiter.exempt(app.view_functions["command_center.cron_refresh_briefings"])
limiter.limit("20 per minute")(app.view_functions["plani_agent.plani_agent"])


def _existing_columns(table_name):
    """Return a set of column names that exist on `table_name`, or an
    empty set if the table doesn't exist. Works on both SQLite and
    Postgres without needing the ORM's metadata to be in sync."""
    from sqlalchemy import inspect as _inspect
    try:
        insp = _inspect(db.engine)
        if table_name not in insp.get_table_names():
            return set()
        return {c["name"] for c in insp.get_columns(table_name)}
    except Exception as e:
        print(f"[migrate] inspect {table_name} failed: {e}")
        return set()


def _migrate_user_columns():
    """Bulletproof ALTER TABLE shim for SQLite/Postgres.

    The previous version used `db.session.execute(...)` per column,
    which on Postgres aborts the WHOLE transaction the moment any one
    of the ALTERs fails (e.g. "column already exists"). Subsequent
    ALTERs in the same transaction then silently no-op and the App
    boots without the new columns — which is the bug that left users
    seeing "Login is briefly unavailable" every time.

    Now we:
      1. Inspect the live schema first and only attempt ALTERs for
         columns that are genuinely missing.
      2. Use a fresh, autocommit-style connection per ALTER so a
         failure on one column can't poison the rest.
      3. Verify after the loop that the new columns landed and log
         the result for production debugging.
    """
    from sqlalchemy import text as _t

    targets = [
        # users — referral / phone / reminder prefs
        ("users", "referral_code", "VARCHAR(16)"),
        ("users", "referred_by_id", "INTEGER"),
        ("users", "phone", "VARCHAR(32)"),
        ("users", "sms_reminders_opt_in", "BOOLEAN DEFAULT FALSE"),
        ("users", "push_reminders_opt_in", "BOOLEAN DEFAULT FALSE"),
        ("users", "reminder_lead_minutes", "INTEGER DEFAULT 60"),
        ("users", "sms_carrier", "VARCHAR(32) DEFAULT 'tmobile'"),
        ("users", "birth_year", "INTEGER"),
        ("users", "parent_email", "VARCHAR(255)"),
        ("users", "parent_consent_granted", "BOOLEAN DEFAULT FALSE"),
        ("users", "parent_consent_token", "VARCHAR(64)"),
        ("users", "lms_preferences", "TEXT DEFAULT '{}'"),
        ("users", "ai_personalization_opt_in", "BOOLEAN DEFAULT FALSE"),
        ("users", "marketing_emails_opt_in", "BOOLEAN DEFAULT FALSE"),
        ("users", "marketing_opt_in_at", "TIMESTAMP"),
        ("users", "role", "VARCHAR(16) DEFAULT 'student'"),
        # users — Active-study focus enforcement
        ("users", "focus_enforcement", "VARCHAR(16) DEFAULT 'off'"),
        ("users", "focus_alarm_file", "VARCHAR(255)"),
        ("users", "focus_grace_seconds", "INTEGER DEFAULT 25"),
        # active_sessions — sparks given up to focus enforcement
        ("active_sessions", "sparks_forfeited", "INTEGER DEFAULT 0"),
        # users — notification preferences. These are listed here as well as
        # in apply_notification_migrations() on purpose: that one runs once,
        # at import, inside an app context. If the database is not reachable
        # at that moment — the gunicorn cold start this whole function exists
        # for — it never runs again, and every User SELECT afterwards fails
        # on the missing column. That is not a broken settings page, it is a
        # 500 on every authenticated request in the app.
        ("users", "email_reminders_opt_in", "BOOLEAN DEFAULT FALSE"),
        ("users", "utc_offset_minutes", "INTEGER DEFAULT 0"),
        ("users", "quiet_hours_enabled", "BOOLEAN DEFAULT TRUE"),
        ("users", "quiet_hours_start", "INTEGER DEFAULT 22"),
        ("users", "quiet_hours_end", "INTEGER DEFAULT 7"),
        ("users", "notification_kinds", "VARCHAR(512)"),
        # manual_tasks provenance for CSV importer / extension scraper
        ("manual_tasks", "import_source", "VARCHAR(32) DEFAULT ''"),
        ("manual_tasks", "import_batch_id", "VARCHAR(64) DEFAULT ''"),
        ("manual_tasks", "external_id", "VARCHAR(128) DEFAULT ''"),
        # saved_schedules — cross-device Interactive View progress
        ("saved_schedules", "progress_json", "TEXT"),
        # user_identities — earlier migration's columns. Without these,
        # _get_or_create_identity() blows up and registration redirects
        # to /onboarding which then 500s.
        ("user_identities", "availability", "TEXT"),
        ("user_identities", "weekly_commitments", "TEXT"),
        ("user_identities", "class_schedule", "TEXT"),
        # notion_integrations
        ("notion_integrations", "auth_type", "VARCHAR(16) DEFAULT 'manual'"),
        ("notion_integrations", "workspace_id", "VARCHAR(64)"),
        ("notion_integrations", "workspace_name", "VARCHAR(256)"),
        ("notion_integrations", "workspace_icon", "VARCHAR(512)"),
        ("notion_integrations", "bot_id", "VARCHAR(64)"),
        ("notion_integrations", "connected_at", "TIMESTAMP"),
        # daily check-in chest tracking
        ("study_points", "last_daily_claim", "VARCHAR(16) DEFAULT ''"),
        # google_integrations — multi-account support
        ("google_integrations", "account_email", "VARCHAR(255)"),
        ("google_integrations", "account_name", "VARCHAR(255)"),
        ("google_integrations", "is_active", "BOOLEAN DEFAULT TRUE"),
        ("google_integrations", "connected_at", "TIMESTAMP"),
        # feature_flags — percentage rollout support
        ("feature_flags", "rollout_percentage", "INTEGER DEFAULT 100"),
        # site_feedback — diagnostics blob attached by the bug-report dialog
        ("site_feedback", "diagnostics", "TEXT DEFAULT ''"),
        # user_streaks — task-completion streak (created by db.create_all,
        # entries here are for columns that may be added after initial deploy)
        ("user_streaks", "nudge_shown_date", "VARCHAR(16) DEFAULT ''"),
        ("user_streaks", "qualified_dates_json", "TEXT DEFAULT '[]'"),
        # plani_pets — virtual creature that grows with site usage
        ("plani_pets", "name", "VARCHAR(40) DEFAULT 'Plani'"),
        ("plani_pets", "xp", "INTEGER DEFAULT 0"),
        ("plani_pets", "last_visit_local_date", "VARCHAR(16) DEFAULT ''"),
        ("plani_pets", "hatched_at", "TIMESTAMP"),
        ("plani_pets", "last_fed_at", "TIMESTAMP"),
        ("plani_pets", "last_played_at", "TIMESTAMP"),
        ("plani_pets", "last_petted_at", "TIMESTAMP"),
        ("plani_pets", "last_studied_at", "TIMESTAMP"),
        ("plani_pets", "last_chest_local_date", "VARCHAR(16) DEFAULT ''"),
        ("plani_pets", "chest_streak_days", "INTEGER DEFAULT 0"),
        ("plani_pets", "perfect_week_paid", "VARCHAR(16) DEFAULT ''"),
    ]

    # Group targets by table so we only inspect each table once.
    by_table = {}
    for t, c, d in targets:
        by_table.setdefault(t, []).append((c, d))

    for table, cols in by_table.items():
        existing = _existing_columns(table)
        if not existing:
            # Table doesn't exist at all — db.create_all() will handle it
            # the next time around. No ALTER needed.
            continue
        for col, decl in cols:
            if col in existing:
                continue
            try:
                # Fresh connection per ALTER, autocommit on, so one failure
                # can't roll back the others.
                with db.engine.connect() as conn:
                    conn.execute(_t(f"ALTER TABLE {table} ADD COLUMN {col} {decl}"))
                    try:
                        conn.commit()
                    except Exception:
                        pass
                print(f"[migrate] added {table}.{col}")
            except Exception as e:
                print(f"[migrate] could not add {table}.{col}: {e}")

    # Verify result so production logs make the state obvious.
    users_cols = _existing_columns("users")
    have_new = {"referral_code", "referred_by_id"}.issubset(users_cols)
    print(f"[migrate] users new columns present: {have_new}  (got: {sorted(users_cols)})")


# Run the schema bootstrap + column migration at IMPORT time so production
# WSGI servers (gunicorn etc.) see the new columns. Previously this was
# gated on __name__=="__main__", which caused every request to surface
# the raw "<h1>Server Error</h1><a href='/'>Home</a>" fallback in prod
# because new columns didn't exist yet and lookups blew up the whole
# Jinja render — including the error page.
_MIGRATION_DONE = False

def _ensure_indexes():
    """Create hot-path indexes on existing tables.

    `db.create_all()` only adds indexes when it creates a table, so an
    already-existing prod table never gets the `index=True` columns we
    added later. These `CREATE INDEX IF NOT EXISTS` statements are
    idempotent and valid on both SQLite and Postgres. They back the
    per-owner lookups that run on every dashboard / command-center load.

    Every statement here backs a query pattern that actually exists in the
    codebase — these are not speculative. Without them each of these lookups
    is a full table scan that grows linearly with total site usage, not with
    the requesting user's own data."""
    from sqlalchemy import text as _t
    statements = [
        "CREATE INDEX IF NOT EXISTS ix_dismissed_user ON dismissed_assignments (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_dismissed_guest ON dismissed_assignments (guest_session_id)",
        "CREATE INDEX IF NOT EXISTS ix_testmarks_user ON test_marks (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_testmarks_guest ON test_marks (guest_session_id)",
        # The single hottest query in the app: "newest active schedule for this
        # owner", run on nearly every scheduler/dashboard/export request.
        # Composite so the filter *and* the ORDER BY ... LIMIT 1 are covered.
        "CREATE INDEX IF NOT EXISTS ix_sched_user_active ON saved_schedules (user_id, is_active, created_at)",
        "CREATE INDEX IF NOT EXISTS ix_sched_guest_active ON saved_schedules (guest_session_id, is_active, created_at)",
        # Task lists — dashboard, scheduler, and the Plani agent all filter
        # by owner and done-state.
        "CREATE INDEX IF NOT EXISTS ix_tasks_user_done ON manual_tasks (user_id, done)",
        "CREATE INDEX IF NOT EXISTS ix_tasks_guest_done ON manual_tasks (guest_session_id, done)",
        # Completion history — read on every schedule generation to build
        # Study DNA, and on every /feedback/predict-time call.
        "CREATE INDEX IF NOT EXISTS ix_feedback_user ON task_feedback (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_feedback_guest ON task_feedback (guest_session_id)",
        # Memories page: owner + date range.
        "CREATE INDEX IF NOT EXISTS ix_archive_user_date ON day_archives (user_id, archive_date)",
        "CREATE INDEX IF NOT EXISTS ix_archive_guest_date ON day_archives (guest_session_id, archive_date)",
        "CREATE INDEX IF NOT EXISTS ix_notes_user ON course_notes (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_notes_guest ON course_notes (guest_session_id)",
        "CREATE INDEX IF NOT EXISTS ix_linked_user ON linked_accounts (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_grades_user ON imported_grades (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_descriptions_user ON custom_descriptions (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_descriptions_guest ON custom_descriptions (guest_session_id)",
        # Spaced-repetition lookups are per-owner and per-question-key.
        "CREATE INDEX IF NOT EXISTS ix_mastery_user ON study_mastery (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_sessions_user ON study_sessions (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_points_user ON study_points (user_id)",
        "CREATE INDEX IF NOT EXISTS ix_pushsub_user ON push_subscriptions (user_id)",
        # The subscribe upsert matches on endpoint; without this it scans
        # every subscription on the site on each notification opt-in.
        "CREATE INDEX IF NOT EXISTS ix_pushsub_endpoint ON push_subscriptions (endpoint)",
        "CREATE INDEX IF NOT EXISTS ix_sessmsg_user ON session_messages (user_id)",
        # Reminder dedupe lookup, run for every user on every cron sweep.
        "CREATE INDEX IF NOT EXISTS ix_reminders_user_key ON reminders_sent (user_id, task_key)",
        "CREATE INDEX IF NOT EXISTS ix_groupmember_group ON study_group_members (group_id)",
        "CREATE INDEX IF NOT EXISTS ix_groupmember_user ON study_group_members (user_id)",
        # Clarification presets: looked up by owner on every schedule request.
        "CREATE INDEX IF NOT EXISTS ix_presets_user_key ON scheduler_presets (user_id, task_key)",
        "CREATE INDEX IF NOT EXISTS ix_presets_guest_key ON scheduler_presets (guest_session_id, task_key)",
        # Client error log: every write is a fingerprint lookup, every admin
        # read is ordered by last_seen. Without these the table degrades into
        # a scan as soon as it holds real traffic.
        "CREATE INDEX IF NOT EXISTS ix_clienterr_fp ON client_error_logs (fingerprint)",
        "CREATE INDEX IF NOT EXISTS ix_clienterr_seen ON client_error_logs (last_seen)",
        # Feedback list is always "mine, newest first".
        "CREATE INDEX IF NOT EXISTS ix_feedback_user_created ON site_feedback (user_id, created_at)",
    ]
    for stmt in statements:
        try:
            db.session.execute(_t(stmt))
        except Exception as _idx_e:
            db.session.rollback()
            print(f"[boot] index create skipped: {_idx_e}")
    db.session.commit()


def _migrate_push_subscription_endpoints():
    """Add push_subscriptions.endpoint and backfill it from the stored JSON.

    Rows written before the column existed have no endpoint, so the upsert
    in /push/subscribe would treat every one of them as a different browser
    and start stacking duplicates. Backfilling is a few hundred rows of
    json.loads at boot, once.
    """
    from sqlalchemy import text as _t

    columns = _existing_columns("push_subscriptions")
    if not columns:
        return
    if "endpoint" not in columns:
        try:
            db.session.execute(_t(
                "ALTER TABLE push_subscriptions ADD COLUMN endpoint VARCHAR(512)"))
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            print(f"[migrate] push endpoint column: {e}")
            return

    try:
        stale = PushSubscription.query.filter(
            (PushSubscription.endpoint.is_(None)) | (PushSubscription.endpoint == "")
        ).all()
        for row in stale:
            try:
                row.endpoint = (json.loads(row.subscription_json) or {}).get("endpoint")
            except Exception:
                # Unparseable JSON was never a usable subscription. Leaving
                # it null keeps it out of the way of the endpoint upsert.
                continue
        if stale:
            db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"[migrate] push endpoint backfill: {e}")


def _run_boot_migration_once():
    """Idempotent migration helper — safe to call repeatedly."""
    global _MIGRATION_DONE
    if _MIGRATION_DONE:
        return
    try:
        db.create_all()
        _migrate_user_columns()
        _migrate_push_subscription_endpoints()
        _ensure_indexes()
        _MIGRATION_DONE = True
    except Exception as _boot_e:
        print(f"[boot] DB bootstrap failed: {_boot_e}")

try:
    with app.app_context():
        _run_boot_migration_once()
except Exception as _boot_e:
    print(f"[boot] App context unavailable at import: {_boot_e}")


@app.route("/health")
def health_check():
    """Diagnostic endpoint. Returns DB-schema state so we can see in
    production whether the migration ran. Public, no secrets exposed."""
    out = {
        "status": "ok",
        "migration_done": bool(_MIGRATION_DONE),
        "users_columns": sorted(_existing_columns("users")),
        "notion_columns": sorted(_existing_columns("notion_integrations")),
    }
    expected = {"id", "email", "password_hash", "referral_code", "referred_by_id"}
    out["users_schema_ok"] = expected.issubset(set(out["users_columns"]))
    return flask.jsonify(out)


@app.before_request
def _ensure_migration_ran():
    """Belt-and-braces: if for any reason the import-time migration
    didn't complete (e.g. the DB wasn't ready yet during gunicorn cold
    start on Railway), retry it on the very first request. This is the
    only way to guarantee the new user.referral_code
    columns exist before any User SELECT hits them."""
    if _MIGRATION_DONE:
        return
    try:
        _run_boot_migration_once()
    except Exception:
        pass


@app.route("/api/syllabus/import", methods=["POST"])
def api_syllabus_import():
    """Extract assignments from an uploaded PDF syllabus using AI.
    Returns a JSON list of {title, due_date, course, description} objects."""
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401

    if "file" not in request.files:
        return jsonify({"status": "error", "message": "No file uploaded"}), 400

    f = request.files["file"]
    if not f.filename or not f.filename.lower().endswith(".pdf"):
        return jsonify({"status": "error", "message": "Only PDF files are supported"}), 400

    raw_text = ""
    try:
        try:
            import pdfplumber
            with pdfplumber.open(f) as pdf:
                raw_text = "\n".join(
                    page.extract_text() or "" for page in pdf.pages[:20]
                )
        except ImportError:
            # Fallback: read raw bytes and try to extract ASCII text
            f.seek(0)
            data = f.read()
            raw_text = data.decode("latin-1", errors="replace")
            raw_text = re.sub(r"[^\x20-\x7E\n\t]", " ", raw_text)
            raw_text = re.sub(r" {3,}", " ", raw_text)
    except Exception as e:
        return jsonify({"status": "error", "message": f"Could not read PDF: {e}"}), 422

    if not raw_text.strip():
        return jsonify({"status": "error", "message": "PDF appears to have no readable text"}), 422

    truncated = raw_text[:6000]
    today_str = utcnow().strftime("%Y-%m-%d")
    prompt = f"""You are a syllabus parser. Today is {today_str}.
Extract every assignment, exam, quiz, project, or deadline from the syllabus text below.
Return ONLY a valid JSON array — no markdown, no explanations — where each item has:
  "title": short assignment name (string),
  "due_date": due date as YYYY-MM-DD if determinable, else null,
  "course": course name/code if visible (string or null),
  "description": one-line description (string or null)

Syllabus text:
---
{truncated}
---
JSON array:"""

    try:
        if not ai_available():
            return jsonify({"status": "error", "message": "AI extraction not configured"}), 503
        parsed = ai_chat_json(
            [{"role": "user", "content": prompt}],
            tier="fast",
            temperature=0.1,
            max_tokens=2000,
        )
        # Model sometimes wraps in {"assignments": [...]}
        if isinstance(parsed, dict):
            for key in ("assignments", "items", "tasks", "deadlines", "data"):
                if isinstance(parsed.get(key), list):
                    parsed = parsed[key]
                    break
        if not isinstance(parsed, list):
            parsed = []
    except Exception as e:
        print(f"[syllabus import] AI parse error: {e}")
        return jsonify({"status": "error", "message": "AI could not parse the syllabus"}), 500

    course_name = (request.form.get("course_name") or "").strip()[:160]
    record_id = None
    try:
        if current_user.is_authenticated or get_guest_session_id():
            rec = SyllabusRecord(
                user_id=current_user.id if current_user.is_authenticated else None,
                guest_session_id=None if current_user.is_authenticated else get_guest_session_id(),
                course_name=course_name,
                filename=f.filename if f else "",
                assignments_json=json.dumps(parsed),
                imported_count=0,
            )
            db.session.add(rec)
            db.session.commit()
            record_id = rec.id
    except Exception as e:
        print(f"[syllabus] save record error: {e}")

    return jsonify({
        "status": "ok",
        "assignments": parsed,
        "count": len(parsed),
        "record_id": record_id,
        "course_name": course_name,
    })


@app.route("/api/syllabus/records", methods=["GET"])
def api_syllabus_records():
    if current_user.is_authenticated:
        rows = SyllabusRecord.query.filter_by(user_id=current_user.id).order_by(SyllabusRecord.created_at.desc()).limit(50).all()
    else:
        gid = get_guest_session_id()
        rows = SyllabusRecord.query.filter_by(guest_session_id=gid).order_by(SyllabusRecord.created_at.desc()).limit(50).all()
    out = []
    for r in rows:
        try:
            items = json.loads(r.assignments_json or "[]")
        except Exception:
            items = []
        out.append({
            "id": r.id,
            "course_name": r.course_name or "",
            "filename": r.filename or "",
            "count": len(items) if isinstance(items, list) else 0,
            "imported_count": r.imported_count or 0,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        })
    return jsonify({"status": "ok", "records": out})


@app.route("/api/syllabus/records/<int:record_id>", methods=["GET"])
def api_syllabus_record_detail(record_id):
    r = SyllabusRecord.query.get(record_id)
    if not r:
        return jsonify({"status": "error", "message": "Not found"}), 404
    if current_user.is_authenticated:
        if r.user_id != current_user.id:
            return jsonify({"status": "error", "message": "Forbidden"}), 403
    elif r.guest_session_id != get_guest_session_id():
        return jsonify({"status": "error", "message": "Forbidden"}), 403
    try:
        assignments = json.loads(r.assignments_json or "[]")
    except Exception:
        assignments = []
    return jsonify({
        "status": "ok",
        "id": r.id,
        "course_name": r.course_name,
        "filename": r.filename,
        "assignments": assignments,
        "imported_count": r.imported_count,
    })


@app.route("/api/settings/lms-sources", methods=["GET", "PATCH"])
def api_lms_sources():
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    if request.method == "GET":
        prefs = get_user_lms_prefs()
        accounts = LinkedAccount.query.filter_by(user_id=current_user.id).all()
        return jsonify({
            "status": "ok",
            "preferences": prefs,
            "linked_accounts": [
                {"login_type": a.login_type, "name": a.name, "is_active": bool(a.is_active)}
                for a in accounts
            ],
            "integrations": {
                "google_classroom": bool(ClassroomIntegration.query.filter_by(user_id=current_user.id).first()),
                "blackboard": bool(BlackboardIntegration.query.filter_by(user_id=current_user.id).first()),
                "moodle": bool(MoodleIntegration.query.filter_by(user_id=current_user.id).first()),
            },
        })
    body = request.get_json(silent=True) or {}
    prefs = get_user_lms_prefs()
    if body.get("grade_source"):
        prefs["grade_source"] = str(body["grade_source"])[:32]
    if isinstance(body.get("assignment_sources"), list):
        prefs["assignment_sources"] = [str(s)[:32] for s in body["assignment_sources"]][:12]
    current_user.lms_preferences = json.dumps(prefs)
    db.session.commit()
    return jsonify({"status": "ok", "preferences": prefs})


@app.route("/api/meetings/saved", methods=["GET", "POST"])
def api_saved_meetings():
    if request.method == "GET":
        if current_user.is_authenticated:
            q = SavedMeeting.query.filter_by(user_id=current_user.id)
        else:
            q = SavedMeeting.query.filter_by(guest_session_id=get_guest_session_id())
        q = q.order_by(SavedMeeting.created_at.desc())

        payload = paginate_query(q, lambda m: {
            "id": m.id, "name": m.name, "url": m.url, "platform": m.platform,
            "schedule_text": m.schedule_text or "", "is_recurring": bool(m.is_recurring),
            "created_at": m.created_at.isoformat() if m.created_at else None,
        }, default_size=25)
        payload["meetings"] = payload["items"]
        return jsonify(payload)
    body = request.get_json(silent=True) or {}
    name = (body.get("name") or "").strip()[:120]
    url = (body.get("url") or "").strip()[:512]
    if not name or not url:
        return jsonify({"status": "error", "message": "name and url required"}), 400
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https"):
        return jsonify({"status": "error", "message": "Invalid URL"}), 400
    m = SavedMeeting(
        user_id=current_user.id if current_user.is_authenticated else None,
        guest_session_id=None if current_user.is_authenticated else get_guest_session_id(),
        name=name,
        url=url,
        platform=(body.get("platform") or "other")[:32],
        schedule_text=(body.get("schedule_text") or body.get("when") or "")[:200],
        is_recurring=bool(body.get("is_recurring")),
    )
    db.session.add(m)
    db.session.commit()
    return jsonify({"status": "ok", "id": m.id})


@app.route("/api/meetings/saved/<int:meeting_id>", methods=["DELETE"])
def api_delete_saved_meeting(meeting_id):
    m = SavedMeeting.query.get(meeting_id)
    if not m:
        return jsonify({"status": "error"}), 404
    if current_user.is_authenticated:
        if m.user_id != current_user.id:
            return jsonify({"status": "error", "message": "Forbidden"}), 403
    elif m.guest_session_id != get_guest_session_id():
        return jsonify({"status": "error", "message": "Forbidden"}), 403
    db.session.delete(m)
    db.session.commit()
    return jsonify({"status": "ok"})


def _session_msg_owner_filter(context_type, context_id):
    q = SessionMessage.query.filter_by(context_type=context_type, context_id=context_id)
    return q.order_by(SessionMessage.created_at.asc()).limit(500)


@app.route("/api/sessions/<context_type>/<int:context_id>/messages", methods=["GET", "POST"])
def api_session_messages(context_type, context_id):
    if context_type not in ("live", "group"):
        return jsonify({"status": "error", "message": "Invalid context"}), 400
    if request.method == "GET":
        rows = _session_msg_owner_filter(context_type, context_id).all()
        return jsonify({"status": "ok", "messages": [{
            "id": m.id,
            "author_name": m.author_name,
            "body": m.body,
            "saved_to_library": bool(m.saved_to_library),
            "created_at": m.created_at.isoformat() if m.created_at else None,
        } for m in rows]})
    body = request.get_json(silent=True) or {}
    text = (body.get("body") or body.get("message") or "").strip()
    if not text:
        return jsonify({"status": "error", "message": "Empty message"}), 400
    if len(text) > 8000:
        text = text[:8000]
    author = (current_user.name or current_user.email or "Student") if current_user.is_authenticated else "Guest"
    msg = SessionMessage(
        context_type=context_type,
        context_id=context_id,
        user_id=current_user.id if current_user.is_authenticated else None,
        author_name=author[:120],
        body=text,
    )
    db.session.add(msg)
    db.session.commit()
    return jsonify({"status": "ok", "message": {
        "id": msg.id,
        "author_name": msg.author_name,
        "body": msg.body,
        "created_at": msg.created_at.isoformat() if msg.created_at else None,
    }})


@app.route("/api/sessions/messages/<int:msg_id>/save", methods=["POST"])
def api_save_session_message(msg_id):
    msg = SessionMessage.query.get(msg_id)
    if not msg:
        return jsonify({"status": "error", "message": "Not found"}), 404
    if not current_user.is_authenticated:
        return jsonify({"status": "error", "message": "login required"}), 401
    msg.saved_to_library = True
    note = CourseNote(
        user_id=current_user.id,
        course_name="Study Session",
        note_date=utcnow().strftime("%Y-%m-%d"),
        title=f"Chat — {msg.author_name}"[:255],
        text_content=msg.body,
    )
    db.session.add(note)
    db.session.commit()
    return jsonify({"status": "ok", "note_id": note.id})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 3000))
    app.run(host="0.0.0.0", port=port)
